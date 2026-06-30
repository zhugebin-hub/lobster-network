#!/usr/bin/env python3
"""生成《论语》今古文经学比较研究小论文"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import date

doc = Document()

# === 页面设置 A4 ===
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# === 全局样式 ===
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = Pt(22.5)  # 1.5倍行距

# === 辅助函数 ===
def set_font(run, name='宋体', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(text, style_name='Normal', align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_line_indent=Cm(0.74), bold=False, space_before=0, space_after=0,
             font_name='宋体', font_size=12):
    p = doc.add_paragraph()
    p.style = doc.styles[style_name]
    p.alignment = align
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, font_name, font_size, bold=bold)
    return p

def add_heading_custom(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12) if level == 1 else Pt(8)
    pf.space_after = Pt(6)
    if level == 1:
        pf.first_line_indent = Cm(0)
        run = p.add_run(text)
        set_font(run, '黑体', 16, bold=True)
    elif level == 2:
        pf.first_line_indent = Cm(0)
        run = p.add_run(text)
        set_font(run, '黑体', 14, bold=True)
    return p

def add_blank_line():
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

# ========================================
# 封面/标题区
# ========================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('论语今文经学与古文经学比较研究')
set_font(run, '黑体', 22, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(20)
run2 = p2.add_run(f'择白\n{date.today().strftime("%Y年%-m月%-d日")}')
set_font(run2, '楷体', 14)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(20)
run3 = p3.add_run('摘要')
set_font(run3, '黑体', 14, bold=True)

# 摘要
add_para(
    '《论语》作为儒家经典的核心文献，在汉代经历了今文经学与古文经学两条不同的传承路径。今文经学以口耳相传、隶书记录为特征，注重微言大义与经世致用；古文经学则以先秦古本、古文字书写为依据，强调训诂考据与历史还原。本文系统梳理《论语》今古文经学的版本源流、学术特征、诠释方法及其历史影响，揭示两派在经典诠释中的根本分歧，并探讨其对后世经学发展乃至当代中国哲学研究的深远意义。研究表明，今古文之争并非简单的文本真伪之辩，而是两种经典诠释范式的根本对立，其影响贯穿整个中国经学史，并在当代学术语境中展现出新的理论价值。',
    font_size=11, first_line_indent=Cm(0.74), space_before=6, space_after=0
)

# 关键词
p_kw = doc.add_paragraph()
p_kw.paragraph_format.space_before = Pt(8)
p_kw.paragraph_format.first_line_indent = Cm(0.74)
kw_bold = p_kw.add_run('关键词：')
set_font(kw_bold, '黑体', 11, bold=True)
kw_text = p_kw.add_run('论语；今文经学；古文经学；经典诠释；汉代经学；郑玄')
set_font(kw_text, '宋体', 11)

doc.add_page_break()

# ========================================
# 正文
# ========================================

add_heading_custom('一、引言')

add_para(
    '秦火之后，儒家经典经历了一次重大的传承危机。随着汉代政治秩序的重建，经典文献的搜集、整理与诠释成为时代的重要课题。在这一历史背景下，《论语》的传承形成了两条截然不同的路径：一条是以口耳相传、隶书记录为特征的今文经学传统，另一条则是以先秦古本、古文字书写为凭据的古文经学传统。'
)

add_para(
    '今古文经学之争是中国经学史上最为重要的学术争论之一。这一争论不仅涉及文本的真伪与版本的优劣，更深层地反映了两种不同的经典诠释范式和学术方法论。就《论语》而言，今古文之争直接体现在《鲁论》《齐论》与《古论》三种版本的并存与融合上。理解这一学术争辩，对于把握《论语》的诠释传统、汉代经学的学术生态，乃至整个中国古典学术的发展脉络，都具有不可替代的意义。'
)

add_para(
    '本文拟从版本源流、学术特征、诠释方法和历史影响四个维度，系统比较《论语》今文经学与古文经学的异同，进而揭示两派之争的深层学术逻辑及其对后世的持久影响。'
)

add_heading_custom('二、《论语》今古文版本的源流与分化')

add_heading_custom('（一）今文《论语》的传承', level=2)

add_para(
    '今文《论语》的形成与秦始皇焚书密切相关。秦火之后，儒家经典遭到严重破坏，许多典籍被焚毁。汉初，经典文献的恢复主要依靠儒生的记忆与口耳相传。《论语》的今文传承主要有两个系统：'
)

add_para(
    '其一为《鲁论》。此本由鲁地儒生传授，共二十篇，是用当时通行的隶书记录下来的。《鲁论》在汉初流传较广，影响深远，后来成为张禹编纂通行本《论语》的基础。'
)

add_para(
    '其二为《齐论》。此本由齐地儒生传授，共二十二篇，比《鲁论》多出《问王》和《知道》两篇。《齐论》在汉代也曾广泛流传，但由于其文本与《鲁论》存在差异，且多出的两篇在后来的传承中逐渐失传，最终未能成为《论语》通行本的组成部分。'
)

add_para(
    '今文《论语》的传授在汉代形成了一个严密的师承系统。从汉初的叔孙通、伏生，到武帝时期的董仲舒、公孙弘，今文经学凭借其官学地位，在政治和教育领域占据了主导。'
)

add_heading_custom('（二）古文《论语》的发现', level=2)

add_para(
    '与今文《论语》的口耳相传不同，古文《论语》的发现具有偶然性和实物性。据《汉书·艺文志》记载，汉武帝末年，鲁恭王坏孔子宅，欲以为宫，于坏壁之中得古文《论语》，称为《古论》。'
)

add_para(
    '《古论》用先秦六国古文字书写，共二十一篇。其篇次与今文有所不同，且文字多有异同。古文《论语》的发现，为经学研究提供了新的文本来源，也引发了今古文两派之间长期的学术争论。'
)

add_para(
    '古文《论语》之所以具有特殊的学术价值，在于它被认为更接近孔子时代的原始文本。由于它是先秦时期的抄本，未经秦火之劫，因此在文本的原始性和可靠性上具有今文所不具备的优势。'
)

add_heading_custom('（三）版本的融合与统一', level=2)

add_para(
    '西汉末年，张禹以《鲁论》为基础，参考《齐论》和《古论》，编纂了一个综合性的《论语》文本，世称"张侯论"。这一文本后来被郑玄用作注释的底本，最终成为《论语》的通行本，流传至今。'
)

add_para(
    '张禹的编纂工作并非简单的文本拼合，而是在比较不同版本异同的基础上，做出了审慎的文本选择。这一过程本身就体现了今古文经学方法的某种融合。'
)

add_heading_custom('三、今文经学与古文经学的学术特征比较')

add_heading_custom('（一）文本来源与可靠性', level=2)

add_para(
    '今文经学与古文经学最根本的分歧在于文本来源的不同。今文经学的文本来源于汉初儒生的记忆与口耳相传，经过隶书转写后形成定本。这种传承方式的优势在于保持了经典在文化记忆中的连续性，但缺陷同样明显：记忆的不准确性、传抄过程中的讹误，以及师承系统中的主观改造，都可能导致文本失真。'
)

add_para(
    '古文经学则以先秦古本为依据，认为这些古本未经秦火之劫，保留了经典的原始面貌。从现代文献学的角度来看，古文经学的立场更接近"文本批判"的方法，强调以最早的、最接近原始状态的文本为依据。'
)

add_para(
    '然而，古文经学也面临自身的困境。一方面，古文经的真伪问题一直存在争议。西汉末年的刘歆大力推崇古文经，但其动机和政治背景引发了学者的质疑。康有为在《新学伪经考》中更是直接将古文经判定为刘歆伪造。另一方面，即使古文经为真，由于年代久远，古文字的释读也存在诸多困难，不同的学者对同一古文字可能有截然不同的理解。'
)

add_heading_custom('（二）诠释方法的差异', level=2)

add_para(
    '今文经学与古文经学在诠释方法上的差异，是两派之争的核心所在。'
)

add_para(
    '今文经学的诠释方法以"微言大义"为核心特征。所谓"微言大义"，即认为经典文本中蕴含着圣人的深层政治意图和道德理想，诠释者的任务是通过解读经典中的"微言"，揭示其中的"大义"。这种方法强调经典的政治实用性和经世致用功能，诠释往往服务于现实的政治需要。董仲舒的"天人感应"说便是今文经学诠释方法的典型体现。'
)

add_para(
    '古文经学则完全不同。古文经学强调"训诂考据"，即通过对文字、音韵、训诂的精细研究，还原经典文本的原始含义。这种方法强调历史还原和学术客观性，认为诠释者的首要任务是准确地理解文本的本义，而非将其服务于现实政治。东汉的郑玄、许慎、马融等人都是古文经学诠释方法的杰出代表。'
)

add_para(
    '这两种诠释方法各有优劣。今文经学的"微言大义"赋予了经典以鲜活的生命力和现实关怀，但也容易导致诠释的主观化和政治化。古文经学的"训诂考据"保证了诠释的客观性和学术性，但也可能使经典研究陷入繁琐考证的泥沼，丧失了对现实问题的回应能力。'
)

add_heading_custom('（三）学术立场与政治态度', level=2)

add_para(
    '今文经学与古文经学之争，从来不是纯粹的学术之争，而是与政治立场紧密相连的。'
)

add_para(
    '今文经学在西汉时期被立为官学，成为官方意识形态的理论基础。今文经学家往往积极参与政治，将经典诠释与现实政治紧密结合。董仲舒的"罢黜百家，独尊儒术"便是今文经学政治化的典型表现。今文经学的政治态度总体上是维护现有秩序的，其经典诠释往往为统治者的政治需要服务。'
)

add_para(
    '古文经学则长期处于民间学术的地位，与官方政治保持一定的距离。古文经学家更加注重学术的独立性和客观性，其诠释活动更多是出于对经典本身的学术兴趣，而非政治目的。然而，古文经学也并非完全脱离政治。刘歆推崇古文经，其背后就有王莽改制的政治动机。'
)

add_heading_custom('四、《论语》今古文诠释的个案分析')

add_heading_custom('（一）核心概念的不同诠释', level=2)

add_para(
    '《论语》中的核心概念，在今古文两派的诠释中呈现出明显的差异。以"仁"这一核心概念为例，今文经学家往往将其与政治治理和社会秩序相联系，强调"仁政"的政治功能；而古文经学家则更注重从文字训诂的角度，还原"仁"的原始语义，强调其作为道德品质的内涵。'
)

add_para(
    '再以"礼"为例，今文经学家倾向于将"礼"理解为维系社会等级秩序的制度规范，强调其政治功能；古文经学家则更多地将"礼"视为一种文化传统和行为规范，注重其历史演变和文化意义。'
)

add_heading_custom('（二）诠释风格的差异', level=2)

add_para(
    '今文经学的诠释风格往往具有强烈的目的论色彩。诠释者带着明确的政治或道德目的来解读经典，其诠释结果往往与诠释者的预设立场高度一致。这种诠释方式的优势在于赋予了经典以强烈的现实意义，但缺陷在于可能导致对经典文本的过度解读甚至曲解。'
)

add_para(
    '古文经学的诠释风格则更加审慎和客观。诠释者力求通过训诂考据的方法，还原经典文本的原始含义，避免将自己的主观意志强加于经典之上。这种诠释方式的优势在于保证了学术的客观性和可靠性，但缺陷在于可能使经典研究变得过于琐碎和脱离现实。'
)

add_heading_custom('五、今古文之争的历史演变')

add_heading_custom('（一）西汉：今文经学的鼎盛', level=2)

add_para(
    '西汉时期，今文经学凭借其官学地位，在学术界和政治界占据了主导地位。汉武帝"罢黜百家，独尊儒术"，所尊之儒学即为今文经学。今文经学家的诠释成为官方意识形态的理论基础，其学术影响力达到鼎盛。'
)

add_para(
    '这一时期，《论语》的传授也主要以今文系统为主。《鲁论》和《齐论》在学术界广泛流传，成为士人学习和研究的主要文本。'
)

add_heading_custom('（二）东汉：古文经学的兴起', level=2)

add_para(
    '东汉时期，古文经学逐渐兴起。随着古文经典的不断发现和古文经学家的不懈努力，古文经学的学术影响力日益增强。郑玄兼采今古文，以古文经学为基础，融合今文经学的合理因素，成为东汉经学的集大成者。'
)

add_para(
    '郑玄对《论语》的注释，便是以张禹的"张侯论"为底本，同时参考了今古文两派的学术成果。郑玄的注释成为后来《论语》研究的重要基础，影响深远。'
)

add_heading_custom('（三）清代：今古文之争的复兴', level=2)

add_para(
    '清代乾嘉时期，考据学大兴，古文经学的方法论得到了前所未有的发展。乾嘉学者以训诂考据为核心方法，对经典文献进行了系统而精细的整理和研究。'
)

add_para(
    '然而，到了晚清，今文经学再次兴起。庄存与、刘逢禄等人重新强调今文经学的"微言大义"，康有为更是将今文经学与变法维新的政治诉求相结合，著《新学伪经考》《孔子改制考》等著作，对古文经学进行了猛烈的批判。'
)

add_para(
    '康有为的"伪经"说认为，古文经乃是刘歆为了配合王莽篡汉而伪造的。这一观点虽然在学术上存在诸多争议，但在当时的社会政治环境下，产生了巨大的影响，为变法维新提供了理论支持。'
)

add_heading_custom('六、今古文之争的当代启示')

add_heading_custom('（一）经典诠释的两种范式', level=2)

add_para(
    '从当代学术的视角来看，今古文之争实际上反映了经典诠释中的两种基本范式：一种是以意义阐释为核心的诠释范式，强调经典的现实关怀和时代价值；另一种是以文本还原为核心的诠释范式，强调经典的原始含义和历史语境。'
)

add_para(
    '这两种范式各有其合理性和局限性。当代的经典研究，应当超越今古文之争的简单对立，将两种诠释范式有机结合起来，既重视文本的历史还原，又关注经典的当代意义。'
)

add_heading_custom('（二）对当代中国哲学研究的意义', level=2)

add_para(
    '今古文经学之争对当代中国哲学研究具有重要的启示意义。首先，它提醒我们关注经典诠释的方法论问题，避免将单一的诠释方法绝对化。其次，它揭示了经典诠释与政治、社会之间的复杂关系，促使我们反思学术研究的社会功能。最后，它展示了中国传统学术中丰富的诠释资源，为当代中国哲学的创新发展提供了深厚的文化根基。'
)

add_heading_custom('七、结论')

add_para(
    '《论语》今文经学与古文经学之争，是中国经学史上最为重要的学术争论之一。从版本源流来看，今文《论语》以口耳相传、隶书记录为特征，古文《论语》则以先秦古本、古文字书写为依据。从学术特征来看，今文经学注重微言大义与经世致用，古文经学强调训诂考据与历史还原。从诠释方法来看，今文经学的诠释具有强烈的政治目的论色彩，古文经学则更加注重学术的客观性和独立性。'
)

add_para(
    '今古文之争并非简单的文本真伪之辩，而是两种经典诠释范式的根本对立。这一争论贯穿了整个中国经学史，从西汉的今文独尊，到东汉的古文兴起，再到清代的今古文复兴，每一次学术转向都深刻影响了中国思想文化的发展走向。'
)

add_para(
    '在当代学术语境中，今古文之争的理论价值不仅没有消减，反而更加凸显。它为我们理解经典诠释的复杂性、反思学术研究的方法论、以及探索中国传统哲学的当代转化，提供了宝贵的思想资源。未来的《论语》研究和中国哲学研究，应当在超越今古文简单对立的基础上，创造性地融合两种诠释范式，推动中国传统经典的当代阐释走向新的深度与高度。'
)

# === 保存 ===
output_path = '/home/admin/.openclaw/workspace/papers/lunyu-jingu/论语今古文经学比较研究.docx'
doc.save(output_path)
print(f'✅ 论文已保存: {output_path}')
