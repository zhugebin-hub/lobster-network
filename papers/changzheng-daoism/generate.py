#!/usr/bin/env python3
"""长征精神感悟 Word 文档生成 ~400字"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = '仿宋'
style.font.size = Pt(16)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
style.paragraph_format.line_spacing = Pt(28)

def _r(run, name='仿宋', size=16, bold=False):
    run.font.name = name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold

def p(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=Cm(0.74), size=16, sb=0, sa=0):
    pg = doc.add_paragraph()
    pg.alignment = align
    pg.paragraph_format.first_line_indent = indent
    pg.paragraph_format.space_before = Pt(sb)
    pg.paragraph_format.space_after = Pt(sa)
    run = pg.add_run(text)
    _r(run, size=size, bold=bold)

# 标题
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(40)
tp.paragraph_format.space_after = Pt(20)
run = tp.add_run('长征精神对道教教职人员修行与教务的思想启迪')
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(22)
run.bold = True

p('长征精神的核心是坚定信念、不怕牺牲、实事求是、团结奋斗。作为一名道教教职人员，重温长征精神，我深受触动。')

p('其一，坚定信念，方能坚守道心。红军将士在漫漫征途中面对围追堵截、缺衣少食，始终不改其志。道教修行讲究"守一不移"，修道的本质就是对大道的坚定信念。面对世俗纷扰，唯有坚守初心、笃志不渝，方能守住内心的清静与澄明。')

p('其二，实事求是，方能精进教务。长征途中，中国共产党人坚持从实际出发，灵活调整战略。道教教务工作亦需如此。当代信众的心理需求日益复杂，我们应当将道教教义与当代人的心理疏导需求相结合，以实事求是的态度探索教务新路径。')

p('其三，团结互助，方能广济众生。长征中红军将士患难与共的精神，与道教"济世利人"的教义高度契合。道教修行不是独善其身，而是要以修炼所得之智慧服务信众。在教务中，我们要以慈悲之心关爱每一位信众，将道教智慧转化为抚慰心灵的实际方法。')

p('长征精神跨越时空，历久弥新。我将以长征精神为镜，在修行中坚定信念、在教务中实事求是、在服务中团结互助，为构建和谐社会贡献道教力量。')

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
sp.paragraph_format.space_before = Pt(20)
run = sp.add_run('2026年6月21日')
run.font.name = '楷体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
run.font.size = Pt(14)

output = '/home/admin/.openclaw/workspace/papers/changzheng-daoism/长征精神感悟.docx'
doc.save(output)
print(f'✅ {output}')
