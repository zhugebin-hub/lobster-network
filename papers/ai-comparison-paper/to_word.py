#!/usr/bin/env python3
"""将论文 markdown 转换为 Word 文档"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# 全局中文字体设置
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 读取 markdown ──
with open('/home/admin/.openclaw/workspace/papers/ai-comparison-paper/小论文_当今各人工智能使用对比_择白.md', 'r') as f:
    lines = f.readlines()

in_table = False
table_lines = []
in_code = False

def flush_table(tbl_lines):
    """解析 markdown 表格并添加到文档"""
    # 找到有效行（去掉空行和分隔线）
    rows = []
    for line in tbl_lines:
        line = line.strip()
        if not line:
            continue
        # 跳过分隔线 |---|---|
        if re.match(r'^\|[\s\-:|]+$', line):
            continue
        # 解析单元格
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_text(text, level=1):
    """添加标题"""
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    elif level == 2:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

def add_body(text):
    """添加正文"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_bold_body(bold_part, normal_part):
    """加粗+正文"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run1 = p.add_run(bold_part)
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.name = '黑体'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run2 = p.add_run(normal_part)
    run2.font.size = Pt(12)
    run2.font.name = '宋体'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 解析并写入 ──
for line in lines:
    raw = line.rstrip('\n')

    # 代码块
    if raw.strip().startswith('```'):
        in_code = not in_code
        continue
    if in_code:
        continue

    # 空行
    if not raw.strip():
        continue

    # 分隔线
    if raw.strip() == '---':
        continue

    # 一级标题（论文标题）
    if raw.startswith('# ') and not raw.startswith('## '):
        title = raw[2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        continue

    # 作者行
    if raw.startswith('**作者：'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(raw.replace('**', '').replace('**', ''))
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        continue

    # 二级标题
    if raw.startswith('## '):
        if in_table:
            flush_table(table_lines)
            table_lines = []
            in_table = False
        text = raw[3:].strip()
        add_heading_text(text, level=2)
        continue

    # 三级标题
    if raw.startswith('### '):
        if in_table:
            flush_table(table_lines)
            table_lines = []
            in_table = False
        text = raw[4:].strip()
        add_heading_text(text, level=3)
        continue

    # 表格
    if '|' in raw and raw.strip().startswith('|'):
        in_table = True
        table_lines.append(raw)
        continue
    elif in_table and not raw.strip().startswith('|'):
        flush_table(table_lines)
        table_lines = []
        in_table = False

    # 摘要
    if raw.startswith('**摘要：**'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('摘  要：')
        run1.bold = True
        run1.font.size = Pt(12)
        run1.font.name = '黑体'
        run1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        content = raw[len('**摘要：**'):]
        run2 = p.add_run(content)
        run2.font.size = Pt(12)
        run2.font.name = '宋体'
        run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        continue

    # 关键词
    if raw.startswith('**关键词：**'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        run1 = p.add_run('关键词：')
        run1.bold = True
        run1.font.size = Pt(12)
        run1.font.name = '黑体'
        run1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        content = raw[len('**关键词：**'):]
        run2 = p.add_run(content)
        run2.font.size = Pt(12)
        run2.font.name = '宋体'
        run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        continue

    # 粗体+正文（**1. xxx：** 内容）
    if raw.startswith('**') and '：**' in raw:
        parts = raw.split('：**', 1)
        label = parts[0].replace('**', '')
        content = parts[1].lstrip() if len(parts) > 1 else ''
        add_bold_body(label + '：', content)
        continue

    # 列表项
    if raw.strip().startswith('- '):
        text = raw.strip()[2:].strip()
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # 处理内部加粗
        if '**' in text:
            segments = text.split('**')
            for idx, seg in enumerate(segments):
                if not seg:
                    continue
                run = p.add_run(seg)
                run.font.size = Pt(12)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if idx % 2 == 1:  # 加粗部分
                    run.bold = True
        else:
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        continue

    # 普通段落
    add_body(raw)

# 处理末尾残留表格
if in_table:
    flush_table(table_lines)

# 保存
output_path = '/home/admin/.openclaw/workspace/papers/ai-comparison-paper/小论文_当今各人工智能使用对比_择白.docx'
doc.save(output_path)
print(f'✅ Word 文档已生成: {output_path}')
