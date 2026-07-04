#!/usr/bin/env python3
"""生成武原中学月亮湖好课开幕式总结发言稿 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# === 页面设置 ===
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# === 默认样式 ===
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# === 标题 ===
title = doc.add_heading(level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('武原中学月亮湖好课开幕式总结发言稿')
run.bold = True
run.font.name = '黑体'
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 0, 0)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 副标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('"数智科研·赋能生长" AI赋能教育科研专场活动')
run.font.name = '楷体'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

doc.add_paragraph()  # 空行

# === 正文段落 ===
def add_heading_custom(text, level=2):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = '黑体'
    run.font.size = Pt(16) if level == 1 else Pt(15)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_after = Pt(8)
    if indent:
        # 首行缩进两字符
        p.paragraph_format.first_line_indent = Pt(28)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(14)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_para_with_bold_prefix(prefix, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Pt(28)
    run1 = p.add_run(prefix)
    run1.bold = True
    run1.font.name = '宋体'
    run1.font.size = Pt(14)
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2 = p.add_run(text)
    run2.font.name = '宋体'
    run2.font.size = Pt(14)
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 称呼
add_para('尊敬的各位领导、各位专家、同仁们：')
add_para('大家下午好！')

# 开场
add_para('今天，我们相聚在美丽的武原中学，共同见证"数智科研·赋能生长"AI赋能教育科研专场活动的圆满开幕。首先，我谨代表武原中学，向莅临本次活动的各位领导、专家和教育同仁表示最热烈的欢迎和最诚挚的感谢！')

# 第一部分
add_heading_custom('一、活动回顾与亮点', level=2)

add_para('今天下午的活动内容丰富、形式多样，既有理论高度，又有实践深度。我们共同见证了以下精彩环节：')

add_para_with_bold_prefix('课题开题，引领方向。', '县级规划课题《"智能批改"背景下初中数学动态分层作业设计与实践研究》和县级微型课题《AI赋能初中地理跨学科主题研学的应用研究》的开题论证，为我们展示了AI技术与学科教学深度融合的创新路径。秦中李杰老师、武中步晓翠老师的课题设计，体现了我们对教育数字化转型的深入思考。')

add_para_with_bold_prefix('专家点评，把脉定向。', '陈梦瑶老师、吴益佳老师的精彩点评，为课题研究提供了宝贵的指导意见，让我们对如何开展好AI赋能教育研究有了更清晰的认识。')

add_para_with_bold_prefix('课题展示，共享智慧。', '陈娟芳老师的《之江汇平台+AI：构建初中生PS课程自主学习新生态》和朱凯琳老师的《智慧中小学平台支持下初中党史课程的三维实施路径研究》，分别从市级数字化专项课题和省级信息化课题的角度，分享了他们在AI赋能教学实践中的探索与成果。')

add_para_with_bold_prefix('专业引领，提升认知。', '市教育技术与数据中心资源建设部主任周雯的专家点评讲座，为我们带来了前沿的教育技术理念和实践经验，让我们对AI赋能教育的未来发展方向有了更深层次的理解。')

# 第二部分
add_heading_custom('二、AI赋能教育的意义', level=2)

add_para('本次活动的主题是"数智科研·赋能生长"，这不仅仅是一个口号，更是我们对未来教育发展的庄严承诺。')

add_para_with_bold_prefix('AI赋能，让教学更精准。', '通过智能批改、动态分层作业设计，我们能够更好地了解每个学生的学习情况，实现个性化教学，让每个孩子都能在自己的节奏中成长。')

add_para_with_bold_prefix('AI赋能，让学习更自主。', '之江汇平台与AI的结合，为学生提供了更加丰富的学习资源和更加灵活的学习方式，培养了学生的自主学习能力和创新思维。')

add_para_with_bold_prefix('AI赋能，让科研更深入。', '数字化课题研究的开展，推动了教育科研的转型升级，让教育研究更加科学化、系统化、实效化。')

# 第三部分
add_heading_custom('三、展望与期待', level=2)

add_para('月亮湖好课项目的启动，标志着我们在AI赋能教育的道路上迈出了坚实的一步。在此，我提出三点期待：')

add_para_with_bold_prefix('第一，深化研究，务求实效。', '希望各课题组能够扎实开展研究工作，将AI技术与学科教学深度融合，形成可复制、可推广的经验和成果。')

add_para_with_bold_prefix('第二，加强交流，共同成长。', '希望各校之间能够建立常态化的交流机制，共享资源、共研问题、共同发展，形成区域教育发展的合力。')

add_para_with_bold_prefix('第三，勇于创新，敢于突破。', '希望各位教师能够保持开放的心态，积极拥抱新技术，勇于尝试新的教学模式，为教育数字化转型贡献智慧和力量。')

# 第四部分
add_heading_custom('四、结语', level=2)

add_para('各位同仁，教育数字化转型的大幕已经拉开，AI赋能教育的时代已经到来。让我们以本次活动为契机，携手并进，共同探索AI赋能教育的新路径，共同开创月亮湖好课的新局面，为培养更多具有创新精神和实践能力的时代新人而不懈努力！')

add_para('最后，再次感谢各位领导和专家的莅临指导，感谢各位同仁的积极参与！祝愿本次活动取得圆满成功！祝愿各位工作顺利、身体健康！')

add_para('谢谢大家！')

# 落款
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(4)
run = p.add_run('发言人：沈利娟')
run.font.name = '宋体'
run.font.size = Pt(14)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p2.paragraph_format.space_after = Pt(4)
run2 = p2.add_run('武原中学')
run2.font.name = '宋体'
run2.font.size = Pt(14)
run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run3 = p3.add_run('2026年5月20日')
run3.font.name = '宋体'
run3.font.size = Pt(14)
run3.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 保存
output_path = '/home/admin/.openclaw/workspace/武原中学月亮湖好课开幕式总结发言稿.docx'
doc.save(output_path)
print(f'✅ Word文档已生成: {output_path}')
