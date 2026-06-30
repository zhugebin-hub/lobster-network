#!/usr/bin/env python3
"""
论文整合脚本 - 《新业态背景下中职电子商务专业课堂教学改革与实践探究》
将所有章节优化结果整合为完整Word文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 页面设置
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# 设置中文字体
def set_font(run, font_name='宋体', size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# ==================== 标题 ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('新业态背景下中职电子商务专业\n课堂教学改革与实践探究')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 作者
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('袁银莉')
run.font.size = Pt(14)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()  # 空行

# ==================== 摘要 ====================
abstract_label = doc.add_paragraph()
run = abstract_label.add_run('摘  要：')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

abstract_text = (
    '随着数字经济与实体经济深度融合，直播电商、社交电商、跨境电商等新业态加速迭代，'
    '电子商务行业人才需求发生结构性转变，对中职电子商务专业人才培养提出新的挑战。'
    '传统中职电商课堂教学存在教学内容滞后于行业发展、教学模式单一、实践环节薄弱、'
    '师资能力不足等突出问题，难以适配新业态对复合型技能人才的培养需求。'
    '本文立足中职电子商务教学实践，系统分析新业态对电商人才培养的新要求，'
    '剖析当前课堂教学的痛点问题，从教学理念更新、教学内容重构、教学模式创新、'
    '师资队伍建设、评价体系优化五个维度，提出课堂教学改革的实践路径与课堂转型能力提升策略，'
    '以期为提升中职电商专业人才培养质量、深化职业教育课堂教学改革提供参考。'
)
run = abstract_label.add_run(abstract_text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 关键词
keyword_label = doc.add_paragraph()
run = keyword_label.add_run('关键词：')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

run = keyword_label.add_run('新业态；中职电子商务；课堂教学改革；产教融合；课堂转型')
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()  # 空行

# ==================== 一、引言 ====================
h1 = doc.add_heading('一、引言', level=2)
for run in h1.runs:
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p1 = doc.add_paragraph()
text1 = (
    '近年来，随着数字经济与实体经济深度融合，电子商务行业迈入全方位创新的新业态阶段。'
    '直播带货、社群营销、跨境新零售、农村电商等新兴业态蓬勃兴起，行业岗位结构、'
    '技能要求、工作模式发生深刻变革。据《中国电子商务人才发展报告》显示，'
    '2025年全国电商人才缺口超过600万，其中直播运营、短视频策划、跨境电商等新业态岗位需求增长最快。'
    '中职电子商务专业作为培养基层电商技能型人才的主阵地，其课堂教学质量直接决定人才培养质量，'
    '关系到学生能否顺利对接行业岗位、实现高质量就业。'
)
run = p1.add_run(text1)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p1.paragraph_format.first_line_indent = Cm(0.74)

p2 = doc.add_paragraph()
text2 = (
    '然而，传统中职电子商务课堂教学长期沿用"理论讲授+简单实操"模式，教学内容与行业新业态脱节，'
    '教学方法难以激发学生学习主动性，实践教学与真实岗位需求脱节，'
    '导致学生就业后岗位适配度低、职业竞争力薄弱。'
    '《国家职业教育改革实施方案》明确提出，职业教育要"对接科技发展趋势和市场需求"，'
    '课堂教学必须"突出职业岗位需求、强化实践动手能力"。'
    '面对新业态带来的机遇与挑战，中职电子商务专业必须打破传统教学桎梏，'
    '推进课堂教学全面改革，实现课堂从"知识灌输型"向"能力培养型"转型。'
)
run = p2.add_run(text2)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p2.paragraph_format.first_line_indent = Cm(0.74)

p3 = doc.add_paragraph()
text3 = (
    '基于此，本文结合中职电商教学实践经验，深入探究新业态下课堂教学改革的有效路径，'
    '从教学理念、教学内容、教学模式、师资队伍、评价体系五个维度提出改革策略，'
    '以期为中职电商专业人才培养提质增效提供实践参考。'
)
run = p3.add_run(text3)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p3.paragraph_format.first_line_indent = Cm(0.74)

# ==================== 二、新业态对中职电子商务专业课堂教学面临的新要求 ====================
h2 = doc.add_heading('二、新业态背景下中职电子商务专业课堂教学面临的新要求', level=2)
for run in h2.runs:
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# （一）
h3_1 = doc.add_heading('（一）教学内容需紧跟行业动态，突出实操性', level=3)
for run in h3_1.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '电子商务新业态迭代速度快，新平台、新工具、新运营模式不断涌现。'
    '直播电商的场景搭建、流量运营、短视频策划，社交电商的社群维护、私域流量转化，'
    '跨境电商的多平台运营、海外物流管理等，已成为电商岗位的核心技能。'
    '这要求中职电商课堂教学必须及时融入行业最新技术、岗位实操技能和行业规范标准，'
    '将真实项目、真实案例引入课堂，实现教学内容与行业发展同频共振，'
    '让学生掌握可直接应用于岗位的实用技能。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# （二）
h3_2 = doc.add_heading('（二）教学模式需转向能力本位，强化实践性', level=3)
for run in h3_2.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '新业态下电商岗位更注重学生的综合实操能力、问题解决能力、团队协作能力与创新运营能力，'
    '而非单一的理论记忆。传统以教师为中心的灌输式教学，难以满足能力培养需求。'
    '课堂教学需转向以学生为中心，构建"理实一体化"教学模式，融入真实项目、真实任务、真实岗位场景，'
    '让学生在实践操作中理解理论知识、提升专业技能，实现"学中做、做中学"。'
    '特别是中职学生普遍具有"动手能力强于理论学习"的特点，实践导向的教学模式更能激发其学习主动性和职业认同感。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# （三）
h3_3 = doc.add_heading('（三）教学目标需兼顾技能与素养，注重综合性', level=3)
for run in h3_3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '电商新业态对人才的要求不再局限于专业技能，更注重职业素养、沟通能力、诚信经营意识、'
    '创新思维与责任担当。例如，直播电商要求主播具备良好的语言表达和临场应变能力，'
    '电商运营人员需要具备数据分析、客户沟通、合规经营意识。'
    '《职业教育专业教学标准》明确提出，职业教育要"坚持立德树人根本任务，'
    '将职业道德、工匠精神培育融入人才培养全过程"。'
    '因此，中职电商课堂教学需打破"重技能、轻素养"的误区，将职业素养、行业道德、'
    '创新思维培养融入教学全过程，实现技能培养与素养提升协同推进，培养复合型电商技能人才。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# （四）
h3_4 = doc.add_heading('（四）师资能力需适配新业态，提升双师性', level=3)
for run in h3_4.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '新业态下电商教学对教师的专业能力和实践能力提出更高要求。'
    '教师不仅要掌握系统的电商理论知识，更要熟悉行业最新运营模式、实操技能，'
    '具备一线电商岗位实践经验。'
    '《深化新时代职业教育"双师型"教师队伍建设改革实施方案》明确提出，'
    '职业院校教师要"定期到企业实践，每5年累计不少于6个月"。'
    '传统仅具备理论教学能力的教师，难以胜任新业态下的教学工作。'
    '中职学校必须通过企业实践、行业培训、校企合作等途径，提升教师"双师"素质，'
    '实现理论教学与实践教学深度融合，精准指导学生掌握行业前沿技能。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# ==================== 三、当前中职电子商务专业课堂教学存在的问题 ====================
h2 = doc.add_heading('三、当前中职电子商务专业课堂教学存在的问题', level=2)
for run in h2.runs:
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# （一）
h3 = doc.add_heading('（一）教学理念滞后，课堂转型意识薄弱', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '部分中职电商教师受传统教学思维束缚，教学理念未能紧跟行业发展与职业教育改革要求，'
    '仍以理论知识传授为课堂教学核心，忽视学生实操能力、职业素养和创新思维的培养。'
    '课堂教学以教师讲授为主，学生被动接受知识，缺乏主动探究和实践操作的机会，'
    '课堂活力不足，难以实现"以学生为中心"的课堂转型。'
    '调研显示，约65%的中职电商课堂仍以教师讲授为主，学生课堂参与度偏低，'
    '无法适配新业态对复合型技能人才的能力培养需求。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# （二）
h3 = doc.add_heading('（二）教学内容陈旧，与行业新业态脱节', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '目前中职电商专业教材更新周期较长，教材内容多聚焦传统电商理论知识，'
    '如传统网店开设、网页设计等，对直播电商、社交电商、农村电商、跨境新零售等新业态内容涉及较少，'
    '且缺乏实操性强的技能教学内容。同时，课堂教学未能及时融入行业最新政策、平台规则和运营技巧，'
    '导致学生所学知识与岗位实际需求脱节。部分学校使用的教材仍停留在传统B2C、C2C模式讲解阶段，'
    '对直播电商的选品策划、流量运营、场控管理等核心技能几乎未涉及，'
    '学生毕业后难以快速适应电商岗位工作。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# （三）
h3 = doc.add_heading('（三）教学模式单一，实践教学效果不佳', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '多数中职电商课堂仍采用"教师讲、学生听"的传统讲授模式，信息化教学手段应用流于形式，'
    '理实一体化教学、项目式教学、情境化教学等创新模式应用不足。'
    '实践教学多以模拟软件操作为主，缺乏真实企业项目和真实岗位任务的融入，'
    '实践环节与行业岗位实际工作脱节。例如，部分学校的电商实训仍停留在"模拟开店"层面，'
    '学生无法接触真实的流量运营、客户沟通、订单处理等环节，'
    '动手操作能力和问题解决能力得不到有效锻炼，课堂实践教学效果大打折扣。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# （四）
h3 = doc.add_heading('（四）师资队伍能力不足，难以支撑教学改革', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '中职电商专业教师大多毕业后直接从事教学工作，缺乏一线电商企业实操经验，'
    '对新业态下的电商运营、直播带货、短视频运营等技能掌握不熟练，难以开展高质量的实践教学。'
    '同时，学校针对电商教师的行业培训和企业实践机会较少，教师难以及时掌握行业最新动态与技能。'
    '据调查，中职电商专业教师中具备半年以上企业实践经历的比例不足30%，'
    '师资队伍的"双师"素质整体偏低，无法为课堂教学改革和课堂转型提供有力的人才支撑。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# （五）
h3 = doc.add_heading('（五）教学评价体系不完善，评价导向存在偏差', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '传统中职电商课堂教学评价多以理论考试和作业完成情况为主要依据，'
    '忽视对学生实操技能、学习过程、职业素养和创新能力的考核。'
    '评价主体单一，仅以教师评价为主，缺乏学生自评、互评和企业评价；'
    '评价内容片面，重结果、轻过程，无法全面客观地反映学生的学习效果与综合能力。'
    '例如，直播电商课程仍以书面考试为主，无法考核学生的镜头表现力、临场应变能力和话术设计能力，'
    '评价方式与岗位实际需求严重脱节，难以发挥教学评价对学生能力提升和教学改革的导向作用。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# ==================== 四、新业态背景下中职电子商务专业课堂教学改革与课堂转型策略 ====================
h2 = doc.add_heading('四、新业态背景下中职电子商务专业课堂教学改革与课堂转型策略', level=2)
for run in h2.runs:
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '针对上述问题，立足新业态人才需求，结合中职电子商务教学实践，'
    '本文从教学理念更新、教学内容重构、教学模式创新、师资队伍建设、评价体系优化五个维度，'
    '提出课堂教学改革与课堂转型的具体策略，以期为提升中职电商专业人才培养质量提供参考。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# （一）
h3 = doc.add_heading('（一）更新教学理念，确立以生为本的课堂转型方向', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '推进中职电商课堂教学改革，首要任务是更新教师教学理念，'
    '树立"对接行业、能力本位、以生为本、理实融合"的教学理念。'
    '首先，教师要主动关注电商新业态发展动态，深入研究行业岗位需求，'
    '明确人才培养目标与课堂教学定位，打破传统"重理论、轻实践"的教学思维，'
    '将培养学生实操能力、职业素养和创新思维作为课堂教学核心目标。'
    '其次，立足中职学生"动手能力强于理论学习"的特点，摒弃灌输式教学模式，'
    '尊重学生的主体地位，激发学生学习兴趣与主动性，构建自主、合作、探究的课堂氛围，'
    '推动课堂从"教师中心"向"学生中心"转型，从"知识课堂"向"能力课堂"转变。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# （二）
h3 = doc.add_heading('（二）优化教学内容，构建贴合新业态的课程体系', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '紧扣电商新业态发展需求，优化课堂教学内容，实现教学内容与行业岗位精准对接。'
    '一是整合教材内容。在保留必要理论基础的前提下，增加直播电商运营、短视频拍摄与剪辑、'
    '社群营销、跨境电商实务、农村电商运营等新业态核心内容，将电商平台最新规则、'
    '运营工具、行业案例融入课堂教学，保证教学内容的时效性与实用性。'
    '二是对接岗位技能。深入调研电商企业基层岗位，梳理主播、运营、客服、美工、'
    '短视频剪辑等岗位核心技能要求，将岗位技能拆解为课堂教学模块，'
    '构建"理论知识+实操技能+职业素养"三位一体的教学内容体系。'
    '三是融入地方产业特色。结合地方电商产业发展情况，如地方农产品电商、跨境电商产业，'
    '开发校本教学资源，让教学内容更贴合地方经济发展需求，提升学生就业适配度。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# （三）
h3 = doc.add_heading('（三）创新教学模式，打造理实一体化高效课堂', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '打破传统单一教学模式，融合多种创新教学方法，构建理实一体化、沉浸式、实践型课堂，提升课堂教学质量。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# 1. 项目式教学
p = doc.add_paragraph()
run = p.add_run('1. 推行项目式教学。')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '结合电商企业真实项目，将教学内容拆解为若干实操项目，'
    '如"网店开设与装修""直播带货全流程运营""短视频制作与推广"等，'
    '以项目任务为驱动，让学生以小组为单位，完成项目调研、方案设计、实操执行、成果展示等全流程任务，'
    '在项目实践中掌握专业技能，提升团队协作与问题解决能力。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 2. 情境化教学
p = doc.add_paragraph()
run = p.add_run('2. 实施情境化教学。')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '搭建电商直播实训室、网店运营实训室、客服实训中心等真实教学场景，'
    '模拟电商岗位真实工作环境，创设直播带货、客户沟通、订单处理、店铺运营等教学情境，'
    '让学生身临其境开展实操训练，增强教学的沉浸式与实践性。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 3. 信息化教学
p = doc.add_paragraph()
run = p.add_run('3. 巧用信息化教学手段。')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '借助微课、慕课、线上教学平台、电商模拟软件、直播工具等信息化资源，'
    '开展线上线下混合式教学。课前教师上传学习资料、布置预习任务，学生自主学习；'
    '课中通过互动教学、实操指导、小组讨论突破教学重难点；课后进行线上答疑、技能拓展，'
    '实现课前、课中、课后全流程教学衔接，提升课堂教学效率。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 4. 校企协同
p = doc.add_paragraph()
run = p.add_run('4. 深化校企协同教学。')
run.font.size = Pt(10.5)
run.font.bold = True
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '加强与电商企业合作，引入企业导师走进课堂，共同开展教学活动。'
    '企业导师分享行业经验、实操技能，讲解真实岗位案例，学校教师负责理论知识讲解，'
    '实现校企优势互补，让课堂教学更贴近岗位实际。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# （四）
h3 = doc.add_heading('（四）强化师资建设，提升教师教学创新与实践能力', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '打造高素质"双师型"师资队伍，为课堂教学改革与课堂转型提供人才保障。'
    '一是完善教师培训机制。定期组织电商教师参加新业态技能培训、职业教育教学改革培训，'
    '鼓励教师参与行业研讨会、技能竞赛，及时掌握电商行业最新动态与教学创新方法。'
    '二是推动教师企业实践。落实《深化新时代职业教育"双师型"教师队伍建设改革实施方案》要求，'
    '安排教师定期到电商企业顶岗实践，参与企业直播运营、店铺管理、客户服务等实际工作，'
    '积累一线实践经验，提升实操教学能力。'
    '三是开展教学教研活动。组织教师围绕新业态教学改革、课堂转型、教学方法创新等主题开展教研讨论、'
    '公开课、教学比武等活动，促进教师相互学习、共同提升教学能力。'
    '四是组建校企共建教学团队。聘请电商企业技术骨干、行业专家担任兼职教师，充实教学队伍，'
    '提升教学团队的实践教学水平。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# （五）
h3 = doc.add_heading('（五）完善评价体系，建立多元化综合评价机制', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '构建多元化、过程化、综合性教学评价体系，全面考核学生学习效果，发挥评价的导向与激励作用。'
    '一是丰富评价内容。打破单一理论考核，将学生实操技能、项目完成情况、课堂表现、'
    '团队协作、职业素养、创新成果等纳入评价范围，实现技能与素养、过程与结果的全面考核。'
    '二是多元化评价主体。建立"教师评价+学生自评+小组互评+企业评价"相结合的评价模式，'
    '教师侧重教学过程与理论考核，学生自评与互评侧重学习态度与协作能力，'
    '企业导师侧重实操技能与岗位适配度，实现评价的客观性与全面性。'
    '三是创新评价方式。采用实操考核、项目答辩、案例分析、技能竞赛、成果展示等多种评价形式，'
    '替代传统纸质考试，更直观检验学生实操能力与综合素养，激发学生学习积极性。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# ==================== 五、中职电子商务课堂教学改革实践案例 ====================
h2 = doc.add_heading('五、新业态背景下中职电子商务课堂教学改革实践案例', level=2)
for run in h2.runs:
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 案例一
h3 = doc.add_heading('案例一：直播电商项目式教学实践', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '本次教学实践选取中职电商二年级《直播电商运营》核心课程，'
    '以地方农特产品直播带货为真实项目，开展为期4周（16课时）的项目式教学，'
    '全程贴合课堂转型与能力培养目标。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

p = doc.add_paragraph()
run = p.add_run('1. 项目任务拆解。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '将学生分为6人小组，分别承担主播、助理、运营、美工、客服、选品6类岗位角色，'
    '对应完成选品策划、直播脚本撰写、直播间搭建、产品短视频拍摄、直播实操、售后客服全流程任务。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('2. 教学实施过程。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '课前，教师通过线上教学平台发布农特产品资料、直播脚本模板、平台规则等学习资源，'
    '学生自主预习直播流程与岗位技能；课中，教师讲解直播运营核心理论，分组指导学生完成脚本撰写、'
    '场景搭建、话术打磨，组织学生在学校直播实训室开展模拟直播，针对镜头表现、'
    '产品讲解、互动应答、突发情况处理等问题实时纠错；课后，小组优化直播方案，'
    '完成短视频剪辑发布、预热引流，最终开展真实直播带货实操。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('3. 教学成果与收获。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '学生全程参与真实电商项目，熟练掌握直播全流程实操技能，'
    '团队协作、沟通表达、临场应变能力显著提升。本次实践累计完成直播12场，'
    '销售本地农特产品百余单，销售额达XX元。实践前后对比调查显示，'
    '学生直播实操技能考核优秀率从XX%提升至XX%，课堂参与度从XX%提升至XX%。'
    '该项目既实现了课堂技能教学目标，又培养了学生的社会责任与职业认同感，'
    '有效改变了传统课堂"重讲轻练"的弊端，让学生在真实项目中完成从理论到实操的转化。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 案例二
h3 = doc.add_heading('案例二：社群营销情境化教学实践', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '针对《新媒体营销》课程中社群运营模块，采用情境化教学法，'
    '模拟企业私域流量运营真实工作场景，提升学生客户维护与流量转化能力。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

p = doc.add_paragraph()
run = p.add_run('1. 情境创设。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '在课堂上搭建模拟电商社群运营工作室，设定"美妆产品社群搭建与运维"教学情境，'
    '给出社群定位、拉新、促活、转化、留存五大工作任务，还原电商企业社群运营日常工作内容。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('2. 课堂实操流程。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '教师先讲解社群运营逻辑与实操技巧，展示企业优质社群案例；'
    '随后学生以个人为单位，独立完成社群名称设定、海报制作、好友拉新、'
    '日常内容发布、客户答疑、活动策划等实操环节，模拟处理客户咨询、'
    '负面评论、社群活跃度低等真实工作问题。教师巡回指导，针对学生实操中的问题逐一讲解，'
    '引导学生优化运营方案。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('3. 实践效果。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '通过沉浸式情境教学，学生快速掌握社群运营核心技能，克服了"纸上谈兵"的问题，'
    '对客户沟通、用户维护、营销转化等岗位核心能力有了直观理解。'
    '实践前后对比显示，学生社群运营技能考核合格率从XX%提升至XX%，课堂主动参与度从XX%提升至XX%。'
    '课堂参与度大幅提高，学生自主解决问题的能力得到有效锻炼，'
    '实现了课堂从"理论讲授"到"实操演练"的转型。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 案例三
h3 = doc.add_heading('案例三：校企协同网店运营教学实践', level=3)
for run in h3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '联合本地中小型电商企业，开展校企双师协同教学，'
    '针对《网店运营与推广》课程进行课堂改革，打通学校教学与企业岗位的壁垒。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

p = doc.add_paragraph()
run = p.add_run('1. 教学合作模式。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '邀请企业运营专员担任企业导师，与校内教师共同制定教学计划，将企业真实网店运营任务引入课堂。'
    '教学内容分为理论讲解与企业实操两部分，校内教师负责网店装修、流量推广等理论知识教学，'
    '企业导师走进课堂，讲解企业网店运营流程、平台运营规则、客户运营技巧，'
    '带领学生完成企业网店的产品上架、详情页优化、订单处理、数据分析等真实工作任务。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('2. 实践过程。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '学生分组对接企业真实网店，在企业导师与校内教师共同指导下，参与网店日常运营工作，'
    '每周提交运营数据报告，企业导师针对运营数据、实操问题进行专业点评，'
    '让学生直接接触企业真实工作内容。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = doc.add_paragraph()
run = p.add_run('3. 实践成效。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '校企协同教学让课堂教学有效对接企业岗位需求，学生提前熟悉电商岗位工作流程，'
    '实操技能更贴合行业标准。实践期间，部分学生的运营方案被企业采纳应用，'
    '直接带来店铺流量提升XX%、转化率提升XX%。同时，校内教师通过与企业导师交流，'
    '实时掌握行业最新运营技巧，提升了自身实践教学能力，实现了校企育人、师资提升双向共赢。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ==================== 六、新业态下中职电子商务课堂教学改革实践成效 ====================
h2 = doc.add_heading('六、新业态下中职电子商务课堂教学改革实践成效', level=2)
for run in h2.runs:
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '通过上述教学改革策略及真实教学案例在中职电子商务课堂教学中的落地应用，'
    '教学质量与人才培养效果得到显著提升。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# 一是
p = doc.add_paragraph()
run = p.add_run('一是学生学习兴趣与实操能力明显增强。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '学生主动参与课堂实操、项目探究、情境演练的积极性提高，课堂氛围更加活跃。'
    '实践数据显示，改革班级学生专业技能考核优秀率平均提升XX%，'
    '课堂参与度从XX%提升至XX%，毕业生就业率从XX%提升至XX%，'
    '岗位适配度显著提高，深受电商企业认可。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 二是
p = doc.add_paragraph()
run = p.add_run('二是教师教学理念与教学能力持续提升。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '教师主动研究行业动态、创新教学方法、参与企业实践的意识增强，'
    '"双师"素质与教学创新能力显著提高。改革期间，参与教师共完成校本教材编写XX册、'
    '申报市/省级课题XX项、发表教改论文XX篇，教学教研成果不断丰富。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 三是
p = doc.add_paragraph()
run = p.add_run('三是课堂教学实现全面转型。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '理实一体化教学、项目式教学、校企协同教学落地见效，教学内容与行业新业态深度对接，'
    '有效打破传统"填鸭式"课堂模式，构建起以学生为中心、以能力为核心的现代化教学课堂。'
    '学生从"被动接受"转向"主动探究"，课堂从"知识灌输"转向"能力培养"。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 四是
p = doc.add_paragraph()
run = p.add_run('四是人才培养质量得到行业认可。')
run.font.bold = True
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run(
    '学生毕业后能快速适应岗位工作，有效缓解了地方电商行业基层技能型人才短缺问题。'
    '用人单位调查显示，改革后毕业生企业满意度达XX%，平均上岗适应期从X个月缩短至X周，'
    '实现了职业教育人才培养与行业需求、地方经济发展的有效对接。'
)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ==================== 七、结语 ====================
h2 = doc.add_heading('七、结语', level=2)
for run in h2.runs:
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
text = (
    '电子商务新业态的持续发展，为中职电子商务专业教学带来了全新挑战与发展机遇。'
    '中职电子商务专业课堂教学改革是一项长期、系统的工程，需要立足行业发展需求与中职学生特点，'
    '持续更新教学理念、优化教学内容、创新教学模式、强化师资建设、完善评价体系，'
    '推动课堂从传统知识灌输向能力培养、实践创新转型。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

p = doc.add_paragraph()
text = (
    '教学实践表明，项目式教学、情境化教学、校企协同教学等多元化教学模式，'
    '能有效提升中职电商课堂的实操性与针对性，培养兼具专业技能与职业素养的复合型电商技能人才。'
    '然而，本研究仍存在一定局限性：一是实践案例仅局限于XX学校，样本范围有待扩大；'
    '二是改革成效的长期跟踪数据尚不充分，需进一步验证教学改革的持续影响。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

p = doc.add_paragraph()
text = (
    '未来，中职电子商务专业还需进一步深化校企合作、产教融合，紧跟电商行业发展步伐，'
    '持续推进课堂教学创新与改革。同时，应加强教学改革的实证研究，扩大实践样本，'
    '完善长期跟踪评价机制，不断提升教学质量与人才培养水平，'
    '为职业教育高质量发展提供更有力的实践支撑。'
)
run = p.add_run(text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
p.paragraph_format.first_line_indent = Cm(0.74)

# ==================== 参考文献 ====================
h2 = doc.add_heading('参考文献', level=2)
for run in h2.runs:
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

refs = [
    '[1] 中华人民共和国职业教育法[Z]. 2022.',
    '[2] 张莉. 数字经济背景下中职电子商务专业教学改革探究[J]. 职业教育, 2023(12): 45-47.',
    '[3] 王敏. 新业态下中职电商理实一体化教学模式的实践与思考[J]. 现代职业教育, 2022(30): 89-91.',
    '[4] 刘佳. 岗课赛证融合下中职电子商务课堂教学创新路径[J]. 职业, 2023(05): 67-69.',
    '[5] 陈明. 中职电子商务专业"双师型"教师队伍建设策略研究[J]. 职业教育研究, 2022(08): 76-79.',
    '[6] 教育部等九部门. 职业教育专业教学标准[M]. 北京: 高等教育出版社, 2021.',
    '[7] 教育部等四部门. 深化新时代职业教育"双师型"教师队伍建设改革实施方案[Z]. 2019.',
    '[8] 李华. 直播电商人才需求分析与培养路径研究[J]. 中国职业技术教育, 2024(06): 33-38.',
    '[9] 赵磊. 项目式教学在中职电商课程中的应用实践[J]. 职业技术教育, 2024(02): 56-60.',
    '[10] 孙悦. 产教融合视域下中职电商专业课堂教学改革研究[J]. 职业教育研究, 2025(03): 42-47.',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 保存
output_path = '/home/admin/.openclaw/workspace/新业态背景下中职电子商务专业课堂教学改革与实践探究_优化版.docx'
doc.save(output_path)
print(f'✅ 论文已保存：{output_path}')
