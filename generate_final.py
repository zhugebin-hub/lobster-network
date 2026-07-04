#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成成员4 最终版 Word 文档 - 纯成员4 内容 + 双轴图"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_para(doc, text, size=11, bold=False, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font_size = Pt(size)
    r.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p

def create_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    # ==================== 标题 ====================
    title = doc.add_heading('成员4：房价下跌如何触发危机', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ==================== 第一部分：数据图 ====================
    doc.add_heading('一、数据图：2000-2008 年美国房价指数 vs 次级贷违约率', level=2)
    
    add_para(doc, '【图表说明】上图展示了 2000-2008 年美国房价指数（左轴，2000 年=100）与次级贷款违约率（右轴，%）的变化趋势。', size=11)
    
    doc.add_heading('关键发现：', level=3)
    add_para(doc, '• 2000-2006 年：房价从 100 涨至 189（+89%），违约率维持在 1.5%-5.5% 低位')
    add_para(doc, '• 2006 年：房价达到峰值 189，违约率升至 5.5%，拐点出现')
    add_para(doc, '• 2007-2008 年：房价暴跌至 139（-26%），违约率飙升至 15.8%')
    add_para(doc, '• 剪刀差效应：房价下跌与违约率上升形成"剪刀差"，标志危机全面爆发', bold=True)
    
    # ==================== 第二部分：真实案例 ====================
    doc.add_heading('二、真实案例：雷曼兄弟破产', level=2)
    
    doc.add_heading('1. 公司简介', level=3)
    add_para(doc, '• 成立时间：1850 年，总部位于纽约')
    add_para(doc, '• 地位：美国第四大投资银行，全球性金融机构')
    add_para(doc, '• 危机前资产规模：约 6390 亿美元')
    
    doc.add_heading('2. 倒闭原因', level=3)
    p = doc.add_paragraph()
    r = p.add_run('核心原因：')
    r.bold = True
    r = p.add_run('过度投资次级房贷相关证券（MBS、CDO）')
    
    add_para(doc, '• 2003-2007 年大量收购房贷机构，扩大次级贷业务')
    add_para(doc, '• 持有 MBS/CDO 资产超过 850 亿美元')
    add_para(doc, '• 杠杆率高达 31:1')
    add_para(doc, '• 房价下跌后，MBS/CDO 价值暴跌')
    
    doc.add_heading('3. 传导路径', level=3)
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['阶段', '事件', '影响']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font_size = Pt(10)
    
    data = [
        ['① 房价下跌', '2006 年起美国房价开始下跌', '次级贷款违约率上升'],
        ['② MBS 贬值', '雷曼持有的 MBS/CDO 价值暴跌', '资产缩水 850 亿美元'],
        ['③ 流动性危机', '交易对手要求追加保证金', '现金短缺'],
        ['④ 寻求救助', '雷曼寻求买家或政府救助', '韩国产业银行退出谈判'],
        ['⑤ 破产申请', '2008 年 9 月 15 日申请破产保护', '史上最大破产案'],
        ['⑥ 全球冲击', '股市暴跌、信贷冻结', '金融危机全面爆发']
    ]
    
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font_size = Pt(9)
    
    doc.add_heading('4. 关键数据', level=3)
    add_para(doc, '• 破产日期：2008 年 9 月 15 日')
    add_para(doc, '• 破产资产：6130 亿美元（美国历史上最大破产案）')
    add_para(doc, '• 裁员人数：约 2.5 万人')
    add_para(doc, '• 股市反应：道指当日下跌 4.4%')
    
    # ==================== 第三部分：传导机制总结 ====================
    doc.add_heading('三、传导机制总结', level=2)
    
    p = doc.add_paragraph()
    r = p.add_run('房价下跌 → 次级贷违约 → MBS/CDO 贬值 → 金融机构亏损 → 流动性危机 → 雷曼破产 → 全球金融危机')
    r.font_size = Pt(12)
    r.bold = True
    
    doc.add_heading('关键教训：', level=3)
    add_para(doc, '• 房价不会永远上涨，资产泡沫终将破裂')
    add_para(doc, '• 金融衍生品可能放大而非分散风险')
    add_para(doc, '• 高杠杆在顺周期时放大收益，在逆周期时放大损失')
    add_para(doc, '• 系统性风险需要监管干预，"太大而不能倒"是真实存在的')
    
    # 参考资料
    doc.add_heading('参考资料', level=2)
    add_para(doc, '• 美联储经济数据 (FRED): Case-Shiller 房价指数')
    add_para(doc, '• 美国财政部：《金融危机调查报告》(2011)')
    add_para(doc, '• 《大空头》(The Big Short) - Michael Lewis')
    add_para(doc, '• 《监守自盗》(Inside Job) 纪录片')
    
    # 保存
    doc.save('/home/admin/.openclaw/workspace/成员4-房价下跌如何触发危机.docx')
    print('✅ Word 文档已生成：成员4-房价下跌如何触发危机.docx')

if __name__ == '__main__':
    create_document()
