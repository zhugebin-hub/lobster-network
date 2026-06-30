#!/usr/bin/env python3
"""Generate 慈云佛学院 开题报告 .docx with proper formatting."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Page setup
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

def set_run_font(run, font_name, size=None, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_para_spacing(para, line_spacing=20, before=0, after=0):
    pf = para.paragraph_format
    pf.line_spacing = Pt(line_spacing)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, '宋体', 16, True)  # 三号=16pt
    set_para_spacing(p, line_spacing=20, after=12)

def add_level1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, '黑体', 14, True)  # 四号=14pt
    set_para_spacing(p, line_spacing=20, before=12, after=6)

def add_level2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, '黑体', 12, True)  # 小四=12pt
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.74)  # 约等于空二格
    set_para_spacing(p, line_spacing=20, before=6, after=3)

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, '宋体', 12, False)  # 小四=12pt
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.74)  # 首行缩进两字符
    set_para_spacing(p, line_spacing=20, before=0, after=0)

def add_sub_item(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, '宋体', 12, False)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.74)
    set_para_spacing(p, line_spacing=20, before=3, after=3)

# === Content ===

add_title('宗教中国化视域下的天台四教四谛研究')

# 一、选题的背景与意义
add_level1('一、选题的背景与意义')

add_level2('（一）选题背景')
add_body('四谛（苦、集、灭、道）是佛教的根本教义，为一切佛法之基石。天台宗智𫖮大师在《四教义》中，以判教的方法将四谛分为四种层次：生灭四谛（藏教）、无生灭四谛（通教）、无量四谛（别教）、无作四谛（圆教），构建了天台宗独特的四谛诠释体系。这一体系不仅是智𫖮判教思想的核心内容，也是理解天台宗"化法四教"（藏、通、别、圆）教理架构的关键。')
add_body('当前，我国宗教工作强调"坚持我国宗教中国化方向"，要求各宗教在教义教理、礼仪制度、组织管理等方面与社会主义社会相适应。佛教作为我国重要的传统宗教之一，其中国化进程既是历史事实，也是当代课题。天台宗作为中国佛教八大宗派之一，其教理体系本身就是印度佛教中国化的典范——智𫖮大师以中国思维的圆融特质，创造性地整合了印度佛教的多种经典与思想，建立了"一念三千""三谛圆融""四教判释"等具有鲜明中国特色的佛学理论。')
add_body('因此，在宗教中国化的视域下重新审视天台四教四谛思想，既有助于深化对天台宗判教理论的理解，也能为当代佛教中国化提供理论资源与实践启示。')

add_level2('（二）选题意义')
add_body('理论意义：第一，深化天台四谛思想研究。现有研究多从哲学或文献学角度探讨智𫖮的四谛思想，较少将其置于宗教中国化的宏观框架中进行系统考察。本研究试图从宗教中国化的视角，重新梳理四教四谛的诠释逻辑与思想特质，揭示其中国化的内在理路。第二，丰富宗教中国化的理论内涵。天台宗的四教判释体系展示了中国佛教徒如何以本土思维整合外来教义，这一历史经验对当代宗教中国化具有重要的理论参考价值。')
add_body('实践意义：第一，为当代佛教中国化提供历史参照。天台宗的成功中国化经验表明，教义教理的创造性诠释是宗教中国化的核心环节。本研究总结的诠释方法，可为当代佛教理论建设提供借鉴。第二，推动天台智慧在当代社会的应用。四教四谛从"生灭"到"无作"的递进，蕴含着对生命痛苦的深刻洞察与超越智慧。将其与当代人的精神困境（如焦虑、异化、意义危机）相结合，有助于推动佛教智慧在心理疏导、心灵关怀等领域的现代转化。')

# 二、研究的基本内容与拟解决的主要问题
add_level1('二、研究的基本内容与拟解决的主要问题')

add_level2('（一）基本内容')
add_body('本研究以智𫖮《四教义》为核心文献，在宗教中国化的视域下，对天台四教（藏、通、别、圆）的四谛思想进行系统研究。具体内容包括四个方面：第一，天台四教四谛的文本与义理分析，以《四教义》为中心，梳理藏教生灭四谛、通教无生灭四谛、别教无量四谛、圆教无作四谛的具体内涵，分析四教对苦、集、灭、道四谛的不同诠释方式。第二，天台四教四谛的判教逻辑，分析智𫖮以四谛为线索构建判教体系的内在理路，探讨"生灭→无生灭→无量→无作"的递进关系及其所体现的中国思维特质。第三，天台四教四谛的中国化特质，从天台四教四谛的诠释方法出发，分析其中国化的具体表现，包括以中国思维的"圆融"特质整合印度佛教的多元教义、以次第与圆顿相结合的方式适应不同根机的中国信众、以"一念三千""三谛圆融"等中国化概念重新诠释四谛。第四，天台四教四谛的当代价值，探讨其在当代宗教中国化进程中的理论价值与实践意义。')

add_level2('（二）拟解决的主要问题')
add_body('本研究拟解决三个主要问题：第一，藏通别圆四教对四谛的诠释差异及其内在逻辑关系是什么？第二，天台四教四谛的诠释方法体现了哪些中国化特质？第三，天台四教四谛思想对当代佛教中国化与心灵关怀有何启示？')

# 三、研究的思路与方法
add_level1('三、研究的思路与方法')

add_level2('（一）研究思路')
add_body('本研究的总体思路是：以宗教中国化为理论视域，以智𫖮《四教义》为文献基础，以天台四教四谛为研究对象，按照"历史溯源—文本分析—逻辑梳理—当代诠释"的路径展开。具体步骤分为四个阶段：第一阶段，历史溯源。梳理印度佛教四谛说的基本内容及其在中国传播过程中面临的诠释挑战，为理解天台四教四谛的中国化语境奠定基础。第二阶段，文本分析。以《四教义》为核心文献，逐教分析藏通别圆四教对苦、集、灭、道四谛的具体诠释，厘清各教四谛的内涵与特征。第三阶段，逻辑梳理。在文本分析的基础上，比较四教四谛的异同，揭示从生灭到无作的递进关系，分析其体现的中国思维特质。第四阶段，当代诠释。将天台四教四谛置于宗教中国化的框架中，总结其中国化的历史经验，探讨其对当代佛教理论建设与心灵关怀的启示。四个阶段环环相扣、逻辑递进，由历史到文本、由文本到逻辑、由逻辑到当代，逐步深入。')

add_level2('（二）研究方法')
add_body('本研究采用以下方法：第一，文献分析法。以《四教义》为核心文献，辅以《法华玄义》《摩诃止观》《维摩经玄疏》等天台典籍，进行文本细读与义理分析，确保研究建立在扎实的文献基础之上。第二，比较研究法。纵向比较印度原始佛教四谛说与天台四教四谛的异同，揭示中国化的具体表现；横向比较藏通别圆四教对四谛诠释的差异，分析判教的内在逻辑。第三，诠释学方法。运用哲学诠释学的方法，分析智𫖮诠释四谛的诠释策略与诠释逻辑，提炼其对中国化的方法论启示。第四，跨学科方法。结合宗教学、哲学、心理学等学科视角，探讨四教四谛的当代价值，特别是其在心理疏导和心灵关怀中的应用可能。')

# 四、研究的总体安排与进度
add_level1('四、研究的总体安排与进度')

add_body('1. 2026年6月～7月，搜集大量相关资料，精读《四教义》及相关天台典籍，完成开题报告；')
add_body('2. 2026年8月～9月，完成第一章（思想渊源与中国化语境）和第二章（四教四谛的义理分析）初稿；')
add_body('3. 2026年10月～11月，完成第三章（递进逻辑与圆融特质）初稿；')
add_body('4. 2026年12月～2027年1月，完成第四章（宗教中国化视域下的当代价值）初稿；')
add_body('5. 2027年2月～3月，完成引言与结语，全文统稿修改，形成二稿；')
add_body('6. 2027年4月～5月，根据导师意见修改完善，定稿并准备答辩。')

# 五、目录
add_level1('五、目录')

add_sub_item('引言')
add_sub_item('第一章 天台四教四谛的思想渊源与中国化语境')
add_sub_item('    第一节 印度佛教四谛说概要')
add_sub_item('    第二节 智𫖮与《四教义》的成书背景')
add_sub_item('    第三节 天台判教的中国化特质')
add_sub_item('第二章 藏通别圆四教四谛的义理分析')
add_sub_item('    第一节 藏教生灭四谛')
add_sub_item('    第二节 通教无生灭四谛')
add_sub_item('    第三节 别教无量四谛')
add_sub_item('    第四节 圆教无作四谛')
add_sub_item('    第五节 四教四谛的比较与总结')
add_sub_item('第三章 四教四谛的递进逻辑与圆融特质')
add_sub_item('    第一节 从生灭到无作：四谛诠释的递进关系')
add_sub_item('    第二节 次第与圆顿：天台判教的双重维度')
add_sub_item('    第三节 三谛圆融与四教四谛的内在关联')
add_sub_item('    第四节 天台四谛诠释的中国思维特质')
add_sub_item('第四章 宗教中国化视域下的天台四教四谛')
add_sub_item('    第一节 天台四教四谛中国化的历史经验')
add_sub_item('    第二节 对当代佛教中国化的启示')
add_sub_item('    第三节 当代价值：对精神困境的回应')
add_sub_item('结语')

# 六、主要参考文献
add_level1('六、主要参考文献')

refs = [
    '[1] 智𫖮. 四教义[M]// 大正藏: 第46册. 台北: 新文丰出版公司, 1983.',
    '[2] 智𫖮. 法华玄义[M]// 大正藏: 第33册. 台北: 新文丰出版公司, 1983.',
    '[3] 智𫖮. 摩诃止观[M]// 大正藏: 第46册. 台北: 新文丰出版公司, 1983.',
    '[4] 智𫖮. 维摩经玄疏[M]// 大正藏: 第38册. 台北: 新文丰出版公司, 1983.',
    '[5] 牟宗三. 佛性与般若[M]. 台北: 学生书局, 1993.',
    '[6] 方立天. 中国佛教哲学要义[M]. 北京: 中国人民大学出版社, 2002.',
    '[7] 楼宇烈. 中国佛教的人文精神[M]. 北京: 宗教文化出版社, 2007.',
    '[8] 慧岳. 天台教学史[M]. 台北: 中华佛教文化馆, 1998.',
    '[9] 潘桂明. 智𫖮评传[M]. 南京: 南京大学出版社, 1998.',
    '[10] 陈英善. 天台四教义研究[M]. 台北: 文津出版社, 2001.',
    '[11] 赖永海. 中国佛教通史[M]. 南京: 江苏人民出版社, 2010.',
    '[12] 释圣严. 天台思想论集[M]. 台北: 法鼓文化, 2002.',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    set_run_font(run, '宋体', 12, False)
    set_para_spacing(p, line_spacing=20, before=0, after=0)

# Save
output_path = '/home/admin/.openclaw/workspace/开题报告-天台四教四谛.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
