#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""为小论文添加脚注并生成新的docx文件"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name='宋体', size=12, bold=False, color=None):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color

def add_footnote(paragraph, text):
    """在段落末尾添加脚注"""
    run = paragraph.add_run()
    footnote_ref = run.add_footnote_ref()
    footnote = footnote_ref.add_footnote()
    footnote_ref_run = footnote.paragraphs[0].add_run(text)
    set_font(footnote_ref_run, '宋体', 9)
    return paragraph

def main():
    doc = Document()
    
    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ========== 标题 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('从"以儒释耶"到"文化润教"')
    set_font(run, '宋体', 22, True)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('——吴雷川的诠释实践与当代镜鉴')
    set_font(run, '宋体', 16, True)
    
    # ========== 作者 ==========
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run('蔡余杭')
    set_font(run, '宋体', 14)
    
    doc.add_paragraph()  # 空行
    
    # ========== 摘要 ==========
    abstract_label = doc.add_paragraph()
    run = abstract_label.add_run('摘 要：')
    set_font(run, '宋体', 12, True)
    
    abstract_text = ('在二十世纪上半叶中国社会剧烈转型与民族救亡图存的宏大历史背景下，'
        '基督教如何摆脱"洋教"标签、实现本土化生存，成为当时知识界与教会面临的紧迫课题。'
        '吴雷川作为兼具晚清进士功名与基督教信仰的"双通"型知识分子，其思想探索具有独特的典范意义。'
        '他不再满足于表层的文化调和，而是运用儒家经学诠释方法，对基督教神学核心进行了系统性的重构。'
        '本文重点梳理吴雷川如何将"上帝"伦理化为"天"，将"耶稣"重塑为"圣贤/革命家"，'
        '将"天国"转化为"大同世界"的诠释实践。在此基础上，文章进一步反思了这种"以儒释耶"路径'
        '在确立文化主体性方面的历史贡献，以及其在消解信仰超越性方面的深刻教训。'
        '通过对这一历史个案的深度剖析，旨在为新时代推进"文化润教"提供从"概念比附"走向"深度内化"的'
        '历史镜鉴，探索外来宗教在中华文明框架下实现创造性转化的可能路径。')
    
    p = doc.add_paragraph()
    run = p.add_run(abstract_text)
    set_font(run, '宋体', 12)
    
    # ========== 关键词 ==========
    keywords = doc.add_paragraph()
    run = keywords.add_run('关 键 词：')
    set_font(run, '宋体', 12, True)
    run = keywords.add_run('吴雷川；以儒释耶；文化润教；基督教中国化')
    set_font(run, '宋体', 12)
    
    # ========== 作者简介 ==========
    bio = doc.add_paragraph()
    run = bio.add_run('作者简介：')
    set_font(run, '宋体', 10, True)
    run = bio.add_run('蔡余杭、玉环市基督教协会副会长')
    set_font(run, '宋体', 10)
    
    doc.add_paragraph()  # 空行
    
    # ========== 引言 ==========
    intro_title = doc.add_paragraph()
    run = intro_title.add_run('引 言')
    set_font(run, '宋体', 16, True)
    intro_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # 引言正文
    intro_text_1 = (
        '在新时代推进"宗教中国化"的理论构建与实践探索中，"以中华优秀传统文化浸润宗教"已然超越了一个单纯的政策口号，'
        '而成为关乎宗教在中国社会能否"行得远、立得住"的核心命题。这一命题的深层逻辑，在于解决外来宗教与本土文明之间的'
        '"水土不服"问题，寻求信仰真理与文化基因的内在契合。回溯历史，二十世纪二三十年代的中国基督教"本色化运动"，'
        '正是这一命题的早期大规模实践场域。'
    )
    p = doc.add_paragraph()
    run = p.add_run(intro_text_1)
    set_font(run, '宋体', 12)
    
    intro_text_2 = (
        '在众多寻求基督教本土化的先贤中，吴雷川（1870—1944）¹的身影尤为卓著，他不仅是燕京大学的首任华人校长²，'
        '更是一位拥有深厚儒家经学造诣的前清翰林。这种独特的"士大夫—基督徒"双重身份，使得他的神学思考迥异于那些具有'
        '西方留学背景的同侪，也区别于草根教会的自发适应。面对"非基督教运动"将基督教斥为"帝国主义侵略先锋"³的严峻挑战，'
        '吴雷川并未采取防御性的护教姿态，而是以一种高度的文化自信与主体自觉，提出了"以儒释耶"的诠释路径。'
        '他试图通过重构基督教的核心教义，证明基督教不仅不与中国文化相悖，反而是儒家"道统"在新时代的延续与补充。'
        '因此，重审吴雷川如何运用孟子"知人论世"与"以意逆志"的方法解读《圣经》，如何在"道一教殊"的框架下处理耶儒关系，'
        '对于我们理解"文化润教"的深层机制——即如何在保持信仰特质的同时确立文化主体性——具有不可替代的学术价值与现实镜鉴意义。'
    )
    p = doc.add_paragraph()
    run = p.add_run(intro_text_2)
    set_font(run, '宋体', 12)
    
    doc.add_paragraph()  # 空行
    
    # ========== 第一章 ==========
    ch1_title = doc.add_paragraph()
    run = ch1_title.add_run('一、"上帝"与"天"：本体论的伦理化融合')
    set_font(run, '宋体', 14, True)
    
    ch1_text = (
        '在神学本体论层面，吴雷川面临的首要难题是如何处理基督教"耶和华"这一位具有强烈位格性、超越性及民族排他性的'
        '"上帝"，与中国传统语境中或为"主宰"或为"自然法则"的"天"之间的巨大张力。吴雷川的解决策略是一种双向的'
        '"理性化祛魅"与"伦理化重构"。'
        '首先，他对基督教的上帝观进行了"去超验化"处理。受当时自由主义神学与科学理性思潮的影响，吴雷川倾向于剥离上帝身上'
        '那些"不合科学"的神秘主义色彩。他将上帝的核心属性提炼为"父"、"公义"、"善"与"灵"⁴。在他看来，上帝并非高高在上、'
        '喜怒无常的人格主宰，而是一种普世性的伦理原则。例如，他将"上帝是父"解释为"人类公共之父"，从而消解了犹太教传统中'
        '上帝仅是以色列人的"选民之神"的狭隘性，使其与儒家"天无私覆"⁵的博爱精神相契合。'
        '其次，他对中国传统的"天"进行了"伦理化"提升。吴雷川指出，儒家经典中的"天"既有"意志之天"（如赏善罚恶的主宰），'
        '也有"自然之天"（如四时行焉的规律）。他巧妙地将《中庸》的"诚"作为连接点，认为《中庸》所言"诚者，天之道也"，'
        '正是对上帝本质的最好注脚。他直言："我以为上帝就是和真理、大自然、最高的原则相等的一种名称"⁶。'
        '在这里，上帝不再是那个在西奈山颁布律法的立法者，而是内化为宇宙运行的最高道德律则——"天道"。'
        '这种诠释实践的得失是显而易见的。其得在于，它极大地降低了基督教与中国知识分子之间的文化隔阂，从本体论上论证了'
        '基督教的"合法性"，使其不再是异质的"洋神"；其失则在于，这种彻底的伦理化几乎完全消解了上帝的位格性与超越性。'
        '当"神人关系"被置换为"天人关系"，基督教原本的"启示"与"救赎"逻辑，便不可避免地滑向了儒家的"修身"与"尽性"逻辑。'
    )
    p = doc.add_paragraph()
    run = p.add_run(ch1_text)
    set_font(run, '宋体', 12)
    
    doc.add_paragraph()  # 空行
    
    # ========== 第二章 ==========
    ch2_title = doc.add_paragraph()
    run = ch2_title.add_run('二、"耶稣"与"圣贤"：基督论的人格重塑与革命化转向')
    set_font(run, '宋体', 14, True)
    
    ch2_text = (
        '如果说在上帝论上吴雷川采取了"温和改良"的路线，那么在基督论上，他的诠释则显得更为激进与彻底。'
        '他运用孟子"知人论世"与"以意逆志"的经学方法，对耶稣的形象进行了颠覆性的重塑，将其从"救世主"改造为"社会革命家"。'
    )
    p = doc.add_paragraph()
    run = p.add_run(ch2_text)
    set_font(run, '宋体', 12)
    
    # 2.1 祛魅与还原
    sub2_1 = doc.add_paragraph()
    run = sub2_1.add_run('1. 祛魅与还原：从"神子"到"完人"')
    set_font(run, '宋体', 12, True)
    
    sub2_1_text = (
        '吴雷川明确表示，对于福音书中关于耶稣童贞女怀孕、水上行走等神迹奇事，即便存在，也并非信仰的核心。'
        '他认为这些神迹记载是后世门徒为了迁就当时人们的迷信心理而附加的，正如儒家经典中对圣人的神化一样。'
        '他主张"去其所蔽"，直探耶稣的人格本质。在他看来，耶稣首先是一个"人"，一个道德完满、具有伟大人格的"圣贤"。'
        '这种"重人伦、轻神迹"的立场，使得耶稣从神坛走向人间，成为中国人可以效法的"榜样"。'
    )
    p = doc.add_paragraph()
    run = p.add_run(sub2_1_text)
    set_font(run, '宋体', 12)
    
    # 2.2 比附与重构
    sub2_2 = doc.add_paragraph()
    run = sub2_2.add_run('2. 比附与重构：纳入"三不朽"谱系')
    set_font(run, '宋体', 12, True)
    
    sub2_2_text = (
        '吴雷川将耶稣的生平与教导纳入儒家"立德、立功、立言"的"三不朽"框架中进行评价⁷。'
        '他将耶稣的登山宝训等言论视为"立言"，将耶稣的谦卑、爱人等品德视为"立德"，'
        '而将耶稣宣传天国、挑战法利赛人权威的行为视为"立功"。通过这种比附，耶稣被成功纳入了中国士大夫所尊崇的圣贤谱系。'
    )
    p = doc.add_paragraph()
    run = p.add_run(sub2_2_text)
    set_font(run, '宋体', 12)
    
    # 2.3 激进的革命化诠释
    sub2_3 = doc.add_paragraph()
    run = sub2_3.add_run('3. 激进的革命化诠释')
    set_font(run, '宋体', 12, True)
    
    sub2_3_text = (
        '随着三十年代民族危机的加深，吴雷川对耶稣的诠释进一步向"革命家"倾斜。'
        '他运用"知人论世"的方法，深入分析犹太民族受罗马帝国压迫的历史背景，认为耶稣是一位怀有强烈民族解放情怀的领袖，'
        '他将耶稣的"登山宝训"解读为一种社会革命的纲领，将耶稣的受难解读为儒家式的"杀身成仁、舍生取义"⁸。'
        '在吴雷川笔下，耶稣不再是一个温良恭俭让的救赎主，而是一个为了改造社会、实现公义而不惜牺牲生命的"革命先锋"。'
        '这种诠释，实际上是以儒家的"内圣外王"之道，重塑了基督教的基督论。'
    )
    p = doc.add_paragraph()
    run = p.add_run(sub2_3_text)
    set_font(run, '宋体', 12)
    
    doc.add_paragraph()  # 空行
    
    # ========== 第三章 ==========
    ch3_title = doc.add_paragraph()
    run = ch3_title.add_run('三、"天国"与"大同"：末世论的历史内转化')
    set_font(run, '宋体', 14, True)
    
    ch3_text = (
        '在末世论层面，吴雷川的诠释最为彻底地体现了其"经世致用"的现实关怀。'
        '他坚决反对教会传统中"死后升天堂"的彼岸式盼望，而是将"天国"完全此世化、政治化，'
        '将其等同于儒家《礼记·礼运》篇中的"大同世界"。'
    )
    p = doc.add_paragraph()
    run = p.add_run(ch3_text)
    set_font(run, '宋体', 12)
    
    # 3.1 空间的此世化
    sub3_1 = doc.add_paragraph()
    run = sub3_1.add_run('1. 空间的此世化')
    set_font(run, '宋体', 12, True)
    
    sub3_1_text = (
        '天国就在人间。吴雷川强调，"天国并不是在这世界之外另有一个世界"⁹，也不是死后灵魂的归宿。'
        '他解释道："愿你的国降临"，本质上就是"愿改造社会的工作顺利进行"，天国就是"将这世界上所有不合仁爱和公义的事全都除去"，'
        '建立一个充满公义与仁爱的新社会。这种诠释，将基督教的盼望从"天上"拉回"地上"，从"来世"拉回"今世"。'
    )
    p = doc.add_paragraph()
    run = p.add_run(sub3_1_text)
    set_font(run, '宋体', 12)
    
    # 3.2 理想的同构化
    sub3_2 = doc.add_paragraph()
    run = sub3_2.add_run('2. 理想的同构化')
    set_font(run, '宋体', 12, True)
    
    sub3_2_text = (
        '吴雷川在《基督教与中国文化》中，系统地将耶稣关于天国的五类训言（如废除私有、凡物公用、爱仇敌等）'
        '与儒家"大同"理想中的"天下为公"、"选贤与能"、"讲信修睦"进行了逐条对勘。'
        '他认为，耶儒两家虽然教法不同，但对理想社会的终极追求是一致的，都是要建立一个没有剥削、没有压迫、人人平等的公义社会。'
    )
    p = doc.add_paragraph()
    run = p.add_run(sub3_2_text)
    set_font(run, '宋体', 12)
    
    # 3.3 手段的革命化
    sub3_3 = doc.add_paragraph()
    run = sub3_3.add_run('3. 手段的革命化')
    set_font(run, '宋体', 12, True)
    
    sub3_3_text = (
        '为了实现这一"天国/大同"，吴雷川在晚年甚至接受了某种形式的暴力革命理论。'
        '他引用耶稣"我来不是叫地上太平，乃是叫人纷争"的话，认为在新旧社会交替之际，革命是不可避免的阵痛。'
        '他将基督教的社会理想与当时流行的社会主义思潮相结合，认为建立天国必须"变更旧有组织"¹⁰，'
        '甚至主张在必要时采取强力手段。这种将宗教末世论转化为政治革命论的尝试，虽然在神学上极具争议，'
        '但在当时却极大地回应了中国社会救亡图存的迫切需求。'
    )
    p = doc.add_paragraph()
    run = p.add_run(sub3_3_text)
    set_font(run, '宋体', 12)
    
    doc.add_paragraph()  # 空行
    
    # ========== 第四章 ==========
    ch4_title = doc.add_paragraph()
    run = ch4_title.add_run('四、反思与启示：对今日"文化润教"的审视')
    set_font(run, '宋体', 14, True)
    
    ch4_intro = (
        '吴雷川的探索是一笔复杂的遗产，他既为基督教中国化提供了宝贵的经验，也留下了深刻的教训。'
        '在新时代推进"文化润教"的今天，我们应当如何汲取其智慧，规避其陷阱？这需要我们从以下三个层面进行深度反思。'
    )
    p = doc.add_paragraph()
    run = p.add_run(ch4_intro)
    set_font(run, '宋体', 12)
    
    # 4.1 根基
    sub4_1 = doc.add_paragraph()
    run = sub4_1.add_run('（一）根基：确立文化主体性，实现从"被动适应"到"深度内化"的跨越')
    set_font(run, '宋体', 12, True)
    
    sub4_1_text = (
        '吴雷川实践的最大贡献在于他确立了鲜明的"中国文化主体性"。与当时许多教会人士盲目崇拜西方神学、'
        '视中国文化为"罪中文化"不同，吴雷川作为儒家道统的天然继承人，他拥有强大的文化自信，'
        '他不是在西方神学的框架内寻找儒家的"预备"，而是以儒家为标准去评判和重构基督教。'
        '这对于今日的"文化润教"具有根本性的启示：真正的文化润教，绝非表面的"符号嫁接"'
        '（如在教堂建筑上加飞檐、在赞美诗中加二胡），而是深层的"精神内化"。'
        '如果宗教界人士自身对中华优秀传统文化缺乏深度的认同与理解，仅仅将文化润教视为一种外在的政治任务，'
        '那么这种"润"必然是肤浅的、甚至是生硬的。新时代的宗教人才应当像吴雷川那样，具备深厚的国学功底，'
        '能够自觉地运用中华文化的思维方式去理解信仰、讲经布道。只有当宗教信仰能够用中华文化的"母语"自然表达出来时，'
        '才能真正融入信众的血脉。'
    )
    p = doc.add_paragraph()
    run = p.add_run(sub4_1_text)
    set_font(run, '宋体', 12)
    
    # 4.2 警醒
    sub4_2 = doc.add_paragraph()
    run = sub4_2.add_run('（二）警醒：警惕"工具化"陷阱，坚守信仰的超越性特质')
    set_font(run, '宋体', 12, True)
    
    sub4_2_text = (
        '然而，吴雷川的尝试也为我们敲响了警钟。为了回应时代的救亡焦虑，他过度强调基督教的"社会功能"与"政治效用"，'
        '导致基督教的许多核心特质被严重稀释。首先是消解神圣为世俗，吴雷川将"上帝"还原为"天道/诚"，'
        '将"圣灵"等同于"仁"，这种做法虽然打通了耶儒的壁垒，但也使得基督教失去了其独特的神圣源头。'
        '赵紫宸曾尖锐地批评吴雷川的基督教是"不拜神、不敬神的宗教"¹¹，这并非全无道理。'
        '其次是消解信仰为伦理，当耶稣被仅仅视为"圣贤"或"革命家"，当"因信称义"被置换为"因德成圣"，'
        '基督教便退化为一种单纯的道德哲学或社会运动，失去了其作为宗教的超越维度。'
        '这提醒我们，在推进"文化润教"时，必须警惕将宗教彻底"工具化"的倾向。'
        '我们不能为了追求外在的适应，而牺牲了基督宗教本身的神圣性与超越性。'
        '文化润教应当是"双向"的：既要用中华文化滋养宗教，也要保留宗教对社会的批判与超越功能。'
        '如果完全抹平了宗教的独特性，使其成为世俗意识形态的附庸，那么"润教"的结果可能不是"共融"，而是"消解"。'
    )
    p = doc.add_paragraph()
    run = p.add_run(sub4_2_text)
    set_font(run, '宋体', 12)
    
    # 4.3 归宿
    sub4_3 = doc.add_paragraph()
    run = sub4_3.add_run('（三）归宿：从"价值共鸣"到"社会担当"的实践升华')
    set_font(run, '宋体', 12, True)
    
    sub4_3_text = (
        '吴雷川最可贵之处在于他强烈的"经世"情怀。他始终认为，考证中西经典、推动耶儒会通，'
        '最终目的都是为了"有益于中国"。他的神学思考始终锚定在"中国向何处去"这一时代之问上。'
        '这对今天的启示在于，新时代的"文化润教"不能仅停留在神学理论的建构或礼仪形式的改良上，'
        '而必须回应现实的社会关切。基督教中国化的深度，取决于其参与社会建设的程度。'
        '在此基础上我们应当挖掘基督教教义中与社会主义核心价值观相契合的内容，例如，'
        '将基督教的"爱邻如己"与社会主义的"友善"相结合，将基督教的"管家职分"与生态文明建设相结合。'
        '同时应该强化社会服务，教会应当引导信众在扶贫济困、养老助残、公益慈善等领域发挥积极作用。'
        '当基督徒在社会中展现出诚实守信、爱国敬业的精神风貌时，当教会成为促进社会和谐的积极力量时，'
        '文化润教才真正实现了其"润物无声"的最高境界。'
    )
    p = doc.add_paragraph()
    run = p.add_run(sub4_3_text)
    set_font(run, '宋体', 12)
    
    doc.add_paragraph()  # 空行
    
    # ========== 结语 ==========
    conclusion_title = doc.add_paragraph()
    run = conclusion_title.add_run('结 语')
    set_font(run, '宋体', 16, True)
    conclusion_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    conclusion_text = (
        '吴雷川的一生，是"士志于道"与"基督信仰"交织的一生。他以"以儒释耶"的大胆实践，'
        '为基督教在中国的生存与发展开辟了一条独特的诠释学路径。他试图在儒家的"道统"中安顿基督教的"信仰"，'
        '在"经世致用"中寻找宗教的"合法性"。站在新的历史方位上，我们推进"文化润教"，'
        '既要继承吴雷川那种"以中国为重心"的文化主体自觉，也要反思他因时代局限而陷入的'
        '"过度政治化"与"伦理化"困境。我们应当追求的，是一种既保持信仰纯正性、又富有中华文化神韵的健康宗教生态。'
        '让基督教不仅在形式上，更在精神实质上，成为建设中华民族现代文明的有机组成部分。'
        '这既是对吴雷川等先贤最好的致敬，也是新时代宗教工作的必然使命。'
    )
    p = doc.add_paragraph()
    run = p.add_run(conclusion_text)
    set_font(run, '宋体', 12)
    
    doc.add_paragraph()  # 空行
    
    # ========== 参考文献 ==========
    ref_title = doc.add_paragraph()
    run = ref_title.add_run('参考文献')
    set_font(run, '宋体', 14, True)
    ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    refs = [
        '[1]《中共中央关于进一步全面深化改革、推进中国式现代化的决定》，《人民日报》2024年7月22日。',
        '[2]张志刚：《构建宗教中国化自主学术话语》，《世界宗教文化》2026年第2期。',
        '[3]李韦：《吴雷川的基督教处境化思想研究》，宗教文化出版社，2010年。',
        '[4]吴雷川：《基督教与中国文化》，商务印书馆，2017年。',
        '[5]吴雷川：《墨翟与耶稣》，青年协会书局，1940年。',
        '[6]赵紫宸：《耶稣为基督——评吴雷川先生之〈基督教与中国文化〉》，载《赵紫宸文集》第三卷，商务印书馆，2007年。',
        '[7]梁慧：《从中国经学传统出发诠释〈圣经〉：吴雷川解读〈圣经〉的立场和方法》，《世界宗教研究》，2013年第2期。',
        '[8]郭清香：《耶儒伦理比较研究：民国时期基督教与儒教伦理思想的冲突与融合》，中国社会科学出版社，2006年。',
    ]
    
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        set_font(run, '宋体', 10)
    
    # 保存文档
    output_path = '/home/admin/.openclaw/workspace/小论文2_带脚注.docx'
    doc.save(output_path)
    print(f'文档已保存至: {output_path}')

if __name__ == '__main__':
    main()
