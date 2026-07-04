#!/usr/bin/env python3
"""Generate a formatted thesis document based on the 浙江省宗教界"双通"人才研修班 template."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re
import sys

def set_cell_font(cell, text, font_name='宋体', font_size=Pt(10.5), bold=False, alignment=None):
    """Set cell text with proper Chinese font."""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml.etree import SubElement
        rFonts = SubElement(rpr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def add_cover_page(doc):
    """Add the cover page with table layout."""
    # Add spacing before the table
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    
    # Title block - repeated title
    title_text = '浙江省宗教界"双通"人才研修班\n毕业论文'
    for i, line in enumerate(title_text.split('\n')):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = '黑体'
        run.font.size = Pt(22)
        run.bold = True
        set_east_asia_font(run, '黑体')
        p.paragraph_format.space_after = Pt(6)
    
    # Paper title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run('新时代道教神仙信仰的传承与中国化路径')
    run.font.name = '黑体'
    run.font.size = Pt(18)
    run.bold = True
    set_east_asia_font(run, '黑体')
    
    # Info table
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    # Row 0: 分类号 / 密级
    set_cell_font(table.cell(0, 0), '分类号:')
    set_cell_font(table.cell(0, 1), '')
    set_cell_font(table.cell(0, 2), '密    级:')
    set_cell_font(table.cell(0, 3), '')
    
    # Row 1: UDC / 编号
    set_cell_font(table.cell(1, 0), 'U D C:')
    set_cell_font(table.cell(1, 1), '')
    set_cell_font(table.cell(1, 2), '编    号:')
    set_cell_font(table.cell(1, 3), '')
    
    # Merge empty cells for labels
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 2).merge(table.cell(0, 3))
    table.cell(1, 0).merge(table.cell(1, 1))
    table.cell(1, 2).merge(table.cell(1, 3))
    
    # Row 2: 作者姓名
    table.cell(2, 0).merge(table.cell(2, 1))
    set_cell_font(table.cell(2, 0), '作    者    姓    名：', bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    table.cell(2, 2).merge(table.cell(2, 3))
    set_cell_font(table.cell(2, 2), '戴泰信')
    
    # Row 3: 学号
    table.cell(3, 0).merge(table.cell(3, 1))
    set_cell_font(table.cell(3, 0), '学    号：', bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    table.cell(3, 2).merge(table.cell(3, 3))
    set_cell_font(table.cell(3, 2), 'XXX')
    
    # Row 4: 指导教师
    table.cell(4, 0).merge(table.cell(4, 1))
    set_cell_font(table.cell(4, 0), '指    导    教    师：', bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    table.cell(4, 2).merge(table.cell(4, 3))
    set_cell_font(table.cell(4, 2), 'XXX')
    
    # Add more rows for 宗教 and 研究方向
    row5 = table.add_row()
    row5.cells[0].merge(row5.cells[1])
    set_cell_font(row5.cells[0], '宗    教：', bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    row5.cells[2].merge(row5.cells[3])
    set_cell_font(row5.cells[2], '道教')
    
    row6 = table.add_row()
    row6.cells[0].merge(row6.cells[1])
    set_cell_font(row6.cells[0], '研    究    方    向：', bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    row6.cells[2].merge(row6.cells[3])
    set_cell_font(row6.cells[2], 'XXX')
    
    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run('2026 年 5 月')
    run.font.name = '宋体'
    run.font.size = Pt(14)
    set_east_asia_font(run, '宋体')
    
    doc.add_page_break()

def add_abstract(doc):
    """Add Chinese and English abstract."""
    # Chinese abstract title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('摘  要')
    run.font.name = '黑体'
    run.font.size = Pt(16)
    run.bold = True
    set_east_asia_font(run, '黑体')
    
    # Empty line
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Chinese abstract content
    abstract_cn = (
        '新时代坚持宗教中国化方向，是以习近平同志为核心的党中央关于宗教工作的重大战略部署。'
        '道教作为中华大地上唯一土生土长的宗教，其中国化进程具有特殊的历史地位与现实意义。'
        '神仙信仰是道教的核心标志与信仰基石，神仙谱系的历史演变即是道教发展史的缩影。'
        '然而，在新时代社会深刻变革的背景下，道教神仙信仰的传承面临信仰理解浅表化、传承主体断层、'
        '科仪商业化倾向以及与当代价值体系对接不足等突出问题。如何在保持信仰本质的同时实现与新时代的适应与融合，'
        '成为道教神仙信仰传承发展亟待破解的核心命题。'
        '本文以"新时代道教神仙信仰的传承与中国化路径"为研究主题，综合运用经典文献分析法、政策文本解读法与案例研究法，'
        '从历史纵深、现实境遇、理论逻辑与实践路径四个维度展开系统研究。'
        '全文核心论点为：道教神仙信仰的中国化并非对传统的否定或消解，而是在"返本开新"中实现信仰的创造性转化与创新性发展。'
        '所谓"返本"，即回归"仙道贵生，无量度人"的信仰本源；所谓"开新"，即在新时代语境下重构神仙信仰的时代内涵、'
        '修持实践、组织形态与传播方式。'
        '研究首先梳理道教神仙信仰从先秦思想渊源到当代存续形态的历史演变脉络，揭示神仙信仰"随方设教"的自我调适传统；'
        '其次考察新时代神仙信仰传承面临的现实困境，论证中国化的必要性；'
        '再次从历史逻辑、经典依据和现实基础三个层面阐释神仙信仰中国化的理论合法性；'
        '最终构建以教义阐释中国化为灵魂、修持实践中国化为关键、组织形态中国化为基础、传播方式中国化为手段的系统性实践路径。'
        '本文的创新之处在于：首次将道教神仙信仰研究与中国化路径问题系统对接，填补了该交叉领域的研究空白；'
        '提出了神仙信仰中国化的"返本开新"理论框架；并以典型案例为支撑，为道教界推进中国化实践提供了学术参考。'
    )
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(abstract_cn)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    set_east_asia_font(run, '宋体')
    
    # Keywords
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run_b = p.add_run('关键词：')
    run_b.bold = True
    run_b.font.name = '宋体'
    run_b.font.size = Pt(12)
    set_east_asia_font(run_b, '宋体')
    run_k = p.add_run('道教神仙信仰；中国化；传承；新时代；返本开新；现代化转型')
    run_k.font.name = '宋体'
    run_k.font.size = Pt(12)
    set_east_asia_font(run_k, '宋体')
    
    doc.add_page_break()
    
    # English abstract
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('Abstract')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    abstract_en = (
        'In the new era, adhering to the Sinicization of religion is a major strategic deployment of the Party Central Committee '
        'with Comrade Xi Jinping at its core regarding religious work. As the only indigenous religion in China, Daoism holds a '
        'special historical position and practical significance in the Sinicization process. Immortal信仰 (Shenxian belief) is the '
        'core symbol and foundation of Daoist faith, and the historical evolution of the immortal pantheon is a microcosm of Daoist '
        'development history. However, against the backdrop of profound social transformation in the new era, the inheritance of '
        'Daoist immortal belief faces prominent challenges including superficial understanding, generational gaps in inheritance '
        'subjects, commercialization tendencies in rituals, and insufficient integration with contemporary value systems.'
        'This paper systematically examines the inheritance and Sinicization path of Daoist immortal belief from four dimensions: '
        'historical depth, contemporary reality, theoretical logic, and practical pathways. The core argument is that the Sinicization '
        'of Daoist immortal belief is not a negation of tradition, but rather a creative transformation and innovative development '
        'through "returning to the roots and opening up the new" (返本开新).'
    )
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(abstract_en)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run_b = p.add_run('Key Words: ')
    run_b.bold = True
    run_b.font.name = 'Times New Roman'
    run_b.font.size = Pt(12)
    run_k = p.add_run('Daoist Immortal Belief; Sinicization; Inheritance; New Era; Returning to Roots and Opening Up the New; Modernization Transformation')
    run_k.font.name = 'Times New Roman'
    run_k.font.size = Pt(12)
    
    doc.add_page_break()

def add_toc(doc):
    """Add table of contents placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('目  录')
    run.font.name = '黑体'
    run.font.size = Pt(16)
    run.bold = True
    set_east_asia_font(run, '黑体')
    
    toc_items = [
        ('摘  要', 'I'),
        ('Abstract', 'II'),
        ('引  言', '1'),
        ('第一章 道教神仙信仰的历史演变与核心内涵', '3'),
        ('  第一节 神仙信仰的思想渊源', '3'),
        ('  第二节 神仙信仰的历史演变', '5'),
        ('  第三节 神仙信仰的核心内涵', '8'),
        ('第二章 新时代道教神仙信仰传承的现实境遇', '10'),
        ('  第一节 新时代社会变迁对神仙信仰的冲击', '10'),
        ('  第二节 神仙信仰传承面临的突出问题', '12'),
        ('  第三节 神仙信仰当代存续形态的考察', '14'),
        ('第三章 道教神仙信仰中国化的理论依据与内在逻辑', '16'),
        ('  第一节 "中国化"的理论内涵', '16'),
        ('  第二节 神仙信仰中国化的历史逻辑', '18'),
        ('  第三节 神仙信仰中国化的经典依据', '20'),
        ('  第四节 神仙信仰中国化的现实依据', '22'),
        ('第四章 新时代道教神仙信仰中国化的实践路径', '24'),
        ('  第一节 教义阐释的中国化', '24'),
        ('  第二节 修持实践的中国化', '27'),
        ('  第三节 组织形态的中国化', '29'),
        ('  第四节 传播方式的中国化', '31'),
        ('  第五节 典型案例分析', '33'),
        ('第五章 神仙信仰传承与中国化的辩证关系与未来展望', '35'),
        ('  第一节 传承与创新的辩证统一', '35'),
        ('  第二节 神仙信仰中国化的基本原则', '37'),
        ('  第三节 未来展望', '38'),
        ('结  语', '40'),
        ('参考文献', '42'),
        ('附  录', '44'),
        ('  一、致  谢', '44'),
        ('  二、论文原创性声明', '45'),
        ('  三、附  件', '46'),
    ]
    
    for title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(title + ' ................................ ' + page)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        set_east_asia_font(run, '宋体')
        if '第' in title and ('章' in title or '目' in title):
            run.bold = True
    
    doc.add_page_break()

def set_para_style(p, font_name='宋体', font_size=Pt(12), line_spacing=1.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=0, first_indent=0):
    """Set paragraph style."""
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = alignment
    if first_indent > 0:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    for run in p.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_chapter_heading(doc, text, level=1):
    """Add chapter/section heading."""
    p = doc.add_paragraph()
    if level == 1:  # 章标题
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(16)
        run.bold = True
        set_east_asia_font(run, '黑体')
    elif level == 2:  # 节标题
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(14)
        run.bold = True
        set_east_asia_font(run, '黑体')
    elif level == 3:  # 一级标题（一、二、三）
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(12)
        run.bold = True
        set_east_asia_font(run, '黑体')
    elif level == 4:  # 二级标题（（一）（二））
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(12)
        run.bold = True
        set_east_asia_font(run, '黑体')

def set_east_asia_font(run, font_name='宋体'):
    """Set eastAsian font attribute properly."""
    from lxml.etree import SubElement
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = SubElement(rpr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)

def add_body_text(doc, text):
    """Add body text paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)  # 约两个字符
    
    # Handle inline bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = '宋体'
            run.font.size = Pt(12)
            set_east_asia_font(run, '宋体')
        elif part:
            run = p.add_run(part)
            run.font.name = '宋体'
            run.font.size = Pt(12)
            set_east_asia_font(run, '宋体')

def add_introduction(doc):
    """Add introduction section."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run('引  言')
    run.font.name = '黑体'
    run.font.size = Pt(16)
    run.bold = True
    set_east_asia_font(run, '黑体')
    
    # Introduction content - condensed version
    intro_text = (
        '党的十八大以来，以习近平同志为核心的党中央高度重视宗教工作，明确提出坚持我国宗教中国化方向的重大战略部署。'
        '2016年全国宗教工作会议强调，积极引导宗教与社会主义社会相适应，坚持我国宗教中国化方向。'
        '2021年全国宗教工作会议进一步指出，要深入推进我国宗教中国化，引导宗教更好与社会主义社会相适应。'
        '党的二十大报告再次强调"坚持我国宗教中国化方向，积极引导宗教与社会主义社会相适应"。'
        '这一系列重要论述为新时代宗教工作提供了根本遵循，也为道教的发展指明了方向。'
    )
    add_body_text(doc, intro_text)
    
    intro_text2 = (
        '道教作为中华大地上唯一土生土长的宗教，其中国化进程具有不同于其他宗教的特殊性。'
        '一方面，道教天然具备中国化的先天优势——从地域性和民族性上看，道教是"中国的宗教"而非"外国宗教在中国"，'
        '其思想渊源直抵中华文明的深层结构，诚如鲁迅所言"中国根柢全在道教"。'
        '另一方面，道教中国化的核心任务并非"本土化"——这一任务已在漫长的历史进程中基本完成——而是"时代化"，'
        '即在新时代条件下实现道教的现代化转型与可持续发展。'
    )
    add_body_text(doc, intro_text2)
    
    intro_text3 = (
        '神仙信仰是道教的核心标志与信仰基石。盖建民在论及道教研究的未来方向时，明确提出要"强化道教神仙信仰以及道教教义学研究"。'
        '李远国在其巨著《中国道教神仙谱系史》中以四卷本五百六十余万字的篇幅论证了一个核心命题：神仙信仰是道教的核心标志，'
        '神仙谱系的历史演变即是道教发展史的缩影。然而，长期以来，道教神仙信仰研究与中国化研究基本处于两条平行线——'
        '前者偏重历史与义理考察，后者侧重政策解读与实践探索，将二者系统结合的专题研究几乎付之阙如。'
    )
    add_body_text(doc, intro_text3)
    
    intro_text4 = (
        '在新时代背景下，道教神仙信仰的传承面临深刻挑战：社会高速发展而信仰阐释相对滞后，科学理性主义对神仙信仰合理性构成冲击，'
        '青年道众与高道人才匮乏导致传承主体断层，宫观科仪的商业化倾向使信仰本质日渐疏离，'
        '神仙信仰与社会主义核心价值观的对接尚不充分。这些问题的实质是传统信仰形态与现代社会需求之间的张力。'
        '如何在保持信仰核心不变的前提下实现神仙信仰的创造性转化与创新性发展，使之更好地与社会主义社会相适应，'
        '成为道教神仙信仰传承发展亟待破解的核心命题。'
    )
    add_body_text(doc, intro_text4)
    
    # Research questions
    add_chapter_heading(doc, '一、研究背景与问题提出', level=3)
    
    rq_text = (
        '基于以上背景，本文提出核心研究问题：新时代道教神仙信仰如何在"返本开新"中实现传承与中国化的辩证统一？'
        '具体而言，需要回答四个递进问题：神仙信仰传承的核心内容是什么？神仙信仰传承面临哪些现实困境？'
        '神仙信仰中国化有何理论依据与内在逻辑？神仙信仰中国化的具体路径是什么？'
        '本研究的意义体现在理论与实践两个层面。在理论层面，本研究有助于填补道教神仙信仰与中国化路径交叉领域的研究空白，'
        '深化对宗教中国化理论的认识，丰富道教学术研究的理论框架。在实践层面，本研究有助于为道教界推进神仙信仰中国化提供学术参考，'
        '为新时代宗教工作提供学理支撑，为道教与社会主义社会相适应提供路径指引。'
    )
    add_body_text(doc, rq_text)
    
    # Research methods
    add_chapter_heading(doc, '二、研究思路与方法', level=3)
    
    rm_text = (
        '本文遵循"传承什么—为何要变—为何能变—如何去变—变与不变的辩证"的逻辑进路，依次展开五个章节的论述。'
        '在研究方法上，本文综合运用以下方法：（一）经典文献分析法。系统解读《道德经》《庄子》《太平经》《抱朴子内篇》'
        '《真灵位业图》《云笈七签》《度人经》等道教经典，梳理神仙信仰的思想渊源与历史演变。'
        '（二）政策文本解读法。深入解读习近平总书记关于宗教工作的重要论述、全国宗教工作会议精神、'
        '《宗教事务条例》《深入推进我国道教中国化五年工作规划纲要（2023-2027年）》等政策文件。'
        '（三）案例研究法。选取武当山"系统推进道教中国化"实践、玄门讲经活动等典型案例。'
        '（四）比较研究法。适度借鉴佛教中国化的历史经验，为道教神仙信仰的中国化路径提供参照。'
    )
    add_body_text(doc, rm_text)
    
    # Core concepts
    add_chapter_heading(doc, '三、核心概念界定', level=3)
    
    cc_text = (
        '（一）"新时代"：本文所指"新时代"，即中国特色社会主义新时代，是以习近平同志为核心的党中央领导下，'
        '中国发展进入新的历史方位的时代。（二）"神仙信仰"：本文所指"神仙信仰"，是道教以神仙为核心对象的信仰体系，'
        '包括对神仙存在（神学层面）、神仙谱系（组织层面）、修仙方法（实践层面）和成仙理想（价值层面）的完整信仰系统。'
        '（三）"中国化"：宗教中国化包含三个维度：地域性、民族性、时代性。对于道教而言，中国化的核心不是"本土化"（已基本完成），'
        '而是"时代化"。（四）"返本开新"：本文的核心理论框架。"返本"即回归"仙道贵生，无量度人"的信仰本源；'
        '"开新"即在新时代语境下重构神仙信仰的时代内涵与实践形态。'
    )
    add_body_text(doc, cc_text)

def add_chapter1(doc):
    """Add Chapter 1."""
    add_chapter_heading(doc, '第一章 道教神仙信仰的历史演变与核心内涵', level=1)
    
    # Section 1
    add_chapter_heading(doc, '第一节 神仙信仰的思想渊源', level=2)
    
    add_chapter_heading(doc, '一、《道德经》与神仙信仰的逻辑起点', level=3)
    
    s1_text = (
        '《道德经》虽未直接论述"神仙"，但其阐发的"长生久视"思想构成了道教神仙信仰的核心理论渊源。'
        '第五十九章明确提出"深根固蒂、长生久视之道"，将养生与治国逻辑打通，核心指向生命长久存续的追求，'
        '为神仙信仰奠定了最根本的经典依据。'
    )
    add_body_text(doc, s1_text)
    
    s1b_text = (
        '第十章"载营魄抱一，能无离乎？专气致柔，能如婴儿乎？"强调精神魂魄合一、聚气致柔的修持状态，'
        '指向"返老还童"的成仙目标。第五十五章以"赤子"为修道理想，认为赤子般纯粹、精气充盈者具备神异属性。'
        '第五十章提出"善摄生者"的超越境界。这些论述虽然旨趣多在修身治国，但其对生命超越性境界的描摹，'
        '为后世道教神仙信仰提供了最原初的经典依据。'
    )
    add_body_text(doc, s1b_text)
    
    s1c_text = (
        '更为关键的是道教对《道德经》的神学化阐释。河上公《老子章句》重点发挥长生思想，注解第五十九章提出'
        '"人能保身中之道，使精气不劳，五神不苦，则可以长久"，被道教视为大典。'
        '《老子想尔注》则直接将"道"等同于至高天神太上老君，注第十章云"一者，道也……一散形为炁，聚形为太上老君，常治昆仑"，'
        '完成了从哲学概念到神学信仰的关键转换。张伯端《悟真篇》更明确宣示："阴符宝字逾三百，道德灵文满五千，'
        '古今上仙无限数，尽于此处达真诠"，将《老子》视为修仙核心典籍。'
    )
    add_body_text(doc, s1c_text)
    
    add_chapter_heading(doc, '二、《庄子》与神仙人格范式的确立', level=3)
    
    s2_text = (
        '《庄子》被道教尊为《南华真经》，其塑造的理想人格与修炼方法，直接构成了道教神仙信仰的人格范式与修持逻辑。'
        '庄子笔下的四类理想人格——真人、至人、神人、圣人——为道教神仙形象提供了经典原型。'
        '《大宗师》描摹真人："古之真人……登高不栗，入水不濡，入火不热……真人之息以踵，众人之息以喉"，'
        '具备超越自然伤害、无生死执念的属性。'
    )
    add_body_text(doc, s2_text)
    
    s2b_text = (
        '《齐物论》刻绘至人："至人神矣！大泽焚而不能热，河汉冱而不能寒……乘云气，骑日月，而游乎四海之外，死生无变于己"。'
        '《逍遥游》塑造神人："藐姑射之山，有神人居焉，肌肤若冰雪，绰约若处子，不食五谷，吸风饮露，乘云气，御飞龙，'
        '而游乎四海之外"，这一形象完全符合道教神仙的典型特征。《天地》篇则明确提出成仙路径："千岁厌世，去而上仙，'
        '乘彼白云，至于帝乡"。'
    )
    add_body_text(doc, s2b_text)
    
    s2c_text = (
        '在修持方法层面，《庄子》同样为道教提供了丰厚的思想资源。《达生》篇"壹其性，养其气，合其德，以通乎物之所造"，'
        '是内丹修炼的核心思想来源。《在宥》篇记载广成子告黄帝"我守其一，以处其和。故我修身千二百岁矣，吾形未常衰"，'
        '"守一"长生法由此确立。心斋与坐忘被道教赋予修炼内涵。由此，《庄子》中的理想人格与修持方法从哲学层面转化为'
        '道教神仙信仰的人格范式与修持逻辑。'
    )
    add_body_text(doc, s2c_text)
    
    add_chapter_heading(doc, '三、先秦神仙信仰与方仙道的融合', level=3)
    
    s3_text = (
        '先秦时期，神仙信仰已在燕齐沿海地区广泛流行，形成了以齐威王、齐宣王、燕昭王遣人入海求仙为标志的方仙道传统。'
        '《史记·封禅书》记载："自威、宣、燕昭使人入海求蓬莱、方丈、瀛洲"，这是中国古代帝王追求成仙的最早记载。'
        '方仙道以"形解销化，依于鬼神之事"为特征，其神仙信仰带有浓厚的神秘主义色彩，为道教神仙信仰的形成提供了社会心理基础和仪式实践资源。'
        '秦汉之际，方仙道与黄老道家逐渐合流，神仙信仰从民间方术上升为有系统教义支撑的宗教信仰体系。'
    )
    add_body_text(doc, s3_text)
    
    # Section 2
    add_chapter_heading(doc, '第二节 神仙信仰的历史演变', level=2)
    
    add_chapter_heading(doc, '一、汉魏晋时期：神仙可学论的确立', level=3)
    
    h1_text = (
        '东汉末年，道教神仙信仰实现了从零散信仰向系统教义的转化。《太平经》首次构建了神仙的等级体系，'
        '将神仙分为六等——神人、真人、仙人、道人、圣人、贤人，并明确其职能分工："神人主天，真人主地，仙人主风雨，'
        '道人主教化吉凶，圣人主治百姓，贤人辅助圣人理万民录也。"同时，《太平经》提出"守一"修仙理论，'
        '指出"古今要道，皆言守一，可长存而不老"。'
    )
    add_body_text(doc, h1_text)
    
    h1b_text = (
        '晋代葛洪《抱朴子内篇》则是道教神仙信仰的系统论证篇。《论仙》篇引《仙经》将神仙分为三等：'
        '"上士举形升虚，谓之天仙；中士游于名山，谓之地仙；下士先死后蜕，谓之尸解仙。"'
        '更重要的是，葛洪在《对俗》篇明确提出"神仙可学"论："仙之可学致，如黍稷之可播种得，甚炳然耳。"'
        '这一论断从根本上确立了神仙信仰的实践性——神仙不是先天的、命定的，而是可以通过后天修学达成的。'
    )
    add_body_text(doc, h1b_text)
    
    add_chapter_heading(doc, '二、南北朝时期：神仙谱系的系统化', level=3)
    
    h2_text = (
        '南北朝时期，道教神仙信仰实现了从散乱到系统的重大转变，其标志是陶弘景《真灵位业图》的问世。'
        '这是道教史上第一个系统的神仙谱系，将神仙分为七个阶层，每一阶层设有主神、佐神和属神，'
        '构成了一个等级分明、秩序井然的神仙世界。元始天尊居于最高层级，统御整个神仙体系。'
        '《真灵位业图》的出现标志着道教神仙信仰从零散的神灵崇拜发展为有系统的神学体系。'
    )
    add_body_text(doc, h2_text)
    
    add_chapter_heading(doc, '三、唐宋时期：从外丹向内丹的转型', level=3)
    
    h3_text = (
        '唐宋时期是道教神仙信仰发展史上的关键转折点。外丹术在唐代达到鼎盛之后，因服食丹药致死事件频发而遭受质疑，'
        '神仙信仰的实践方式遂从外丹转向内丹。这一转型的核心是：神仙信仰的实现路径从依赖外在丹药转向依靠内在心性修炼，'
        '神仙观念也从"肉体飞升"转向"精神超越"。'
    )
    add_body_text(doc, h3_text)
    
    h3b_text = (
        '全真道的兴起是这一转型的集中体现。王重阳创立全真道，明确提出"性命双修"的修持纲领，'
        '将神仙信仰的核心从肉体的长生不老转向心性的超越解脱。全真道所追求的"真仙"，不再是乘云驾鹤的超自然存在，'
        '而是心性觉悟、道德圆满的修行者。这一转型深刻地改变了道教神仙信仰的内涵。'
    )
    add_body_text(doc, h3b_text)
    
    add_chapter_heading(doc, '四、明清至今：神仙信仰的世俗化与地方化', level=3)
    
    h4_text = (
        '明清以降，道教神仙信仰呈现出明显的世俗化与地方化趋势。一方面，神仙信仰与民间信仰深度交融，'
        '大量地方性神灵被纳入道教神仙体系（如关帝、妈祖、城隍、土地等），神仙信仰的社会基础从道门内部扩展到广大民间信众。'
        '另一方面，神仙信仰的功能从超越性的终极追求逐渐向现世性的祈福禳灾偏移，神仙更多地被视为满足现实需求的超自然力量。'
        '在组织形态上，明清道教的正一、全真两大派系格局趋于稳固，正一派以符箓斋醮为主、全真派以内丹修持为主，'
        '两派在神仙信仰的侧重上也有所不同。'
    )
    add_body_text(doc, h4_text)
    
    # Section 3
    add_chapter_heading(doc, '第三节 神仙信仰的核心内涵', level=2)
    
    add_chapter_heading(doc, '一、"道—气—神"的三层结构', level=3)
    
    c1_text = (
        '詹石窗在《论道教神仙形象与易学符号之关系》中揭示了道教神仙信仰的基本结构：道教神仙信仰体现了抽象与具体的统一——'
        '"道"是抽象的，各路神仙是具体的，没有具体神仙大道无以彰显，没有抽象的"道"神仙则找不到根本。'
        '这一论述精辟地概括了"道—气—神"三层结构的内在逻辑：道为根本，是宇宙万物的本源与终极实在；'
        '气为中介，是道化生万物的机制与媒介；神为显现，是道在具体信仰层面的呈现与人格化。'
    )
    add_body_text(doc, c1_text)
    
    add_chapter_heading(doc, '二、"仙道贵生，无量度人"的基本宗教精神', level=3)
    
    c2_text = (
        '《度人经》提出的"仙道贵生，无量度人"，是道教神仙信仰的基本宗教精神。"仙道贵生"确立了生命至上的价值取向——'
        '神仙信仰的核心不是否定现世、追求来世，而是肯定生命、提升生命、超越生命的有限性。'
        '"无量度人"则将个体修仙与济度众生统一起来——神仙不是独善其身的超然存在，而是兼济天下的道德典范。'
    )
    add_body_text(doc, c2_text)
    
    add_chapter_heading(doc, '三、修仙实践与道德修养的统一', level=3)
    
    c3_text = (
        '道教神仙信仰从来不是纯粹的理论构想，而是与修仙实践紧密相连的信仰体系。更为重要的是，道教始终坚持修仙实践与道德修养的统一。'
        '《太平经》提出"积善不止"方能成仙，《抱朴子》强调"欲求仙者，要当以忠孝和顺仁信为本"，'
        '全真道更以"真功真行"为修仙的基本要求——内修"真功"（心性修炼），外修"真行"（济世利人）。'
    )
    add_body_text(doc, c3_text)
    
    # Chapter summary
    add_chapter_heading(doc, '本章小结', level=3)
    
    summary1 = (
        '道教神仙信仰从先秦思想渊源到当代存续形态，经历了一个不断演变、调适的漫长历史过程。'
        '这一过程的根本特征是"随方设教"——即对特定地域和时代具有适应性、随顺性。'
        '从这一历史考察中，可以得出两个对于本文后续论证至关重要的结论：第一，神仙信仰的核心——"道"的体悟与生命超越的追求——具有恒定性。'
        '第二，神仙信仰的形式——神仙形象、谱系结构、修持方式——具有时代性。恒定性与时代性的辩证统一，构成了神仙信仰"传承"与"创新"的内在张力。'
    )
    add_body_text(doc, summary1)

def add_chapter2(doc):
    """Add Chapter 2."""
    add_chapter_heading(doc, '第二章 新时代道教神仙信仰传承的现实境遇', level=1)
    
    # Section 1
    add_chapter_heading(doc, '第一节 新时代社会变迁对神仙信仰的冲击', level=2)
    
    add_chapter_heading(doc, '一、高速发展的社会生活与道教信仰滞后的矛盾', level=3)
    
    cs1_text = (
        '陈耀庭在对道教中国化面临的挑战进行深刻分析时指出，当代道教发展面临的第一个深层困难，就是'
        '"高速发展的社会生活和科技文化提出的各种人类生存问题，同道教百年未变的信仰滞后之间的矛盾，'
        '其后果就是道教信仰对于社会认识丧失了话语权"。这一论断切中了问题的要害。'
        '当代中国社会的现代化进程以前所未有的速度推进，而道教神仙信仰的理论阐释和信仰表达却长期停留在传统框架之内，'
        '未能有效回应时代提出的生命伦理、生态危机、精神焦虑等重大问题。'
    )
    add_body_text(doc, cs1_text)
    
    add_chapter_heading(doc, '二、科学理性主义对神仙信仰合理性的挑战', level=3)
    
    cs2_text = (
        '新时代是科学理性高扬的时代。现代科学的世界观和方法论深刻影响了社会公众的认知方式，'
        '神仙信仰中关于超越自然规律的超验描述（如白日飞升、分身变化、辟谷不食等）面临着前所未有的合理性拷问。'
        '传统的"神仙可学论"以服食丹药、修炼方术为实践路径，在现代科学视野下显然难以自圆其说；'
        '而内丹心性学的现代阐释又远远不够充分。神仙信仰的合理性危机，本质上是传统信仰的表述方式与当代认知方式之间的断裂。'
    )
    add_body_text(doc, cs2_text)
    
    add_chapter_heading(doc, '三、信息化时代信众信仰方式的变化', level=3)
    
    cs3_text = (
        '互联网和移动通讯技术的普及，深刻改变了宗教信仰的传播方式和信众的信仰行为。'
        '传统的宫观参拜、斋醮法会、师徒传承不再是信众接触神仙信仰的唯一渠道，网络空间中的道教资讯、线上祈福、虚拟社区正在重塑神仙信仰的传播生态。'
        '信息化为神仙信仰的传播提供了前所未有的广阔平台，但碎片化、浅表化的网络传播也容易导致神仙信仰的误解与曲解。'
    )
    add_body_text(doc, cs3_text)
    
    # Section 2
    add_chapter_heading(doc, '第二节 神仙信仰传承面临的突出问题', level=2)
    
    add_chapter_heading(doc, '一、信仰理解浅表化：神仙沦为功能性的祈福对象', level=3)
    
    cp1_text = (
        '神仙信仰的核心内涵在当代信仰实践中不同程度地被消解或遮蔽。大量信众对神仙的理解停留在功能性层面：'
        '拜财神求财运、拜药王求健康、拜文昌求功名、拜月老求姻缘，神仙更多地被视为满足现实需求的超自然工具，'
        '而非"道"的显现与生命超越的典范。这种功能性的信仰理解不仅弱化了神仙信仰的精神深度，'
        '也使神仙信仰与民间巫术、封建迷信的界限变得模糊。'
    )
    add_body_text(doc, cp1_text)
    
    add_chapter_heading(doc, '二、传承主体断层：青年道众与高道人才的匮乏', level=3)
    
    cp2_text = (
        '神仙信仰的传承依赖于道众的修持实践与道长的人格感召。然而，当代道教面临着严重的传承主体断层问题。'
        '一方面，青年道众的数量不足，且文化素养参差不齐，难以承担神仙信仰义理阐释和信仰引导的重任；'
        '另一方面，具有深厚修持造诣和人格感召力的高道大德稀缺，道教人才的培养速度远远落后于宫观恢复和信众增长的速度。'
    )
    add_body_text(doc, cp2_text)
    
    add_chapter_heading(doc, '三、宫观科仪的商业化倾向与信仰本质的疏离', level=3)
    
    cp3_text = (
        '近年来，部分宫观在商业化浪潮中逐渐偏离了神仙信仰的本质。烧高香、撞吉祥钟、挂牌祈福等商业性服务项目大行其道，'
        '斋醮科仪沦为明码标价的"宗教消费"，宫观成为旅游经济的组成部分而非信仰修持的清静之地。'
        '商业化倾向不仅损害了道教的社会形象，更严重地侵蚀了神仙信仰的超越性本质。'
    )
    add_body_text(doc, cp3_text)
    
    add_chapter_heading(doc, '四、神仙信仰与当代价值体系的对接不足', level=3)
    
    cp4_text = (
        '新时代以社会主义核心价值观为引领的价值体系，是全社会共同的价值基础。道教神仙信仰中蕴含的"仙道贵生""济世利人""道法自然""天人合一"等思想，'
        '与社会主义核心价值观具有内在的契合性。然而，这种契合性目前主要停留在宏观论述层面，缺乏系统的学理论证和有效的转化机制。'
    )
    add_body_text(doc, cp4_text)
    
    # Section 3
    add_chapter_heading(doc, '第三节 神仙信仰当代存续形态的考察', level=2)
    
    cs3a_text = (
        '在挑战与困境之中，道教神仙信仰仍在当代社会中以多种形态存续和发展。'
        '宫观是道教神仙信仰传承的实体空间，目前全国道教宫观约9000余处，供奉着从三清四御到地方神灵的庞大神仙体系。'
        '道教神仙在民间信仰中有着广泛而深刻的影响，关帝、妈祖、城隍、土地、灶神等道教神仙早已融入民间信仰的日常生活。'
        '互联网正在成为道教神仙信仰传播的新兴阵地，微信公众号、短视频平台、网络社区中出现了大量道教文化内容。'
        '尽管面临诸多挑战，道教神仙信仰在当代社会精神生活中仍然具有不可替代的潜在价值，'
        '神仙信仰所追求的生命超越和精神安顿，为现代人提供了一种不同于物质主义的生命观和价值观。'
    )
    add_body_text(doc, cs3a_text)
    
    # Chapter summary
    add_chapter_heading(doc, '本章小结', level=3)
    
    summary2 = (
        '新时代道教神仙信仰的传承面临深刻挑战：社会变迁对信仰的宏观冲击、信仰传承的具体困境、以及信仰存续形态的复杂分化，'
        '共同构成了神仙信仰"为什么要中国化"的现实依据。这些挑战的实质并非信仰本身的危机，而是信仰表达方式与时代需求之间的断裂。'
        '中国化不是外在强加于道教的政治要求，而是道教自身"返本开新"传统的延续，是神仙信仰在新时代条件下的"随方设教"。'
    )
    add_body_text(doc, summary2)

def add_chapter3(doc):
    """Add Chapter 3."""
    add_chapter_heading(doc, '第三章 道教神仙信仰中国化的理论依据与内在逻辑', level=1)
    
    # Section 1
    add_chapter_heading(doc, '第一节 "中国化"的理论内涵', level=2)
    
    cs1_text = (
        '宗教中国化包含三个维度：一是地域性，即宗教是中国的宗教，而不是外国宗教在中国的简单移植；'
        '二是民族性，即宗教是中华民族的宗教，体现中华民族的文化特征和精神气质；'
        '三是时代性，即宗教要与时俱进、与时偕行，与社会主义社会相适应。'
        '道教作为中华大地上唯一土生土长的宗教，其中国化具有不同于外来宗教的特殊性。'
        '对于道教来说，本土化的任务基本完成，现在主要要完成的是道教的现代化转型与可持续发展。'
        '所以坚持道教中国化方向就是在爱国爱教的基础上，以道教基本教理教义为核心，构建新时代的教义思想体系，更好地服务社会、利益人群。'
    )
    add_body_text(doc, cs1_text)
    
    # Section 2
    add_chapter_heading(doc, '第二节 神仙信仰中国化的历史逻辑', level=2)
    
    ch1_text = (
        '陈寅恪在论及道教的历史特性时，提出了"道教之真精神"这一深刻概念。卢国龙在此基础上进一步阐释，'
        '将"道教之真精神"归结为两个方面：一方面尽量吸收各种外来思想（开放精神），另一方面不忘本来民族之地位（守护民族文化主体性）。'
        '六朝隋唐时期，道教吸收佛学"缘起性空"思想方法，形成"重玄学"，正是"道教之真精神"的历史典范。'
    )
    add_body_text(doc, ch1_text)
    
    ch2_text = (
        '"随方设教"是道教神仙信仰自我调适的历史传统。道教"常道"虽大全永恒，但对"常道"的认知总有偏颇与局限，'
        '一切教化方式只能是"随方设教"——即对特定地域和时代具有适应性、随顺性。'
        '从历史考察中可以发现，神仙信仰的每一次重大转型都是"随方设教"的体现：'
        '《太平经》构建神仙等级体系是对汉代官僚体制的映射，陶弘景《真灵位业图》系统化神仙谱系是对南北朝门阀制度的对应，'
        '唐宋内丹转型是对儒释道三教融合趋势的回应，明清神仙世俗化是对民间信仰需求的吸纳。'
    )
    add_body_text(doc, ch2_text)
    
    # Section 3
    add_chapter_heading(doc, '第三节 神仙信仰中国化的经典依据', level=2)
    
    cj1_text = (
        '《道德经》第四十章云："反者道之动，弱者道之用。"此论揭示了"道"的运动的根本规律：一切事物的发展都包含着向对立面转化的趋势。'
        '"反"既意味着"返回"（复归本源），也意味着"相反"（对立转化）。这一辩证逻辑为神仙信仰的中国化提供了深刻的经典依据。'
    )
    add_body_text(doc, cj1_text)
    
    cj2_text = (
        '《度人经》"仙道贵生，无量度人"是道教神仙信仰的基本宗教精神。"仙道贵生"确立了生命至上的价值取向，'
        '这一取向与当代社会的生命关怀形成了深刻的契合。"无量度人"则将个体修仙与济度众生统一起来，'
        '与当代社会公益、志愿服务、扶贫济困等实践形成了价值共鸣。'
    )
    add_body_text(doc, cj2_text)
    
    cj3_text = (
        '"济世利人"是道教神仙信仰中修仙实践与道德修养统一的核心体现。《度人经》"无量度人"、《抱朴子》"欲求仙者，要当以忠孝和顺仁信为本"、'
        '全真道"真功真行"——道教经典中蕴含着丰富的利他主义和道德修养思想。道教教理教义体系中"济世利人"等思想与社会主义核心价值观具有内在的契合性。'
        '这种内在契合性，是神仙信仰中国化最坚实的价值基础。'
    )
    add_body_text(doc, cj3_text)
    
    # Section 4
    add_chapter_heading(doc, '第四节 神仙信仰中国化的现实依据', level=2)
    
    cr1_text = (
        '党的十八大以来，党中央高度重视宗教工作，坚持宗教中国化方向已成为宗教工作的重大战略部署。'
        '《宗教事务条例》（2017年修订）为宗教活动的规范化管理提供了法律保障，全国宗教工作会议精神为宗教中国化指明了方向。'
        '道教界对自身中国化的自觉意识，是神仙信仰中国化最直接的内在动力。'
        '中国道教协会发布的《深入推进我国道教中国化五年工作规划纲要（2023-2027年）》，是道教界系统推进中国化的纲领性文件。'
        '当代道教信众的构成和需求正在发生深刻变化，年轻信众不仅关注祈福禳灾等传统功能，更关心心灵安顿、生命意义、生态伦理等精神层面的议题。'
    )
    add_body_text(doc, cr1_text)
    
    # Chapter summary
    add_chapter_heading(doc, '本章小结', level=3)
    
    summary3 = (
        '道教神仙信仰的中国化具有三重理论依据：历史逻辑表明，神仙信仰从来不是一个封闭的体系，而是在"随方设教"的传统中不断自我调适、自我更新；'
        '经典依据表明，道教经典中蕴含着丰富的自我更新思想资源；现实依据表明，新时代宗教政策提供了制度保障，道教界的自觉意识提供了内在动力。'
        '三重依据相互支撑，共同论证了一个核心命题：神仙信仰的中国化不是外在强加的，而是内在于神仙信仰发展逻辑的必然要求。'
    )
    add_body_text(doc, summary3)

def add_chapter4(doc):
    """Add Chapter 4."""
    add_chapter_heading(doc, '第四章 新时代道教神仙信仰中国化的实践路径', level=1)
    
    # Section 1
    add_chapter_heading(doc, '第一节 教义阐释的中国化——神仙信仰的义理重构', level=2)
    
    p1_text = (
        '教义阐释的中国化是神仙信仰中国化的灵魂。信仰的核心具有恒定性，不可动摇；但信仰的阐释框架必须与时俱进。'
        '义理重构不是改宗变教，而是在回归本源中激活传统的当代意义。'
        '以"仙道贵生"为核心，重构神仙信仰的时代内涵：从"个体长生"到"生命尊严"，从"出世超脱"到"入世度人"，从"道法自然"到"生态文明"。'
    )
    add_body_text(doc, p1_text)
    
    p1b_text = (
        '针对当前神仙信仰理解浅表化的问题，义理重构的核心任务之一是引导信众从功能化的信仰理解回归义理化的信仰认知。'
        '拜财神不只是求财运，更是体悟"利而不害""为而不争"的财富伦理；拜药王不只是求健康，更是践行"贵生""养生"的生命智慧。'
        '每一尊神仙都不是功能性的祈福工具，而是"道"在特定领域的显现，蕴含着相应的价值理念和精神追求。'
    )
    add_body_text(doc, p1b_text)
    
    # Section 2
    add_chapter_heading(doc, '第二节 修持实践的中国化——神仙信仰的实践转化', level=2)
    
    p2_text = (
        '义理重构提供了方向指引，但神仙信仰的中国化最终必须落实到修持实践的层面。'
        '内丹修持与心理健康、生命教育的结合：内丹修持中的"守一""坐忘""心斋"等修持方法，在去除宗教神秘性外衣后，'
        '与现代正念冥想、心理调适等方法具有功能上的相似性，可以为当代人的精神健康提供有益的修持资源。'
        '斋醮科仪的规范化与简洁化：在保护传承的前提下推进科仪的规范化整理，在保持神圣性的前提下推进科仪的简洁化，'
        '加强科仪中的义理阐释。'
    )
    add_body_text(doc, p2_text)
    
    p2b_text = (
        '玄门讲经：以社会主义核心价值观引领经典阐释。中国道教协会持续组织的玄门讲经巡讲活动，'
        '围绕社会主义核心价值观选择讲解经典，设定年度主题，组织道众进行深入阐释。修道生活与现代生活方式的调适：'
        '在保持修道生活核心要素（诵经、打坐、科仪、济世）的前提下，适当吸收现代生活的合理成分。'
    )
    add_body_text(doc, p2b_text)
    
    # Section 3
    add_chapter_heading(doc, '第三节 组织形态的中国化——神仙信仰的制度保障', level=2)
    
    p3_text = (
        '宫观管理的民主化与现代化：完善民主管理委员会制度，建立规范的财务管理制度，消除商业化倾向，'
        '推进宫观服务的标准化和规范化，建设"文化道观""生态道观"。'
        '道教人才培养体系的完善：扩大道教学院的办学规模，优化课程体系，创新培养模式，将学院教育与宫观实修、师徒传承有机结合。'
        '信众服务模式的分层化与规范化：针对不同层次的信众提供差异化的服务，建立统一的服务标准和行为规范。'
    )
    add_body_text(doc, p3_text)
    
    # Section 4
    add_chapter_heading(doc, '第四节 传播方式的中国化——神仙信仰的当代表达', level=2)
    
    p4_text = (
        '道教文化的数字化传播与新媒体运用：建设道教文化数字资源库，开发道教文化新媒体内容，培养道教新媒体传播人才，'
        '建立道教网络传播规范。神仙信仰的美学表达与文化创意产业：将道教神仙信仰蕴含的丰富美学资源转化为当代文化产品。'
        '道教文化的国际传播与文明互鉴：以神仙信仰中"和而不同""万物并育而不相害"的包容精神，为不同文明之间的对话与互鉴贡献道教智慧。'
    )
    add_body_text(doc, p4_text)
    
    # Section 5
    add_chapter_heading(doc, '第五节 典型案例分析', level=2)
    
    case_text = (
        '武当山"系统推进道教中国化"实践：将真武大帝信仰与爱国主义教育相结合，推进宫观管理的规范化建设，开展生态道观建设，'
        '利用新媒体传播武当道教文化。玄门讲经活动的教义当代阐释：以"仙道贵生"为主题阐释生命的尊严与价值，'
        '以"济世利人"为主题阐释道教的社会责任，以"道法自然"为主题阐释生态文明理念。'
        '地方道协推进中国化的创新实践：上海市道协组织编写《道教常识》等通俗读本，四川省道协推进"文化道观"建设，'
        '浙江省道协推进"生态道观"建设，湖北省道协推行科仪规范化改革，广东省道协利用新媒体传播道教文化。'
    )
    add_body_text(doc, case_text)
    
    # Chapter summary
    add_chapter_heading(doc, '本章小结', level=3)
    
    summary4 = (
        '新时代道教神仙信仰中国化的实践路径，是一个以"返本开新"为核心理念的系统性工程，包括四个相互支撑的维度：'
        '教义阐释的中国化是灵魂，修持实践的中国化是关键，组织形态的中国化是基础，传播方式的中国化是手段。'
        '四条路径相互支撑、缺一不可，典型案例验证了路径的可行性和有效性——中国化不是理论构想，而是正在进行的实践探索。'
    )
    add_body_text(doc, summary4)

def add_chapter5(doc):
    """Add Chapter 5."""
    add_chapter_heading(doc, '第五章 神仙信仰传承与中国化的辩证关系与未来展望', level=1)
    
    # Section 1
    add_chapter_heading(doc, '第一节 传承与创新的辩证统一', level=2)
    
    d1_text = (
        '在神仙信仰的传承与中国化问题上，存在两种可能的偏向：一种偏向是片面强调"传承"，将神仙信仰视为不可更改的固定体系；'
        '另一种偏向是片面强调"创新"，以"中国化"为名对神仙信仰进行过度改造，甚至消解信仰的核心内涵。'
        '这两种偏向都未能把握传承与创新的辩证统一关系。'
    )
    add_body_text(doc, d1_text)
    
    d2_text = (
        '传承是创新之根：传承的核心在于守护神仙信仰中具有恒定性的要素——"道"的体悟与生命超越的追求。'
        '这些要素是道教区别于其他宗教的根本标志，也是道教作为宗教存在的身份基础。'
        '"仙道贵生，无量度人"的基本宗教精神、"道—气—神"的三层信仰结构、修仙实践与道德修养的统一——这些核心要素必须被坚守。'
    )
    add_body_text(doc, d2_text)
    
    d3_text = (
        '创新是传承之翼：如果神仙信仰的阐释框架和实践方式始终停留在传统形态，无法回应当代社会提出的精神需求和价值问题，'
        '那么信仰就会在现实中被边缘化。历史已经反复证明：神仙信仰的每一次重大转型都不是对传统的否定，而是在新的时代条件下激活传统的生命力。'
    )
    add_body_text(doc, d3_text)
    
    d4_text = (
        '"返本开新"是传承与创新辩证统一的集中体现。"返本"不是简单的复古，而是对信仰核心的重新发现和深刻体认；'
        '"开新"不是随意的创造，而是在信仰核心的引导下，根据时代条件探索新的阐释框架、实践方式和传播形式。'
        '"返本"与"开新"是同一过程中的辩证两面，正如《道德经》所言"反者道之动"，回归本源与开创新局是"道"的运动规律的内在要求。'
    )
    add_body_text(doc, d4_text)
    
    # Section 2
    add_chapter_heading(doc, '第二节 神仙信仰中国化的基本原则', level=2)
    
    bp_text = (
        '第一，坚持爱国爱教的政治方向。坚持爱国爱教，是神仙信仰中国化的根本政治前提。'
        '道教自古就有"助国保民"的优良传统，新时代神仙信仰的中国化，必须在政治上坚持拥护中国共产党的领导。'
        '第二，坚守信仰核心不偏移。"仙道贵生，无量度人"的基本宗教精神等核心要素是神仙信仰的"根"与"魂"，任何中国化的实践都不能触及这一底线。'
        '第三，坚持与时俱进的时代自觉。神仙信仰的中国化是一个持续的历史进程，不是一劳永逸的阶段性任务。'
        '第四，坚持以人为本的服务导向。中国化的成效，不是由理论上的自洽性来衡量，而是由信众的获得感和社会的认可度来检验。'
    )
    add_body_text(doc, bp_text)
    
    # Section 3
    add_chapter_heading(doc, '第三节 未来展望', level=2)
    
    future_text = (
        '神仙信仰与中华优秀传统文化传承发展的深度融合：神仙信仰作为中华传统文化的重要组成部分，其"道法自然""天人合一""仙道贵生"等思想资源，'
        '不仅是道教自身的信仰财富，更是中华优秀传统文化的精神瑰宝。'
        '道教神仙信仰在构建人类命运共同体中的文化贡献：神仙信仰所蕴含的"天人合一"的宇宙观，为超越人类中心主义、构建人与自然和谐共生的全球生态伦理提供了思想资源。'
        '道教中国化的持续推进与长效机制建设：在制度层面完善道教中国化的制度保障体系，在学术层面持续推进神仙信仰的义理研究和当代阐释，'
        '在实践层面及时总结和推广中国化的典型经验，在人才层面培养一批既具深厚信仰修养又具时代眼光的道教人才队伍。'
    )
    add_body_text(doc, future_text)

def add_conclusion(doc):
    """Add conclusion."""
    add_chapter_heading(doc, '结  语', level=1)
    
    conc_text = (
        '新时代坚持宗教中国化方向，是以习近平同志为核心的党中央关于宗教工作的重大战略部署。'
        '道教作为中华大地上唯一土生土长的宗教，其中国化进程具有特殊的历史地位与现实意义。'
        '本文以"新时代道教神仙信仰的传承与中国化路径"为研究主题，从历史纵深、现实境遇、理论逻辑与实践路径四个维度展开系统研究，得出以下主要结论：'
    )
    add_body_text(doc, conc_text)
    
    conc2 = (
        '第一，道教神仙信仰的核心——"道"的体悟与生命超越的追求——具有恒定性，而神仙信仰的形式具有时代性。'
        '恒定性与时代性的辩证统一，是神仙信仰"传承"与"创新"的内在张力，也是"中国化"命题的深层依据。'
        '第二，新时代道教神仙信仰的传承面临深刻挑战，这些挑战的实质是传统信仰表达方式与现代社会需求之间的断裂。'
        '中国化不是外在强加的政治要求，而是道教自身"返本开新"传统的延续。'
        '第三，神仙信仰的中国化具有深厚的历史逻辑、经典依据和现实基础，共同为中国化提供了三重理论合法性。'
        '第四，神仙信仰的中国化应以"返本开新"为核心理念，从教义阐释、修持实践、组织形态、传播方式四个维度构建系统性实践路径。'
    )
    add_body_text(doc, conc2)
    
    conc3 = (
        '本文的主要贡献在于：首次将道教神仙信仰研究与中国化路径问题系统对接，填补了该交叉领域的研究空白；'
        '提出了神仙信仰中国化的"返本开新"理论框架；从四个维度构建了系统性的中国化实践路径；以典型案例为支撑，验证了路径的可行性和有效性。'
        '展望未来，道教神仙信仰的中国化是一个持续的历史进程。在新时代的伟大征程中，道教神仙信仰必将在"返本开新"的道路上，'
        '实现信仰的创造性转化与创新性发展，为中华优秀传统文化的传承发展、为社会主义社会的和谐进步、为人类命运共同体的文明互鉴，贡献道教智慧和道教力量。'
    )
    add_body_text(doc, conc3)

def add_references(doc):
    """Add references section."""
    add_chapter_heading(doc, '参考文献', level=1)
    
    refs = [
        '《道德经》.',
        '《庄子》（《南华真经》）.',
        '《太平经》.',
        '葛洪. 《抱朴子内篇》.',
        '陶弘景. 《真灵位业图》.',
        '张君房. 《云笈七签》.',
        '《度人经》.',
        '李远国. 中国道教神仙谱系史（四卷本）[M]. 成都: 四川大学出版社/成都时代出版社, 2023-2024.',
        '刘屹. 神格与地域：汉唐间道教信仰世界研究[M]. 上海: 上海人民出版社, 2011.',
        '姜生. 汉帝国的遗产：汉鬼考[M]. 北京: 科学出版社, 2017.',
        '孙亦平. 唐宋道教的转型[M]. 北京: 中华书局, 2018.',
        '丁常云（主编）. 道教中国化研究[M]. 上海: 上海三联书店, 2020.',
        '陈耀庭. 道教神学概论（修订本）[M]. 北京: 宗教文化出版社, 2016.',
        '余平. 神仙信仰现象学引论[M]. 成都: 四川大学出版社, 2015.',
        '高丽杨. 坚持道教中国化方向的五个着力点[J]. 世界宗教文化, 2023(5).',
        '曹玉华. 道教"仙真"形态演变论[J]. 四川大学学报（哲学社会科学版）, 2005(3).',
        '宫哲兵. 当代道教主神的结构体系论[J]. 哲学研究, 2008(4).',
        '程乐松. 内外之间与古今之际——信仰的义理建构与道教的文化价值[J]. 道家文化研究, 2023(1).',
        '张志刚. 试论道教义理的现代建构[J]. 世界宗教文化, 2025(1).',
        '班泰勇, 楼剑涛, 李蕾. 系统推进道教中国化走深走实的几点思考[J]. 中国宗教, 2025(4).',
        '刘傲然. 道教现代人才培养在20世纪的探索及实践成果[J]. 中国道教, 2025.',
        '卢国龙. 发挥"道教之真精神"，推进道教的"中国化"进程[N]. 中国民族报, 2019-04-09.',
        '陆文荣. 道教中国化与当代道教发展的思考[J/OL]. 道教之音, 2023.',
        '盖建民. 构建道教研究的"中国学派"话语体系的思考[J/OL]. 中国社会科学网, 2020.',
        '习近平. 在全国宗教工作会议上的重要讲话[R]. 2016, 2021.',
        '国务院. 宗教事务条例（国务院令第686号）[Z]. 2017.',
        '中国道教协会. 深入推进我国道教中国化五年工作规划纲要（2023-2027年）[Z]. 2023.',
        '李光富. 以社会主义核心价值观引领道教的赓续与创新[Z]. 全国政协十四届常委会第九次会议书面发言, 2024.',
        '章伟文. 芻议道教教义的传承、创新之维[J]. 北京师范大学学术成果.',
        '康立坤, 张广保. 性与命——全真教中国化的根基[J]. 世界宗教文化, 2023(3).',
    ]
    
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(f'[{i}] {ref}')
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        set_east_asia_font(run, '宋体')

def add_appendix(doc):
    """Add appendix section."""
    add_chapter_heading(doc, '附  录', level=1)
    
    # Acknowledgments
    add_chapter_heading(doc, '一、致  谢', level=2)
    
    ack_text = (
        '（此处填写致谢内容）'
        '在本文的撰写过程中，得到了导师的悉心指导和无私帮助，在此表示诚挚的感谢。'
        '同时，感谢浙江省宗教界"双通"人才研修班提供的学习平台，感谢各位授课老师的精彩讲授，'
        '感谢同窗学友的相互交流与启发。本文的研究也离不开道教界各位道长的大力支持与分享，'
        '在此一并致以诚挚的谢意。'
    )
    add_body_text(doc, ack_text)
    
    doc.add_page_break()
    
    # Originality statement
    add_chapter_heading(doc, '二、论文原创性声明', level=2)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    
    statement_text = (
        '本人郑重声明：所呈交的毕业论文，是本人在导师的指导下，独立进行研究工作所取得的成果。'
        '除文中已经注明引用的内容外，本论文不包含任何其他个人或集体已经发表或撰写过的作品成果。'
        '对本文的研究作出重要贡献的个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的法律和学术责任由本人承担。'
    )
    add_body_text(doc, statement_text)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('毕业论文作者签名：')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    set_east_asia_font(run, '宋体')
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('日    期：      年    月    日')
    run.font.name = '宋体'
    run.font.size = Pt(12)
    set_east_asia_font(run, '宋体')

def main():
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    set_east_asia_font_style(style, '宋体')
    style.paragraph_format.line_spacing = 1.5
    
    # Page setup
    section = doc.sections[0] if doc.sections else doc.add_section()
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    
    # Generate content
    add_cover_page(doc)
    add_abstract(doc)
    add_toc(doc)
    add_introduction(doc)
    add_chapter1(doc)
    add_chapter2(doc)
    add_chapter3(doc)
    add_chapter4(doc)
    add_chapter5(doc)
    add_conclusion(doc)
    add_references(doc)
    add_appendix(doc)
    
    output_path = '/home/admin/.openclaw/workspace/新时代道教神仙信仰的传承与中国化路径.docx'
    doc.save(output_path)
    print(f'✅ 论文已生成: {output_path}')

if __name__ == '__main__':
    main()
