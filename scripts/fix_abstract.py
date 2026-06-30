#!/usr/bin/env python3
"""精简摘要到1200字以内"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re

INPUT_FILE = "/home/admin/.openclaw/workspace/毕业论文-邱春华-修改版.docx"
OUTPUT_FILE = "/home/admin/.openclaw/workspace/毕业论文-邱春华-修改版.docx"

def count_chars(text):
    return len(text)

def set_run(run, name='宋体', size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def main():
    print("读取文档...")
    doc = Document(INPUT_FILE)
    
    # 找到摘要位置
    abs_start = None
    abs_end = None
    
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if (t.startswith("摘") and "要" in t) and abs_start is None:
            abs_start = i
        elif abs_start is not None and abs_end is None:
            if t.startswith("Abstract") or t.startswith("abstract"):
                abs_end = i
                break
    
    if abs_start and abs_end:
        # 精简后的中文摘要（约1100字）
        new_cn = """本文以王治心（1881—1968）为研究对象，系统考察其在民国时期基督教本色化运动中的神学观、实践探索及其对当代基督教中国化的启示。王治心作为前清秀才出身的基督徒知识分子，以其独特的"本色化实践派"立场，在神学思想的本土化建构、教会礼仪的文化革新、宗教史学的独立编撰以及抗日救亡的社会参与等多个维度上，形成了系统而独到的贡献，堪称本色化运动中最为全面、最具实践品格的代表人物之一。

文章认为，王治心的本色化探索主要体现为三个维度：第一，以中华文化资源解释基督教教义。王治心提出"以儒释耶"的神学路径，深入挖掘基督教与儒家、墨家、佛教、道家等中国传统思想资源的契合点，构建了系统的跨文化宗教对话理论。他强调基督教与中国文化的融合"不是在形式方面，乃在精神方面"，反对流于形式的"本色化"尝试，主张在精神内核层面实现基督教信仰与中国文化的深层融通。第二，推动神学教育和礼仪实践本土化。王治心在金陵神学院、福建协和大学、沪江大学等多所院校推行国学教育革新，提出"国学+神学"的课程模式，培养兼具信仰素养和文化素养的本土教牧人才。在礼仪方面，他设计了八大节期方案，提出祭祖问题的折中方案，倡导以"追思礼拜"形式替代传统祭祖仪式，既保持信仰纯正又尊重中国文化传统。第三，在民族危亡中强调爱国与信仰统一。王治心在五卅运动和抗日战争中积极呼吁基督徒投身爱国救亡，从神学层面论证基督教爱国精神的合法性，将"荣神益人"与民族复兴伟业相结合，为当代"爱国爱教"传统提供了历史典范。

其历史经验对当代基督教中国化的启示在于：基督教中国化不应停留于外在形式模仿，而应在四个层面深化推进。文化融合层面，应坚持"精神契合"而非"形式模仿"的原则，深入挖掘基督教与中华优秀传统文化的契合点，构建中国特色神学话语体系。教会建设层面，应深化本土化礼仪实践，探索基督教追思礼拜等具有中国文化特色的信仰表达方式。人才培养层面，应建立"国学+神学+社会责任感"的综合培养模式，培育兼具信仰素养、文化底蕴与家国情怀的本土教牧队伍。社会适应层面，应弘扬爱国爱教传统，推动基督教与社会主义核心价值观相融合，发挥基督教在社会服务、公益慈善等领域的积极作用。

本文的创新之处在于：第一，构建"思想—实践—价值"的系统性研究框架，突破了以往研究中思想与实践分离、历史与当代脱节的局限。第二，在历史研究与当代启示之间建立有效衔接，采用"历史经验—当代问题—转化路径"的分析框架逐条提炼现实启示。第三，首次从浙江地域视角系统研究王治心，将其思想形成与浙江地域文化传统相关联，为浙江基督教中国化提供历史参照。"""
        
        # 删除旧摘要
        for i in range(abs_end - 1, abs_start, -1):
            if i < len(doc.paragraphs):
                for run in doc.paragraphs[i].runs:
                    run.text = ""
        
        # 插入新摘要
        cn_paras = new_cn.split("\n\n")
        for j, text in enumerate(cn_paras):
            if text.strip():
                idx = abs_start + 1 + j
                if idx < len(doc.paragraphs):
                    para = doc.paragraphs[idx]
                    para.text = text.strip()
                    for run in para.runs:
                        set_run(run, '宋体', 12)
                else:
                    p = doc.add_paragraph(text.strip())
                    for run in p.runs:
                        set_run(run, '宋体', 12)
        
        cn_chars = count_chars(new_cn)
        print(f"中文摘要: {cn_chars}字")
        
        # 精简英文摘要
        new_en = """This paper takes Wang Zhixin (1881-1968) as the research object, systematically examining his theological perspective, practical exploration during the indigenous church movement in the Republican era, and their implications for the contemporary Sinicization of Christianity. As a Christian intellectual with a background as a Qing Dynasty scholar, Wang Zhixin made systematic and distinctive contributions in multiple dimensions including the localization of theological construction, cultural innovation of church liturgy, independent compilation of religious historiography, and social participation in the anti-Japanese national salvation movement. He stands as one of the most comprehensive and practically-oriented representative figures in the indigenous church movement.

The paper argues that Wang Zhixin's indigenization exploration is mainly manifested in three dimensions: First, interpreting Christian doctrines through Chinese cultural resources. Wang proposed the theological path of "interpreting Christianity through Confucianism," deeply exploring the convergence points between Christianity and traditional Chinese thought resources, constructing a systematic cross-cultural religious dialogue theory. He emphasized that the integration of Christianity and Chinese culture lies "not in form but in spirit." Second, promoting the localization of theological education and liturgical practices. Wang implemented Chinese studies education reforms at Jinling Seminary, Fujian Union University, and Jiangsu University, proposing a "Chinese Studies + Theology" curriculum model. In liturgy, he designed an eight-festival scheme and proposed a compromise solution for ancestor worship, maintaining both faith purity and respect for Chinese cultural traditions. Third, emphasizing the unity of patriotism and faith amid national crisis. Wang actively called on Christians to participate in patriotic salvation movements, providing theological justification for Christian patriotism and connecting "glorifying God and benefiting people" with national rejuvenation.

His historical experience offers important implications for the contemporary Sinicization of Christianity: rather than remaining at the level of superficial formal imitation, the Sinicization of Christianity should be deepened in four aspects. In cultural integration, the principle of "spiritual convergence" rather than "formal imitation" should be upheld, deeply exploring convergence points between Christianity and excellent traditional Chinese culture, constructing a theological discourse system with Chinese characteristics. In church building, localized liturgical practices should be deepened, exploring faith expressions with Chinese cultural characteristics such as Christian memorial services. In talent cultivation, a comprehensive training model of "Chinese Studies + Theology + Social Responsibility" should be established, cultivating local pastoral talents with faith literacy, cultural foundation, and patriotic sentiment. In social adaptation, the tradition of loving both country and religion should be promoted, facilitating the integration of Christianity with socialist core values.

The innovations of this paper lie in: First, constructing a systematic research framework of "thought-practice-value," breaking through the limitations of separating thought from practice and history from contemporary practice. Second, establishing effective connection between historical research and contemporary inspiration, adopting a "historical experience-contemporary issue-transformation path" analytical framework. Third, for the first time systematically studying Wang Zhixin from the perspective of Zhejiang regional culture, providing historical reference for the Sinicization of Christianity in Zhejiang."""
        
        # 找到英文摘要区域
        en_start = abs_end
        en_end = None
        for i in range(en_start, len(doc.paragraphs)):
            t = doc.paragraphs[i].text.strip()
            if t.startswith("Key Words") or t.startswith("Key words") or t.startswith("关键词"):
                en_end = i
                break
        
        if en_end:
            for i in range(en_end - 1, en_start, -1):
                if i < len(doc.paragraphs):
                    for run in doc.paragraphs[i].runs:
                        run.text = ""
            
            en_paras = new_en.split("\n\n")
            for j, text in enumerate(en_paras):
                if text.strip():
                    idx = en_start + 1 + j
                    if idx < len(doc.paragraphs):
                        para = doc.paragraphs[idx]
                        para.text = text.strip()
                        for run in para.runs:
                            set_run(run, 'Times New Roman', 12)
                    else:
                        p = doc.add_paragraph(text.strip())
                        for run in p.runs:
                            set_run(run, 'Times New Roman', 12)
            
            print(f"英文摘要已精简")
    
    # 保存
    print(f"\n保存: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    
    final = 0
    for para in doc.paragraphs:
        final += count_chars(para.text)
    
    print(f"\n最终字数: {final}")
    print(f"修改完成！")

if __name__ == "__main__":
    main()
