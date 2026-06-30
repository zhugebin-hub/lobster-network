#!/usr/bin/env python3
"""生成修改后的毕业论文"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

OUTPUT_FILE = "/home/admin/.openclaw/workspace/毕业论文-邱春华-修改版.docx"

def set_cell_text(cell, text, font_name='宋体', font_size=12, bold=False, alignment=None):
    """设置单元格文本"""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if alignment:
        p.alignment = alignment

def add_heading_text(doc, text, level=1, font_size=None):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    if font_size:
        for run in heading.runs:
            run.font.size = Pt(font_size)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return heading

def add_para(doc, text, font_name='宋体', font_size=12, bold=False, alignment=None, line_spacing=None, space_before=0, space_after=0):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    pf = p.paragraph_format
    if alignment:
        p.alignment = alignment
    if line_spacing:
        pf.line_spacing = line_spacing
    if space_before:
        pf.space_before = Pt(space_before)
    if space_after:
        pf.space_after = Pt(space_after)
    
    return p

def main():
    doc = Document()
    
    # 设置文档默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ==================== 封面 ====================
    # 空行
    for _ in range(6):
        doc.add_paragraph()
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("浙江省宗教界"双通"人才研修班")
    run.font.name = '宋体'
    run.font.size = Pt(16)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_paragraph()
    
    # 论文题目（修改后）
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("王治心神学观与基督教中国化实践的当代价值探析")
    run.font.name = '宋体'
    run.font.size = Pt(22)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 信息表格
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    
    info = [
        ["作  者  姓  名：", "邱春华", "学      号：", "ST20241144"],
        ["指  导  教  师：", "商  琳", "研究方向：", "基督教中国化"],
        ["答  辩  日  期：", "2026年6月", "成      绩：", ""],
        ["浙江工商大学", "宗教中国化研究所", "", ""],
    ]
    
    for i, row_data in enumerate(info):
        if i < len(table.rows):
            for j, text in enumerate(row_data):
                if j < len(table.columns):
                    set_cell_text(table.rows[i].cells[j], text, font_size=12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_para.add_run("2026年6月")
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ==================== 中文摘要（扩充版） ====================
    doc.add_page_break()
    
    add_heading_text(doc, "摘  要", level=1, font_size=16)
    doc.add_paragraph()
    
    # 扩充后的中文摘要（约1800字，满足全文5%要求）
    abstract_cn = [
        "本文以王治心（1881—1968）为研究对象，系统考察其在民国时期基督教本色化运动中的神学观、实践探索及其对当代基督教中国化的启示。王治心作为前清秀才出身的基督徒知识分子，以其独特的"本色化实践派"立场，在神学思想的本土化建构、教会礼仪的文化革新、宗教史学的独立编撰以及抗日救亡的社会参与等多个维度上，形成了系统而独到的贡献，堪称本色化运动中最为全面、最具实践品格的代表人物之一。",
        
        "文章认为，王治心的本色化探索主要体现为三个维度：第一，以中华文化资源解释基督教教义。王治心提出"以儒释耶"的神学路径，深入挖掘基督教与儒家、墨家、佛教、道家等中国传统思想资源的契合点，构建了系统的跨文化宗教对话理论。他强调基督教与中国文化的融合"不是在形式方面，乃在精神方面"，反对流于形式的"本色化"尝试，主张在精神内核层面实现基督教信仰与中国文化的深层融通。第二，推动神学教育和礼仪实践本土化。王治心在金陵神学院、福建协和大学、沪江大学等多所院校推行国学教育革新，提出"国学+神学"的课程模式，培养兼具信仰素养和文化素养的本土教牧人才。在礼仪方面，他设计了八大节期方案，提出祭祖问题的折中方案，倡导以"追思礼拜"形式替代传统祭祖仪式，既保持信仰纯正又尊重中国文化传统。第三，在民族危亡中强调爱国与信仰统一。王治心在五卅运动和抗日战争中积极呼吁基督徒投身爱国救亡，从神学层面论证基督教爱国精神的合法性，将"荣神益人"与民族复兴伟业相结合，为当代"爱国爱教"传统提供了历史典范。",
        
        "其历史经验对当代基督教中国化的启示在于：基督教中国化不应停留于外在形式模仿，而应在四个层面深化推进。文化融合层面，应坚持"精神契合"而非"形式模仿"的原则，深入挖掘基督教与中华优秀传统文化的契合点，构建中国特色神学话语体系，以"文化润教"推动宗教中国化的深入发展。教会建设层面，应深化本土化礼仪实践，探索基督教追思礼拜等具有中国文化特色的信仰表达方式，推进教会空间的中国化建设。人才培养层面，应建立"国学+神学+社会责任感"的综合培养模式，将中华优秀传统文化课程纳入神学教育核心模块，培育兼具信仰素养、文化底蕴与家国情怀的本土教牧队伍。社会适应层面，应弘扬爱国爱教传统，推动基督教与社会主义核心价值观相融合，发挥基督教在社会服务、公益慈善等领域的积极作用。",
        
        "本文的创新之处在于：第一，构建"思想—实践—价值"的系统性研究框架，将王治心的神学观、本色化实践与当代价值置于统一的学术视野中进行考察，突破了以往研究中思想与实践分离、历史与当代脱节的局限。第二，在历史研究与当代启示之间建立有效衔接，实现历史经验与当代实践的双向贯通，使历史研究具有了当代价值。第三，首次从浙江地域视角系统研究王治心，将其思想形成与浙江地域文化传统相关联，为浙江基督教中国化提供历史镜鉴，拓展了浙江基督教人物研究的学术空间。本文的研究成果不仅对丰富基督教中国化历史研究维度、深化跨文化宗教融合理论认知具有重要学术价值，更为当代基督教摆脱"洋教"困境、培养本土教牧人才、推进中国化实践提供了可资借鉴的历史资源。"
    ]
    
    for text in abstract_cn:
        add_para(doc, text, font_size=12, line_spacing=1.5, space_after=6)
    
    # 关键词
    keywords_para = doc.add_paragraph()
    run = keywords_para.add_run("关键词：")
    run.font.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2 = keywords_para.add_run("王治心；本色化运动；以儒释耶；基督教中国化；爱国爱教")
    run2.font.name = '宋体'
    run2.font.size = Pt(12)
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ==================== 英文摘要（扩充版） ====================
    doc.add_paragraph()
    add_heading_text(doc, "Abstract", level=1, font_size=16)
    doc.add_paragraph()
    
    abstract_en = [
        "This paper takes Wang Zhixin (1881-1968) as the research object, systematically examining his theological perspective, practical exploration during the indigenous church movement in the Republican era, and their implications for the contemporary Sinicization of Christianity. As a Christian intellectual with a background as a Qing Dynasty xiucai (scholar), Wang Zhixin, with his unique "indigenization practice school" stance, made systematic and distinctive contributions in multiple dimensions including the localization of theological construction, cultural innovation of church liturgy, independent compilation of religious historiography, and social participation in the anti-Japanese national salvation movement. He stands as one of the most comprehensive and practically-oriented representative figures in the indigenous church movement.",
        
        "The paper argues that Wang Zhixin's indigenization exploration is mainly manifested in three dimensions: First, interpreting Christian doctrines through Chinese cultural resources. Wang proposed the theological path of "interpreting Christianity through Confucianism," deeply exploring the convergence points between Christianity and traditional Chinese thought resources such as Confucianism, Mohism, Buddhism, and Daoism, constructing a systematic cross-cultural religious dialogue theory. He emphasized that the integration of Christianity and Chinese culture lies "not in form but in spirit," opposing superficial "indigenization" attempts and advocating for deep integration at the spiritual core level. Second, promoting the localization of theological education and liturgical practices. Wang implemented Chinese studies education reforms at Jinling Seminary, Fujian Union University, and Jiangsu University, proposing a "Chinese Studies + Theology" curriculum model to cultivate local pastoral talents with both faith literacy and cultural literacy. In liturgy, he designed an eight-festival scheme, proposed a compromise solution for ancestor worship issues, and advocated replacing traditional ancestor worship rituals with "memorial services" that maintain faith purity while respecting Chinese cultural traditions. Third, emphasizing the unity of patriotism and faith amid national crisis. Wang actively called on Christians to participate in patriotic salvation movements during the May 30th Incident and the Anti-Japanese War, providing theological justification for Christian patriotism and connecting "glorifying God and benefiting people" with the great cause of national rejuvenation, providing a historical model for the contemporary tradition of "loving both country and religion.",
        
        "His historical experience offers important implications for the contemporary Sinicization of Christianity: rather than remaining at the level of superficial formal imitation, the Sinicization of Christianity should be deepened in four aspects. In cultural integration, the principle of "spiritual convergence" rather than "formal imitation" should be upheld, deeply exploring the convergence points between Christianity and excellent traditional Chinese culture, constructing a theological discourse system with Chinese characteristics, and promoting religious Sinicization through "cultural nourishment." In church building, localized liturgical practices should be deepened, exploring faith expressions with Chinese cultural characteristics such as Christian memorial services, and promoting the Sinicization of church spaces. In talent cultivation, a comprehensive training model of "Chinese Studies + Theology + Social Responsibility" should be established, integrating excellent traditional Chinese culture courses into the core module of theological education to cultivate local pastoral talents with faith literacy, cultural foundation, and patriotic sentiment. In social adaptation, the tradition of loving both country and religion should be promoted, facilitating the integration of Christianity with socialist core values, and leveraging Christianity's positive role in social services and public welfare.",
        
        "The innovations of this paper lie in: First, constructing a systematic research framework of "thought-practice-value," examining Wang Zhixin's theological perspective, indigenization practice, and contemporary value within a unified academic perspective, breaking through the limitations of separating thought from practice and history from contemporary practice in previous research. Second, establishing effective connection between historical research and contemporary inspiration, achieving bidirectional integration of historical experience and contemporary practice, giving historical research contemporary value. Third, for the first time systematically studying Wang Zhixin from the perspective of Zhejiang regional culture, connecting his thought formation with Zhejiang's regional cultural traditions, providing historical reference for the Sinicization of Christianity in Zhejiang, and expanding the academic space of Zhejiang Christian figure research. The research results of this paper not only have important academic value for enriching the historical research dimensions of Christianity Sinicization and deepening the theoretical cognition of cross-cultural religious integration, but also provide valuable historical resources for contemporary Christianity to escape the "foreign religion" dilemma, cultivate local pastoral talents, and promote Sinicization practices."
    ]
    
    for text in abstract_en:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = 6
    
    # Key Words
    kw_para = doc.add_paragraph()
    run = kw_para.add_run("Key Words: ")
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = kw_para.add_run("Wang Zhixin; indigenous church movement; interpreting Christianity through Confucianism; Sinicization of Christianity; patriotism and religious faith")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    
    # ==================== 目录 ====================
    doc.add_page_break()
    add_heading_text(doc, "目  录", level=1, font_size=16)
    doc.add_paragraph()
    
    toc_content = [
        "摘  要 .................................................................................................................... 3",
        "Abstract .................................................................................................................... 3",
        "引  言 .................................................................................................................... 5",
        "  第一节 选题背景与研究意义 ............................................................................. 6",
        "  第二节 国内外研究现状述评 ............................................................................. 7",
        "  第三节 研究思路与研究方法 ........................................................................... 10",
        "  第四节 研究创新点与不足 ............................................................................... 12",
        "第一章 王治心生平与二十世纪基督教中国化的历史语境 .............................. 13",
        "  第一节 从儒家秀才到基督徒知识分子：双重文化身份的形成 ................... 13",
        "  第二节 20世纪基督教中国化的历史动因与时代背景 ................................... 16",
        "  第三节 时代文化思潮对王治心神学观形成的影响 ....................................... 18",
        "第二章 王治心神学观的核心内涵与理论特质 .................................................. 20",
        "  第一节 "以儒释耶"与神学本土化建构及本土上帝观论证 .......................... 20",
        "  第二节 跨文化宗教互补思想 ........................................................................... 25",
        "  第三节 爱国与信仰统一的神学立场 ............................................................... 28",
        "  第四节 宗教本土化学理范式辨析 ................................................................... 30",
        "第三章 王治心基督教中国化的实践路径与历史成效 ...................................... 32",
        "  第一节 神学教育的本土化改造 ....................................................................... 32",
        "  第二节 教会礼仪与文化的中国化尝试 ........................................................... 35",
        "  第三节 宗教著作编撰、文化传播与本色神学家比较 ................................... 38",
        "  第四节 社会参与实践与历史成效局限分析 ................................................... 42",
        "第四章 王治心神学观与实践的当代价值与现实启示 ...................................... 46",
        "  第一节 当代基督教中国化的现状、挑战与本色化思想的时代局限 ........... 46",
        "  第二节 文化融合层面的启示 ........................................................................... 50",
        "  第三节 教会建设层面的启示 ........................................................................... 53",
        "  第四节 人才培养与社会适应层面的启示 ....................................................... 56",
        "结  语 .................................................................................................................... 60",
        "参考文献 ................................................................................................................ 64",
        "附  录 .................................................................................................................... 69",
    ]
    
    for line in toc_content:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = '宋体'
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.space_after = 2
    
    print("封面、摘要、目录已生成")
    
    # ==================== 正文内容 ====================
    # 由于正文内容较长，我们直接从原文档复制并修改
    # 这里先保存当前进度
    
    doc.save(OUTPUT_FILE)
    print(f"文档已保存到: {OUTPUT_FILE}")
    print("注意：这只是封面+摘要+目录部分，正文需要从原文档修改后追加")

if __name__ == "__main__":
    main()
