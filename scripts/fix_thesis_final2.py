#!/usr/bin/env python3
"""最终修改毕业论文"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re

INPUT_FILE = "/home/admin/.openclaw/media/inbound/2f0729c6-1a17-4fb4-9c79-5bedbeb2b7c1.docx"
OUTPUT_FILE = "/home/admin/.openclaw/workspace/毕业论文-邱春华-修改版.docx"

def count_chars(text):
    return len(text)

def set_run(run, name='宋体', size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def main():
    print("读取文档...")
    doc = Document(INPUT_FILE)
    
    # 统计原文字数
    orig_chars = 0
    for para in doc.paragraphs:
        orig_chars += count_chars(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                orig_chars += count_chars(cell.text)
    
    print(f"原文总字数: {orig_chars}")
    
    # 目标: 中文摘要 ≥ 5% of (orig_chars + abstract_added)
    # N ≥ 0.05 * (orig_chars + N) => N ≥ 0.05*orig_chars / 0.95
    min_abstract = int(orig_chars * 0.05 / 0.95)
    print(f"中文摘要需≥{min_abstract}字")
    
    # ==================== 1. 全局替换 ====================
    print("\n1. 替换 '王治心神学思想' → '王治心神学观'...")
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if "王治心神学思想" in run.text:
                run.text = run.text.replace("王治心神学思想", "王治心神学观")
                count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "王治心神学思想" in run.text:
                            run.text = run.text.replace("王治心神学思想", "王治心神学观")
                            count += 1
    print(f"   替换了 {count} 处")
    
    # ==================== 2. 整合章节 ====================
    print("\n2. 整合章节...")
    
    delete_patterns = [
        r"^第二节\s*本土上帝观的学术论证",
        r"^第三节\s*宗教本土化的三种学理范式辨析",
        r"^第四节\s*与民国本色神学家的比较",
        r"^第六节\s*历史成效与局限分析",
        r"^第二节\s*本色化思想的内在张力与时代局限",
        r"^第六节\s*社会适应层面的启示",
    ]
    
    deleted = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for pat in delete_patterns:
            if re.match(pat, text) and i > 30:
                for run in para.runs:
                    run.text = ""
                deleted += 1
                print(f"   删除: {text[:40]} (段落{i})")
    
    print(f"   共删除 {deleted} 个节标题")
    
    # ==================== 3. 扩充摘要 ====================
    print("\n3. 扩充摘要...")
    
    abs_start = None
    abs_end = None
    
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if (t.startswith("摘") and "要" in t) and abs_start is None:
            abs_start = i
        elif abs_start is not None and abs_end is None:
            if t.startswith("Abstract") or t.startswith("abstract"):
                abs_end = i
                break
    
    if abs_start and abs_end:
        # 扩充后的中文摘要（约3200字，确保≥5%）
        new_cn = """本文以王治心（1881—1968）为研究对象，系统考察其在民国时期基督教本色化运动中的神学观、实践探索及其对当代基督教中国化的启示。王治心作为前清秀才出身的基督徒知识分子，以其独特的"本色化实践派"立场，在神学思想的本土化建构、教会礼仪的文化革新、宗教史学的独立编撰以及抗日救亡的社会参与等多个维度上，形成了系统而独到的贡献，堪称本色化运动中最为全面、最具实践品格的代表人物之一。王治心是浙江湖州人，其思想的形成与浙江地域文化有着深层的内在关联，湖州深厚的儒学传统、江南独特的人文精神、浙江务实的区域品格，都对王治心的思想与实践产生了深刻的影响。

文章认为，王治心的本色化探索主要体现为三个维度：第一，以中华文化资源解释基督教教义。王治心提出"以儒释耶"的神学路径，深入挖掘基督教与儒家、墨家、佛教、道家等中国传统思想资源的契合点，构建了系统的跨文化宗教对话理论。他强调基督教与中国文化的融合"不是在形式方面，乃在精神方面"，反对流于形式的"本色化"尝试，主张在精神内核层面实现基督教信仰与中国文化的深层融通。他的"以儒释耶"不是简单的概念比附，而是在两种思想传统之间建立深层的道德哲学对话——"仁"与"爱"的互释、"天""帝"观与上帝观的融通、"大同"与天国观的比较，都是在精神层面寻找契合点的努力。在上帝观建构方面，他援引《诗经》《尚书》等先秦核心文献，系统论证上古中国文化体系中早已存在对应终极主宰的至高神灵崇拜传统，将基督教上帝观与中国传统"昊天上帝"信仰展开深度文本对话与价值对接，有效消解了"基督教为纯粹外来信仰、与中国文化全然对立"的片面认知。在基督论方面，他强调耶稣的伦理教训和人格典范，认为"登山宝训"与儒家道德理想高度一致，"爱人如己"与"仁者爱人"可以相互诠释。在救赎论方面，他借用儒家"尽心知性""反身而诚"等修养功夫来阐释基督教的悔改重生，强调信仰的内在体验和道德实践的重要性，体现了鲜明的伦理化诠释取向。在跨文化宗教互补思想方面，他以耶儒对话为核心、以耶墨对话为补充、以耶佛对话为借鉴、以道家思想为辅助，形成了多维度的对话体系。第二，推动神学教育和礼仪实践本土化。王治心在金陵神学院、福建协和大学、沪江大学等多所院校推行国学教育革新，提出"国学+神学"的课程模式，要求神学生"必须熟读四书，略通诸子"，培养兼具信仰素养和文化素养的本土教牧人才。在礼仪方面，他设计了八大节期方案，将中国传统节日与基督教信仰内涵相结合——如元旦对应感恩节、上元节对应灯节、清明节对应追思礼拜、端午节对应纪念救恩、中秋节对应天国盼望、重阳节对应敬老教训、冬至节对应圣诞前预备期、除夕对应守望礼拜，并非简单嫁接传统节日，而是将基督教信仰内核与中国人的生活节律、情感诉求相结合。他提出祭祖问题的折中方案，区分祭祖的"精神内核"与"表现形式"，倡导以"追思礼拜"形式替代传统祭祖仪式，用献花代替焚香、用鞠躬代替跪拜，既保持信仰纯正又尊重中国文化传统。在婚丧礼改革方面，他主张在保留基督教婚姻神学的基础上融入中国婚礼的合理元素，以基督教的追思礼拜取代传统丧礼仪式，体现了在保持信仰纯正前提下对传统习俗的尊重和包容。第三，在民族危亡中强调爱国与信仰统一。王治心在五卅运动和抗日战争中积极呼吁基督徒投身爱国救亡，从神学层面论证基督教爱国精神的合法性，将"荣神益人"与民族复兴伟业相结合。在1932年出版的《孙文主义与耶稣主义》中，他将三民主义的"民族独立、民权平等、民生幸福"与基督教的"自由、平等、博爱"相对接，系统论证了孙中山的三民主义与耶稣精神在价值追求上的内在一致性。在1938年出版的《战时教会讲话》中，他鼓励信徒在民族危亡中坚持信仰、服务社会，以实际行动支持抗战，体现了将基督教信仰与民族命运相结合的坚定立场，为当代"爱国爱教"传统提供了历史典范。

其历史经验对当代基督教中国化的启示在于：基督教中国化不应停留于外在形式模仿，而应在四个层面深化推进。文化融合层面，应坚持"精神契合"而非"形式模仿"的原则，深入挖掘基督教与中华优秀传统文化的契合点，在哲学层面探索创造论与"天生烝民""道法自然"的对话空间，在伦理层面深化"爱人如己"与"仁者爱人""兼爱天下"的呼应，在社会理想层面比较"上帝之国"与"大同世界"的异曲同工，在修养层面比较灵修传统与"坐忘""心斋""禅定"的可比性，构建中国特色神学话语体系，以"文化润教"推动宗教中国化的深入发展。教会建设层面，应深化本土化礼仪实践，探索基督教追思礼拜等具有中国文化特色的信仰表达方式，在礼拜音乐方面不仅改编传统圣诗更要创作具有中国音乐风格的原创圣诗，在礼拜空间方面探索具有中国建筑美学特征（如飞檐、斗拱、庭院等）的教堂设计，在讲道语言方面使用更贴近中国信徒文化心理的表述方式，在教会治理方面建立符合中国社会治理体系的教会管理制度，推进教会空间的中国化建设。人才培养层面，应建立"国学+神学+社会责任感"的综合培养模式，将中华优秀传统文化课程纳入神学教育核心模块而非"补充"课程，加大国学内容的比重和学时保障，培育兼具信仰素养、文化底蕴与家国情怀的本土教牧队伍，在人才选拔上注重文化素养，在培养上注重实践能力，在使用上注重地域特色，充分发挥本地教牧人员对地域文化的理解优势。社会适应层面，应弘扬爱国爱教传统，在国家层面将"荣神益人"思想与"富强、民主、文明、和谐"价值目标相融合，在社会层面将"大同"与"天国"比较思想与"自由、平等、公正、法治"社会取向相对接，在个人层面将"仁""爱"互释思想与"爱国、敬业、诚信、友善"公民道德规范相呼应，推动基督教与社会主义核心价值观深度融合，发挥基督教在社会服务、公益慈善、养老服务、社区建设等领域的积极作用。

本文的创新之处在于：第一，构建"思想—实践—价值"的系统性研究框架，将王治心的神学观、本色化实践与当代价值置于统一的学术视野中进行考察，突破了以往研究中思想与实践分离、历史与当代脱节的局限。第二，在历史研究与当代启示之间建立有效衔接，采用"历史经验—当代问题—转化路径"的分析框架逐条提炼现实启示，实现历史经验与当代实践的双向贯通，使历史研究具有了当代价值。第三，首次从浙江地域视角系统研究王治心，将其思想形成与浙江地域文化传统（湖州儒学传统、江南人文精神、浙江务实品格）相关联，为浙江基督教中国化提供最具地域亲和力的历史参照，拓展了浙江基督教人物研究的学术空间。本文的研究成果不仅对丰富基督教中国化历史研究维度、深化跨文化宗教融合理论认知具有重要学术价值，更为当代基督教摆脱"洋教"困境、培养本土教牧人才、推进中国化实践提供了可资借鉴的历史资源，也为浙江省宗教界"双通"人才深耕基督教中国化、传承爱国爱教传统提供了历史参考与实践思路。本文在研究方法上综合运用历史文献研究法、比较研究法、跨文化诠释学方法，以王治心的原著为核心文献，结合民国时期基督教期刊的相关文献，通过文本细读与历史语境还原，对王治心的思想与实践进行了深入解读，确保了研究的学术严谨性和论证的充分性。"""
        
        # 删除旧摘要
        for i in range(abs_end - 1, abs_start, -1):
            if i < len(doc.paragraphs):
                for run in doc.paragraphs[i].runs:
                    run.text = ""
        
        # 插入新摘要
        cn_paras = new_cn.split("\n\n")
        for j, text in enumerate(cn_paras):
            if text.strip():
                idx = abs_start + 1 + j
                if idx < len(doc.paragraphs):
                    para = doc.paragraphs[idx]
                    para.text = text.strip()
                    for run in para.runs:
                        set_run(run, '宋体', 12)
                else:
                    p = doc.add_paragraph(text.strip())
                    for run in p.runs:
                        set_run(run, '宋体', 12)
        
        cn_chars = count_chars(new_cn)
        print(f"   中文摘要: {cn_chars}字 (需≥{min_abstract}字)")
        
        # 扩充英文摘要
        new_en = """This paper takes Wang Zhixin (1881-1968) as the research object, systematically examining his theological perspective, practical exploration during the indigenous church movement in the Republican era, and their implications for the contemporary Sinicization of Christianity. As a Christian intellectual with a background as a Qing Dynasty scholar, Wang Zhixin, with his unique "indigenization practice school" stance, made systematic and distinctive contributions in multiple dimensions including the localization of theological construction, cultural innovation of church liturgy, independent compilation of religious historiography, and social participation in the anti-Japanese national salvation movement. He stands as one of the most comprehensive and practically-oriented representative figures in the indigenous church movement.

The paper argues that Wang Zhixin's indigenization exploration is mainly manifested in three dimensions: First, interpreting Christian doctrines through Chinese cultural resources. Wang proposed the theological path of "interpreting Christianity through Confucianism," deeply exploring the convergence points between Christianity and traditional Chinese thought resources such as Confucianism, Mohism, Buddhism, and Daoism, constructing a systematic cross-cultural religious dialogue theory. He emphasized that the integration of Christianity and Chinese culture lies "not in form but in spirit," opposing superficial "indigenization" attempts and advocating for deep integration at the spiritual core level. In terms of God concept construction, he cited pre-Qin core texts such as the Book of Songs and the Book of Documents, systematically demonstrating that ancient Chinese cultural system had long existed a tradition of supreme deity worship corresponding to the ultimate ruler, conducting deep textual dialogue and value docking between the Christian concept of God and the traditional Chinese "Haotian Shangdi" faith. In Christology, he emphasized the ethical teachings and personal example of Jesus, believing that the "Sermon on the Mount" was highly consistent with Confucian moral ideals. In soteriology, he borrowed Confucian cultivation methods such as "extending the mind and knowing nature" and "reflecting on oneself with sincerity" to interpret Christian repentance and rebirth, reflecting a distinctive ethical interpretation orientation. Second, promoting the localization of theological education and liturgical practices. Wang implemented Chinese studies education reforms at Jinling Seminary, Fujian Union University, and Jiangsu University, proposing a "Chinese Studies + Theology" curriculum model to cultivate local pastoral talents with both faith literacy and cultural literacy. In liturgy, he designed an eight-festival scheme integrating Chinese traditional festivals with Christian faith connotations—such as New Year corresponding to Thanksgiving, Qingming corresponding to Memorial Service, Dragon Boat Festival corresponding to Salvation Commemoration, Mid-Autumn Festival corresponding to Kingdom of Heaven hope, etc. These were not simple grafts of traditional festivals, but combined Christian faith core with Chinese people's life rhythms and emotional appeals. He proposed a compromise solution for ancestor worship, distinguishing between the "spiritual core" and "expression form" of ancestor worship, advocating "memorial services" to replace traditional rituals, using flower offerings instead of incense burning and bowing instead of kneeling, maintaining both faith purity and respect for Chinese cultural traditions. Third, emphasizing the unity of patriotism and faith amid national crisis. Wang actively called on Christians to participate in patriotic salvation movements during the May 30th Incident and the Anti-Japanese War, providing theological justification for Christian patriotism and connecting "glorifying God and benefiting people" with national rejuvenation. In his 1932 work "Sun Yat-sen's Principles and Jesus' Principles," he connected the Three Principles of the People's "national independence, democratic equality, and people's livelihood" with Christianity's "freedom, equality, and fraternity," systematically demonstrating the intrinsic consistency between Sun Yat-sen's Three Principles and the spirit of Jesus in value pursuit, providing a historical model for the contemporary tradition of "loving both country and religion."

His historical experience offers important implications for the contemporary Sinicization of Christianity: rather than remaining at the level of superficial formal imitation, the Sinicization of Christianity should be deepened in four aspects. In cultural integration, the principle of "spiritual convergence" rather than "formal imitation" should be upheld, deeply exploring convergence points between Christianity and excellent traditional Chinese culture—exploring dialogue spaces between creation theory and "Heaven produces all people" and "Dao follows nature" at the philosophical level, deepening the resonance between "love your neighbor as yourself" and "the benevolent loves others" and "universal love" at the ethical level, comparing the "Kingdom of God" with "Great Harmony" at the social ideal level, and comparing spiritual traditions with "sitting in oblivion," "fasting of the mind," and "meditation" at the cultivation level—constructing a theological discourse system with Chinese characteristics and promoting religious Sinicization through "cultural nourishment." In church building, localized liturgical practices should be deepened, exploring faith expressions with Chinese cultural characteristics such as Christian memorial services, not only adapting traditional hymns but also creating original hymns with Chinese musical style in worship music, exploring church designs with Chinese architectural aesthetic features (such as flying eaves, Dougong brackets, courtyards, etc.) in worship spaces, using expressions closer to Chinese believers' cultural psychology in preaching language, and promoting the Sinicization of church spaces. In talent cultivation, a comprehensive training model of "Chinese Studies + Theology + Social Responsibility" should be established, integrating excellent traditional Chinese culture courses into the core module rather than "supplementary" courses of theological education, increasing the proportion and class hours of Chinese studies content, cultivating local pastoral talents with faith literacy, cultural foundation, and patriotic sentiment, emphasizing cultural literacy in talent selection, practical ability in cultivation, and regional characteristics in usage. In social adaptation, the tradition of loving both country and religion should be promoted, integrating "glorifying God and benefiting people" with "prosperity, democracy, civility, and harmony" at the national level, connecting "Great Harmony" and "Kingdom of God" comparisons with "freedom, equality, justice, and rule of law" at the social level, and resonating "benevolence" and "love" mutual interpretation with "patriotism, dedication, integrity, and friendliness" civic moral norms at the personal level, facilitating deep integration of Christianity with socialist core values and leveraging Christianity's positive role in social services, public welfare, elderly care, and community building.

The innovations of this paper lie in: First, constructing a systematic research framework of "thought-practice-value," examining Wang Zhixin's theological perspective, indigenization practice, and contemporary value within a unified academic perspective, breaking through the limitations of separating thought from practice and history from contemporary practice in previous research. Second, establishing effective connection between historical research and contemporary inspiration, adopting a "historical experience-contemporary issue-transformation path" analytical framework to extract practical implications item by item, achieving bidirectional integration of historical experience and contemporary practice, giving historical research contemporary value. Third, for the first time systematically studying Wang Zhixin from the perspective of Zhejiang regional culture, connecting his thought formation with Zhejiang's regional cultural traditions (Huzhou Confucian tradition, Jiangnan humanistic spirit, Zhejiang pragmatic character), providing the most regionally-relevant historical reference for the Sinicization of Christianity in Zhejiang, and expanding the academic space of Zhejiang Christian figure research. The research results of this paper not only have important academic value for enriching the historical research dimensions of Christianity Sinicization and deepening the theoretical cognition of cross-cultural religious integration, but also provide valuable historical resources for contemporary Christianity to escape the "foreign religion" dilemma, cultivate local pastoral talents, and promote Sinicization practices, while also providing historical reference and practical guidance for Zhejiang Province's religious talents to deepen Christianity Sinicization and inherit the tradition of loving both country and religion."""
        
        # 找到英文摘要区域
        en_start = abs_end
        en_end = None
        for i in range(en_start, len(doc.paragraphs)):
            t = doc.paragraphs[i].text.strip()
            if t.startswith("Key Words") or t.startswith("Key words") or t.startswith("关键词"):
                en_end = i
                break
        
        if en_end:
            for i in range(en_end - 1, en_start, -1):
                if i < len(doc.paragraphs):
                    for run in doc.paragraphs[i].runs:
                        run.text = ""
            
            en_paras = new_en.split("\n\n")
            for j, text in enumerate(en_paras):
                if text.strip():
                    idx = en_start + 1 + j
                    if idx < len(doc.paragraphs):
                        para = doc.paragraphs[idx]
                        para.text = text.strip()
                        for run in para.runs:
                            set_run(run, 'Times New Roman', 12)
                    else:
                        p = doc.add_paragraph(text.strip())
                        for run in p.runs:
                            set_run(run, 'Times New Roman', 12)
            
            print(f"   英文摘要已扩充")
    
    # ==================== 保存 ====================
    print(f"\n保存: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    
    final = 0
    for para in doc.paragraphs:
        final += count_chars(para.text)
    
    print(f"\n最终字数: {final}")
    print(f"5%阈值: {int(final * 0.05)}")
    print(f"中文摘要: {cn_chars}字")
    print(f"修改完成！")

if __name__ == "__main__":
    main()
