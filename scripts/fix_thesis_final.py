#!/usr/bin/env python3
"""完整修改毕业论文：1.扩充摘要≥5% 2.每章≤4节 3.神学思想→神学观"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

INPUT_FILE = "/home/admin/.openclaw/media/inbound/2f0729c6-1a17-4fb4-9c79-5bedbeb2b7c1.docx"
OUTPUT_FILE = "/home/admin/.openclaw/workspace/毕业论文-邱春华-修改版.docx"

def count_chars(text):
    return len(text)

def set_run(run, name='宋体', size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def main():
    print("读取文档...")
    doc = Document(INPUT_FILE)
    
    # 统计全文
    total_chars = 0
    for para in doc.paragraphs:
        total_chars += count_chars(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_chars += count_chars(cell.text)
    
    print(f"全文总字数: {total_chars}")
    min_abstract = int(total_chars * 0.05)
    print(f"摘要需≥{min_abstract}字")
    
    # ==================== 1. 全局替换 ====================
    print("\n1. 替换 '王治心神学思想' → '王治心神学观'...")
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if "王治心神学思想" in run.text:
                run.text = run.text.replace("王治心神学思想", "王治心神学观")
                count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "王治心神学思想" in run.text:
                            run.text = run.text.replace("王治心神学思想", "王治心神学观")
                            count += 1
    print(f"   替换了 {count} 处")
    
    # ==================== 2. 整合章节 ====================
    print("\n2. 整合章节...")
    
    # 需要删除的节标题（只删标题行，内容保留并归入上一节）
    delete_patterns = [
        # 第二章
        r"^第二节\s*本土上帝观的学术论证",
        r"^第三节\s*宗教本土化的三种学理范式辨析",
        # 第三章
        r"^第四节\s*与民国本色神学家的比较",
        r"^第六节\s*历史成效与局限分析",
        # 第四章
        r"^第二节\s*本色化思想的内在张力与时代局限",
        r"^第六节\s*社会适应层面的启示",
    ]
    
    deleted = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for pat in delete_patterns:
            if re.match(pat, text):
                # 只删除正文中的节标题，不删除目录中的
                # 判断是否在正文区域（段落索引>30大概是正文开始）
                if i > 30:
                    for run in para.runs:
                        run.text = ""
                    deleted += 1
                    print(f"   删除: {text[:40]}... (段落{i})")
    
    print(f"   共删除 {deleted} 个节标题")
    
    # ==================== 3. 扩充摘要 ====================
    print("\n3. 扩充摘要...")
    
    # 找到摘要区域
    abs_start = None
    abs_end = None  # Abstract标题
    
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if (t.startswith("摘") and "要" in t) and abs_start is None:
            abs_start = i
        elif abs_start is not None and abs_end is None:
            if t.startswith("Abstract") or t.startswith("abstract"):
                abs_end = i
                break
    
    if abs_start and abs_end:
        # 计算当前摘要
        cur_chars = sum(count_chars(doc.paragraphs[i].text) for i in range(abs_start+1, abs_end))
        print(f"   当前摘要: {cur_chars}字, 需要≥{min_abstract}字")
        
        # 新摘要内容（约2800字中文 + 约2800字英文）
        new_cn = """本文以王治心（1881—1968）为研究对象，系统考察其在民国时期基督教本色化运动中的神学观、实践探索及其对当代基督教中国化的启示。王治心作为前清秀才出身的基督徒知识分子，以其独特的"本色化实践派"立场，在神学思想的本土化建构、教会礼仪的文化革新、宗教史学的独立编撰以及抗日救亡的社会参与等多个维度上，形成了系统而独到的贡献，堪称本色化运动中最为全面、最具实践品格的代表人物之一。

文章认为，王治心的本色化探索主要体现为三个维度：第一，以中华文化资源解释基督教教义。王治心提出"以儒释耶"的神学路径，深入挖掘基督教与儒家、墨家、佛教、道家等中国传统思想资源的契合点，构建了系统的跨文化宗教对话理论。他强调基督教与中国文化的融合"不是在形式方面，乃在精神方面"，反对流于形式的"本色化"尝试，主张在精神内核层面实现基督教信仰与中国文化的深层融通。他的"以儒释耶"不是简单的概念比附，而是在两种思想传统之间建立深层的道德哲学对话——"仁"与"爱"的互释、"天""帝"观与上帝观的融通、"大同"与天国观的比较，都是在精神层面寻找契合点的努力。第二，推动神学教育和礼仪实践本土化。王治心在金陵神学院、福建协和大学、沪江大学等多所院校推行国学教育革新，提出"国学+神学"的课程模式，培养兼具信仰素养和文化素养的本土教牧人才。在礼仪方面，他设计了八大节期方案，将中国传统节日与基督教信仰内涵相结合；提出祭祖问题的折中方案，倡导以"追思礼拜"形式替代传统祭祖仪式，既保持信仰纯正又尊重中国文化传统。第三，在民族危亡中强调爱国与信仰统一。王治心在五卅运动和抗日战争中积极呼吁基督徒投身爱国救亡，从神学层面论证基督教爱国精神的合法性，将"荣神益人"与民族复兴伟业相结合，在1932年出版的《孙文主义与耶稣主义》中系统论证了三民主义与耶稣精神在价值追求上的内在一致性，为当代"爱国爱教"传统提供了历史典范。

其历史经验对当代基督教中国化的启示在于：基督教中国化不应停留于外在形式模仿，而应在四个层面深化推进。文化融合层面，应坚持"精神契合"而非"形式模仿"的原则，深入挖掘基督教与中华优秀传统文化的契合点，在哲学层面探索创造论与"天生烝民"、"道法自然"的对话空间，在伦理层面深化"爱人如己"与"仁者爱人""兼爱天下"的呼应，在社会理想层面比较"上帝之国"与"大同世界"的异曲同工，构建中国特色神学话语体系，以"文化润教"推动宗教中国化的深入发展。教会建设层面，应深化本土化礼仪实践，探索基督教追思礼拜等具有中国文化特色的信仰表达方式，在礼拜音乐方面创作具有中国音乐风格的原创圣诗，在礼拜空间方面探索具有中国建筑美学特征的教堂设计，推进教会空间的中国化建设。人才培养层面，应建立"国学+神学+社会责任感"的综合培养模式，将中华优秀传统文化课程纳入神学教育核心模块而非"补充"课程，培育兼具信仰素养、文化底蕴与家国情怀的本土教牧队伍，优先选拔具有中国文化根基的信徒进入教牧队伍。社会适应层面，应弘扬爱国爱教传统，在国家层面将"荣神益人"思想与"富强、民主、文明、和谐"价值目标相融合，在社会层面将"大同"与"天国"比较思想与"自由、平等、公正、法治"社会取向相对接，在个人层面将"仁""爱"互释思想与"爱国、敬业、诚信、友善"公民道德规范相呼应，推动基督教与社会主义核心价值观深度融合，发挥基督教在社会服务、公益慈善、养老服务、社区建设等领域的积极作用。

本文的创新之处在于：第一，构建"思想—实践—价值"的系统性研究框架，将王治心的神学观、本色化实践与当代价值置于统一的学术视野中进行考察，突破了以往研究中思想与实践分离、历史与当代脱节的局限。第二，在历史研究与当代启示之间建立有效衔接，采用"历史经验—当代问题—转化路径"的分析框架逐条提炼现实启示，实现历史经验与当代实践的双向贯通。第三，首次从浙江地域视角系统研究王治心，将其思想形成与浙江地域文化传统（湖州儒学传统、江南人文精神、浙江务实品格）相关联，为浙江基督教中国化提供最具地域亲和力的历史参照，拓展了浙江基督教人物研究的学术空间。本文的研究成果不仅对丰富基督教中国化历史研究维度、深化跨文化宗教融合理论认知具有重要学术价值，更为当代基督教摆脱"洋教"困境、培养本土教牧人才、推进中国化实践提供了可资借鉴的历史资源，也为浙江省宗教界"双通"人才深耕基督教中国化、传承爱国爱教传统提供了历史参考与实践思路。"""
        
        # 删除旧摘要内容
        for i in range(abs_end - 1, abs_start, -1):
            if i < len(doc.paragraphs):
                for run in doc.paragraphs[i].runs:
                    run.text = ""
        
        # 插入新摘要
        cn_paras = new_cn.split("\n\n")
        for j, text in enumerate(cn_paras):
            if text.strip():
                idx = abs_start + 1 + j
                if idx < len(doc.paragraphs):
                    para = doc.paragraphs[idx]
                    para.text = text.strip()
                    for run in para.runs:
                        set_run(run, '宋体', 12)
                else:
                    p = doc.add_paragraph(text.strip())
                    for run in p.runs:
                        set_run(run, '宋体', 12)
        
        cn_chars = count_chars(new_cn)
        print(f"   中文摘要: {cn_chars}字")
        
        # 扩充英文摘要
        new_en = """This paper takes Wang Zhixin (1881-1968) as the research object, systematically examining his theological perspective, practical exploration during the indigenous church movement in the Republican era, and their implications for the contemporary Sinicization of Christianity. As a Christian intellectual with a background as a Qing Dynasty scholar, Wang Zhixin, with his unique "indigenization practice school" stance, made systematic and distinctive contributions in multiple dimensions including the localization of theological construction, cultural innovation of church liturgy, independent compilation of religious historiography, and social participation in the anti-Japanese national salvation movement. He stands as one of the most comprehensive and practically-oriented representative figures in the indigenous church movement.

The paper argues that Wang Zhixin's indigenization exploration is mainly manifested in three dimensions: First, interpreting Christian doctrines through Chinese cultural resources. Wang proposed the theological path of "interpreting Christianity through Confucianism," deeply exploring the convergence points between Christianity and traditional Chinese thought resources such as Confucianism, Mohism, Buddhism, and Daoism, constructing a systematic cross-cultural religious dialogue theory. He emphasized that the integration of Christianity and Chinese culture lies "not in form but in spirit," opposing superficial "indigenization" attempts and advocating for deep integration at the spiritual core level. His approach was not a simple conceptual analogy, but rather established a deep moral-philosophical dialogue between two intellectual traditions—the mutual interpretation of "ren" (benevolence) and "love," the integration of "Tian" (Heaven) and "Di" (Lord) concepts with the concept of God, and the comparison of "Datong" (Great Harmony) with the Kingdom of God—all representing efforts to find convergence at the spiritual level. Second, promoting the localization of theological education and liturgical practices. Wang implemented Chinese studies education reforms at Jinling Seminary, Fujian Union University, and Jiangsu University, proposing a "Chinese Studies + Theology" curriculum model to cultivate local pastoral talents with both faith literacy and cultural literacy. In liturgy, he designed an eight-festival scheme integrating Chinese traditional festivals with Christian faith connotations, proposed a compromise solution for ancestor worship advocating "memorial services" as alternatives to traditional rituals, maintaining both faith purity and respect for Chinese cultural traditions. Third, emphasizing the unity of patriotism and faith amid national crisis. Wang actively called on Christians to participate in patriotic salvation movements during the May 30th Incident and the Anti-Japanese War, providing theological justification for Christian patriotism and connecting "glorifying God and benefiting people" with national rejuvenation. In his 1932 work "Sun Yat-sen's Principles and Jesus' Principles," he systematically demonstrated the内在 consistency between the Three Principles of the People and the spirit of Jesus in value pursuit, providing a historical model for the contemporary tradition of "loving both country and religion."

His historical experience offers important implications for the contemporary Sinicization of Christianity: rather than remaining at the level of superficial formal imitation, the Sinicization of Christianity should be deepened in four aspects. In cultural integration, the principle of "spiritual convergence" rather than "formal imitation" should be upheld, deeply exploring convergence points between Christianity and excellent traditional Chinese culture—exploring dialogue spaces between creation theory and "Tian sheng zheng min" and "Dao fa zi ran" at the philosophical level, deepening the resonance between "love your neighbor as yourself" and "ren zhe ai ren" and "jian ai tian xia" at the ethical level, comparing the "Kingdom of God" with "Datong" at the social ideal level—constructing a theological discourse system with Chinese characteristics and promoting religious Sinicization through "cultural nourishment." In church building, localized liturgical practices should be deepened, exploring faith expressions with Chinese cultural characteristics such as Christian memorial services, creating original hymns with Chinese musical style in worship music, exploring church designs with Chinese architectural aesthetics in worship spaces, and promoting the Sinicization of church spaces. In talent cultivation, a comprehensive training model of "Chinese Studies + Theology + Social Responsibility" should be established, integrating excellent traditional Chinese culture courses into the core module rather than "supplementary" courses of theological education, cultivating local pastoral talents with faith literacy, cultural foundation, and patriotic sentiment, and prioritizing the selection of believers with Chinese cultural roots for pastoral work. In social adaptation, the tradition of loving both country and religion should be promoted, integrating "glorifying God and benefiting people" with "prosperity, democracy, civility, and harmony" at the national level, connecting "Datong" and "Kingdom of God" comparisons with "freedom, equality, justice, and rule of law" at the social level, and resonating "ren" and "love" mutual interpretation with "patriotism, dedication, integrity, and friendliness"公民道德 norms at the personal level, facilitating deep integration of Christianity with socialist core values and leveraging Christianity's positive role in social services, public welfare, elderly care, and community building.

The innovations of this paper lie in: First, constructing a systematic research framework of "thought-practice-value," examining Wang Zhixin's theological perspective, indigenization practice, and contemporary value within a unified academic perspective, breaking through the limitations of separating thought from practice and history from contemporary practice in previous research. Second, establishing effective connection between historical research and contemporary inspiration, adopting a "historical experience—contemporary issue—transformation path" analytical framework to extract practical implications item by item, achieving bidirectional integration of historical experience and contemporary practice. Third, for the first time systematically studying Wang Zhixin from the perspective of Zhejiang regional culture, connecting his thought formation with Zhejiang's regional cultural traditions (Huzhou Confucian tradition, Jiangnan humanistic spirit, Zhejiang pragmatic character), providing the most regionally-relevant historical reference for the Sinicization of Christianity in Zhejiang, and expanding the academic space of Zhejiang Christian figure research. The research results of this paper not only have important academic value for enriching the historical research dimensions of Christianity Sinicization and deepening the theoretical cognition of cross-cultural religious integration, but also provide valuable historical resources for contemporary Christianity to escape the "foreign religion" dilemma, cultivate local pastoral talents, and promote Sinicization practices, while also providing historical reference and practical guidance for Zhejiang Province's religious "dual通" talents to deepen Christianity Sinicization and inherit the tradition of loving both country and religion."""
        
        # 找到英文摘要区域
        en_start = abs_end
        en_end = None
        for i in range(en_start, len(doc.paragraphs)):
            t = doc.paragraphs[i].text.strip()
            if t.startswith("Key Words") or t.startswith("Key words") or t.startswith("关键词"):
                en_end = i
                break
        
        if en_end:
            # 删除旧英文摘要
            for i in range(en_end - 1, en_start, -1):
                if i < len(doc.paragraphs):
                    for run in doc.paragraphs[i].runs:
                        run.text = ""
            
            # 插入新英文摘要
            en_paras = new_en.split("\n\n")
            for j, text in enumerate(en_paras):
                if text.strip():
                    idx = en_start + 1 + j
                    if idx < len(doc.paragraphs):
                        para = doc.paragraphs[idx]
                        para.text = text.strip()
                        for run in para.runs:
                            set_run(run, 'Times New Roman', 12)
                    else:
                        p = doc.add_paragraph(text.strip())
                        for run in p.runs:
                            set_run(run, 'Times New Roman', 12)
            
            en_chars = count_chars(new_en)
            print(f"   英文摘要: {en_chars}字")
    
    # ==================== 保存 ====================
    print(f"\n保存: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    
    # 最终统计
    final = 0
    for para in doc.paragraphs:
        final += count_chars(para.text)
    
    print(f"\n最终字数: {final}")
    print(f"摘要需≥{int(final * 0.05)}字")
    print(f"修改完成！")

if __name__ == "__main__":
    main()
