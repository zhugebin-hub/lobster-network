#!/usr/bin/env python3
"""按导师意见全面修改论文 v3"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re

INPUT_FILE = "/home/admin/.openclaw/workspace/毕业论文-邱春华-修改版.docx"
OUTPUT_FILE = "/home/admin/.openclaw/workspace/毕业论文-邱春华-导师修改版.docx"

def set_run(run, name='宋体', size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def main():
    print("读取文档...")
    doc = Document(INPUT_FILE)
    
    # 找到正文开始（引言内容开始处）
    body_start = None
    for i, para in enumerate(doc.paragraphs):
        if '基督教中国化是当前中国宗教工作的核心议题之一' in para.text:
            body_start = i
            break
    
    print(f"正文开始: 段落{body_start}")
    
    # ==================== 第一章修改 ====================
    print("\n=== 修改第一章 ===")
    
    # 找到第一章正文范围
    ch1_start = None
    ch1_end = None
    for i in range(body_start, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if t.startswith('第一章') and ch1_start is None:
            ch1_start = i
        elif ch1_start is not None and ch1_end is None:
            if t.startswith('第二章'):
                ch1_end = i
                break
    
    if ch1_start and ch1_end:
        print(f"  第一章范围: 段落{ch1_start}到{ch1_end}")
        
        # 找到各节位置
        sections = []
        for i in range(ch1_start, ch1_end):
            t = doc.paragraphs[i].text.strip()
            if re.match(r'^第[一二三四五六七八九十]+节\s', t):
                sections.append(i)
        
        print(f"  找到{len(sections)}节")
        for idx, sec in enumerate(sections):
            print(f"    第{idx+1}节: 段落{sec} - {doc.paragraphs[sec].text.strip()[:40]}")
        
        # 修改"佛教中国化"相关节标题为"其他宗教本土化经验的参照意义"
        for i in range(ch1_start, ch1_end):
            t = doc.paragraphs[i].text.strip()
            if '佛教中国化' in t and ('节' in t or i in sections):
                doc.paragraphs[i].text = '第四节 其他宗教本土化经验的参照意义'
                for run in doc.paragraphs[i].runs:
                    set_run(run, '宋体', 14, bold=True)
                print(f"  修改节标题: 其他宗教本土化经验的参照意义 (段落{i})")
                break
        
        # 压缩佛教中国化相关内容
        for i in range(ch1_start, ch1_end):
            t = doc.paragraphs[i].text
            if '佛教中国化的成功经验为基督教中国化提供了参照' in t or ('佛教中国化' in t and '王治心' in t and len(t) > 100):
                new_text = '佛教中国化的成功经验为基督教中国化提供了重要参照。王治心在《基督徒之佛学研究》中深入分析了佛教中国化的历程，指出佛教之所以能够在中国扎根并发展，关键在于它成功地与中国传统文化相融合。但他也强调，基督教与佛教在教义上存在根本差异，基督教的一神信仰、创世论、救赎论等核心教义不能改变。这一认识对王治心的本色化思想产生了深远影响，他强调基督教与中国文化的融合"不是在形式方面，乃在精神方面"。'
                doc.paragraphs[i].text = new_text
                for run in doc.paragraphs[i].runs:
                    set_run(run, '宋体', 12)
                print(f"  压缩佛教中国化内容 (段落{i})")
                break
    
    # ==================== 第二章修改 ====================
    print("\n=== 修改第二章 ===")
    
    # 找到第二章正文范围
    ch2_start = None
    ch2_end = None
    for i in range(body_start, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if t.startswith('第二章') and ch2_start is None:
            ch2_start = i
        elif ch2_start is not None and ch2_end is None:
            if t.startswith('第三章'):
                ch2_end = i
                break
    
    if ch2_start and ch2_end:
        print(f"  第二章范围: 段落{ch2_start}到{ch2_end}")
        
        # 找到各节位置
        sections = []
        for i in range(ch2_start, ch2_end):
            t = doc.paragraphs[i].text.strip()
            if re.match(r'^第[一二三四五六七八九十]+节\s', t):
                sections.append(i)
        
        print(f"  找到{len(sections)}节")
        for idx, sec in enumerate(sections):
            print(f"    第{idx+1}节: 段落{sec} - {doc.paragraphs[sec].text.strip()[:40]}")
        
        # 在每节结尾添加"当代转化提示"
        transformation_tips = [
            '当代转化提示：王治心"以儒释耶"的神学路径启示我们，当代基督教中国化的神学建设应当深入挖掘中华优秀传统文化资源，在保持信仰核心的前提下实现文化层面的创造性转化。但需注意，儒学与基督教在终极关怀层面存在根本差异，不能简单等同，而应在对话中寻求互补。',
            '当代转化提示：王治心以"天""上帝"等中国传统概念对接基督教上帝观的尝试，为当代神学表达的本土化提供了重要参照。但需要注意概念错位的神学风险，在借鉴传统术语时应明确其与基督教位格性上帝之间的差异，避免简单比附。',
            '当代转化提示：王治心以耶儒对话为核心、以耶墨对话为补充的跨文化宗教思想，为当代宗教对话提供了多维框架。当代基督教中国化可以借鉴其"精神融合"的对话原则，在挖掘传统文化资源时应突出最能支撑中国化主线的儒家和墨家思想，佛道思想可作为辅助性资源。',
            '当代转化提示：王治心在民族救亡中论证爱国与信仰统一的神学立场，为当代"爱国爱教"传统提供了历史典范。当代基督教中国化应继续弘扬这一传统，在保持信仰纯正的同时，积极参与社会服务、推动基督教与社会主义核心价值观的融合。'
        ]
        
        # 在各节结尾添加提示
        for idx in range(len(sections)):
            if idx < len(transformation_tips):
                sec_start = sections[idx]
                sec_end = sections[idx + 1] if idx + 1 < len(sections) else ch2_end
                
                # 找到该节最后一个非空段落
                insert_pos = sec_end - 1
                for j in range(sec_end - 1, sec_start, -1):
                    if doc.paragraphs[j].text.strip():
                        insert_pos = j + 1
                        break
                
                if insert_pos < len(doc.paragraphs):
                    para = doc.paragraphs[insert_pos]
                    para.text = transformation_tips[idx]
                    for run in para.runs:
                        set_run(run, '宋体', 12)
                    print(f"  在第{idx+1}节结尾添加当代转化提示 (段落{insert_pos})")
    
    # ==================== 第三章修改 ====================
    print("\n=== 修改第三章 ===")
    
    # 找到第三章正文范围
    ch3_start = None
    ch3_end = None
    for i in range(body_start, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if t.startswith('第三章') and ch3_start is None:
            ch3_start = i
        elif ch3_start is not None and ch3_end is None:
            if t.startswith('第四章'):
                ch3_end = i
                break
    
    if ch3_start and ch3_end:
        print(f"  第三章范围: 段落{ch3_start}到{ch3_end}")
        
        # 找到各节位置
        sections = []
        for i in range(ch3_start, ch3_end):
            t = doc.paragraphs[i].text.strip()
            if re.match(r'^第[一二三四五六七八九十]+节\s', t):
                sections.append(i)
        
        print(f"  找到{len(sections)}节")
        for idx, sec in enumerate(sections):
            print(f"    第{idx+1}节: 段落{sec} - {doc.paragraphs[sec].text.strip()[:40]}")
        
        # 在各节开头添加判断句
        judgment_sentences = [
            '王治心推动神学教育本土化，核心不是增加几门国学课程，而是试图改变基督教人才对中国文化的疏离状态。',
            '王治心的礼仪改革实践表明，基督教中国化不能停留在外在形式的改变，而应当在精神内核层面实现信仰与中国文化的深度融合。',
            '王治心通过宗教著作的编撰，彰显了中国人自主书写基督教历史与思想的主体意识，打破了西方学者对中国基督教史的学术垄断。',
            '王治心的社会参与实践体现了爱国爱教与宗教社会责任的统一，证明了基督教信仰与民族命运可以相互成就而非彼此对立。'
        ]
        
        for idx in range(len(sections)):
            if idx < len(judgment_sentences):
                sec_start = sections[idx]
                
                # 在节标题后第一段开头添加判断句
                if sec_start + 1 < len(doc.paragraphs):
                    para = doc.paragraphs[sec_start + 1]
                    current_text = para.text
                    if current_text and not current_text.startswith(judgment_sentences[idx]):
                        new_text = judgment_sentences[idx] + '\n\n' + current_text
                        para.text = new_text
                        for run in para.runs:
                            set_run(run, '宋体', 12)
                        print(f"  在第{idx+1}节开头添加判断句 (段落{sec_start+1})")
    
    # ==================== 第四章修改 ====================
    print("\n=== 修改第四章 ===")
    
    # 找到第四章正文范围
    ch4_start = None
    ch4_end = None
    for i in range(body_start, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if t.startswith('第四章') and ch4_start is None:
            ch4_start = i
        elif ch4_start is not None and ch4_end is None:
            if t.startswith('结'):
                ch4_end = i
                break
    
    if ch4_start and ch4_end:
        print(f"  第四章范围: 段落{ch4_start}到{ch4_end}")
        
        # 找到各节位置
        sections = []
        for i in range(ch4_start, ch4_end):
            t = doc.paragraphs[i].text.strip()
            if re.match(r'^第[一二三四五六七八九十]+节\s', t):
                sections.append(i)
        
        print(f"  找到{len(sections)}节")
        for idx, sec in enumerate(sections):
            print(f"    第{idx+1}节: 段落{sec} - {doc.paragraphs[sec].text.strip()[:40]}")
        
        # 在各节开头添加"历史经验—当代问题—转化路径"结构化内容
        chapter4_tips = [
            '历史经验：王治心强调基督教与中国文化的融合"不是在形式方面，乃在精神方面"。当代问题：当前一些教会的"中国化"尝试往往停留在加装中国元素等形式层面。转化路径：坚持"精神契合"而非"形式模仿"的原则，深入挖掘基督教与中华优秀传统文化的契合点。',
            '历史经验：王治心设计了八大节期方案，提出祭祖问题的折中方案。当代问题：当前礼仪中国化多停留在形式层面，祭祖问题在农村教会中仍是难题。转化路径：在保持信仰纯正的前提下，探索具有中国文化特色的"基督教追思礼拜"。',
            '历史经验：王治心自身"国学+神学"的知识结构证明了培养兼具文化素养和信仰素养的本土教牧人才是可行的。当代问题：当前神学院校国学课程往往作为"补充"而非"核心"存在。转化路径：建立"国学+神学+社会责任感"的综合培养模式。',
            '历史经验：王治心在五卅运动和抗日战争中的爱国表现，论证了基督教信仰与爱国情感可以在"为正义、真理奋斗"的神学框架中实现有机统一。当代问题：当代基督教需要进一步融入社会主义核心价值观，发挥社会服务功能。转化路径：弘扬爱国爱教传统，推动基督教与社会主义核心价值观相融合。'
        ]
        
        for idx in range(len(sections)):
            if idx < len(chapter4_tips):
                sec_start = sections[idx]
                
                # 在节标题后添加结构化内容
                if sec_start + 1 < len(doc.paragraphs):
                    para = doc.paragraphs[sec_start + 1]
                    current_text = para.text
                    if current_text and not current_text.startswith('历史经验'):
                        new_text = chapter4_tips[idx] + '\n\n' + current_text
                        para.text = new_text
                        for run in para.runs:
                            set_run(run, '宋体', 12)
                        print(f"  在第{idx+1}节添加结构化内容 (段落{sec_start+1})")
    
    # ==================== 结语修改 ====================
    print("\n=== 修改结语 ===")
    
    # 找到结语正文范围
    conclusion_start = None
    conclusion_end = None
    for i in range(body_start, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if t.startswith('结') and '语' in t and conclusion_start is None:
            conclusion_start = i
        elif conclusion_start is not None and conclusion_end is None:
            if t.startswith('参考'):
                conclusion_end = i
                break
    
    if conclusion_start and conclusion_end:
        print(f"  结语范围: 段落{conclusion_start}到{conclusion_end}")
        
        # 简化结语为三点
        new_conclusion = """结  语

一、研究结论

本文通过对王治心神学观的核心内涵、基督教中国化的实践路径及其当代价值的系统考察，得出以下三条主要结论：

第一，王治心的本色化探索表明，基督教中国化并非外在装饰，而是信仰表达、文化解释、人才培养和社会责任的整体转化。其"以儒释耶"的神学路径、"精神融合"的本土化原则以及四维一体的实践体系，为当代基督教中国化提供了至今仍有启发意义的思想范式。

第二，王治心思想与实践的当代价值不在于提供可直接照搬的具体方案，而在于坚持以中华文化为思想资源、以国家认同为基本立场、以社会责任为实践路径，推动基督教在中国社会中实现更深层次的适应。对当代基督教中国化而言，其启示不在于复制民国时期的礼仪方案，而在于把握"精神契合"而非"形式模仿"的根本原则。

第三，本文从浙江地域视角考察王治心，揭示了其思想形成与浙江地域文化的深层关联，为浙江基督教中国化提供了具有地域亲和力的历史参照。浙江是王治心的故乡，湖州的儒学传统、江南的人文精神、浙江的务实品格都对其本色化探索产生了深刻影响。这一地域视角的引入，丰富了王治心研究的维度，也为浙江基督教中国化的实践提供了历史资源与文化自信。

二、研究不足与展望

本文研究仍存在较为明显的文献短板与研究局限，有待后续进一步深化完善。在文献搜集层面，本文主要依托王治心已公开出版的专著、期刊文章及学界现有研究成果展开分析，但王治心在金陵神学院任教期间撰写的多部核心讲义均为未刊馆藏文献，尚未公开出版发行，现有公开资料难以完整覆盖其思想全貌。后续研究可通过专业档案渠道，系统搜集、整理、释读相关未刊文献，补齐现有文献短板，实现对王治心思想体系的全方位、深层次解读。

除文献局限外，本文对王治心与民国其他本色神学家的横向比较仍有拓展空间，未能完全厘清不同神学流派的核心差异与谱系特征；同时，对其"儒耶融合"思想的学理范式辨析、思想内在张力的深挖仍不够透彻。后续可立足文化翻译、文化对话、文化融合的学理框架，精准界定其思想范式，并结合当代基督教中国化的具体实践，细化本土化落地路径，进一步提升研究的理论深度与现实价值。

本研究也期望为浙江省宗教界"双通"人才深耕基督教中国化、传承爱国爱教传统提供历史参考与实践思路。"""
        
        # 替换结语内容
        for i in range(conclusion_end - 1, conclusion_start, -1):
            if i < len(doc.paragraphs):
                for run in doc.paragraphs[i].runs:
                    run.text = ""
        
        conclusion_paras = new_conclusion.split('\n\n')
        for j, text in enumerate(conclusion_paras):
            if text.strip():
                idx = conclusion_start + 1 + j
                if idx < len(doc.paragraphs):
                    para = doc.paragraphs[idx]
                    para.text = text.strip()
                    for run in para.runs:
                        set_run(run, '宋体', 12)
                else:
                    p = doc.add_paragraph(text.strip())
                    for run in p.runs:
                        set_run(run, '宋体', 12)
        
        print(f"  结语已简化为三点")
    
    # ==================== 保存 ====================
    print(f"\n保存: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    print(f"修改完成！")

if __name__ == "__main__":
    main()
