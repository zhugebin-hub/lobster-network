#!/usr/bin/env python3
"""Build the complete formatted thesis with cover page + pandoc content."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml.etree import SubElement
import copy

def set_rfonts(run, font_name='宋体'):
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = SubElement(rpr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)

def fix_run(run, font_name='宋体', font_size=Pt(12), bold=False):
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    set_rfonts(run, font_name)
    # Remove any color
    color_elem = run._element.find(qn('w:color'))
    if color_elem is not None:
        run._element.remove(color_elem)

def add_run(p, text, font_name='宋体', font_size=Pt(12), bold=False):
    run = p.add_run(text)
    fix_run(run, font_name, font_size, bold)
    return run

def add_para(doc, text, font_name='宋体', font_size=Pt(12), bold=False, align=None, spacing=1.5, indent=0.74, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    if indent > 0:
        p.paragraph_format.first_line_indent = Cm(indent)
    if text:
        add_run(p, text, font_name, font_size, bold)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_run(p, part[2:-2], '宋体', Pt(12), True)
        elif part:
            add_run(p, part, '宋体', Pt(12), False)
    return p

def process_paragraph(p):
    """Process a paragraph from the pandoc doc and apply formatting."""
    text = p.text.strip()
    if not text:
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        return
    
    p.paragraph_format.line_spacing = 1.5
    
    is_chapter = False
    is_section = False
    is_subsection = False
    
    # Chapter titles
    if re.match(r'^第[一二三四五六七八九十]+章\s', text) or text.startswith('参考文献') or text == '结  语' or text == '结语':
        is_chapter = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
        p.paragraph_format.first_line_indent = 0
    elif re.match(r'^第[一二三四五六七八九十]+节\s', text):
        is_section = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.first_line_indent = 0
    elif re.match(r'^[一二三四五六七八九十]+[、\.]', text):
        is_subsection = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = 0
    elif re.match(r'^[（(][一二三四五六七八九十]+[）)]', text):
        is_subsection = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = 0
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0.74)
    
    for run in p.runs:
        if is_chapter:
            fix_run(run, '黑体', Pt(16), True)
        elif is_section:
            fix_run(run, '黑体', Pt(14), True)
        elif is_subsection:
            fix_run(run, '黑体', Pt(12), True)
        else:
            run_text = run.text.strip()
            if run_text and all(c.isascii() for c in run_text) and len(run_text) > 3:
                fix_run(run, 'Times New Roman', Pt(12), run.bold)
            else:
                fix_run(run, '宋体', Pt(12), run.bold)

import re

def main():
    # First, convert markdown to docx using pandoc
    import subprocess
    md_path = '/home/admin/.openclaw/workspace/论文_新时代道教神仙信仰的传承与中国化路径.md'
    temp_docx = '/tmp/thesis_temp.docx'
    output_path = '/home/admin/.openclaw/workspace/新时代道教神仙信仰的传承与中国化路径_完整版.docx'
    
    subprocess.run(['pandoc', md_path, '-o', temp_docx, '--from', 'markdown', '--to', 'docx', '--standalone'], check=True)
    
    # Now build the final document
    final = Document()
    
    # Page setup
    section = final.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    
    # === COVER PAGE ===
    for _ in range(6):
        final.add_paragraph().paragraph_format.space_after = Pt(12)
    
    add_para(final, '浙江省宗教界"双通"人才研修班', '黑体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 12, 6)
    add_para(final, '毕业论文', '黑体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 12, 24)
    add_para(final, '新时代道教神仙信仰的传承与中国化路径', '黑体', Pt(18), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 0, 36)
    add_para(final, '作    者：戴泰信', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 0, 6)
    add_para(final, '指导教师：XXX', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 0, 6)
    add_para(final, '宗    教：道教', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 0, 6)
    add_para(final, '2026 年 5 月', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 0, 0)
    
    final.add_page_break()
    
    # === CHINESE ABSTRACT ===
    add_para(final, '摘  要', '黑体', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 12, 12)
    
    abstract_cn = (
        '新时代坚持宗教中国化方向，是以习近平同志为核心的党中央关于宗教工作的重大战略部署。'
        '道教作为中华大地上唯一土生土长的宗教，其中国化进程具有特殊的历史地位与现实意义。'
        '神仙信仰是道教的核心标志与信仰基石，神仙谱系的历史演变即是道教发展史的缩影。'
        '然而，在新时代社会深刻变革的背景下，道教神仙信仰的传承面临信仰理解浅表化、传承主体断层、'
        '科仪商业化倾向以及与当代价值体系对接不足等突出问题。如何在保持信仰本质的同时实现与新时代的适应与融合，'
        '成为道教神仙信仰传承发展亟待破解的核心命题。'
        '本文以"新时代道教神仙信仰的传承与中国化路径"为研究主题，综合运用经典文献分析法、政策文本解读法与案例研究法，'
        '从历史纵深、现实境遇、理论逻辑与实践路径四个维度展开系统研究。'
        '全文核心论点为：道教神仙信仰的中国化并非对传统的否定或消解，而是在"返本开新"中实现信仰的创造性转化与创新性发展。'
        '所谓"返本"，即回归"仙道贵生，无量度人"的信仰本源；所谓"开新"，即在新时代语境下重构神仙信仰的时代内涵、'
        '修持实践、组织形态与传播方式。'
        '研究首先梳理道教神仙信仰从先秦思想渊源到当代存续形态的历史演变脉络，揭示神仙信仰"随方设教"的自我调适传统；'
        '其次考察新时代神仙信仰传承面临的现实困境，论证中国化的必要性；'
        '再次从历史逻辑、经典依据和现实基础三个层面阐释神仙信仰中国化的理论合法性；'
        '最终构建以教义阐释中国化为灵魂、修持实践中国化为关键、组织形态中国化为基础、传播方式中国化为手段的系统性实践路径。'
        '本文的创新之处在于：首次将道教神仙信仰研究与中国化路径问题系统对接，填补了该交叉领域的研究空白；'
        '提出了神仙信仰中国化的"返本开新"理论框架；并以典型案例为支撑，为道教界推进中国化实践提供了学术参考。'
    )
    add_body(final, abstract_cn)
    
    p = final.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_run(p, '关键词：', '宋体', Pt(12), True)
    add_run(p, '道教神仙信仰；中国化；传承；新时代；返本开新；现代化转型', '宋体', Pt(12), False)
    
    final.add_page_break()
    
    # === ENGLISH ABSTRACT ===
    add_para(final, 'Abstract', 'Times New Roman', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 12, 12)
    
    p = final.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, 'In the new era, adhering to the Sinicization of religion is a major strategic deployment of the Party Central Committee with Comrade Xi Jinping at its core regarding religious work. As the only indigenous religion in China, Daoism holds a special historical position and practical significance in the Sinicization process. Immortal belief (Shenxian belief) is the core symbol and foundation of Daoist faith, and the historical evolution of the immortal pantheon is a microcosm of Daoist development history. However, against the backdrop of profound social transformation in the new era, the inheritance of Daoist immortal belief faces prominent challenges including superficial understanding, generational gaps in inheritance subjects, commercialization tendencies in rituals, and insufficient integration with contemporary value systems. How to achieve adaptation and integration with the new era while maintaining the essence of faith has become a core issue for the inheritance and development of Daoist immortal belief. This paper takes "the Inheritance and Sinicization Path of Daoist Immortal Belief in the New Era" as the research theme, comprehensively using classical literature analysis, policy text interpretation, and case study methods to conduct systematic research from four dimensions: historical depth, contemporary reality, theoretical logic, and practical pathways.', 'Times New Roman', Pt(12), False)
    
    p = final.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_run(p, 'Key Words: ', 'Times New Roman', Pt(12), True)
    add_run(p, 'Daoist Immortal Belief; Sinicization; Inheritance; New Era; Returning to Roots and Opening Up the New; Modernization Transformation', 'Times New Roman', Pt(12), False)
    
    final.add_page_break()
    
    # === TOC ===
    add_para(final, '目  录', '黑体', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 12, 12)
    toc_items = [
        '摘  要', 'Abstract', '引  言',
        '第一章 道教神仙信仰的历史演变与核心内涵',
        '  第一节 神仙信仰的思想渊源',
        '  第二节 神仙信仰的历史演变',
        '  第三节 神仙信仰的核心内涵',
        '第二章 新时代道教神仙信仰传承的现实境遇',
        '  第一节 新时代社会变迁对神仙信仰的冲击',
        '  第二节 神仙信仰传承面临的突出问题',
        '  第三节 神仙信仰当代存续形态的考察',
        '第三章 道教神仙信仰中国化的理论依据与内在逻辑',
        '  第一节 "中国化"的理论内涵',
        '  第二节 神仙信仰中国化的历史逻辑',
        '  第三节 神仙信仰中国化的经典依据',
        '  第四节 神仙信仰中国化的现实依据',
        '第四章 新时代道教神仙信仰中国化的实践路径',
        '  第一节 教义阐释的中国化',
        '  '