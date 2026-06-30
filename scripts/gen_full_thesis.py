#!/usr/bin/env python3
"""Generate a fully formatted thesis from the complete markdown source."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml.etree import SubElement
import re

def set_rfonts(run, font_name='宋体'):
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = SubElement(rpr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)

def add_run(p, text, font_name='宋体', font_size=Pt(12), bold=False):
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    set_rfonts(run, font_name)
    return run

def add_para(doc, text, font_name='宋体', font_size=Pt(12), bold=False, align=None, spacing=1.5, indent=0.74, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    if indent > 0:
        p.paragraph_format.first_line_indent = Cm(indent)
    if text:
        add_run(p, text, font_name, font_size, bold)
    return p

def add_body(doc, text, indent=0.74):
    """Add body text with **bold** support."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent > 0:
        p.paragraph_format.first_line_indent = Cm(indent)
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_run(p, part[2:-2], '宋体', Pt(12), True)
        elif part:
            add_run(p, part, '宋体', Pt(12), False)
    return p

# ===== Cover Page =====
def add_cover(doc):
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    
    add_para(doc, '浙江省宗教界"双通"人才研修班', '黑体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 12, 6)
    add_para(doc, '毕业论文', '黑体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 12, 24)
    add_para(doc, '新时代道教神仙信仰的传承与中国化路径', '黑体', Pt(18), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 0, 36)
    add_para(doc, '作    者：戴泰信', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 0, 6)
    add_para(doc, '指导教师：XXX', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 0, 6)
    add_para(doc, '宗    教：道教', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 0, 6)
    add_para(doc, '2026 年 5 月', '宋体', Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 0, 0)

# ===== Abstract =====
def add_abstract(doc):
    add_para(doc, '摘  要', '黑体', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 12, 12)
    
    abstract_text = (
        '新时代坚持宗教中国化方向，是以习近平同志为核心的党中央关于宗教工作的重大战略部署。'
        '道教作为中华大地上唯一土生土长的宗教，其中国化进程具有特殊的历史地位与现实意义。'
        '神仙信仰是道教的核心标志与信仰基石，神仙谱系的历史演变即是道教发展史的缩影。'
        '然而，在新时代社会深刻变革的背景下，道教神仙信仰的传承面临信仰理解浅表化、传承主体断层、'
        '科仪商业化倾向以及与当代价值体系对接不足等突出问题。如何在保持信仰本质的同时实现与新时代的适应与融合，'
        '成为道教神仙信仰传承发展亟待破解的核心命题。'
        '本文以"新时代道教神仙信仰的传承与中国化路径"为研究主题，综合运用经典文献分析法、政策文本解读法与案例研究法，'
        '从历史纵深、现实境遇、理论逻辑与实践路径四个维度展开系统研究。'
        '全文核心论点为：道教神仙信仰的中国化并非对传统的否定或消解，而是在"返本开新"中实现信仰的创造性转化与创新性发展。'
        '所谓"返本"，即回归"仙道贵生，无量度人"的信仰本源；所谓"开新"，即在新时代语境下重构神仙信仰的时代内涵、'
        '修持实践、组织形态与传播方式。'
        '研究首先梳理道教神仙信仰从先秦思想渊源到当代存续形态的历史演变脉络，揭示神仙信仰"随方设教"的自我调适传统；'
        '其次考察新时代神仙信仰传承面临的现实困境，论证中国化的必要性；'
        '再次从历史逻辑、经典依据和现实基础三个层面阐释神仙信仰中国化的理论合法性；'
        '最终构建以教义阐释中国化为灵魂、修持实践中国化为关键、组织形态中国化为基础、传播方式中国化为手段的系统性实践路径。'
        '本文的创新之处在于：首次将道教神仙信仰研究与中国化路径问题系统对接，填补了该交叉领域的研究空白；'
        '提出了神仙信仰中国化的"返本开新"理论框架；并以典型案例为支撑，为道教界推进中国化实践提供了学术参考。'
    )
    add_body(doc, abstract_text)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_run(p, '关键词：', '宋体', Pt(12), True)
    add_run(p, '道教神仙信仰；中国化；传承；新时代；返本开新；现代化转型', '宋体', Pt(12), False)

def add_eng_abstract(doc):
    add_para(doc, 'Abstract', 'Times New Roman', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 12, 12)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, 'In the new era, adhering to the Sinicization of religion is a major strategic deployment of the Party Central Committee with Comrade Xi Jinping at its core regarding religious work. As the only indigenous religion in China, Daoism holds a special historical position and practical significance in the Sinicization process. Immortal belief (Shenxian belief) is the core symbol and foundation of Daoist faith, and the historical evolution of the immortal pantheon is a microcosm of Daoist development history. However, against the backdrop of profound social transformation in the new era, the inheritance of Daoist immortal belief faces prominent challenges including superficial understanding, generational gaps in inheritance subjects, commercialization tendencies in rituals, and insufficient integration with contemporary value systems. How to achieve adaptation and integration with the new era while maintaining the essence of faith has become a core issue亟待破解 for the inheritance and development of Daoist immortal belief.', 'Times New Roman', Pt(12), False)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_run(p, 'Key Words: ', 'Times New Roman', Pt(12), True)
    add_run(p, 'Daoist Immortal Belief; Sinicization; Inheritance; New Era; Returning to Roots and Opening Up the New; Modernization Transformation', 'Times New Roman', Pt(12), False)

# ===== TOC =====
def add_toc(doc):
    add_para(doc, '目  录', '黑体', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 12, 12)
    toc = [
        '摘  要', 'Abstract', '引  言',
        '第一章 道教神仙信仰的历史演变与核心内涵',
        '  第一节 神仙信仰的思想渊源',
        '  第二节 神仙信仰的历史演变',
        '  第三节 神仙信仰的核心内涵',
        '第二章 新时代道教神仙信仰传承的现实境遇',
        '  第一节 新时代社会变迁对神仙信仰的冲击',
        '  第二节 神仙信仰传承面临的突出问题',
        '  第三节 神仙信仰当代存续形态的考察',
        '第三章 道教神仙信仰中国化的理论依据与内在逻辑',
        '  第一节 "中国化"的理论内涵',
        '  第二节 神仙信仰中国化的历史逻辑',
        '  第三节 神仙信仰中国化的经典依据',
        '  第四节 神仙信仰中国化的现实依据',
        '第四章 新时代道教神仙信仰中国化的实践路径',
        '  第一节 教义阐释的中国化',
        '  第二节 修持实践的中国化',
        '  第三节 组织形态的中国化',
        '  第四节 传播方式的中国化',
        '  第五节 典型案例分析',
        '第五章 神仙信仰传承与中国化的辩证关系与未来展望',
        '  第一节 传承与创新的辩证统一',
        '  第二节 神仙信仰中国化的基本原则',
        '  第三节 未来展望',
        '结  语', '参考文献', '附  录',
    ]
    for item in toc:
        add_para(doc, item, '宋体', Pt(12), False, None, 1, 0, 2, 2)

# ===== Process markdown =====
def process_markdown(doc, md_text):
    """Parse markdown and add to document."""
    lines = md_text.split('\n')
    i = 0
    in_refs = False
    ref_section = ''
    in_block = False
    
    while i < len(lines):
        line = lines[i]
        
        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue
        
        # Skip title line
        if line.startswith('# 新时代道教'):
            i += 1
            continue
        
        # Skip abstract/keywords section (already added)
        if line.startswith('## 摘要') or line.startswith('## 绪论'):
            i += 1
            continue
        
        # Chapter heading (## 第X章)
        if line.startswith('## '):
            title = line.replace('## ', '').strip()
            if title == '参考文献':
                in_refs = True
                i += 1
                continue
            if title == '结语':
                add_para(doc, '结  语', '黑体', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 24, 18)
                i += 1
                continue
            # Chapter title
            add_para(doc, title, '黑体', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 24, 18)
            i += 1
            continue
        
        # Section heading (### 第X节)
        if line.startswith('### '):
            title = line.replace('### ', '').strip()
            add_para(doc, title, '黑体', Pt(14), True, None, 1, 0, 18, 12)
            i += 1
            continue
        
        # Sub-section (#### )
        if line.startswith('#### '):
            title = line.replace('#### ', '').strip()
            add_para(doc, title, '黑体', Pt(12), True, None, 1, 0, 12, 6)
            i += 1
            continue
        
        # Bold inline heading (**一、** or **一、xxx**)
        if line.startswith('**') and ('、' in line or '、' in line):
            # Extract bold text
            match = re.match(r'\*\*(.+?)\*\*', line)
            if match:
                title = match.group(1)
                add_para(doc, title, '黑体', Pt(12), True, None, 1, 0, 12, 6)
                # Check if there's text after the bold part
                rest = line[match.end():].strip()
                if rest:
                    add_body(doc, rest)
                i += 1
                continue
        
        # Reference section
        if in_refs:
            if line.startswith('### '):
                ref_section = line.replace('### ', '').strip()
                add_para(doc, ref_section, '黑体', Pt(14), True, None, 1, 0, 18, 12)
                i += 1
                continue
            if line.strip().startswith('[') and ']' in line:
                # Reference line
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Cm(0.74)
                ref_text = line.strip()
                # Extract number from [1] format
                match = re.match(r'^\[(\d+)\]\s*(.*)', ref_text)
                if match:
                    add_run(p, f'[{match.group(1)}] ', '宋体', Pt(10.5), False)
                    add_run(p, match.group(2), '宋体', Pt(10.5), False)
                else:
                    add_run(p, ref_text, '宋体', Pt(10.5), False)
                i += 1
                continue
            if line.strip() == '':
                i += 1
                continue
            i += 1
            continue
        
        # Appendix
        if line.startswith('## 附'):
            add_para(doc, '附  录', '黑体', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 24, 18)
            i += 1
            continue
        
        # Empty line
        if line.strip() == '':
            i += 1
            continue
        
        # Regular paragraph text
        if line.strip():
            # Collect consecutive non-heading lines
            para_lines = []
            while i < len(lines) and lines[i].strip() != '' and not lines[i].startswith('#') and lines[i].strip() != '---':
                para_lines.append(lines[i].strip())
                i += 1
            if para_lines:
                text = ' '.join(para_lines)
                # Skip reference lines at end of chapters
                if re.match(r'^\[\d+\]', text):
                    continue
                add_body(doc, text)
        else:
            i += 1

def main():
    # Read markdown
    with open('/home/admin/.openclaw/workspace/论文_新时代道教神仙信仰的传承与中国化路径.md', 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    
    # Default style
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    # Cover
    add_cover(doc)
    doc.add_page_break()
    
    # Abstract
    add_abstract(doc)
    doc.add_page_break()
    
    # English abstract
    add_eng_abstract(doc)
    doc.add_page_break()
    
    # TOC
    add_toc(doc)
    doc.add_page_break()
    
    # Process main content
    process_markdown(doc, md_text)
    
    # Appendix placeholder
    add_para(doc, '附  录', '黑体', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, 1, 0, 24, 18)
    add_para(doc, '一、致  谢', '黑体', Pt(14), True, None, 1, 0, 18, 12)
    add_body(doc, '在本文的撰写过程中，得到了导师的悉心指导和无私帮助，在此表示诚挚的感谢。同时，感谢浙江省宗教界"双通"人才研修班提供的学习平台，感谢各位授课老师的精彩讲授，感谢同窗学友的相互交流与启发。本文的研究也离不开道教界各位道长的大力支持与分享，在此一并致以诚挚的谢意。')
    
    add_para(doc, '二、论文原创性声明', '黑体', Pt(14), True, None, 1, 0, 18, 12)
    add_body(doc, '本人郑重声明：所呈交的毕业论文，是本人在导师的指导下，独立进行研究工作所取得的成果。除文中已经注明引用的内容外，本论文不包含任何其他个人或集体已经发表或撰写过的作品成果。对本文的研究作出重要贡献的个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的法律和学术责任由本人承担。')
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    add_run(p, '毕业论文作者签名：', '宋体', Pt(12), False)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    add_run(p, '日    期：      年    月    日', '宋体', Pt(12), False)
    
    output_path = '/home/admin/.openclaw/workspace/新时代道教神仙信仰的传承与中国化路径_完整版.docx'
    doc.save(output_path)
    print(f'✅ 论文已生成: {output_path}')

if __name__ == '__main__':
    main()
