#!/usr/bin/env python3
"""生成法治演讲稿 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# === 页面设置 ===
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# === 样式设置 ===
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = Pt(28)
pf.space_before = Pt(0)
pf.space_after = Pt(6)

# === 辅助函数 ===
def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.font.size = Pt(14 if level == 1 else 13)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h

def add_para(text, bold=False, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.space_after = Pt(6)
    if indent:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_quote(text):
    """引用段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Cm(1.5)
    pf.right_indent = Cm(1.0)
    pf.line_spacing = Pt(26)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '楷体'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.italic = True
    return p

def add_bold_normal(text):
    """混合加粗段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0.74)
    # 简单处理：整段不加粗，由调用方传入纯文本
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_mixed_para(segments, indent=True):
    """混合加粗段落，segments = [(text, bold), ...] 或 [text, ...]"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.space_after = Pt(6)
    if indent:
        pf.first_line_indent = Cm(0.74)
    for seg in segments:
        if isinstance(seg, tuple):
            text, bold_val = seg
            bold = bool(bold_val)
        else:
            text, bold = seg, False
        run = p.add_run(text)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.bold = bold
    return p

# ==================== 正文开始 ====================

# 标题
title = doc.add_heading('以戒为师，依法护心', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = subtitle.paragraph_format
pf.space_after = Pt(12)
run = subtitle.add_run('——从佛教戒律精神阐释"法治"核心价值观')
run.font.name = '黑体'
run.font.size = Pt(14)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 开场白
add_para('各位老师、各位同学，大家好。')
add_para('今天我想和大家探讨的，是社会主义核心价值观中的"法治"。')
add_para('在座各位来自五大宗教，每个人都有自己的信仰传统和戒律规范。有人可能会问：我是学佛的，讲戒律就够了，为什么要讲法治？同样，道教的同学会讲"道法自然"，伊斯兰教的同学会讲"沙里亚"的秩序精神，天主教会讲自然法，基督教会讲"顺服掌权者"——每个宗教都有自己关于行为规范、遵守法律的教导。')
add_para('那么，我们为什么要专门来谈"法治"？宗教的戒律、教规，与国家的法律之间，到底是什么关系？')
add_para('今天，我从佛教视角出发，尝试给出部分答案。全文分五个部分：')

# 目录
add_mixed_para([('一、制度传承，高于个人', True)])
add_mixed_para([('二、规则面前，人人平等', True)])
add_mixed_para([('三、慈悲生起，自觉持戒', True)])
add_mixed_para([('四、德法相济，闭环成序', True)])
add_mixed_para([('五、尊法守戒，从我做起', True)])

add_para('好，让我们正式开始。')

# ========== 第一部分 ==========
add_heading_custom('一、制度传承，高于个人', level=1)

add_para('首先，回答一个根本问题：制度传承，为什么高于个人？')

add_para('佛陀临涅槃时，阿难问了一个关键问题："佛在世时，以佛为师；佛灭度后，以谁为师？"佛陀的回答在座的诸位法师都知道——以戒为师。')

add_quote('"汝等比丘，于我灭后，当尊重珍敬波罗提木叉，如暗遇明、贫人得宝。当知此则是汝等大师，若我住世无异此也。" ——《佛遗教经》')

add_para('"波罗提木叉"就是戒律。佛陀没有说"你们听阿难的"，也没有说"你们听大迦叶的"。他说：以戒为师。')

add_para('这意味着佛陀做出了一个极其重要的选择——制度传承，而非个人传承。佛陀把自己创立的僧团，交给了一套规则体系来治理。这个选择告诉我们三件事：')

add_mixed_para([('第一，制度的权威高于个人的权威。', True), (' 谁都不能凌驾于规则之上。就像一座灯塔，它不属于任何船长，却为所有航船指明方向。')])
add_mixed_para([('第二，规则的作用是"如暗遇明"。', True), (' 戒律就是那条清晰的边界：什么该做，什么不该做，一目了然。')])
add_mixed_para([('第三，制度的有效性靠共同信守。', True), (' 戒律在，就等于佛陀在。这不是说戒律代替了佛，而是说，共同信守的规则，才是僧团存在的根基。')])

add_para('各位同学，两千五百年前佛陀的这个选择，核心就是一个字：法。这和今天"全面依法治国"的精神——让规则而非个人成为治理的根基——在思想层面上，跨越时空，遥相呼应。')

add_para('历史上，许多王朝依靠明君、贤相，一时兴盛，却人亡政息。而佛陀选择了制度建设。这不正是"法治"最古老的智慧之光吗？')

add_mixed_para([('所以，从佛教的视角看：', ''), ('戒律就是佛门的"法律"，法律就是世间的戒律。', True), (' 两者的底层逻辑完全相同——用明确的规则，约束行为、保护权益、维护秩序。')])

add_para('那么，戒律和法律之间，除了这个共同的底层逻辑，还有更深层的关系吗？带着这个问题，我们进入第二部分。')

# ========== 第二部分 ==========
add_heading_custom('二、规则面前，人人平等', level=1)

add_para('规则面前，为什么人人平等？这可不是一句空话。')

add_para('大乘佛教有一个非常重要的原则，叫"四依四不依"，第一条就是：')

add_quote('"依法不依人。" ——《大宝积经》')

add_para('这句话的意思很直白：要依止法（规则），而不是依止某个人。')

add_para('有人会问：规则不也是人定的吗？如果规则本身有问题，我们还要遵守吗？')

add_para('对此，佛陀在律藏中给出了一个极有智慧的原则，叫"随方毗尼"。《五分律》中佛陀这样开示：')

add_quote('"虽是我所制，而于余方不以为清净者，皆不应用；虽非我所制，而于余方必应行者，皆不得不行。"')

add_para('什么意思？佛陀说，我在印度制定的某条戒律，如果到了另一个国家、另一种文化环境中，与当地的风俗法律相冲突，不再被视为合适的，那么就可以不执行。反过来，我虽然没有制定某条规则，但如果那个地方的法律风俗要求必须这样做，那么僧团也不得不遵守。')

add_mixed_para([('从这里我们可以看出，佛陀留出了因地制宜的空间，但有一个底线从未动摇——', ''), ('规则面前人人平等的原则，不可动摇。', True)])

add_para('佛陀本人就是最好的示范。')

add_para('经中有这样一个故事，关于佛陀的堂弟提婆达多。他野心勃勃，想要取代佛陀领导僧团。他提出更严格的苦行标准，拉拢了一部分比丘，公开分裂僧团。换作一般的领袖，可能会动用权威，直接驱赶。但佛陀怎么做的？')

add_para('他没有动用个人权威去压制，而是按照戒律程序，召集全体僧团，依法羯磨——也就是集体表决、公开讨论，依规则作出决定。')

add_para('这就是"依法不依人"最生动的实践。在佛教的制度设计里，没有一个人可以凌驾于规则之上。这不正是现代法治"法律面前人人平等"的原则，在宗教治理中最古老的体现吗？')

add_para('规则平等的问题解决了。但下一个问题来了：人为什么要守规则？是因为怕惩罚吗？是因为怕坐牢、怕下地狱吗？如果有一天惩罚不在了，是不是就可以不守了？')

add_para('我们进入第三部分。')

# ========== 第三部分 ==========
add_heading_custom('三、慈悲生起，自觉持戒', level=1)

add_para('佛陀给出了比"怕惩罚"更深的答案。这个答案，直指人心。')

add_quote('"一切男子是我父，一切女子是我母，我生生无不从之受生，故六道众生皆是我父母。" ——《梵网经》')

add_para('当一个人真正生起这种慈悲心，看一切众生都如同自己的父母——那么，持戒就不再是"不敢做"的外在约束，而是不忍做的内在觉醒。')

add_para('请大家用心体会这两个词的区别。')

add_mixed_para([('"不敢做"，', True), ('是因为外面有警察、有监控、有惩罚。')])
add_mixed_para([('"不忍做"，', True), ('是因为你的心告诉你：那样做会伤害到别人，而那个人，如同你的父母。')])

add_para('我们用佛教最基本的五戒来一一对照：')

# 五戒表格
table = doc.add_table(rows=6, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['五戒', '不是怕…', '而是不忍…', '对应的法治价值']
data = [
    ['不杀生', '怕犯戒', '伤害生命', '对生命权的敬畏'],
    ['不偷盗', '怕被抓', '侵犯他人劳动成果', '对财产权的尊重'],
    ['不邪淫', '怕名声受损', '破坏他人家庭幸福', '对婚姻制度的守护'],
    ['不妄语', '怕失去信用', '欺骗信任自己的人', '对诚信体系的维护'],
    ['不饮酒', '怕失态', '因失去理性而伤害他人', '对自律精神的持守'],
]

for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = '黑体'
            run.font.size = Pt(10)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for row_idx, row_data in enumerate(data):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.cell(row_idx + 1, col_idx)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = '宋体'
                run.font.size = Pt(10)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_para('各位同学，对比一下：')

add_mixed_para([('法律', True), ('是全体公民必须遵守的最低行为规范，具有国家强制力；')])
add_mixed_para([('戒律', True), ('是宗教信仰者内心自觉遵循的修行准则，其力量来自信仰与慈悲。')])

add_para('法律为戒律提供了不可逾越的外部边界，戒律则从内在帮助信众更好地遵守法律。两者不是对立关系，而是法律为根本、戒律为辅助的关系。')

add_para('我相信，其他宗教的朋友也能找到类似的教导：道教讲"天道承负"，伊斯兰教讲"敬畏安拉"，天主教讲"良心的指引"，基督新教讲"顺服掌权者"。各宗教的戒律与教规，最终都指向同一个实践要求——在遵守国家法律的前提下，以信仰的力量提升自律。')

add_para('那么，"不敢"与"不忍"之间，是互相替代，还是互相补充？如果一个人有了"不忍"，是不是就不需要"不敢"了？反过来，法律能不能取代信仰的教化？')

add_para('我们进入第四部分。')

# ========== 第四部分 ==========
add_heading_custom('四、德法相济，闭环成序', level=1)

add_para('这个问题，古人早就给出了精辟的回答。')

add_para('两千多年前，孔子在《论语》中这样说：')

add_quote('"道之以政，齐之以刑，民免而无耻；道之以德，齐之以礼，有耻且格。" ——《论语》')

add_para('意思很清晰：用政令引导、用刑罚规范，老百姓虽然能避免犯罪，但内心没有羞耻感；用道德教化、用礼制约束，老百姓不但有羞耻心，而且能够自觉归正。')

add_para('这句话放在今天依然振聋发聩。')

add_mixed_para([('单靠法律', True), ('，人只是不敢犯——一旦有机会、有漏洞，还会犯。')])
add_mixed_para([('加上道德教化、信仰引导', True), ('，人才能从内心深处建立起"不忍犯"的自觉。')])

add_para('宗教的戒律、教规，可以成为道德教化的重要组成部分，帮助信众从内心认同法律的正当性。但必须明确：法律是底线，是根本；宗教的道德教化是在法律框架内发挥辅助作用。任何宗教的戒律、教规，如果与国家法律相抵触，都必须以法律为准绳。')

add_mixed_para([('这正是中国特色社会主义法治道路的鲜明特色——', ''), ('法治为根本，德治为辅助，德法相济，共同维护社会秩序。', True)])

add_para('有人担心：强调法治，会不会与宗教教义冲突？')

add_mixed_para([('不会。', True), (' 佛教经典中就有明确的护国利民思想。《仁王护国般若波罗蜜经》中，佛陀教导国王：')])

add_quote('"国王应正法治世，不以非法。若王以非法治世，则天龙鬼神不护其国。"')

add_para('什么意思？佛陀明确要求君主以正法、以公正的法律治理国家，而不是以非法的手段统治。如果一个国家不以法治国，连天龙鬼神都不护佑。')

add_para('在佛教中国化的历史进程中，历代祖师反复强调："不违国法"是持戒的基本前提。')

add_para('一个有序、公正、稳定的社会，是所有宗教健康发展的土壤。国家法律保护了宗教活动场所的合法财产，保护了信众的信仰自由，保护了宗教界的合法权益。因此，遵纪守法，是每一个宗教信众应尽的基本义务，也是信仰实践的前提条件。')

add_para('从佛陀的制度选择，到规则面前人人平等；从外在约束到内在觉醒，再到德法共治的完整链条——道理讲完了。回到现实，也是最关键的一步：我们该怎么做？')

add_para('我们进入第五部分。')

# ========== 第五部分 ==========
add_heading_custom('五、尊法守戒，从我做起', level=1)

add_mixed_para([('作为双通班学员，我们不是旁观者，不是评论员，而是参与者和建设者。', True), ('怎么践行？我提三点，供大家参考。')])

add_mixed_para([('第一，在个人修学上：学法懂法，做知法守法的表率。', True)])

add_para('无论哪个宗教，如果教职人员或信众法律意识淡薄，出现网上造谣传谣、借教敛财、非法集资等行为，既违反国法，也违背各自的教规。这样的案例，我们或多或少都听说过。')

add_para('一个人违法，受伤害的不仅是他自己，还有他所在的宗教的声誉，以及广大信众的感情。所以，我们首先要知法懂法。守法，本身就是信仰实践的一部分。')

add_mixed_para([('第二，在信众引导上：以各自的戒律传统，守护法治底线。', True)])

add_para('我们每个宗教都有丰富的经典依据，教导信众遵守国家法律：')

add_mixed_para([('佛教', True), ('有"以戒为师"，强调"不犯国法，不违毗尼"；')])
add_mixed_para([('道教', True), ('有"积功行善"，强调"助国化民"；')])
add_mixed_para([('伊斯兰教', True), ('有"命人行善，止人作恶"，强调遵守所在国的法律是信仰的一部分；')])
add_mixed_para([('天主教和基督教', True), ('强调"顺服掌权者，行善守法"，因为法律是为维护正义而设立的。')])

add_para('每个宗教都有自己教导信众遵守法律的经典依据。我们要善于运用这些教义，向信众讲清楚：守法不违背信仰，守法本身就是信仰的实践。')

add_mixed_para([('第三，在社会参与上：运用法律维护正当权益，积极投身社会建设。', True)])

add_para('国家鼓励宗教界在法律框架内开展活动，保护宗教界的合法权益。我们要善于运用法律武器，维护宗教活动场所的合法财产、正常的宗教活动秩序。同时，积极投身公益慈善、生态环保、扶贫济困，为"中国式现代化"贡献宗教界的力量。')

add_para('这正是各个宗教"庄严国土、利乐有情""济世利人""服务社会"的本怀。')

# ========== 结语 ==========
add_heading_custom('结语：双通学员的使命', level=1)

add_para('各位同学，今天我们走了一条清晰的路：')

add_para('从佛陀的一个选择出发——制度传承，以戒为师；')
add_para('进入三重递进——')
add_para('第一层：规则大于个人，依法不依人；')
add_para('第二层：从"不敢"到"不忍"，从外在约束到内在觉醒，让慈悲成为守法的动力；')
add_para('第三层：法治与德治结合，法律与戒律互补，共同构成从约束到自觉的完整链条。')

add_mixed_para([('最后落到三点践行——', ''), ('知法守法、以信仰护法、服务社会。', True)])

add_quote('"当尊重珍敬波罗提木叉，如暗遇明、贫人得宝。" ——《佛遗教经》')

add_mixed_para([('今天，我们也可以说：', ''), ('让我们尊重珍敬国家的法律，如暗遇明、贫人得宝。', True), (' 因为法治，是守护每一个人的庄严，也是守护我们信仰实践环境的根本保障。')])

add_mixed_para([('作为双通班学员，我们的使命就是——', ''), ('让两束光，信仰的光明与法治的阳光，在我们手中交相辉映，照亮宗教中国化的前行之路。', True)])

add_para('最后，以一句话结束今天的分享：')

# 结尾金句 - 居中加粗
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format
pf.space_before = Pt(12)
pf.space_after = Pt(12)
run = p.add_run('以戒为师，是信仰的定力；依法护心，是时代的担当。')
run.font.name = '黑体'
run.font.size = Pt(14)
run.font.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

add_para('请各位老师、各位同学批评指正。谢谢大家！')

# 保存
output_path = '/home/admin/.openclaw/workspace/以戒为师依法护心-法治核心价值观演讲.docx'
doc.save(output_path)
print(f'✅ 文档已生成: {output_path}')
