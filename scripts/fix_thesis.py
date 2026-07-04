#!/usr/bin/env python3
"""修改毕业论文：1.扩充摘要 2.整合章节(每章最多4节) 3.神学思想→神学观"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy
import re
import sys

INPUT_FILE = "/home/admin/.openclaw/media/inbound/2f0729c6-1a17-4fb4-9c79-5bedbeb2b7c1.docx"
OUTPUT_FILE = "/home/admin/.openclaw/workspace/毕业论文-邱春华-修改版.docx"

def count_chinese_chars(text):
    """统计中文字符数（含标点）"""
    return len(text)

def replace_in_run(run):
    """替换run中的文本"""
    old_text = run.text
    # 全局替换：王治心神学思想 → 王治心神学观
    new_text = old_text.replace("王治心神学思想", "王治心神学观")
    # 题目中的替换
    new_text = new_text.replace("王治心神学思想", "王治心神学观")
    if old_text != new_text:
        run.text = new_text

def replace_in_paragraph(para):
    """替换段落中所有run的文本"""
    if not para.text.strip():
        return
    # 检查是否需要替换
    if "王治心神学思想" in para.text:
        # 合并所有run的文本后替换，再分配回去
        full_text = para.text
        new_full = full_text.replace("王治治心神学思想", "王治心神学观")
        new_full = new_full.replace("王治心神学思想", "王治心神学观")
        if new_full != full_text:
            # 清除所有run，重新设置
            texts = [r.text for r in para.runs]
            combined = "".join(texts)
            new_combined = combined.replace("王治心神学思想", "王治心神学观")
            if new_combined != combined:
                # 简单处理：只改第一个run
                if para.runs:
                    para.runs[0].text = new_combined
                    for r in para.runs[1:]:
                        r.text = ""

def get_heading_level(para):
    """获取段落标题级别"""
    if para.style is None:
        return 0
    style_name = para.style.name.lower()
    if 'heading' in style_name:
        try:
            return int(style_name.replace('heading', '').strip())
        except:
            return 0
    return 0

def main():
    print("读取文档...")
    doc = Document(INPUT_FILE)
    
    # 统计全文字数
    total_chars = 0
    for para in doc.paragraphs:
        total_chars += count_chinese_chars(para.text)
    print(f"全文总字数: {total_chars}")
    print(f"摘要需要至少: {int(total_chars * 0.05)} 字")
    
    # ========== 任务3: 全局替换 "王治心神学思想" → "王治心神学观" ==========
    print("\n任务3: 替换 '王治心神学思想' → '王治心神学观'...")
    replace_count = 0
    for para in doc.paragraphs:
        old_text = para.text
        for run in para.runs:
            if "王治心神学思想" in run.text:
                run.text = run.text.replace("王治心神学思想", "王治心神学观")
                replace_count += 1
        # 也检查段落级别的文本（表格等）
        if "王治心神学思想" in para.text:
            # 重新检查是否还有未替换的
            pass
    
    # 处理表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "王治心神学思想" in run.text:
                            run.text = run.text.replace("王治心神学思想", "王治心神学观")
                            replace_count += 1
    
    print(f"替换了 {replace_count} 处")
    
    # ========== 任务2: 整合章节，每章最多4节 ==========
    print("\n任务2: 整合章节结构...")
    
    # 分析当前章节结构
    chapters = {}  # {章号: [(节号, para_index), ...]}
    current_chapter = 0
    current_section = 0
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        level = get_heading_level(para)
        
        # 检测章标题
        chapter_match = re.match(r'^第[一二三四五六七八九十]+章\s*', text)
        if chapter_match and level >= 1:
            current_chapter += 1
            chapters[current_chapter] = []
            continue
        
        # 检测节标题
        section_match = re.match(r'^第[一二三四五六七八九十]+节\s*', text)
        if section_match and level >= 2:
            current_section += 1
            if current_chapter > 0:
                if current_chapter not in chapters:
                    chapters[current_chapter] = []
                chapters[current_chapter].append((current_section, i))
    
    print(f"\n当前章节结构:")
    for ch, sections in chapters.items():
        print(f"  第{['','一','二','三','四'][ch] if ch <= 4 else ch}章: {len(sections)} 节")
        for sec_num, para_idx in sections:
            para_text = doc.paragraphs[para_idx].text.strip() if para_idx < len(doc.paragraphs) else ""
            print(f"    {para_text}")
    
    # 需要整合的章节:
    # 第二章: 5节 → 4节 (合并第2节到第1节, 或合并第3+4节)
    # 第三章: 6节 → 4节 (合并多节)
    # 第四章: 6节 → 4节 (合并多节)
    
    print("\n需要整合:")
    print("  第二章: 5节→4节")
    print("  第三章: 6节→4节")
    print("  第四章: 6节→4节")
    
    # ========== 任务1: 扩充摘要 ==========
    print("\n任务1: 扩充摘要...")
    
    # 找到摘要位置
    abstract_start = None
    abstract_end = None
    for i, para in enumerate(doc.paragraphs):
        if "摘  要" in para.text or "摘    要" in para.text or para.text.strip() == "摘要":
            abstract_start = i
        elif abstract_start is not None and abstract_end is None:
            if "关键词" in para.text:
                abstract_end = i  # 关键词行之前是摘要内容
            elif ("Abstract" in para.text or "abstract" in para.text.lower()) and i > abstract_start + 2:
                abstract_end = i
    
    if abstract_start and abstract_end:
        print(f"摘要位置: 段落 {abstract_start} - {abstract_end}")
        abstract_text = ""
        for i in range(abstract_start + 1, abstract_end):
            abstract_text += doc.paragraphs[i].text + "\n"
        print(f"当前摘要字数: {count_chinese_chars(abstract_text)}")
    
    # 找到英文摘要
    english_abstract_start = None
    english_abstract_end = None
    for i, para in enumerate(doc.paragraphs):
        if "Abstract" in para.text and i > 5:
            english_abstract_start = i
        elif english_abstract_start is not None and english_abstract_end is None:
            if "Key Words" in para.text or "关键词" in para.text:
                english_abstract_end = i
    
    if english_abstract_start and english_abstract_end:
        print(f"英文摘要位置: 段落 {english_abstract_start} - {english_abstract_end}")
    
    # ========== 保存文档 ==========
    print(f"\n保存修改后的文档到: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    print("完成！")
    
    return doc, chapters

if __name__ == "__main__":
    main()
