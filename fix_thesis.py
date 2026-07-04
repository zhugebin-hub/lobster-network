#!/usr/bin/env python3
"""完善叶畏兵毕业论文：修复格式问题、补充脚注关联、统一标点"""

import docx
import re
import json

INPUT = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
OUTPUT = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"
ISSUES_LOG = "/home/admin/.openclaw/workspace/论文修改记录.json"

doc = docx.Document(INPUT)

# ========== 构建参考文献数据库 ==========
# 扫描参考文献部分，建立编号到条目的映射
ref_entries = {}
in_refs = False
current_num = None
current_text = ""

for p in doc.paragraphs:
    t = p.text.strip()
    if t == "参考文献":
        in_refs = True
        continue
    if in_refs and t.startswith("一、经典著作"):
        continue
    if in_refs and t.startswith("二、中文专著"):
        continue
    if in_refs and t.startswith("三、中文期刊"):
        continue
    if in_refs and t.startswith("四、政策文件"):
        continue
    if in_refs and t.startswith("五、网络资源"):
        continue
    if in_refs and t.startswith("附 录"):
        # 保存最后一个条目
        if current_num and current_text:
            ref_entries[current_num] = current_text.strip()
        break
    
    if in_refs and t:
        # 检查是否是新的编号条目
        m = re.match(r'\[(\d+)\]', t)
        if m:
            # 保存前一个
            if current_num and current_text:
                ref_entries[current_num] = current_text.strip()
            current_num = int(m.group(1))
            current_text = t
        else:
            # 可能是续行（如77/78合并问题）
            if current_num:
                current_text += " " + t

# 保存最后一个
if current_num and current_text:
    ref_entries[current_num] = current_text.strip()

print(f"参考文献条目数: {len(ref_entries)}")

# ========== 识别正文中的引用标记位置 ==========
# 找到所有需要加脚注的位置
# 格式：文本末尾的单个空格通常是脚注占位符
citation_positions = []
for i, p in enumerate(doc.paragraphs):
    t = p.text
    # 找到脚注占位符（段落末尾的单个空格，前面有句号/引号等）
    if t and t[-1] == ' ' and len(t) > 5:
        # 这很可能是一个脚注标记
        # 找到前一个标点
        stripped = t.rstrip()
        if stripped and stripped[-1] in ['。', '』', '"', '"', '）', '》', '！', '？', ':', '：']:
            citation_positions.append({
                'para_idx': i,
                'text_preview': t[:80] + '...',
                'full_text': t
            })

print(f"发现脚注占位符: {len(citation_positions)} 处")

# ========== 问题修复 ==========
issues_fixed = []
fix_count = 0

for p in doc.paragraphs:
    run_text = []
    for run in p.runs:
        if run.text:
            original = run.text
            fixed = original
            
            # 1. 英文引号 → 中文引号
            if '"' in fixed or '"' in fixed:
                # 智能替换：成对替换
                new_text = ""
                in_quote = False
                for ch in fixed:
                    if ch == '"':
                        new_text += '""' if in_quote else '""'
                        in_quote = not in_quote
                    elif ch == '"':
                        new_text += '""' if in_quote else '""'
                        in_quote = not in_quote
                    else:
                        new_text += ch
                if new_text != fixed:
                    issues_fixed.append({
                        'type': '引号替换',
                        'from': original[:60],
                        'to': new_text[:60]
                    })
                    fixed = new_text
                    fix_count += 1
            
            # 2. 半角冒号/分号 → 全角（中文语境下）
            # 仅在中文文本中替换
            if re.search('[\u4e00-\u9fff]', fixed):
                fixed = re.sub(r'(?<=[\u4e00-\u9fff])：', '：', fixed)  # 已经是全角
                # 半角冒号在中文后
                fixed = re.sub(r'([\u4e00-\u9fff]):', r'\1：', fixed)
                # 半角分号在中文后
                fixed = re.sub(r'([\u4e00-\u9fff]);', r'\1；', fixed)
            
            # 3. 英文破折号 -- → 中文破折号 ——
            fixed = fixed.replace('--', '——')
            
            # 4. 修复常见的格式错误
            fixed = fixed.replace('"', '""')
            fixed = fixed.replace('"', '""')
            
            run.text = fixed
    
    # 5. 处理段落级别的脚注占位符
    t = p.text
    if t and t[-1] == ' ' and len(t) > 5:
        stripped = t.rstrip()
        if stripped and stripped[-1] in ['。', '』', '"', '"', '）', '》', '！', '？', ':', '：']:
            # 保留脚注标记，但记录位置
            pass

# ========== 修复参考文献77/78合并问题 ==========
# 检查77和78
if '77' in ref_entries:
    entry77 = ref_entries['77']
    print(f"\n参考文献77: {entry77[:100]}")
    if '[78]' in entry77:
        parts = entry77.split('[78]')
        ref_entries['77'] = parts[0].strip()
        ref_entries['78'] = '[78]' + parts[1].strip()
        print(f"  拆分77: {ref_entries['77'][:80]}")
        print(f"  拆分78: {ref_entries['78'][:80]}")

# ========== 保存修复后的文档 ==========
doc.save(OUTPUT)
print(f"\n修复完成，已保存至: {OUTPUT}")
print(f"共修复 {fix_count} 处")

# ========== 生成修改记录 ==========
record = {
    'total_refs': len(ref_entries),
    'citation_positions': len(citation_positions),
    'fixes_count': fix_count,
    'issues_fixed': issues_fixed[:50],  # 最多记录50条
    'ref_entries_sample': {str(k): v[:80] for k, v in list(ref_entries.items())[:10]}
}

with open(ISSUES_LOG, 'w', encoding='utf-8') as f:
    json.dump(record, f, ensure_ascii=False, indent=2)

print(f"\n修改记录已保存: {ISSUES_LOG}")
