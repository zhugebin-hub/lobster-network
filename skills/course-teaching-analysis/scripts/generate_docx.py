#!/usr/bin/env python3
"""Generate Word document from course analysis data."""

import sys
import os
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def create_report(doc, data):
    """Create a course analysis report in Word format."""
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # Title
    title = doc.add_heading(data.get('title', '课程教学分析报告'), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta info
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'📅 {data.get("date", "")}').font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Content sections
    for section in data.get('sections', []):
        doc.add_heading(section['title'], level=1)
        for para in section.get('paragraphs', []):
            doc.add_paragraph(para)
        for table_data in section.get('tables', []):
            create_table(doc, table_data)
    
    return doc

def create_table(doc, data):
    """Create a table in the document."""
    rows = len(data)
    cols = len(data[0]) if rows > 0 else 0
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    run.font.bold = (row_idx == 0)
                    if row_idx == 0:
                        run.font.color.rgb = RGBColor(255, 255, 255)
    
    return table

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python3 generate_docx.py <data.json> [output.docx]")
        sys.exit(1)
    
    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    output = sys.argv[2] if len(sys.argv) > 2 else 'output.docx'
    doc = Document()
    create_report(doc, data)
    doc.save(output)
    print(f"Word document saved: {output}")
