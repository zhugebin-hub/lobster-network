#!/usr/bin/env python3
"""
Markdown → DOCX 转换脚本
支持中文排版、表格、图片、样式设置

用法:
  python generate_docx.py input.md -o output.docx
  python generate_docx.py input.md -o output.docx --title "标题" --author "作者"
  python generate_docx.py input.md -o output.docx --style academic
  echo "# Hello" | python generate_docx.py - -o output.docx
"""

import argparse
import sys
import os
from pathlib import Path

try:
    import pandoc
except ImportError:
    pandoc = None

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT


def set_chinese_font(doc, font_name='宋体', font_size=Pt(12)):
    """设置默认中文字体"""
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = font_size
    style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 标题字体
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = '黑体'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        sizes = {1: Pt(22), 2: Pt(16), 3: Pt(14)}
        heading_style.font.size = sizes.get(level, Pt(14))


def set_page_layout(doc, style='default'):
    """设置页面布局"""
    layouts = {
        'default': {
            'width': Cm(21), 'height': Cm(29.7),
            'top': Cm(2.54), 'bottom': Cm(2.54),
            'left': Cm(3.18), 'right': Cm(3.18),
        },
        'academic': {
            'width': Cm(21), 'height': Cm(29.7),
            'top': Cm(2.5), 'bottom': Cm(2.5),
            'left': Cm(3.0), 'right': Cm(3.0),
        },
        'narrow': {
            'width': Cm(21), 'height': Cm(29.7),
            'top': Cm(2.0), 'bottom': Cm(2.0),
            'left': Cm(2.0), 'right': Cm(2.0),
        },
    }
    layout = layouts.get(style, layouts['default'])
    section = doc.sections[0] if doc.sections else doc.add_section()
    section.page_width = layout['width']
    section.page_height = layout['height']
    section.top_margin = layout['top']
    section.bottom_margin = layout['bottom']
    section.left_margin = layout['left']
    section.right_margin = layout['right']
    return section


def md_to_docx(input_path, output_path, title=None, author=None, date=None,
               style='default', font_name='宋体', font_size=Pt(12)):
    """使用 pandoc 将 Markdown 转换为 DOCX"""
    import subprocess
    import json

    # 读取输入
    if input_path == '-':
        content = sys.stdin.read()
        input_file = None
    else:
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        input_file = input_path

    # 构建 pandoc 命令
    cmd = ['pandoc', '-f', 'gfm+yaml_metadata_block', '-t', 'docx']

    # 元数据
    metadata_yaml = []
    if title:
        metadata_yaml.append(f'title: "{title}"')
    if author:
        metadata_yaml.append(f'author: "{author}"')
    if date:
        metadata_yaml.append(f'date: "{date}"')
    if metadata_yaml:
        # 将元数据插入到内容前面
        yaml_block = '---\n' + '\n'.join(metadata_yaml) + '\n---\n\n'
        # 如果内容已有 yaml 块，合并
        if content.strip().startswith('---'):
            parts = content.split('---', 2)
            if len(parts) >= 3:
                existing = parts[1].strip()
                for line in metadata_yaml:
                    key = line.split(':')[0]
                    if not any(l.startswith(key + ':') for l in existing.split('\n')):
                        existing += '\n' + line
                content = '---\n' + existing + '\n---\n' + parts[2]
            else:
                content = yaml_block + content
        else:
            content = yaml_block + content

    # 使用临时文件
    import tempfile
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        cmd.extend([tmp_path, '-o', output_path])
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            print(f"pandoc error: {result.stderr}", file=sys.stderr)
            return False
        print(f"✅ 已生成: {output_path}")
        return True
    finally:
        os.unlink(tmp_path)


def create_simple_docx(output_path, content_lines, title=None, author=None,
                       style='default', font_name='宋体', font_size=Pt(12)):
    """用 python-docx 直接创建 DOCX（不依赖 pandoc）"""
    doc = Document()
    set_chinese_font(doc, font_name, font_size)
    set_page_layout(doc, style)

    # 标题页
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(author)
        run.font.size = Pt(14)

    if title:
        doc.add_page_break()

    # 逐行解析
    for line in content_lines:
        stripped = line.rstrip()

        # 空行
        if not stripped:
            continue

        # 标题
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
            continue
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
            continue

        # 水平线
        if stripped.startswith('---') or stripped.startswith('***'):
            doc.add_page_break()
            continue

        # 无序列表
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
            continue
        if stripped.startswith('-'):
            p = doc.add_paragraph(stripped[1:].strip(), style='List Bullet')
            continue

        # 有序列表
        import re
        ol_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if ol_match:
            doc.add_paragraph(ol_match.group(2), style='List Number')
            continue

        # 引用
        if stripped.startswith('> '):
            p = doc.add_paragraph(stripped[2:])
            p.paragraph_format.left_indent = Cm(1.27)
            run = p.runs[0] if p.runs else p.add_run(stripped[2:])
            run.italic = True
            continue

        # 普通段落
        doc.add_paragraph(stripped)

    doc.save(output_path)
    print(f"✅ 已生成: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Markdown → Word 文档生成器')
    parser.add_argument('input', nargs='?', default='-',
                        help='输入文件路径（默认从 stdin 读取）')
    parser.add_argument('-o', '--output', required=True,
                        help='输出 .docx 文件路径')
    parser.add_argument('--title', help='文档标题')
    parser.add_argument('--author', help='文档作者')
    parser.add_argument('--date', help='文档日期')
    parser.add_argument('--style', choices=['default', 'academic', 'narrow'],
                        default='default', help='页面样式')
    parser.add_argument('--font', default='宋体', help='正文字体')
    parser.add_argument('--font-size', type=float, default=12,
                        help='字体大小（磅）')
    parser.add_argument('--method', choices=['pandoc', 'python'], default='pandoc',
                        help='生成方法（默认 pandoc）')

    args = parser.parse_args()

    font_size = Pt(args.font_size)

    if args.method == 'pandoc':
        success = md_to_docx(
            args.input, args.output,
            title=args.title, author=args.author, date=args.date,
            style=args.style, font_name=args.font, font_size=font_size
        )
        sys.exit(0 if success else 1)
    else:
        # 读取内容
        if args.input == '-':
            lines = sys.stdin.readlines()
        else:
            with open(args.input, 'r', encoding='utf-8') as f:
                lines = f.readlines()
        create_simple_docx(
            args.output, lines,
            title=args.title, author=args.author,
            style=args.style, font_name=args.font, font_size=font_size
        )


if __name__ == '__main__':
    main()
