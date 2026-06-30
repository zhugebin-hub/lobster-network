#!/usr/bin/env python3
"""叶畏兵论文完善 - 干净重写版"""
import docx, re
from copy import deepcopy
from lxml import etree
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
DST = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"

d = docx.Document(SRC)

# ── 1. 引号 ──
qpos = []
for pi,p in enumerate(d.paragraphs):
    for ri,r in enumerate(p.runs):
        if r.text:
            for ci,ch in enumerate(r.text):
                if ch=='"': qpos.append((pi,ri,ci))
for i,(pi,ri,ci) in enumerate(qpos):
    r=d.paragraphs[pi].runs[ri]; c=list(r.text)
    c[ci]='\u201c' if i%2==0 else '\u201d'; r.text=''.join(c)
print(f"引号: {len(qpos)} → ✅")

# ── 2. 77/78 ──
for i in range(len(d.paragraphs)):
    p=d.paragraphs[i]
    if '[77]' in p.text and '[78]' in p.text:
        si=p.text.index('[78]'); a,b=p.text[:si].rstrip(),p.text[si:].strip()
        p.clear(); p.add_run(a)
        np=d.add_paragraph()
        if p.style: np.style=p.style
        np.add_run(b); p._p.addnext(np._p)
        print("77/78 → ✅"); break

# ── 3. 脚注上标 ──
# 段落号→[ref号,...] (None=无脚注)
FM={
124:[59], 126:[56,62], 127:[63,54,51],
138:[16,4], 140:[70,52], 142:[10,6],
143:[23,25,15], 144:[8,19], 146:[63,36,40],
147:[30,37,41,42], 148:[29,33,45], 150:[47,66,50],
151:[55,53,48],
184:[16], 185:[4,17], 189:[2], 190:[70], 191:[65], 192:[30,53], 196:[60],
215:[2], 216:[9], 220:[8], 221:[19], 223:[7],
229:[55], 230:[52], 245:[53], 246:[11], 254:[10],
265:[27], 269:[52], 270:[9], 275:[61], 283:[48], 293:[48],
311:[2], 320:[67], 321:[67], 322:[64], 324:[46,35],
340:[79], 351:[59],
}

def sup(para,cp,num):
    pos=0
    for run in para.runs:
        if not run.text: continue
        rl=len(run.text)
        if pos<=cp<pos+rl:
            ia=cp-pos+1; bf=run.text[:ia]; af=run.text[ia:]
            run.text=bf
            r=etree.SubElement(run._element.getparent(),run._element.tag)
            rPr=run._element.find(qn('w:rPr'))
            if rPr: r.append(deepcopy(rPr))
            rp=r.find(qn('w:rPr'))
            if rp is None: rp=OxmlElement(qn('w:rPr')); r.append(rp)
            va=OxmlElement(qn('w:vertAlign'));va.set(qn('w:val'),'superscript');rp.append(va)
            sz=OxmlElement(qn('w:sz'));sz.set(qn('w:val'),'16');rp.append(sz)
            sz2=OxmlElement(qn('w:szCs'));sz2.set(qn('w:val'),'16');rp.append(sz2)
            t=OxmlElement(qn('w:t'));t.text=f'[{num}]';t.set(qn('xml:space'),'preserve');r.append(t)
            if af:
                r2=etree.SubElement(r.getparent(),run._element.tag)
                if rPr: r2.append(deepcopy(rPr))
                t2=OxmlElement(qn('w:t'));t2.text=af;t2.set(qn('xml:space'),'preserve');r2.append(t2)
            return True
        pos+=rl
    return False

n=0
for pi,refs in FM.items():
    if pi>=len(d.paragraphs): continue
    p=d.paragraphs[pi]; t=p.text
    pos=[m.start() for m in re.finditer(r'([。》"）]) (?=[\u4e00-\u9fffA-Z])',t)]
    for s,po in enumerate(pos):
        if s<len(refs) and refs[s]:
            if sup(p,po,refs[s]): n+=1
print(f"脚注: {n} → ✅")

d.save(DST)
print(f"\n✅ {DST}")
v=docx.Document(DST)
print(f"ASCII引号: {sum(1 for p in v.paragraphs for c in p.text if c=='\"')}")
print(f"脚注标记: {sum(len(re.findall(r'\[\d+\]',p.text)) for p in v.paragraphs)}")
