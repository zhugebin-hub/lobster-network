#!/usr/bin/env python3
"""
叶畏兵毕业论文 - 全面完善脚本 v6
"""
import docx
import re

INPUT = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
OUTPUT = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"

doc = docx.Document(INPUT)

def smart_fix_quotes(text):
    """
    将段落中混合的引号统一为中文引号。
    策略：先收集所有开闭引号位置，然后成对替换。
    """
    # 先把所有中文引号也统一标记
    # 找到所有引号字符的位置
    chars = list(text)
    positions = []
    for i, ch in enumerate(chars):
        if ch in ('"', '"', '"', '"'):
            positions.append(i)
    
    # 成对替换
    for idx, pos in enumerate(positions):
        if idx % 2 == 0:
            chars[pos] = '"'  # 开引号
        else:
            chars[pos] = '"'  # 闭引号
    
    return ''.join(chars)

# ==========================================
# 修复1：所有引号
# ==========================================
print("=== 修复1：引号统一为中文引号 ===")
quote_fixes = 0

for p in doc.paragraphs:
    # 合并所有run的文本，统一处理
    full_text = ''.join(run.text for run in p.runs)
    if '"' not in full_text:
        continue
    
    fixed = smart_fix_quotes(full_text)
    if fixed != full_text:
        # 只修改第一个run，清空后续run
        if p.runs:
            p.runs[0].text = fixed
            for run in p.runs[1:]:
                run.text = ''
            quote_fixes += 1

print(f"引号修复: {quote_fixes} 处")

# ==========================================
# 修复2：参考文献77/78合并
# ==========================================
print("\n=== 修复2：参考文献77/78合并 ===")
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if '[77]' in p.text and '[78]' in p.text:
        split_pos = p.text.index('[78]')
        part77 = p.text[:split_pos].rstrip()
        part78 = p.text[split_pos:].strip()
        
        p.text = part77
        print(f"  77: {part77}")
        
        # 在77后面插入新段落
        new_p = doc.add_paragraph()
        if p.style:
            new_p.style = p.style
        run = new_p.add_run(part78)
        p._p.addnext(new_p._p)
        print(f"  78: {part78}")
        break

# ==========================================
# 修复3：半角标点（中文段落中）
# ==========================================
print("\n=== 修复3：标点规范 ===")
punct_fixes = 0

for p in doc.paragraphs:
    for run in p.runs:
        if not run.text:
            continue
        old = run.text
        if re.search(r'[\u4e00-\u9fff]', old):
            new = old
            new = re.sub(r'([\u4e00-\u9fff]):', r'\1：', new)
            new = re.sub(r'([\u4e00-\u9fff]);', r'\1；', new)
            new = re.sub(r'([\u4e00-\u9fff]),(?=[\u4e00-\u9fff\u3000-\u303f\uff01-\uff5f])', r'\1，', new)
            if new != old:
                run.text = new
                punct_fixes += 1

print(f"标点修复: {punct_fixes} 处")

# ==========================================
# 保存
# ==========================================
doc.save(OUTPUT)
print(f"\n✅ 完成! 输出: {OUTPUT}")

# 验证
doc2 = docx.Document(OUTPUT)
ascii_quotes = sum(1 for p in doc2.paragraphs if '"' in p.text)
print(f"验证: 仍有ASCII引号的段落 {ascii_quotes} 处")
