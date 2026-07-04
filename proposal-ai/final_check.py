#!/usr/bin/env python3
from docx import Document
doc = Document("/home/ubuntu/毕业论文_修订版.docx")

print("=== 最终残留检查 ===")
keywords = ["17 类", "17类", "17 个章节", "17个章节", "1/17", "Elasticsearch", "Monaco",
            "Zustand", "三栏", "OpenAI、Claude", "TableOfContents", "目录域"]
found = False
for i, para in enumerate(doc.paragraphs):
    for kw in keywords:
        if kw in para.text:
            # 排除误报：第5.2节提到的"17 年"、"17岁"等可能的误判
            if kw == "17" and "1/17" not in para.text:
                continue
            print(f"[残留P{i}] ({kw}): {para.text[:100]}")
            found = True

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for kw in keywords:
                if kw in cell.text:
                    print(f"[残留表{ti}行{ri}列{ci}] ({kw}): {cell.text[:60]}")
                    found = True

if not found:
    print("✓ 无残留关键词！")

print("\n=== 测试用例表格 ===")
for ti, table in enumerate(doc.tables):
    row_texts = [" | ".join(c.text for c in row.cells) for row in table.rows]
    has_create = any("项目创建" in t and "通过" in t for t in row_texts)
    has_export = any("Word 导出" in t or "Word导出" in t for t in row_texts)
    if has_create and has_export:
        print(f"表{ti}行数: {len(table.rows)}")
        for ri, rt in enumerate(row_texts):
            print(f"  行{ri}: {rt[:100]}")
        break

print("\n=== 关键段落快照 ===")
for i in [159, 161, 163, 287, 181]:
    if i < len(doc.paragraphs):
        text = doc.paragraphs[i].text
        print(f"[P{i}] {text[:150]}...")
