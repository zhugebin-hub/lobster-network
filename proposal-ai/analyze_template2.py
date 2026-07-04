from docx import Document
from docx.oxml.ns import qn

doc = Document('/tmp/template_original.docx')

print('=== ALL PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f'[{i}] style={p.style.name!r} | {text[:150]!r}')

print()
print(f'=== TOTAL TABLES: {len(doc.tables)} ===')
print()

# Show only first 3 tables in detail (cover page tables)
for ti in range(min(3, len(doc.tables))):
    t = doc.tables[ti]
    print(f'--- Table {ti} ({len(t.rows)} rows x {len(t.columns)} cols) ---')
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f'  [{ri},{ci}] {txt[:100]!r}')
    print()
