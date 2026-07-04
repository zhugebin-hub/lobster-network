#!/usr/bin/env python3
"""第二轮修改：处理段落级改写、表格行增删、数据库字段补全等复杂修改"""

from docx import Document
from docx.shared import Pt
from copy import deepcopy

DST = "/home/ubuntu/毕业论文_修订版.docx"
doc = Document(DST)

modified_count = 0

# ============ 第一步：扫描所有段落，定位关键段落 ============
for i, para in enumerate(doc.paragraphs):
    text = para.text

    # === 摘要中"17类"描述 ===
    if "17 类" in text or "17类" in text:
        new_text = text.replace("17 类", "14 类").replace("17类", "14类")
        if new_text != text:
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            modified_count += 1
            print(f"[P{i}] 17类→14类: {text[:60]}...")

    # === 第3.1.1节 Elasticsearch 文献库描述 ===
    if "Elasticsearch" in text or "elasticsearch" in text:
        new_text = text
        # 整段替换关于 Elasticsearch 的内容
        if "文献库" in text and ("Elasticsearch" in text or "全文检索" in text):
            # 这段是"文献库基于 Elasticsearch..."的整段描述
            new_text = "文献检索功能作为未来扩展方向暂未实现，本期系统的数据层聚焦于关系数据库和配置与模板库两部分。"
        else:
            new_text = text.replace("Elasticsearch", "关系数据库").replace("elasticsearch", "关系数据库")
        if new_text != text:
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            modified_count += 1
            print(f"[P{i}] Elasticsearch段落改写: {text[:60]}... → {new_text[:60]}...")

    # === 第3.1.2节 LLM网关 多模型 ===
    if "LLM 网关" in text or "LLM网关" in text or "多模型路由" in text:
        new_text = text
        new_text = new_text.replace("对接 OpenAI、Claude 等多个大模型，后面配上多模型路由",
                                    "调用平台内置 LLM 服务")
        new_text = new_text.replace("多模型路由", "平台内置模型调用")
        new_text = new_text.replace("OpenAI、Claude 等", "平台内置")
        if new_text != text:
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            modified_count += 1
            print(f"[P{i}] LLM网关描述改写")

    # === 第3.1.2节 知识库服务 ===
    if "知识库服务" in text and "知识图谱" in text:
        new_text = "知识库服务作为未来扩展方向，可集成学术知识图谱和文献检索能力，本期系统暂未实现。"
        if para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""
        modified_count += 1
        print(f"[P{i}] 知识库服务改写")

    # === 第3.1.2节 质量评估服务 ===
    if "质量评估服务" in text and "评估" in text:
        new_text = text.replace("质量评估服务", "质量评估服务（规划中）")
        if new_text != text:
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
            modified_count += 1
            print(f"[P{i}] 质量评估服务标注规划中")

    # === proposals表字段补全 ===
    if "updatedAt（更新时间）" in text and "proposals" not in text:
        # 在此段落之后插入字段补充说明——跳过，通过后续段落定位
        pass

# ============ 第二步：处理表格 ============
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                text = para.text
                if not text.strip():
                    continue

                new_text = text
                # 17 → 14
                if "17" in text and any(k in text for k in ["章节", "/17"]):
                    new_text = text.replace("1/17 = 6%", "1/14 ≈ 7%")
                    new_text = new_text.replace("1/17", "1/14")
                    new_text = new_text.replace("17 个章节", "14 个章节")
                    new_text = new_text.replace("17个章节", "14个章节")

                # Elasticsearch 行（技术栈表格）
                if "Elasticsearch" in text:
                    new_text = ""  # 清空此单元格，稍后整行处理

                # Zustand 状态管理行
                if "Zustand" in text:
                    new_text = text.replace("Zustand", "tRPC React Query")

                if new_text != text:
                    if para.runs:
                        para.runs[0].text = new_text
                        for r in para.runs[1:]:
                            r.text = ""
                    else:
                        para.text = new_text
                    modified_count += 1
                    print(f"[表{ti}行{ri}列{ci}] {text[:40]}... → {new_text[:40]}...")

# ============ 第三步：处理技术栈表格——删除 Elasticsearch 行 ============
# 定位技术栈表格（含"前端框"、"后端框"的表格）
for ti, table in enumerate(doc.tables):
    # 检查是否是技术栈表格
    is_tech_stack = False
    for row in table.rows[:3]:
        cell_texts = [c.text for c in row.cells]
        if any("React" in t or "前端框" in t for t in cell_texts):
            is_tech_stack = True
            break

    if is_tech_stack:
        # 找到并移除Elasticsearch行
        rows_to_remove = []
        for ri, row in enumerate(table.rows):
            row_text = " ".join(c.text for c in row.cells)
            if "Elasticsearch" in row_text or "elasticsearch" in row_text:
                rows_to_remove.append(ri)

        # 从后往前移除（避免索引变化）
        for ri in reversed(rows_to_remove):
            tbl = table._tbl
            tr = table.rows[ri]._tr
            tbl.remove(tr)
            modified_count += 1
            print(f"[技术栈表格] 删除第{ri}行（Elasticsearch）")
        break

# ============ 保存 ============
doc.save(DST)
print(f"\n第二轮共修改 {modified_count} 处")
print(f"已保存到: {DST}")
