#!/usr/bin/env python3
"""第四轮：数据库字段补全 + Word导出段落改写 + 测试表格扩充"""

from docx import Document
from copy import deepcopy

DST = "/home/ubuntu/毕业论文_修订版.docx"
doc = Document(DST)

modified_count = 0

def replace_para_text(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.text = new_text

# ============ 第一步：Word导出段落改写 ============
# 查找"在文档最前面"+"TableOfContents"/"目录域"的段落
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if ("目录域" in text and "TableOfContents" in text) or \
       ("在文档最前面" in text and "TableOfContents" in text):
        new_text = (
            "在文档最前面，会生成封面页（含申报编号、项目名称、申报单位、项目负责人等信息），"
            "紧跟其后是填报说明、项目基本信息表（含申报单位、推荐单位、项目负责人、项目联系人、"
            "课题分解等含复杂合并单元格的表格），然后是申报项目简介和各已确认章节正文。"
            "整体结构和排版细节与《国家重点研发计划项目申报书模板》保持一致，生成完成后"
            "会把 Document 对象转成 buffer 返回给前端，浏览器端通过 atob+Uint8Array 解码 base64 "
            "数据后直接触发下载，整个导出流程走一趟下来，文档格式是统一的，也便于后续按不同"
            "模板进行定制。"
        )
        replace_para_text(para, new_text)
        modified_count += 1
        print(f"[P{i}] Word导出段落改写完成")

    # proposals表字段补全——在"updatedAt（更新时间）"后面的段落添加说明
    if "updatedAt（更新时间）" in text and "proposalId" not in text:
        # 此段是proposals表字段最后一行，判断是否是proposals的"updatedAt"
        # 检查前面几段是否有"proposals 表"
        is_proposals_table = False
        for j in range(max(0, i-15), i):
            if "proposals 表" in doc.paragraphs[j].text or "proposals表" in doc.paragraphs[j].text:
                is_proposals_table = True
                break
            if "sections 表" in doc.paragraphs[j].text or "operationLogs" in doc.paragraphs[j].text:
                is_proposals_table = False
                break

        if is_proposals_table:
            # 在此段落内容中补充字段说明
            new_text = text + "\n（注：实际实现中还包含申报单位、推荐单位、项目负责人、项目联系人、预算信息等扩展字段，以支持完整的申报书模板填充，此处仅列出核心字段。）"
            replace_para_text(para, new_text)
            modified_count += 1
            print(f"[P{i}] 数据库字段补全说明")

# ============ 第二步：处理技术栈表格——删除Elasticsearch行、Zustand行 ============
for ti, table in enumerate(doc.tables):
    is_tech_stack = False
    for row in table.rows[:3]:
        cell_texts = [c.text for c in row.cells]
        if any("前端框" in t or "React 19" in t for t in cell_texts):
            is_tech_stack = True
            break

    if is_tech_stack:
        rows_to_remove = []
        for ri, row in enumerate(table.rows):
            row_text = " ".join(c.text for c in row.cells)
            if "Elasticsearch" in row_text or "elasticsearch" in row_text:
                rows_to_remove.append(ri)

        for ri in reversed(rows_to_remove):
            tbl = table._tbl
            tr = table.rows[ri]._tr
            tbl.remove(tr)
            modified_count += 1
            print(f"[技术栈表格] 删除第{ri}行（Elasticsearch）")
        break

# ============ 第三步：测试用例表格扩充——在Word导出行之前插入list和delete两行 ============
for ti, table in enumerate(doc.tables):
    # 寻找含"项目创建"和"Word 导出"的表格
    row_texts = []
    for row in table.rows:
        row_texts.append(" ".join(c.text for c in row.cells))

    has_create = any("项目创建" in t and "通过" in t for t in row_texts)
    has_export = any("Word 导出" in t or "Word导出" in t for t in row_texts)

    if has_create and has_export:
        # 找到"Word 导出"行的索引
        export_idx = None
        for ri, row in enumerate(table.rows):
            row_text = " ".join(c.text for c in row.cells)
            if "Word 导出" in row_text or "Word导出" in row_text:
                export_idx = ri
                break

        if export_idx is not None:
            # 在 export_idx 之前插入两行：项目列表、项目删除
            # 复制现有行作为模板
            template_row = table.rows[export_idx]._tr

            new_rows_data = [
                ("项目列表", "✓ 通过", "成功查询用户名下所有项目"),
                ("项目删除", "✓ 通过", "成功删除项目及关联的章节与操作日志"),
            ]

            for data in new_rows_data:
                # 创建新行（复制现有行的结构）
                new_tr = deepcopy(template_row)
                # 清空原有文本，写入新文本
                new_row_cells = new_tr.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc")
                for idx, cell_xml in enumerate(new_row_cells):
                    if idx < len(data):
                        # 清空cell内所有文字，保留格式
                        # 先移除所有现有的 <w:p> 段落的 <w:r>
                        for p in cell_xml.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                            for r in p.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"):
                                for t in r.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                                    t.text = data[idx]
                                    data_set = True
                                    break
                                break
                            break
                # 插入新行到 export_idx 位置之前
                template_row.addprevious(new_tr)
                modified_count += 1

            print(f"[测试表格] 已插入'项目列表'和'项目删除'两行")
        break

doc.save(DST)
print(f"\n第四轮共修改 {modified_count} 处")
