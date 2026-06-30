#!/usr/bin/env python3
from docx import Document

doc = Document("/home/ubuntu/毕业论文_修订版2.docx")

# 从段落249开始，往后读60段
for i in range(249, min(320, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    style = para.style.name
    if 'toc' in style.lower():
        continue
    print(f"[{i}] [{style}] {text[:120]}")
    # 遇到4.2.4停止
    if '4.2.4' in text or '4.3' in text:
        break
