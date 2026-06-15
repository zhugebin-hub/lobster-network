#!/usr/bin/env python3
"""第三轮：针对性修改剩余问题"""

from docx import Document
from copy import deepcopy

DST = "/home/ubuntu/毕业论文_修订版.docx"
doc = Document(DST)

modified_count = 0

def replace_para_text(para, new_text):
    """安全替换段落文本，保留首个run的格式"""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.text = new_text

# ============ 精准修改 ============
for i, para in enumerate(doc.paragraphs):
    text = para.text

    # P161 服务层整段改写
    if "LLM 网关服务" in text and "OpenAI、Claude" in text:
        new_text = (
            "服务层负责把核心逻辑和算法能力打包好然后封装。LLM 网关服务主要是把对外的模型调用收拢到一个口子，"
            "通过平台内置的 LLM 服务统一对接底层大模型，由平台负责请求限流、结果缓存和异常重试，"
            "这样上层就不用关心到底在跟哪个模型打交道。内容管理服务管的是章节内容的版本怎么控制、"
            "谁在编辑时锁定、多人协同发生冲突怎么解决，以及怎么回退到历史版本，都在这一层处理。"
            "知识库服务和学术知识图谱集成作为未来扩展方向，本期系统暂未实现。"
            "质量评估服务（规划中）则负责检查生成内容的学术规范性、字数符合度等，为后续迭代留下接口。"
        )
        replace_para_text(para, new_text)
        modified_count += 1
        print(f"[P{i}] 服务层段落重写完成")

    # P163 应用层"三栏"改为"两栏"
    elif "前端控制台做成了三栏布局" in text or ("三栏布局" in text and "章节导航" in text):
        new_text = text.replace(
            "前端控制台做成了三栏布局：左边是项目章节导航，中间是实时编辑和预览区，右边集中显示进度、日志和操作面板",
            "前端控制台做成了两栏布局：左边是项目章节导航和进度显示，右边是内容编辑和预览区，操作按钮集成在页面顶部"
        )
        if new_text != text:
            replace_para_text(para, new_text)
            modified_count += 1
            print(f"[P{i}] 应用层两栏布局改写")

    # P166 工作流"17个章节"
    elif "17 个章节" in text or "17个章节" in text:
        new_text = text.replace("17 个章节", "14 个章节").replace("17个章节", "14 个章节")
        replace_para_text(para, new_text)
        modified_count += 1
        print(f"[P{i}] 工作流章节数量")

    # 4.3.1 主工作区——三栏描述修改
    elif "我们把它设计成了三栏" in text or "三栏。左侧是章节树" in text:
        new_text = text.replace("设计成了三栏", "设计成了两栏")
        new_text = new_text.replace("三栏", "两栏")
        new_text = new_text.replace("章节树", "章节导航")
        replace_para_text(para, new_text)
        modified_count += 1
        print(f"[P{i}] 主工作区两栏描述")

# ============ 第二步：再扫一次，确保无遗漏 ============
REMAINING_REPLACEMENTS = [
    ("17 类", "14 类"),
    ("17类", "14类"),
    ("17 个章节", "14 个章节"),
    ("17个章节", "14 个章节"),
    ("1/17 = 6%", "1/14 ≈ 7%"),
    ("1/17", "1/14"),
    ("OpenAI、Claude 等大模型", "平台内置 LLM 服务"),
    ("OpenAI、Claude 等多个大模型", "平台内置 LLM 服务"),
    ("对接 OpenAI、Claude", "通过平台内置 LLM 服务对接"),
    ("多模型路由、请求限流", "由平台负责请求限流"),
    ("Elasticsearch", "关系数据库"),
    ("Monaco 编辑器", "Streamdown 组件"),
    ("Monaco编辑器", "Streamdown组件"),
    ("Zustand store", "React Query 缓存"),
    ("Zustand", "tRPC React Query"),
    ("三栏布局", "两栏布局"),
    ("做成了三栏", "做成了两栏"),
    ("设计成了三栏", "设计成了两栏"),
]

for i, para in enumerate(doc.paragraphs):
    text = para.text
    new_text = text
    for find, replace in REMAINING_REPLACEMENTS:
        if find in new_text:
            new_text = new_text.replace(find, replace)
    if new_text != text:
        replace_para_text(para, new_text)
        modified_count += 1
        print(f"[P{i}]兜底替换: {text[:50]}... → {new_text[:50]}...")

# 表格也兜底一遍
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                text = para.text
                new_text = text
                for find, replace in REMAINING_REPLACEMENTS:
                    if find in new_text:
                        new_text = new_text.replace(find, replace)
                if new_text != text:
                    replace_para_text(para, new_text)
                    modified_count += 1
                    print(f"[表{ti}行{ri}列{ci}]兜底: {text[:30]}... → {new_text[:30]}...")

doc.save(DST)
print(f"\n第三轮共修改 {modified_count} 处")
