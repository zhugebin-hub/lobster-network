#!/usr/bin/env python3
"""修改毕业论文：根据工程实际情况调整论文内容"""

from docx import Document
import shutil
import re

SRC = "/home/ubuntu/upload/毕业论文final.docx"
DST = "/home/ubuntu/毕业论文_修订版.docx"

# 复制一份，在副本上修改
shutil.copy(SRC, DST)
doc = Document(DST)

# ============ 全文文本替换规则 ============
# (find, replace, description)
REPLACEMENTS = [
    # 一、17章节 → 14章节
    ("17类官方章节", "14类官方章节"),
    ("17 类官方章节", "14 类官方章节"),
    ("17个章节", "14个章节"),
    ("17 个章节", "14 个章节"),
    ("17个官方章节", "14个官方章节"),
    ("17 个官方章节", "14 个官方章节"),
    ("1/17 = 6%", "1/14 ≈ 7%"),
    ("1/17", "1/14"),
    ("17 个子章节", "14 个子章节"),
    ("17个子章节", "14个子章节"),
    # 二、删除/改写 Elasticsearch 相关描述
    ("文献检索的活则交给 Elasticsearch，全文检索能力\n摆在那里，关键词搜索和语义搜索都指望它。",
     "文献检索功能作为未来扩展方向，本期暂未实现。"),
    ("文献检索的活则交给 Elasticsearch，全文检索能力摆在那里，关键词搜索和语义搜索都指望它。",
     "文献检索功能作为未来扩展方向，本期暂未实现。"),
    # 三、Monaco 编辑器 → Streamdown
    ("集成了 Monaco 编辑器，支持 Markdown 的实时预览",
     "使用 Streamdown 组件渲染 Markdown 内容，支持实时预览"),
    ("集成了Monaco编辑器", "使用Streamdown组件"),
    ("Monaco 编辑器", "Streamdown 组件"),
    ("Monaco编辑器", "Streamdown组件"),
    # 四、Zustand → tRPC React Query
    ("用一个共享的 Zustand store 来传\n递状态",
     "通过 tRPC React Query 的缓存机制来共享状态"),
    ("用一个共享的 Zustand store 来传递状态",
     "通过 tRPC React Query 的缓存机制来共享状态"),
    ("store 里的\nappendContent 方法", "React Query 的 setData 方法"),
    ("store 里的 appendContent 方法", "React Query 的 setData 方法"),
    ("Zustand store", "React Query 缓存"),
    ("状态管理选了 Zustand，足\n量轻量", "状态管理使用 tRPC React Query，足够轻量"),
    ("状态管理选了 Zustand，足\n够轻量", "状态管理使用 tRPC React Query，足够轻量"),
    ("状态管理选了 Zustand", "状态管理使用 tRPC React Query"),
    # 五、三栏 → 两栏
    ("我们把它设计成了三栏", "我们把它设计成了两栏"),
    ("把它设计成了三栏", "把它设计成了两栏"),
    ("三栏。左侧是章节树", "两栏。左侧是章节导航"),
    # 六、LLM 多模型 → 平台内置LLM
    ("支持多个 LLM 提供商", "调用平台内置 LLM 服务"),
    ("对接 OpenAI、Claude 等多个大模型",
     "调用平台内置 LLM 服务（封装了底层大模型调用）"),
    # 七、Word 导出描述更新（目录域 → 模板格式）
    ("在文档最前面，会插入一个\n目录域，也就是 TableOfContents，打开 Word 之后它会自动刷出目录。页眉那\n边放上项目名称，页码也一并生成。",
     "在文档最前面，会生成封面页（含申报编号、项目名称、申报单位、项目负责人等信息），紧跟其后是填报说明、项目基本信息表（含申报单位、推荐单位、项目负责人、项目联系人、课题分解等复杂合并单元格表格），然后是申报项目简介和各已确认章节正文。整体结构与《国家重点研发计划项目申报书模板》保持一致。"),
    ("在文档最前面，会插入一个目录域，也就是 TableOfContents，打开 Word 之后它会自动刷出目录。页眉那边放上项目名称，页码也一并生成。",
     "在文档最前面，会生成封面页（含申报编号、项目名称、申报单位、项目负责人等信息），紧跟其后是填报说明、项目基本信息表（含申报单位、推荐单位、项目负责人、项目联系人、课题分解等复杂合并单元格表格），然后是申报项目简介和各已确认章节正文。整体结构与《国家重点研发计划项目申报书模板》保持一致。"),
    # 八、系统访问URL（保持原地址不变，已正确）
    # 九、测试数量调整
    ("设计了六个从头到尾覆盖全流程的用例",
     "设计了八个从头到尾覆盖全流程的用例"),
    ("设计了 六 个从头到尾覆盖全流程的用例",
     "设计了八个从头到尾覆盖全流程的用例"),
]

# ============ 执行段落替换 ============
modified_count = 0
for para in doc.paragraphs:
    original = para.text
    new_text = original
    for find, replace in REPLACEMENTS:
        if find in new_text:
            new_text = new_text.replace(find, replace)
    if new_text != original:
        # 保留段落格式：清空runs后，在第一个run中写入新文本
        if para.runs:
            # 保留第一个run的格式，清除其他run
            first_run = para.runs[0]
            first_run.text = new_text
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.text = new_text
        modified_count += 1
        print(f"[段落替换] {original[:50]}... → {new_text[:50]}...")

# ============ 执行表格单元格替换 ============
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                original = para.text
                new_text = original
                for find, replace in REPLACEMENTS:
                    if find in new_text:
                        new_text = new_text.replace(find, replace)
                if new_text != original:
                    if para.runs:
                        first_run = para.runs[0]
                        first_run.text = new_text
                        for run in para.runs[1:]:
                            run.text = ""
                    else:
                        para.text = new_text
                    modified_count += 1
                    print(f"[表格替换] {original[:50]}... → {new_text[:50]}...")

print(f"\n共修改 {modified_count} 处")

# ============ 保存 ============
doc.save(DST)
print(f"\n已保存到: {DST}")
