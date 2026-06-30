#!/usr/bin/env python3
"""
基于模板生成Word文档的Python脚本
将AI生成的内容插入到国家重点研发计划模板的对应位置
"""

import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_from_template(template_path, output_path, proposal_data):
    """
    基于模板生成Word文档
    
    Args:
        template_path: 模板文件路径
        output_path: 输出文件路径
        proposal_data: 包含项目信息和章节内容的字典
    """
    
    # 加载模板
    doc = Document(template_path)
    
    # 查找"申报项目简介"位置
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        if "申报项目简介" in para.text:
            insert_index = i
            break
    
    if insert_index is None:
        # 如果找不到标记，在最后一个表格后插入
        insert_index = len(doc.paragraphs) - 1
    
    # 获取要插入的段落对象
    target_para = doc.paragraphs[insert_index]
    
    # 在目标位置后插入内容
    sections = proposal_data.get('sections', [])
    
    for section in sections:
        # 插入章节标题
        new_para = target_para.insert_paragraph_before(section['title'])
        new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 手动设置格式
        for run in new_para.runs:
            run.bold = True
            run.font.size = Pt(14)
        
        # 插入章节内容
        content_lines = section['content'].split('\n')
        for line in content_lines:
            if line.strip():
                new_para = target_para.insert_paragraph_before(line.strip())
                new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # 设置首行缩进
                new_para.paragraph_format.first_line_indent = Inches(0.5)
                # 设置字体大小
                for run in new_para.runs:
                    run.font.size = Pt(12)
    
    # 保存文档
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("用法: python docxTemplateGenerator.py <template_path> <output_path> <json_data>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    output_path = sys.argv[2]
    json_data = json.loads(sys.argv[3])
    
    generate_from_template(template_path, output_path, json_data)
