#!/usr/bin/env python3
from docx import Document
doc = Document("/home/ubuntu/毕业论文_修订版.docx")

print("=== 关键段落检查 ===\n")
for i, para in enumerate(doc.paragraphs):
    if 155 <= i <= 170:
        text = para.text
        if text.strip():
            print(f"[P{i}] {text[:200]}")
            print("---")

print("\n=== 搜索残留的未修改关键词 ===")
keywords = ["17 类", "17类", "17 个章节", "17个章节", "1/17", "Elasticsearch", "Monaco", "Zustand",
            "三栏", "OpenAI、Claude"]
for i, para in enumerate(doc.paragraphs):
    for kw in keywords:
        if kw in para.text:
            print(f"[残留P{i}] ({kw}): {para.text[:100]}")

print("\n=== 表格中关键词 ===")
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for kw in keywords:
                if kw in cell.text:
                    print(f"[残留表{ti}行{ri}列{ci}] ({kw}): {cell.text[:80]}")
