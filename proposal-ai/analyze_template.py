from docx import Document
from docx.oxml.ns import qn
import json

doc = Document('/tmp/template_original.docx')

print('=== PARAGRAPHS (first 80) ===')
for i, p in enumerate(doc.paragraphs[:80]):
    text = p.text.strip()
    if text:
        print(f'[{i}] style={p.style.name!r} | {text[:120]!r}')

print()
print(f'Total tables: {len(doc.tables)}')
print()
for ti, t in enumerate(doc.tables):
    print(f'--- Table {ti} ({len(t.rows)} rows x {len(t.columns)} cols) ---')
    for ri, row in enumerate(t.rows):
        row_data = []
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            row_data.append(f'[{ci}]={txt[:60]!r}')
        print(f'  Row {ri}: ' + ' | '.join(row_data))
    print()
