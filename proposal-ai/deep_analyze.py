#!/usr/bin/env python3
"""深度解析模板所有表格结构和格式细节"""
from docx import Document
from docx.shared import Pt, Cm
import json

doc = Document('/tmp/template.docx')

print("=" * 80)
print("模板总体信息")
print("=" * 80)
print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

# 页面设置
section = doc.sections[0]
print(f"页面宽度: {section.page_width.cm:.1f}cm")
print(f"页面高度: {section.page_height.cm:.1f}cm")
print(f"左边距: {section.left_margin.cm:.1f}cm")
print(f"右边距: {section.right_margin.cm:.1f}cm")
print(f"上边距: {section.top_margin.cm:.1f}cm")
print(f"下边距: {section.bottom_margin.cm:.1f}cm")

print("\n" + "=" * 80)
print("所有段落（前100个）")
print("=" * 80)
for i, p in enumerate(doc.paragraphs[:100]):
    text = p.text.strip()
    if text:
        style = p.style.name
        align = str(p.alignment)
        # 获取字体信息
        font_info = ""
        if p.runs:
            r = p.runs[0]
            if r.font.size:
                font_info = f" 字号:{r.font.size.pt:.0f}pt"
            if r.font.bold:
                font_info += " 粗体"
            if r.font.name:
                font_info += f" 字体:{r.font.name}"
        print(f"P{i:03d} [{style}] [{align}]{font_info}: {text[:80]}")

print("\n" + "=" * 80)
print("所有表格详细结构")
print("=" * 80)
for ti, table in enumerate(doc.tables):
    print(f"\n{'='*60}")
    print(f"表格 {ti+1}: {len(table.rows)}行 x {len(table.columns)}列")
    
    # 表格宽度
    if table.columns:
        try:
            widths = []
            for col in table.columns:
                try:
                    w = col.width
                    widths.append(f"{w.cm:.1f}cm" if w else "auto")
                except:
                    widths.append("?")
            print(f"列宽: {widths}")
        except:
            pass
    
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()[:60]
            # 检查合并
            tc = cell._tc
            grid_span = tc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
            v_merge = tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
            
            merge_info = ""
            if grid_span and int(grid_span) > 1:
                merge_info += f" [横跨{grid_span}列]"
            if v_merge is not None:
                val = v_merge.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'continue')
                merge_info += f" [纵向:{val}]"
            
            # 背景色
            shading = tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            bg = ""
            if shading is not None:
                fill = shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if fill and fill != 'auto' and fill != 'FFFFFF':
                    bg = f" [背景:{fill}]"
            
            if text or merge_info:
                print(f"  [{ri},{ci}]{merge_info}{bg}: {text}")
