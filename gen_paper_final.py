#!/usr/bin/env python3
"""Final Word generation for 叶畏兵 paper."""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = '仿宋_GB2312'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
pf = style.paragraph_format
pf.first_line_indent = Cm(0.74)
pf.alignment = 2
pf.line_spacing = Pt(25)
pf.space_before = Pt(0)
pf.space_after = Pt(0)

def blank():
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(12)

def title(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = 1
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(16); r.font.bold = True
    r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def author(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = 1
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = '仿宋_GB2312'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

def meta(label, content):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0); pf.alignment = 2
    pf.line_spacing = Pt(25); pf.space_before = Pt(0); pf.space_after = Pt(0)
    r1 = p.add_run(label)
    r1.font.size = Pt(12); r1.font.bold = True
    r1.font.name = '黑体'; r1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r2 = p.add_run(content)
    r2.font.size = Pt(12)
    r2.font.name = '仿宋_GB2312'; r2.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

def h1(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0); pf.alignment = 2
    pf.space_before = Pt(12); pf.space_after = Pt(6)
    pf.line_spacing = Pt(25)
    r = p.add_run(text)
    r.font.size = Pt(14); r.font.bold = True
    r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def h2(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0); pf.alignment = 2
    pf.space_before = Pt(6); pf.space_after = Pt(6)
    pf.line_spacing = Pt(25)
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.bold = True
    r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def para(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.74); pf.alignment = 2
    pf.line_spacing = Pt(25); pf.space_before = Pt(0); pf.space_after = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = '仿宋_GB2312'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

def ref_para(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0); pf.alignment = 2
    pf.line_spacing = Pt(25); pf.space_before = Pt(0); pf.space_after = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = '仿宋_GB2312'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

# Read the md file and parse it
with open('/home/admin/.openclaw/workspace/叶畏兵_眼泪神学_去AI最终版.md', 'r') as f:
    lines = f.readlines()

# Skip front matter, start from first ##
skip = True
for line in lines:
    s = line.strip()
    if not s:
        continue
    if skip:
        if s.startswith('## 一、'):
            skip = False
            h1(s[3:])
        continue

    if not s:
        continue

    if s.startswith('## '):
        ht = s[3:].strip()
        if any(ht.startswith(x) for x in ['一、','二、','三、','四、','五、','结语','参考文献']):
            h1(ht)
        else:
            h2(ht)
    else:
        para(s)

# References
refs = [
    '[1]《圣经》（和合本），南京：中国基督教协会，2000年。',
    '[2]Brueggemann, Walter. The Message of the Psalms: A Theological Commentary. Minneapolis: Augsburg Publishing House, 1984.',
    '[3]Williamson, H.G.M. Ezra, Nehemiah. Word Biblical Commentary, Vol. 16. Waco: Word Books, 1985.',
    '[4]Yamauchi, Edwin M. "Nehemiah." In The Expositor\'s Bible Commentary, Vol. 3, edited by Frank E. Gaebelein, 637-724. Grand Rapids: Zondervan, 1990.',
    '[5]Roberts, J.J.M. "Nehemiah." In The New Interpreter\'s Bible, Vol. III, 815-918. Nashville: Abingdon Press, 1999.',
    '[6]丁光训：《丁光训文集》，南京：译林出版社，1998年。',
    '[7]丁光训：《神学思考的足迹》，北京：宗教文化出版社，2005年。',
    '[8]赵晓阳编：《基督教中国化研究》，北京：社会科学文献出版社，2013年。',
    '[9]卓新平：《当代西方新教神学》，上海：上海三联书店，1998年。',
    '[10]许志伟：《基督教神学基础》，北京：宗教文化出版社，2006年。',
    '[11]尤思德：《和而不同：基督教神学与中国文化的对话》，香港：汉语圣经协会，2009年。',
]

for ref in refs:
    ref_para(ref)

out = '/home/admin/.openclaw/workspace/叶畏兵_眼泪的神学_最终版.docx'
doc.save(out)
print(f'Saved to {out}')
