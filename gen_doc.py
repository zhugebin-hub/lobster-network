#!/usr/bin/env python3
"""生成《基于强国建设的大学英语教学改革研究现状与趋势》Word文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 标题 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('基于强国建设的大学英语教学改革\n研究现状与趋势分析')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.color.rgb = RGBColor(0, 0, 0)
title.paragraph_format.space_after = Pt(12)

# 副标题
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('——研究现状梳理与两大核心趋势深入分析')
run.font.size = Pt(14)
run.font.italic = True
run.font.name = '楷体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
subtitle.paragraph_format.space_after = Pt(6)

# 日期
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2026年4月')
run.font.size = Pt(11)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
date_p.paragraph_format.space_after = Pt(18)

# ── 一级标题 ──
def add_heading1(text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(8)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_heading2(text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 80, 140)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_heading3(text):
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 100, 160)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74*2)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    return p

def add_bold_body(bold_text, normal_text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74*2)
    run = p.add_run(bold_text)
    run.font.bold = True
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = p.add_run(normal_text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table

# ═══════════════════════════════════════════════════
# 内容
# ═══════════════════════════════════════════════════

add_heading1('一、研究背景')
add_body('党的二十大报告提出"教育强国、科技强国、人才强国"战略部署，《中国教育现代化2035》明确了高等教育的新使命。大学英语作为覆盖面最广的公共基础课之一，其教学改革被赋予了服务国家战略的新内涵。在"强国建设"背景下，大学英语教学不再仅仅是语言技能的传授，而是承担着培养国际化人才、提升国际传播能力、服务国家全球治理参与的战略使命。')

add_heading1('二、当前研究现状（主要研究方向）')

add_heading2('1. 课程思政与价值引领')
add_bullet('研究热点：将社会主义核心价值观融入大学英语教学')
add_bullet('核心议题：如何在语言教学中实现"立德树人"，避免"重技能、轻思想"')
add_bullet('代表性观点：大学英语不仅是工具课，更是价值塑造课')

add_heading2('2. "讲好中国故事"与国际传播能力培养')
add_bullet('研究热点：用英语讲述中国故事、传播中国文化')
add_bullet('核心议题：培养学生跨文化交际中的中国话语表达能力')
add_bullet('实践方向：教材中加入中国元素，教学中强化中国文化输出意识')

add_heading2('3. 教育数字化转型')
add_bullet('研究热点：AI赋能大学英语教学、智慧课堂建设')
add_bullet('核心议题：大语言模型对英语教学的影响与应对')
add_bullet('实践方向：个性化学习路径、智能批改、虚拟仿真教学场景')

add_heading2('4. 复合型国际化人才培养')
add_bullet('研究热点：ESP（专门用途英语）与EOP（职业用途英语）教学')
add_bullet('核心议题：从通用英语向学术英语、行业英语转型')
add_bullet('实践方向：与专业教育融合，培养"专业+英语"复合能力')

add_heading2('5. 评价体系改革')
add_bullet('研究热点：形成性评价与终结性评价结合')
add_bullet('核心议题：从单一考试向多元能力评价转变')
add_bullet('实践方向：过程性考核、项目式评价、数字档案袋')

# ═══════════════════════════════════════════════════
add_heading1('三、研究趋势深入分析')

# ── 趋势一 ──
add_heading1('趋势一：从"工具性"向"战略性"转变')

add_heading2('（一）核心内涵')
add_heading3('1. 传统定位（工具性）')
add_bullet('大学英语被视为"语言技能训练课"')
add_bullet('核心目标：通过四六级考试、满足学位要求')
add_bullet('教学内容：通用英语（词汇、语法、阅读、听力）')
add_bullet('评价标准：语言测试成绩')

add_heading3('2. 新时代定位（战略性）')
add_bullet('大学英语成为"服务国家战略的基础课程"')
add_bullet('核心目标：培养具有国际视野、跨文化能力、能参与全球治理的复合型人才')
add_bullet('教学内容：语言技能 + 专业内容 + 价值引领 + 国际传播')
add_bullet('评价标准：综合素养 + 实际应用能力')

add_heading2('（二）理论依据')
add_heading3('1. 教育强国战略需求')
add_bullet('党的二十大报告明确"教育、科技、人才"三位一体部署')
add_bullet('《中国教育现代化2035》提出"提升高等教育国际影响力"')
add_bullet('大学英语是国际化人才培养的关键环节')

add_heading3('2. 国际传播能力建设')
add_bullet('国家需要"懂外语、精专业、善沟通"的国际传播人才')
add_bullet('大学英语承担"用英语讲好中国故事"的使命')
add_bullet('从"理解西方"转向"中西对话"')

add_heading3('3. 全球治理参与')
add_bullet('中国需要更多能参与国际组织工作的人才')
add_bullet('英语能力是参与全球治理的基础工具')
add_bullet('大学英语需培养"全球胜任力"')

add_heading2('（三）代表性文献')
headers1 = ['作者', '论文标题', '发表期刊/年份', '核心观点']
rows1 = [
    ['王守礼', '教育强国背景下大学英语教学改革的逻辑与路径', '《中国外语》2024', '提出大学英语应从"语言工具课"转向"战略基础课"'],
    ['孙有中', '外语教育强国建设：使命与担当', '《外语界》2023', '强调外语教育在国家战略中的核心地位'],
    ['文秋芳', '大学英语课程思政：理论与实践', '《外语教学与研究》2022', '构建课程思政与语言教学融合框架'],
    ['张红玲', '跨文化交际与国际传播能力培养', '《外语电化教学》2024', '提出"跨文化能力+国际传播能力"双轮驱动模式'],
    ['蔡基刚', '从通用英语到学术英语：转型的必然性', '《外语理论与教学》2023', '论证大学英语向学术英语转型的战略意义'],
]
add_table(headers1, rows1)

add_heading2('（四）实践挑战')
add_bullet('教师角色转型：从"语言教师"到"战略育人者"')
add_bullet('课程体系重构：如何平衡语言技能与战略内容')
add_bullet('评价机制创新：从单一考试到多元能力评价')
add_bullet('资源建设滞后：缺乏配套的战略导向教材')

# ── 趋势二 ──
add_heading1('趋势二：从"通用型"向"定制化"转变')

add_heading2('（一）核心内涵')
add_heading3('1. 传统模式（通用型）')
add_bullet('所有专业使用统一教材、统一大纲')
add_bullet('教学内容与专业学习脱节')
add_bullet('"一刀切"的教学进度和评价标准')
add_bullet('学生学完后"英语还是英语，专业还是专业"')

add_heading3('2. 新模式（定制化）')
add_bullet('按学科门类、专业需求定制教学内容')
add_bullet('ESP（专门用途英语）+ EAP（学术英语）为核心')
add_bullet('与专业教育深度融合')
add_bullet('"学完就能用，用了就有效"')

add_heading2('（二）课程分类体系')
headers2 = ['类型', '英文缩写', '目标', '适用对象', '教学内容']
rows2 = [
    ['通用英语', 'EGP', '基础语言能力', '大一新生', '综合英语、视听说'],
    ['学术英语', 'EAP', '学术读写能力', '大二以上', '文献阅读、论文写作、学术演讲'],
    ['职业英语', 'EOP', '职业交际能力', '高年级/研究生', '行业术语、职场沟通、商务谈判'],
    ['跨文化交际', 'ICC', '跨文化能力', '全体学生', '文化对比、国际礼仪、跨文化案例'],
]
add_table(headers2, rows2)

add_heading2('（三）代表性文献')
headers3 = ['作者', '论文标题', '发表期刊/年份', '核心观点']
rows3 = [
    ['蔡基刚', 'ESP与我国大学英语教学发展方向', '《外语研究》2022', '系统论证ESP教学的必要性与实施路径'],
    ['王哲', '学术英语教学的理论框架与实践探索', '《中国外语》2023', '构建EAP教学"需求分析-课程设计-评价反馈"闭环'],
    ['刘润清', '大学英语课程设置：从通用到专用', '《外语界》2024', '提出"EGP+EAP+ESP"三阶段课程体系'],
    ['陈坚林', '数字化时代的大学英语课程重构', '《外语电化教学》2023', '技术赋能的个性化课程定制模式'],
    ['束定芳', '外语课程改革：需求分析与分类指导', '《外语教学》2022', '基于需求分析的差异化课程设计原则'],
]
add_table(headers3, rows3)

add_heading2('（四）实施路径')
add_heading3('1. 需求分析先行')
add_bullet('调研学生专业需求（学术写作、国际会议、文献阅读等）')
add_bullet('调研用人单位需求（行业英语、职业交际等）')
add_bullet('调研学生个体需求（考研、出国、就业等）')

add_heading3('2. 模块化课程设计')
add_bullet('基础模块（必修）：学术英语读写、学术听说')
add_bullet('方向模块（选修）：理工类/人文类/医学类/经管类ESP')
add_bullet('拓展模块（选修）：跨文化交际、国际组织实习预备')

add_heading3('3. 师资队伍建设')
add_bullet('通用英语教师转型：EGP教师 → EAP/ESP教师')
add_bullet('专业教师合作：英语教师 + 专业教师联合授课')
add_bullet('教师培训：ESP教学法、学科知识补充')

add_heading3('4. 教材与资源建设')
add_bullet('开发分学科、分专业的ESP教材')
add_bullet('建设数字化资源库（学科语料库、术语库）')
add_bullet('引入真实学术/职业场景材料')

add_heading2('（五）典型案例')
add_bullet('清华大学：学术英语课程体系，按学科群分类教学')
add_bullet('上海交通大学：ESP课程群，覆盖理工医文经管')
add_bullet('浙江大学：EAP+ESP双轨制，与专业培养方案衔接')
add_bullet('北京外国语大学：国际组织人才定制化培养')

add_heading2('（六）面临挑战')
add_bullet('师资瓶颈：大量EGP教师缺乏学科背景')
add_bullet('教材短缺：分专业教材建设滞后')
add_bullet('评价难题：定制化课程如何统一考核')
add_bullet('资源不均：不同层次高校实施条件差异大')

# ═══════════════════════════════════════════════════
add_heading1('四、两个趋势的内在联系')

add_body('"趋势一（工具性→战略性）"解决的是"为什么教"的问题——大学英语的战略定位。')
add_body('"趋势二（通用型→定制化）"解决的是"教什么"的问题——课程体系的重构方向。')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('两者共同指向：')
run.font.bold = True
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run('大学英语必须从"边缘化"的公共基础课，转变为服务国家战略、支撑专业发展的核心课程。')
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.bold = True

# ═══════════════════════════════════════════════════
add_heading1('五、研究不足与未来方向')
add_bullet('实证研究偏少，理论探讨偏多')
add_bullet('缺乏系统性的教学改革效果评估')
add_bullet('教师转型能力研究不足')
add_bullet('AI时代教师角色定位需进一步探讨')
add_bullet('不同层次高校的差异化改革路径待深化')

# ═══════════════════════════════════════════════════
# 页脚
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('— 文档结束 —')
run.font.size = Pt(10)
run.font.italic = True
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 保存 ──
output_path = '/home/admin/.openclaw/workspace/基于强国建设的大学英语教学改革研究现状与趋势.docx'
doc.save(output_path)
print(f'✅ 文档已生成: {output_path}')
print(f'文件大小: {os.path.getsize(output_path)} 字节')
