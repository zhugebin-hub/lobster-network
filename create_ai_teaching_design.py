"""
生成AI赋能版教学设计 Word 文档
课题：七年级上册 第四章 第4节 地球板块的缓慢运动（第1课时）
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

def set_cell_font(run, font_name='仿宋', size=10.5, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_cell_text(cell, text, font_name='仿宋', size=10.5, bold=False, alignment=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    set_cell_font(run, font_name, size, bold)
    # 行距
    pPr = p._element.get_or_add_pPr()
    spc = OxmlElement('w:spacing')
    spc.set(qn('w:line'), '276')
    spc.set(qn('w:lineRule'), 'auto')
    pPr.append(spc)

def set_col_widths(table, widths_cm):
    """设置每列宽度（厘米）"""
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)

def set_cell_span(cell, rows=1, cols=1):
    """合并单元格"""
    tc = cell._tc
    if cols > 1:
        tcgrid = tc.get_or_add_tcPr()
        hMerge = OxmlElement('w:hMerge')
        hMerge.set(qn('w:val'), 'restart')
        tcgrid.append(hMerge)
    if rows > 1:
        tcgrid = tc.get_or_add_tcPr()
        vMerge = OxmlElement('w:vMerge')
        vMerge.set(qn('w:val'), 'restart')
        tcgrid.append(vMerge)

def merge_cells_horizontal(cell_start, cell_end):
    cell_start.merge(cell_end)

def merge_cells_vertical(table, col_idx, start_row, end_row):
    for r in range(start_row, end_row + 1):
        tc = table.cell(r, col_idx)._tc
        tcPr = tc.get_or_add_tcPr()
        vMerge = OxmlElement('w:vMerge')
        if r == start_row:
            vMerge.set(qn('w:val'), 'restart')
        else:
            vMerge.set(qn('w:val'), 'continue')
        tcPr.append(vMerge)

# ============ 创建文档 ============
doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# 样式设置
style = doc.styles['Normal']
font = style.font
font.name = '仿宋'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

# ============ 标题 ============
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.space_before = Pt(10)
title_p.space_after = Pt(15)
run = title_p.add_run('AI赋能的教学设计')
run.font.name = '黑体'
run.font.size = Pt(22)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_p.space_after = Pt(15)
run = subtitle_p.add_run('七年级上册 第四章 第4节 地球板块的缓慢运动（第1课时）')
run.font.name = '楷体'
run.font.size = Pt(14)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

# ============ 基本信息表格 ============
# 表1：基本信息
info_table = doc.add_table(rows=2, cols=6)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'

# 合并课题单元格
cell0 = info_table.cell(0, 0)
add_cell_text(cell0, '课题', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
cell1 = info_table.cell(0, 1)
add_cell_text(cell1, '地球板块的缓慢运动\n（第1课时）', alignment=WD_ALIGN_PARAGRAPH.CENTER)
merge_cells_horizontal(info_table.cell(0, 1), info_table.cell(0, 3))

cell4 = info_table.cell(0, 4)
add_cell_text(cell4, '课型', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
cell5 = info_table.cell(0, 5)
add_cell_text(cell5, '新授课', alignment=WD_ALIGN_PARAGRAPH.CENTER)

cell6 = info_table.cell(1, 0)
add_cell_text(cell6, '教材版本', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
cell7 = info_table.cell(1, 1)
add_cell_text(cell7, '浙教版科学七年级上册', alignment=WD_ALIGN_PARAGRAPH.CENTER)
merge_cells_horizontal(info_table.cell(1, 1), info_table.cell(1, 3))

cell8 = info_table.cell(1, 4)
add_cell_text(cell8, '课时', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
cell9 = info_table.cell(1, 5)
add_cell_text(cell9, '1课时（45分钟）', alignment=WD_ALIGN_PARAGRAPH.CENTER)

set_col_widths(info_table, [2.0, 4.5, 4.5, 4.5, 2.0, 2.5])

doc.add_paragraph()  # 空行

# ============ 教学目标 ============
obj_title = doc.add_paragraph()
run = obj_title.add_run('一、教学目标')
run.font.name = '黑体'
run.font.size = Pt(14)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
obj_title.space_after = Pt(8)

obj_table = doc.add_table(rows=5, cols=2)
obj_table.style = 'Table Grid'
obj_table.alignment = WD_TABLE_ALIGNMENT.CENTER

obj_headers = ['维度', '目标描述']
for i, h in enumerate(obj_headers):
    add_cell_text(obj_table.cell(0, i), h, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

obj_data = [
    ['科学观念',
     '1.初步形成"地球岩石圈由六大板块构成"的核心认知，理解大陆漂移说、海底扩张说、板块构造学说的递进关系与科学内涵。\n'
     '2.明确板块运动（碰撞、张裂）是地形地貌演化的重要动力，掌握珠穆朗玛峰"长高"的本质是印度洋板块与亚欧板块持续挤压的结果。\n'
     '3.能运用板块构造学说解释喜马拉雅山隆起、地中海演化等自然现象，建立"地形与地质运动"的关联观念。\n'
     '4.借助AI三维地球模型，直观感知板块的空间分布与运动方向，形成动态的地球系统观念。\n'
     '核心概念：地球系统\n'
     '学习内容：地球内部圈层和地壳运动'],
    ['科学思维',
     '1.通过拼合大陆轮廓、撕纸类比等活动，发展类比推理与实证思维，学会依据证据推导科学结论。\n'
     '2.借助AI模拟实验和动态图示分析，提升模型认知能力，能用板块运动模型解释抽象的地质现象。\n'
     '3.梳理三大科学假说的演进逻辑，培养逻辑串联与归纳概括能力，理解科学理论"猜想—证据—完善"的发展过程。\n'
     '4.利用AI学习助手进行证据辨析训练，提升批判性思维和多角度分析能力。'],
    ['探究实践',
     '1.参与拼合大陆轮廓、模拟板块碰撞张裂等动手活动，掌握"观察—猜想—验证—结论"的探究流程。\n'
     '2.通过小组合作分析科学证据、讨论地形演化趋势，提升合作探究与表达交流能力。\n'
     '3.运用AI地质模拟平台预测地中海、红海的演化方向，实现知识迁移与实践应用。\n'
     '4.利用AI生成个性化探究任务，培养自主学习和解决实际地理问题的能力。'],
    ['态度责任',
     '1.激发对地球科学的探究兴趣，树立"像科学家一样思考"的主动探究意识。\n'
     '2.尊重科学理论的演进规律，体会科学家敢于猜想、严谨求证的科学精神，培养严谨务实的治学态度。\n'
     '3.认识地球地形的动态演化特征，树立"善待地球、敬畏自然"的可持续发展观念，增强对地理学科的价值认同。\n'
     '4.通过AI技术感受现代科技在科学研究中的应用，理解科技与科学的相互促进关系。'],
]

for r, (dim, desc) in enumerate(obj_data):
    add_cell_text(obj_table.cell(r+1, 0), dim, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(obj_table.cell(r+1, 1), desc)

set_col_widths(obj_table, [2.5, 16.5])

doc.add_paragraph()

# ============ 重点难点 ============
kd_title = doc.add_paragraph()
run = kd_title.add_run('二、重点与难点')
run.font.name = '黑体'
run.font.size = Pt(14)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
kd_title.space_after = Pt(8)

kd_table = doc.add_table(rows=3, cols=2)
kd_table.style = 'Table Grid'
kd_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_cell_text(kd_table.cell(0, 0), '项目', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(kd_table.cell(0, 1), '内容', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_cell_text(kd_table.cell(1, 0), '重点', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(kd_table.cell(1, 1),
    '1.掌握大陆漂移说、海底扩张说、板块构造学说的核心内涵及递进逻辑，明确板块运动（碰撞、张裂）的基本形式。\n'
    '2.理解珠穆朗玛峰"长高"的地质本质——印度洋板块与亚欧板块持续挤压，能运用板块构造学说解释喜马拉雅山隆起、红海扩张等自然现象。\n'
    '3.掌握"观察—猜想—验证—结论"的探究流程，初步形成基于证据推导科学结论的实证思维。\n'
    '4.借助AI三维地球模型和模拟平台，直观理解板块运动与地形演化的关系。')

add_cell_text(kd_table.cell(2, 0), '难点', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(kd_table.cell(2, 1),
    '1.理解海底扩张说的动力机制（地幔对流推动洋壳扩张），以及其作为大陆漂移动力来源的逻辑关联，突破抽象地质过程的认知障碍。\n'
    '2.建立板块运动与地形演化的空间逻辑，准确区分板块碰撞（消亡边界）、张裂（生长边界）对应的不同地形类型（隆起山脉、裂谷、海洋）。\n'
    '3.理解科学理论"猜想—证据—完善"的演进规律，将三大假说的逻辑串联起来，形成对地质科学发展的系统性认知。\n'
    '4.利用AI地质模拟平台将抽象地质过程可视化，帮助学生突破空间想象和动态推理的认知难点。')

set_col_widths(kd_table, [2.0, 17.0])

doc.add_paragraph()

# ============ 教学资源 ============
res_title = doc.add_paragraph()
run = res_title.add_run('三、教学资源（含AI赋能资源）')
run.font.name = '黑体'
run.font.size = Pt(14)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
res_title.space_after = Pt(8)

res_table = doc.add_table(rows=2, cols=2)
res_table.style = 'Table Grid'
res_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_cell_text(res_table.cell(0, 0), '类别', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(res_table.cell(0, 1), '资源清单', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_cell_text(res_table.cell(1, 0), '教学资源',
    '传统资源：\n'
    '1.PPT课件（含板块分布图、海底年龄图）\n'
    '2.证据资料卡（古生物分布、地层构造对比）\n'
    '3.活动材料：世界地图拼图、纸张、两本书、泡沫板、土豆糊\n\n'
    'AI赋能资源：\n'
    '1.AI生成视频：《魏格纳的发现》（AI动画还原发现过程）\n'
    '2.AI模拟视频：海底扩张动态模拟（AI生成地幔对流可视化）\n'
    '3.AI三维地球模型：交互式板块分布与运动演示\n'
    '4.AI学习助手：课堂实时问答与证据辨析辅导\n'
    '5.AI地质模拟平台：板块碰撞/张裂交互模拟\n'
    '6.AI测评系统：随堂练习自动批改与学情分析')

add_cell_text(res_table.cell(1, 1),
    '【AI视频1】《魏格纳的发现》——AI动画还原1910年魏格纳在病床上观察世界地图、提出大陆漂移说的故事，以沉浸式叙事激发学生兴趣。\n\n'
    '【AI视频2】海底扩张模拟——AI生成地幔对流推动洋壳向两侧扩张的动态过程，将抽象的地幔对流机制可视化。\n\n'
    '【AI三维地球模型】交互式3D地球仪，学生可旋转查看六大板块分布、板块边界类型、板块运动方向，支持AR叠加显示地震带与火山分布。\n\n'
    '【AI学习助手】嵌入课堂的智能问答系统，学生可随时提问（如"为什么魏格纳的猜想最初不被接受？"），AI即时反馈并引导思考。\n\n'
    '【AI地质模拟平台】学生通过拖拽板块边界、设置运动方向，实时观察地形变化（山脉隆起、裂谷形成、海洋扩张），实现"做中学"。\n\n'
    '【AI测评系统】随堂3-5道选择题，AI自动批改并生成班级正确率热力图，教师据此调整教学节奏。')

set_col_widths(res_table, [4.5, 14.5])

doc.add_paragraph()

# ============ 教学过程 ============
tp_title = doc.add_paragraph()
run = tp_title.add_run('四、教学过程')
run.font.name = '黑体'
run.font.size = Pt(14)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
tp_title.space_after = Pt(8)

# 教学过程表格
tp_table = doc.add_table(rows=7, cols=4)
tp_table.style = 'Table Grid'
tp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
headers = ['环节', '教师活动', '学生活动', '设计意图（含AI赋能说明）']
for i, h in enumerate(headers):
    add_cell_text(tp_table.cell(0, i), h, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# 导入
add_cell_text(tp_table.cell(1, 0), '导入\n（5分钟）', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(tp_table.cell(1, 1),
    '1.播放AI生成视频《如何精准测量珠穆朗玛峰的海拔高度》，引导学生关注珠峰海拔持续缓慢升高的事实。\n'
    '2.提问："珠峰为什么会\'长高\'？"\n'
    '3.利用AI学习助手发起即时投票：学生通过平板选择自己的初步猜测（地震抬升/板块挤压/火山堆积/其他），AI实时统计并投影结果。\n'
    '4.引出课题《揭秘珠穆朗玛峰"会长高"》。')
add_cell_text(tp_table.cell(1, 2),
    '1.观看AI生成视频，感受珠峰测量的科学过程。\n'
    '2.思考问题，通过平板参与投票。\n'
    '3.查看全班投票结果，产生认知冲突和探究欲望。')
add_cell_text(tp_table.cell(1, 3),
    '【AI赋能】\n'
    '1.AI生成视频以沉浸式叙事还原珠峰测量过程，比传统视频更具吸引力。\n'
    '2.AI即时投票系统实现全员参与，教师快速掌握学情前概念，为后续教学提供数据支撑。\n'
    '3.以真实情境和悬念激发探究兴趣，搭建"生活现象→科学问题"的联结。')

# 任务一
add_cell_text(tp_table.cell(2, 0), '任务一\n大陆漂移说\n（10分钟）', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(tp_table.cell(2, 1),
    '1.播放AI动画视频《魏格纳的发现》，介绍魏格纳的猜想：大陆原来是连在一起的，后来漂移分开。\n'
    '2.组织活动：学生拼合南美洲和非洲大陆轮廓图，并尝试拼合被撕开的纸张（类比推理）。\n'
    '3.引导思考：魏格纳需要寻找哪些证据？利用AI学习助手，学生可向其提问并获取证据线索。\n'
    '4.出示证据卡：轮廓吻合、古生物亲缘关系（舌羊齿化石）、地质构造相似（地层、冰川沉积连续分布）。\n'
    '5.利用AI地质模拟平台展示大陆漂移的动态过程。\n'
    '6.提问："能否用大陆漂移说解释珠峰长高？"（不能，缺少动力机制）')
add_cell_text(tp_table.cell(2, 2),
    '1.观看AI动画视频，了解魏格纳的发现过程。\n'
    '2.小组拼图活动，感受大陆轮廓的吻合性。\n'
    '3.与AI学习助手互动，获取证据线索并辨析证据有效性。\n'
    '4.在AI平台上观察大陆漂移的动态模拟。\n'
    '5.记录大陆漂移说的核心观点及证据。')
add_cell_text(tp_table.cell(2, 3),
    '【AI赋能】\n'
    '1.AI动画视频以叙事化方式还原科学发现过程，增强情境代入感。\n'
    '2.AI学习助手作为"智能学伴"，支持差异化学习——基础薄弱学生可获得更多提示，学有余力学生可深入探究。\n'
    '3.AI地质模拟平台将抽象的大陆漂移过程可视化，帮助学生建立动态空间认知。\n'
    '4.通过动手操作感知假说的直观依据，培养观察能力；借助类比推理和证据辨析，发展实证思维。')

# 任务二
add_cell_text(tp_table.cell(3, 0), '任务二\n海底扩张说\n（10分钟）', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(tp_table.cell(3, 1),
    '1.过渡：大陆漂移说因缺少动力来源未被认可。播放AI生成视频《海底扩张说》。\n'
    '2.展示大洋中脊及洋底岩石年龄分布图，提问："越远离大洋中脊，岩石年龄越老，这说明了什么？"\n'
    '3.模拟实验：加热土豆糊（模拟地幔物质上涌），观察泡沫板（模拟洋壳）向两边移动。\n'
    '4.利用AI地质模拟平台，学生自主调节地幔对流速度，观察洋壳扩张速率的变化。\n'
    '5.小结：地幔物质上涌推动洋壳向两侧扩张，为大陆漂移提供了动力。\n'
    '6.再次提问："现在能解释珠峰长高吗？"（仍不能，需要板块碰撞）')
add_cell_text(tp_table.cell(3, 2),
    '1.观看AI生成的海底扩张模拟视频。\n'
    '2.读图分析，发现洋底年龄对称分布规律。\n'
    '3.观察模拟实验，描述现象。\n'
    '4.在AI平台上自主调节参数，观察不同对流速度下的洋壳扩张效果。\n'
    '5.讨论海底扩张说对大陆漂移说的支持作用。')
add_cell_text(tp_table.cell(3, 3),
    '【AI赋能】\n'
    '1.AI生成视频将地幔对流这一不可见的地下过程可视化，突破认知难点。\n'
    '2.AI地质模拟平台支持参数调节（对速度、板块厚度等），学生通过"试错—观察—归纳"自主建构知识，实现探究式学习。\n'
    '3.利用多媒体和模拟实验将抽象地质过程具象化，引导学生分析实验与原理的关联，提升模型认知能力。')

# 任务三
add_cell_text(tp_table.cell(4, 0), '任务三\n板块构造学说\n（12分钟）', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(tp_table.cell(4, 1),
    '1.展示AI三维地球模型，引导学生观察：岩石圈被海沟、海岭分割成六大板块。\n'
    '2.识别板块名称及边界类型，重点标注印度洋板块与亚欧板块。\n'
    '3.模拟活动：两本书相向碰撞，观察"隆起"现象。\n'
    '4.讲解珠峰长高原理：印度洋板块与亚欧板块持续碰撞挤压，使喜马拉雅山脉缓慢抬升。\n'
    '5.利用AI地质模拟平台，学生自主设置板块运动方式（碰撞/张裂），观察地形变化。\n'
    '6.呈现地中海、红海板块分布示意图，提问："未来两地会如何演化？"引导学生运用碰撞（消亡）和张裂（生长）规律分析。\n'
    '7.利用AI测评系统发布3道随堂练习题，即时检测学习效果。')
add_cell_text(tp_table.cell(4, 2),
    '1.操作AI三维地球模型，旋转查看六大板块分布。\n'
    '2.参与碰撞模拟，描述现象。\n'
    '3.在AI平台上自主设置板块运动方式，观察地形变化并记录结果。\n'
    '4.解释珠峰长高原因。\n'
    '5.预测地中海缩小、红海扩大，并说明依据。\n'
    '6.完成AI随堂练习，查看即时反馈。')
add_cell_text(tp_table.cell(4, 3),
    '【AI赋能】\n'
    '1.AI三维地球模型支持AR叠加显示地震带与火山分布，帮助学生建立板块运动与地质灾害的空间关联。\n'
    '2.AI地质模拟平台实现"做中学"——学生通过自主操作探索板块运动规律，而非被动接受结论。\n'
    '3.AI测评系统即时批改并生成班级正确率热力图，教师据此精准调整教学节奏，实现数据驱动的教学决策。\n'
    '4.建立板块运动的空间认知，通过模拟实验深化"碰撞→隆起"的逻辑关联；开展知识迁移。')

# 小结与升华
add_cell_text(tp_table.cell(5, 0), '小结与升华\n（5分钟）', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(tp_table.cell(5, 1),
    '1.利用AI思维导图工具，师生共同梳理三大假说的演进逻辑：猜想（魏格纳）→ 证据（古生物、地质）→ 动力（海底扩张）→ 完善（板块构造）。\n'
    '2.强调科学假说需要实践检验，随着新证据出现会不断修正和发展。\n'
    '3.引导学生思考：AI技术如何帮助我们更好地理解地球科学？\n'
    '4.布置课后任务。')
add_cell_text(tp_table.cell(5, 2),
    '1.回顾学习内容，在AI思维导图工具上补充完善知识结构。\n'
    '2.体会科学理论的动态发展过程。\n'
    '3.思考AI技术在科学研究中的应用价值。')
add_cell_text(tp_table.cell(5, 3),
    '【AI赋能】\n'
    '1.AI思维导图工具支持协作式知识建构，学生可实时查看他人补充的内容，形成集体智慧。\n'
    '2.引导学生反思AI技术在科学研究中的作用，培养科技素养和数字意识。\n'
    '3.帮助学生建立系统性认知，理解科学本质，激发进一步探究的兴趣。')

# 板书设计
add_cell_text(tp_table.cell(6, 0), '板书设计', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(tp_table.cell(6, 1),
    '第四章 第4节 地球板块的缓慢运动\n\n'
    '一、大陆漂移说（魏格纳）\n'
    '  猜想：大陆原为整体，后漂移分开\n'
    '  证据：轮廓吻合、古生物化石、地质构造相似\n'
    '  缺陷：动力来源不明\n\n'
    '二、海底扩张说（赫斯等）\n'
    '  动力：地幔物质上涌，推动洋壳向两侧扩张\n'
    '  证据：洋底岩石年龄对称分布\n\n'
    '三、板块构造学说\n'
    '  六大板块：亚欧、非洲、印度洋、太平洋、美洲、南极洲\n'
    '  板块运动：碰撞（消亡边界）→ 山脉、海沟\n'
    '           张裂（生长边界）→ 裂谷、海洋\n'
    '  应用：珠峰长高（印度洋板块与亚欧板块碰撞）\n'
    '       地中海缩小（亚欧与非洲碰撞）\n'
    '       红海扩大（印度洋与非洲张裂）\n\n'
    '科学假说的发展：猜想 → 证据 → 完善',
    size=9.5)

# 合并后三列
merge_cells_horizontal(tp_table.cell(6, 1), tp_table.cell(6, 3))

set_col_widths(tp_table, [2.5, 5.5, 4.5, 6.5])

doc.add_paragraph()

# ============ 作业设计 ============
hw_title = doc.add_paragraph()
run = hw_title.add_run('五、作业设计（含AI赋能作业）')
run.font.name = '黑体'
run.font.size = Pt(14)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
hw_title.space_after = Pt(8)

hw_table = doc.add_table(rows=4, cols=3)
hw_table.style = 'Table Grid'
hw_table.alignment = WD_TABLE_ALIGNMENT.CENTER

add_cell_text(hw_table.cell(0, 0), '类型', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(hw_table.cell(0, 1), '作业内容', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(hw_table.cell(0, 2), 'AI赋能说明', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

hw_data = [
    ['基础作业',
     '完成作业本中相关练习题。',
     '利用AI测评系统在线提交，系统自动批改并生成个人错题本，推送针对性巩固练习。'],
    ['实践作业',
     '利用橡皮泥或泡沫板，模拟板块碰撞和张裂过程，拍摄30秒短视频并解释地形变化。',
     '学生上传视频至AI学习平台，AI自动识别视频中的板块运动类型并给出评分建议；学生也可使用AI视频编辑工具添加字幕和解说。'],
    ['拓展作业',
     '查阅资料，了解东非大裂谷的形成与板块运动的关系，写一段100字左右的说明。',
     '可使用AI搜索助手筛选可靠资料，利用AI写作助手辅助修改语言表达；AI自动生成东非大裂谷的板块运动示意图供参考。'],
]

for r, (typ, content, ai_desc) in enumerate(hw_data):
    add_cell_text(hw_table.cell(r+1, 0), typ, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(hw_table.cell(r+1, 1), content)
    add_cell_text(hw_table.cell(r+1, 2), ai_desc)

set_col_widths(hw_table, [2.5, 7.0, 9.5])

doc.add_paragraph()

# ============ 教学反思 ============
rf_title = doc.add_paragraph()
run = rf_title.add_run('六、教学反思')
run.font.name = '黑体'
run.font.size = Pt(14)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
rf_title.space_after = Pt(8)

rf_text = (
    '本课以"珠峰长高"为主线，串联三大假说，AI技术贯穿教学全过程，学生探究兴趣浓厚。'
    '成功之处：\n'
    '1.AI动画视频《魏格纳的发现》以沉浸式叙事还原科学发现过程，比传统讲授更具吸引力，学生注意力高度集中。\n'
    '2.AI三维地球模型支持学生自主旋转、缩放、叠加地震带和火山分布，空间认知效果显著优于平面地图。\n'
    '3.AI地质模拟平台的参数调节功能激发了学生的探索热情，通过"试错—观察—归纳"自主建构知识，体现了探究式学习的价值。\n'
    '4.AI即时投票和测评系统为教师提供了实时学情数据，使教学决策更加精准，实现了数据驱动的教学。\n'
    '5.拼图活动和撕纸类比有效帮助学生理解大陆漂移的直观证据；海底扩张模拟实验（土豆糊+泡沫板）生动形象，突破了动力机制难点。\n\n'
    '待改进之处：\n'
    '1.部分学生对AI平台操作不够熟练，首次使用时花费了较多时间，后续教学需提前进行简短培训。\n'
    '2.AI学习助手的问答内容需要进一步丰富，尤其是针对学生常见误解的引导性回复。\n'
    '3.大洋中脊岩石年龄图应使用标注更清晰的版本，便于学生读图分析。\n'
    '4.海底扩张说的逻辑呈现可以更递进：先展示年龄分布图，再播放模拟实验，最后关联大陆漂移，避免跳跃。\n'
    '5.在解释珠峰长高前，可先以大西洋为例（板块张裂形成海洋），让学生先理解张裂运动，再对比碰撞运动，效果更佳。\n'
    '6.学生模拟活动时，使用希沃投屏展示不同小组的碰撞效果，对比更直观。\n\n'
    'AI赋能反思：\n'
    '1.AI技术不是替代传统教学，而是增强和补充——模拟实验、动手操作等传统方法仍然不可替代。\n'
    '2.AI学习助手和测评系统有效支持了差异化教学，但需要教师精心设计提示词和反馈策略。\n'
    '3.后续教学可增加"板块运动与地震、火山分布"的AI模拟关联，进一步拓展知识的应用价值。'
)

rf_p = doc.add_paragraph()
run = rf_p.add_run(rf_text)
run.font.name = '仿宋'
run.font.size = Pt(10.5)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

# ============ 保存 ============
output_path = '/home/admin/.openclaw/workspace/AI赋能教学设计_地球板块的缓慢运动.docx'
doc.save(output_path)
print(f'✅ AI赋能教学设计已生成：{output_path}')
