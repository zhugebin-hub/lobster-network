#!/usr/bin/env python3
"""修复论文中的英文引号和遗漏字"""
import docx

SRC = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
DST = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"

doc = docx.Document(SRC)

# 收集所有ASCII引号位置，成对替换为中文引号
qpos = []
for pi, p in enumerate(doc.paragraphs):
    for ri, run in enumerate(p.runs):
        if run.text:
            for ci, ch in enumerate(run.text):
                if ch == '"':
                    qpos.append((pi, ri, ci))

for idx, (pi, ri, ci) in enumerate(qpos):
    run = doc.paragraphs[pi].runs[ri]
    chars = list(run.text)
    chars[ci] = '\u201c' if idx % 2 == 0 else '\u201d'
    run.text = ''.join(chars)

# 修复段落341遗漏的"教"字
for p in doc.paragraphs:
    t = p.text
    if '以慈善传"' in t and '以慈善传教"' not in t:
        # 替换
        for run in p.runs:
            if run.text and '以慈善传"' in run.text:
                run.text = run.text.replace('以慈善传"', '以慈善传教"')

# 参考文献77/78拆分
import copy
from docx.oxml.ns import qn
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if '[77]' in p.text and '[78]' in p.text:
        si = p.text.index('[78]')
        p77 = p.text[:si].rstrip()
        p78 = p.text[si:].strip()
        p.clear()
        p.add_run(p77)
        np = doc.add_paragraph()
        if p.style: np.style = p.style
        np.add_run(p78)
        p._p.addnext(np._p)
        break

doc.save(DST)

# 验证
v = docx.Document(DST)
aq = sum(1 for p in v.paragraphs for c in p.text if c == '"')
print(f"ASCII引号剩余: {aq}")
# 检查遗漏字修复
fixed = False
for p in v.paragraphs:
    if '以慈善传教"' in p.text:
        fixed = True
print(f"遗漏字修复: {'✅' if fixed else '❌'}")
