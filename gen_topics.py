#!/usr/bin/env python3
"""生成省级重点教改项目课题名称文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 标题 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('省级重点教改项目课题名称建议')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title.paragraph_format.space_after = Pt(12)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('基于强国建设背景的大学英语教学改革')
run.font.size = Pt(14)
run.font.italic = True
run.font.name = '楷体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
subtitle.paragraph_format.space_after = Pt(18)

# ── 辅助函数 ──
def add_heading1(text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(8)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_heading2(text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 80, 140)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_heading3(text):
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 100, 160)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74*2)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    return p

def add_bold_body(bold_text, normal_text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74*2)
    run = p.add_run(bold_text)
    run.font.bold = True
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = p.add_run(normal_text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table

# ═══════════════════════════════════════════════════
# 课题一
# ═══════════════════════════════════════════════════
add_heading1('课题一')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(10)
run = p.add_run('教育强国视域下大学英语"课程思政+国际传播能力"双轮驱动教学模式构建与实践')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

add_heading2('一、选题依据')
add_body('党的二十大报告明确提出"加强国际传播能力建设，全面提升国际传播效能"，教育强国战略对大学英语教学提出了新的使命要求。当前大学英语教学存在"重语言技能、轻价值引领"和"重西方文化输入、轻中国话语输出"两大痛点。本课题旨在构建"课程思政+国际传播能力"双轮驱动的教学模式，将价值塑造与能力培养有机融合，培养既能理解世界又能讲好中国故事的复合型人才。')

add_heading2('二、研究目标')
add_bullet('构建"课程思政+国际传播能力"双轮驱动的教学模式框架')
add_bullet('开发融入中国元素的大学英语系列教学资源')
add_bullet('建立双轮驱动模式下的多元评价体系')
add_bullet('在2-3所兄弟院校开展实践验证，形成可推广方案')

add_heading2('三、研究内容')
add_bullet('大学英语课程思政与国际化传播能力融合的理论框架研究')
add_bullet('"讲好中国故事"主题单元设计与教学实践')
add_bullet('基于真实场景的国际传播能力训练体系构建')
add_bullet('双轮驱动模式下的形成性评价工具开发')
add_bullet('教学效果实证研究与模式优化')

add_heading2('四、创新点')
add_bullet('视角创新：首次将"课程思政"与"国际传播能力"整合为统一的教学框架')
add_bullet('内容创新：开发"中国主题+语言技能"融合的模块化教学资源')
add_bullet('方法创新：采用"项目式学习+真实传播场景"的情境化教学方法')
add_bullet('评价创新：构建语言能力、价值认同、传播效能三维评价体系')

add_heading2('五、预期成果')
add_bullet('研究报告1份（含实证数据）')
add_bullet('教学案例集1部（20-30个课程思政+国际传播教学案例）')
add_bullet('发表教改论文2-3篇（核心期刊）')
add_bullet('形成可推广的教学模式方案1套')
add_bullet('建设在线开放课程1门')

add_heading2('六、操作性说明')
add_bullet('研究周期：2年（适合省级重点教改项目周期）')
add_bullet('实施范围：可从1-2个教学班试点，逐步推广到全学院')
add_bullet('团队构成：英语教师+思政教师+传播学教师跨学科合作')
add_bullet('经费预算：教学资源开发、调研差旅、论文发表、会议交流')
add_bullet('风险可控：有成熟理论支撑（课程思政理论、跨文化交际理论），实施路径清晰')

# ═══════════════════════════════════════════════════
# 课题二
# ═══════════════════════════════════════════════════
add_heading1('课题二')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(10)
run = p.add_run('面向教育强国建设的"EGP→EAP→ESP"三阶递进式大学英语课程体系重构与实践')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

add_heading2('一、选题依据')
add_body('《大学英语教学指南（2020版）》明确提出大学英语课程应包含通用英语、学术英语和专门用途英语三大模块。然而，当前多数高校仍停留在"通用英语一统天下"的阶段，教学内容与专业学习严重脱节，学生学完后无法有效阅读专业文献、撰写学术论文或进行专业交流。本课题针对这一痛点，构建"EGP→EAP→ESP"三阶递进式课程体系，实现大学英语从"通用型"向"定制化"的根本转变。')

add_heading2('二、研究目标')
add_bullet('构建符合本校学科特色的三阶递进式课程体系')
add_bullet('开发各阶段衔接的课程标准与教学大纲')
add_bullet('建设分学科方向的ESP教学资源库')
add_bullet('探索EGP教师向EAP/ESP教师转型的有效路径')
add_bullet('建立三阶课程的质量保障与评价机制')

add_heading2('三、研究内容')
add_bullet('基于需求分析的三阶课程体系设计（学生需求+专业需求+社会需求）')
add_bullet('EGP与EAP衔接机制研究：从通用语言技能到学术语言能力的过渡')
add_bullet('ESP课程群建设：覆盖本校主要学科门类（理工/人文/经管/医学等）')
add_bullet('模块化课程设计与学分配置方案')
add_bullet('教师转型培训机制与教学团队建设')
add_bullet('三阶课程学习成效追踪与评价体系')

add_heading2('四、创新点')
add_bullet('体系创新：构建完整的"基础→学术→专业"三阶递进体系，打通大学英语教学全链条')
add_bullet('衔接创新：设计EGP到EAP的平滑过渡机制，避免课程断层')
add_bullet('资源创新：建设分学科语料库和术语库，支撑ESP教学')
add_bullet('机制创新：建立"英语教师+专业教师"双师协同教学模式')

add_heading2('五、预期成果')
add_bullet('三阶递进式大学英语课程体系方案1套（含课程标准、教学大纲）')
add_bullet('ESP教材/讲义3-5本（覆盖主要学科方向）')
add_bullet('发表教改论文2-3篇')
add_bullet('教师转型培训方案1份')
add_bullet('学生学习成效追踪报告1份')

add_heading2('六、操作性说明')
add_bullet('研究周期：2-3年')
add_bullet('实施路径：第一年试点（EGP→EAP衔接）→第二年扩展（ESP课程群建设）→第三年完善（体系优化推广）')
add_bullet('团队构成：大学英语教师团队+各学院专业教师代表+教务处')
add_bullet('资源需求：教材开发经费、教师培训经费、语料库建设技术支持')
add_bullet('实施保障：可与学校专业培养方案修订同步推进，获得制度支持')

# ═══════════════════════════════════════════════════
# 课题三
# ═══════════════════════════════════════════════════
add_heading1('课题三')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(10)
run = p.add_run('AI赋能教育强国：大学英语"智能教学+精准评价"数字化转型模式研究')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

add_heading2('一、选题依据')
add_body('以大语言模型为代表的人工智能技术正在深刻改变教育生态。教育强国建设要求高等教育主动拥抱数字化转型。大学英语作为大班授课、个性化需求强烈的课程，是AI技术应用的天然场景。然而，当前AI在大学英语教学中的应用仍处于"辅助工具"层面，缺乏系统性的教学模式重构。本课题旨在构建AI深度融入大学英语教学的完整模式，实现从"技术辅助"到"教学重构"的跨越。')

add_heading2('二、研究目标')
add_bullet('构建AI赋能大学英语教学的完整模式框架')
add_bullet('开发基于大语言模型的智能教学应用场景')
add_bullet('建立基于学习分析的精准评价体系')
add_bullet('探索AI时代英语教师角色转型与能力提升路径')
add_bullet('制定AI辅助教学的伦理规范与使用指南')

add_heading2('三、研究内容')
add_bullet('AI技术在大学英语教学中的应用场景设计与开发')
add_bullet('基于大语言模型的个性化学习路径生成机制')
add_bullet('智能写作批改与口语训练系统的教学应用研究')
add_bullet('基于学习行为数据的精准评价与预警机制')
add_bullet('AI辅助教学与教师主导教学的协同模式')
add_bullet('AI教学应用的伦理风险识别与规范制定')
add_bullet('师生AI素养现状调查与提升策略')

add_heading2('四、创新点')
add_bullet('模式创新：构建"AI智能辅助+教师价值引领"的协同教学新模式')
add_bullet('技术整合创新：整合多种AI工具（大语言模型、语音识别、学习分析）形成教学闭环')
add_bullet('评价创新：基于多模态学习数据实现精准评价和个性化反馈')
add_bullet('规范创新：率先制定AI辅助外语教学的伦理规范和使用指南')

add_heading2('五、预期成果')
add_bullet('AI赋能大学英语教学模式方案1套')
add_bullet('智能教学应用场景案例集1部（15-20个应用场景）')
add_bullet('发表教改论文2-3篇（含1篇核心期刊）')
add_bullet('AI辅助教学伦理规范/使用指南1份')
add_bullet('师生AI素养调查报告1份')
add_bullet('建设AI辅助教学示范课程1-2门')

add_heading2('六、操作性说明')
add_bullet('研究周期：2年')
add_bullet('技术门槛：利用现有AI工具（ChatGPT、文心一言、讯飞等），无需自主开发AI模型')
add_bullet('实施范围：可从写作、口语等单一场景切入，逐步扩展到阅读、听力等')
add_bullet('团队构成：英语教师+教育技术教师+数据分析人员')
add_bullet('经费预算：AI工具订阅、教师培训、调研差旅、论文发表')
add_bullet('时效性强：AI+教育是当前最热方向，易获得评审关注和立项支持')

# ═══════════════════════════════════════════════════
# 对比表格
# ═══════════════════════════════════════════════════
add_heading1('四、课题对比与选择建议')

headers = ['维度', '课题一', '课题二', '课题三']
rows = [
    ['核心关键词', '课程思政+国际传播', 'EGP→EAP→ESP', 'AI赋能+数字化转型'],
    ['对应趋势', '工具性→战略性', '通用型→定制化', '技术辅助→深度融合'],
    ['理论成熟度', '高（课程思政理论成熟）', '高（ESP/EAP理论成熟）', '中（AI教育应用较新）'],
    ['实施难度', '中等', '较高（需课程体系重构）', '中等（利用现有AI工具）'],
    ['创新亮点', '价值引领+能力培养融合', '完整三阶课程体系', 'AI深度融入教学全链条'],
    ['评审关注度', '高（思政是热点）', '中高（课程体系改革）', '很高（AI是顶流热点）'],
    ['推广价值', '强（可跨校复制）', '较强（需结合本校学科）', '强（通用性强）'],
    ['适合团队', '英语+思政跨学科团队', '大学英语教师团队', '英语+教育技术团队'],
]
add_table(headers, rows)

add_heading1('五、选题建议')
add_bold_body('如果团队有思政教学基础：', '推荐课题一。思政是当前教育领域的核心热点，与"讲好中国故事"结合后既有政治高度又有学术深度，立项成功率高。')
add_bold_body('如果团队希望系统性改革课程体系：', '推荐课题二。这是大学英语教学的根本性问题，研究价值高，但需要较强的组织协调能力和较长的实施周期。')
add_bold_body('如果团队希望快速出成果、赶热点：', '推荐课题三。AI+教育是当前最热门的方向，技术门槛不高（利用现成AI工具），成果形式多样，容易获得关注。')

# ═══════════════════════════════════════════════════
# 页脚
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('— 文档结束 —')
run.font.size = Pt(10)
run.font.italic = True
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 保存 ──
output_path = '/home/admin/.openclaw/workspace/省级重点教改项目课题名称建议.docx'
doc.save(output_path)
print(f'✅ 文档已生成: {output_path}')
print(f'文件大小: {os.path.getsize(output_path)} 字节')
