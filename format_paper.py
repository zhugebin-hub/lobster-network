#!/usr/bin/env python3
"""
将利玛窦论文整理为标准学术论文格式（.docx）。
内容完全不变，使用 Word 原生脚注（每页底部，从1开始重新编号）。
方法：先生成基础文档，然后手动修改 XML 添加脚注。
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import zipfile, shutil, os, tempfile, re
from collections import Counter

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
w = lambda tag: '{%s}%s' % (W, tag)
fn_nsmap = {'w': W, 'r': R}

# ============================================================
# 1. 读入原始全文
# ============================================================
with open('/home/admin/.openclaw/media/inbound/b6aef016-21bc-4db5-bd2a-82e52e349ada', 'r', encoding='utf-8') as f:
    raw = f.read()

SEP = '脚  注'
idx = raw.index(SEP)
body_text = raw[:idx]
fn_text = raw[idx:]

# ============================================================
# 2. 解析脚注
# ============================================================
sup_digits = '⁰¹²³⁴⁵⁶⁷⁸⁹'
sup_map = {c: str(i) for i, c in enumerate(sup_digits)}

def sup_to_num(s):
    return int(''.join(sup_map.get(c, c) for c in s if c in sup_map))

footnote_entries = {}
for line in fn_text.split('\n'):
    line = line.strip()
    if not line or line.startswith('===') or line.startswith('脚'):
        continue
    if line[0] in sup_digits:
        chars = []
        for c in line:
            if c in sup_digits: chars.append(c)
            else: break
        num = sup_to_num(''.join(chars))
        footnote_entries[num] = line[len(''.join(chars)):].strip()

print(f"解析到 {len(footnote_entries)} 条脚注")

# ============================================================
# 3. 去重脚注
# ============================================================
seen = {}
dedup_map = {}
new_num = 0

for old_n in sorted(footnote_entries.keys()):
    content = footnote_entries[old_n]
    key = re.sub(r'第\d+页', '', content).strip()
    if key in seen:
        dedup_map[old_n] = seen[key]
    else:
        new_num += 1
        seen[key] = new_num
        dedup_map[old_n] = new_num

dedup_footnotes = {}
for old_n, new_n in dedup_map.items():
    dedup_footnotes[new_n] = footnote_entries[old_n]

counts = Counter(dedup_map.values())
merged = {n: c for n, c in counts.items() if c > 1}
print(f"去重后: {new_num} 条（原 {len(footnote_entries)} 条，合并 {len(merged)} 组）")

# ============================================================
# 4. 正文脚注引用替换
# ============================================================
def replace_refs(text):
    def repl(m):
        old_n = sup_to_num(m.group(1))
        return f"[{dedup_map[old_n]}]" if old_n in dedup_map else m.group(0)
    text = re.sub(r'([⁰-⁹]+)', repl, text)
    while re.search(r'\[(\d+)\]\[\d+\]', text):
        text = re.sub(r'\[(\d+)\]\[\d+\]', r'[\1]', text)
    return text

body_text = replace_refs(body_text)

# ============================================================
# 5. 按"第X节"拆分
# ============================================================
pattern = r'(第[一二三四五六七八九十]+节\s+[^\n]{1,80})'
parts = re.split(pattern, body_text)

abstract_raw = parts[0] if parts else ''
sections = []
for i in range(1, len(parts), 2):
    title = parts[i].strip()
    content = parts[i + 1] if i + 1 < len(parts) else ''
    sections.append((title, content))

print(f"拆分为 {len(sections)} 节")

# ============================================================
# 6. 生成 Word 文档（基础版，不含脚注）
# ============================================================
doc = Document()

for sec in doc.sections:
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

normal = doc.styles['Normal']
normal.font.name = '宋体'
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = Pt(22)

for lv in [1, 2, 3]:
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.size = [18, 16, 14][lv - 1]
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)

def fnt(run, size=12, bold=False, name='宋体'):
    run.font.size = Pt(size)
    run.font.name = name
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_para_with_markers(doc, text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
    """添加段落，[n] 用占位文本标记，后续替换为脚注引用"""
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    
    parts = re.split(r'(\[\d+\])', text)
    for part in parts:
        if not part:
            continue
        m = re.match(r'^\[(\d+)\]$', part)
        if m:
            fn_id = int(m.group(1))
            # 使用占位符文本，后面替换
            run = p.add_run(f'__FN{fn_id}__')
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
        else:
            run = p.add_run(part)
            fnt(run, size=size)
    return p

# 写入内容
# 标题页
for _ in range(8):
    doc.add_paragraph()

for text, size, bold, space in [
    ('利玛窦的译文实践', 26, True, 12),
    ('——"太极""理"的翻译与概念挪用', 18, False, 30),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space)
    run = p.add_run(text)
    fnt(run, size, bold, '黑体')

for text, size, space in [('戴建华', 14, 60), ('2026年6月', 12, 0)]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space)
    run = p.add_run(text)
    fnt(run, size, False, '宋体')

doc.add_page_break()

# 摘要
p = doc.add_heading('摘  要', level=1)
for r in p.runs: fnt(r, 16, True, '黑体')

if '关键词' in abstract_raw:
    abstract_only = abstract_raw[:abstract_raw.index('关键词')].strip()
    keywords_part = abstract_raw[abstract_raw.index('关键词'):].strip()
    add_para_with_markers(doc, abstract_only, indent=False)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(40)
    run = p.add_run(keywords_part)
    fnt(run, 12, True, '黑体')
else:
    add_para_with_markers(doc, abstract_raw.strip(), indent=False)

# 正文
for sec_title, sec_content in sections:
    p = doc.add_heading(sec_title, level=2)
    for r in p.runs: fnt(r, 16, True, '黑体')
    
    for para in re.split(r'\n{2,}', sec_content):
        para = para.strip()
        if para:
            add_para_with_markers(doc, para)

# 保存基础文档
base_path = '/tmp/base_doc.docx'
doc.save(base_path)
print(f"基础文档已保存: {base_path}")

# ============================================================
# 7. 解压、修改 XML、重新打包
# ============================================================
tmpdir = tempfile.mkdtemp()
with zipfile.ZipFile(base_path, 'r') as z:
    z.extractall(tmpdir)

# 7a. 创建 footnotes.xml
fn_root = etree.Element(w('footnotes'), nsmap=fn_nsmap)

# separator
sep = etree.SubElement(fn_root, w('footnote'))
sep.set('{%s}type' % W, 'separator')
sep.set('{%s}id' % W, '-1')
sep_p = etree.SubElement(sep, w('p'))
sep_r = etree.SubElement(sep_p, w('r'))
etree.SubElement(sep_r, w('separator'))

# continuationSeparator
cont = etree.SubElement(fn_root, w('footnote'))
cont.set('{%s}type' % W, 'continuationSeparator')
cont.set('{%s}id' % W, '0')
cont_p = etree.SubElement(cont, w('p'))
cont_r = etree.SubElement(cont_p, w('r'))
etree.SubElement(cont_r, w('continuationSeparator'))

# 添加所有脚注内容
for fn_id in sorted(dedup_footnotes.keys()):
    fn_text = dedup_footnotes[fn_id]
    fn = etree.SubElement(fn_root, w('footnote'))
    fn.set('{%s}id' % W, str(fn_id))
    
    p = etree.SubElement(fn, w('p'))
    pPr = etree.SubElement(p, w('pPr'))
    
    ind = etree.SubElement(pPr, w('ind'))
    ind.set('{%s}left' % W, '420')
    ind.set('{%s}hanging' % W, '420')
    
    sp = etree.SubElement(pPr, w('spacing'))
    sp.set('{%s}line' % W, '240')
    sp.set('{%s}lineRule' % W, 'auto')
    
    r = etree.SubElement(p, w('r'))
    rPr = etree.SubElement(r, w('rPr'))
    sz = etree.SubElement(rPr, w('sz'))
    sz.set('{%s}val' % W, '18')
    szCs = etree.SubElement(rPr, w('szCs'))
    szCs.set('{%s}val' % W, '18')
    rf = etree.SubElement(rPr, w('rFonts'))
    rf.set('{%s}ascii' % W, '宋体')
    rf.set('{%s}hAnsi' % W, '宋体')
    rf.set('{%s}eastAsia' % W, '宋体')
    
    t = etree.SubElement(r, w('t'))
    t.text = fn_text

fn_blob = etree.tostring(fn_root, encoding='utf-8', xml_declaration=True)
with open(os.path.join(tmpdir, 'word', 'footnotes.xml'), 'wb') as f:
    f.write(fn_blob)

# 7b. 修改 document.xml.rels 添加关系
rels_path = os.path.join(tmpdir, 'word', '_rels', 'document.xml.rels')
rels_tree = etree.parse(rels_path)
rels_root = rels_tree.getroot()

max_rid = 0
for rel in rels_root:
    rid = rel.get('Id', '')
    if rid.startswith('rId'):
        num = int(rid[3:])
        if num > max_rid:
            max_rid = num

new_rid = f'rId{max_rid + 1}'
new_rel = etree.SubElement(rels_root, '{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
new_rel.set('Id', new_rid)
new_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')
new_rel.set('Target', 'footnotes.xml')

with open(rels_path, 'wb') as f:
    f.write(etree.tostring(rels_tree, xml_declaration=True, encoding='UTF-8'))

# 7c. 修改 document.xml，将占位符替换为脚注引用
doc_path = os.path.join(tmpdir, 'word', 'document.xml')
doc_tree = etree.parse(doc_path)
doc_root = doc_tree.getroot()

body = doc_root.find(w('body'))
if body is not None:
    for p in body.findall(w('p')):
        for r in p.findall(w('r')):
            for t in r.findall(w('t')):
                if t.text and t.text.startswith('__FN') and t.text.endswith('__'):
                    try:
                        fn_id = int(t.text[4:-2])
                        # 清除占位符文本
                        t.text = ''
                        # 添加上标格式
                        rPr = r.get_or_add_rPr()
                        va = etree.SubElement(rPr, w('vertAlign'))
                        va.set('{%s}val' % W, 'superscript')
                        sz = etree.SubElement(rPr, w('sz'))
                        sz.set('{%s}val' % W, '16')
                        # 添加脚注引用
                        fn_ref = etree.SubElement(r, w('footnoteReference'))
                        fn_ref.set('{%s}id' % W, str(fn_id))
                    except:
                        pass

with open(doc_path, 'wb') as f:
    f.write(etree.tostring(doc_tree, xml_declaration=True, encoding='UTF-8'))

# 7d. 修改 settings.xml 设置每页重新编号
settings_path = os.path.join(tmpdir, 'word', 'settings.xml')
settings_tree = etree.parse(settings_path)
settings_root = settings_tree.getroot()

footnotePr = settings_root.find('.//w:footnotePr', fn_nsmap)
if footnotePr is None:
    footnotePr = etree.SubElement(settings_root, w('footnotePr'))

nr = etree.SubElement(footnotePr, w('numRestart'))
nr.set('{%s}val' % W, 'eachPage')

nf = etree.SubElement(footnotePr, w('numFmt'))
nf.set('{%s}val' % W, 'decimal')

sp = etree.SubElement(footnotePr, w('start'))
sp.set('{%s}val' % W, '1')

with open(settings_path, 'wb') as f:
    f.write(etree.tostring(settings_tree, xml_declaration=True, encoding='UTF-8'))

# 7e. 修改 [Content_Types].xml 添加 footnotes.xml 类型
ct_path = os.path.join(tmpdir, '[Content_Types].xml')
ct_tree = etree.parse(ct_path)
ct_root = ct_tree.getroot()

ct_ns = 'http://schemas.openxmlformats.org/package/2006/content-types'
override = etree.SubElement(ct_root, '{%s}Override' % ct_ns)
override.set('PartName', '/word/footnotes.xml')
override.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml')

with open(ct_path, 'wb') as f:
    f.write(etree.tostring(ct_tree, xml_declaration=True, encoding='UTF-8'))

# 7f. 重新打包
output = '/home/admin/.openclaw/workspace/利玛窦译文实践_标准论文格式.docx'
if os.path.exists(output):
    os.remove(output)

shutil.make_archive('/tmp/final_doc', 'zip', tmpdir)
os.rename('/tmp/final_doc.zip', output)

print(f"\n✅ 已保存: {output}")
print(f"脚注: {len(dedup_footnotes)} 条（每页底部显示，从1开始重新编号）")

# 清理
shutil.rmtree(tmpdir)
if os.path.exists(base_path):
    os.remove(base_path)
