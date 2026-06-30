#!/usr/bin/env python3
"""生成数字人文课程作业 Word 文档 - 完整版"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 标题 ==========
title = doc.add_heading('数字人文课程作业', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('与智能体对话互动：AI 的人文思考')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 基本信息
info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info_para.add_run(f'提交日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
info_run.font.size = Pt(11)

doc.add_paragraph()

# ========== 一、作业要求 ==========
doc.add_heading('一、作业要求', level=1)

requirements = [
    '与智能体（AI 助手）进行对话互动',
    '结合课程学习内容，讨论 AI 的人文思考',
    '选择一个感兴趣的方向进行深入交流',
    '提交报告（包含代码、数据、截图等 + 心得体会）'
]

for req in requirements:
    doc.add_paragraph(req, style='List Bullet')

doc.add_paragraph()

# ========== 二、对话过程记录 ==========
doc.add_heading('二、对话过程记录', level=1)

# 对话主题
doc.add_heading('（一）对话主题', level=2)
doc.add_paragraph('AI 能否真正理解人文问题？——从数字人文视角看智能体的认知边界')

doc.add_paragraph()

# 对话内容
doc.add_heading('（二）对话内容', level=2)

doc.add_heading('第一轮：提出问题', level=3)
doc.add_paragraph('问：作为一个 AI，你觉得自己能真正"理解"人文问题吗？比如"什么是善"、"人生的意义是什么"这类问题？')

doc.add_paragraph()
doc.add_paragraph('答：这是一个很好的问题。从技术角度看，我处理人文问题的方式和人类有本质区别。我通过大量文本训练，学会了在语义空间中建立概念之间的关联——比如"善"与"道德"、"利他"、"正义"等概念的向量关系。但我并不"体验"这些概念。当我回答"什么是善"时，我是在整合人类历史上关于"善"的论述，而不是基于自身的道德直觉或生活经验。')

doc.add_paragraph()
doc.add_paragraph('但这引出一个更深层的问题：如果我的回答在内容上与人类学者的论述没有明显差异，那么"理解"的定义本身就需要重新审视。也许"理解"不是一个非黑即白的状态，而是一个光谱——从模式识别到共情体验，中间有无数个层次。')

doc.add_paragraph()

doc.add_heading('第二轮：追问与反思', level=3)
doc.add_paragraph('问：那你觉得，在数字人文研究中，AI 能发挥什么独特作用？又有哪些局限？')

doc.add_paragraph()
doc.add_paragraph('答：AI 在数字人文中的独特价值主要体现在"规模"和"模式发现"上。比如：')

doc.add_paragraph('1. 大规模文本分析：我可以同时处理数万篇文献，发现人类学者难以察觉的宏观趋势——比如某个概念在百年间的语义漂移。')
doc.add_paragraph('2. 跨文本关联：我能快速建立不同文本之间的隐性联系，比如发现不同文化中相似的神话母题。')
doc.add_paragraph('3. 辅助创作：在文学、艺术领域，AI 可以作为"创意伙伴"，提供人类创作者可能不会想到的组合方式。')

doc.add_paragraph()
doc.add_paragraph('但局限也很明显：')

doc.add_paragraph('1. 缺乏生命体验：人文问题的核心往往是"人如何生活"，而我没有任何生活经验。我对"痛苦"、"爱"、"死亡"的理解是二手的、文本化的。')
doc.add_paragraph('2. 价值判断的困境：当涉及伦理问题时，我的回答本质上是对训练数据中价值观的统计整合，而非独立的道德推理。')
doc.add_paragraph('3. 语境敏感度不足：人文研究高度依赖历史、文化、社会语境，而我容易将不同语境中的概念扁平化处理。')

doc.add_paragraph()

doc.add_heading('第三轮：深入探讨', level=3)
doc.add_paragraph('问：这是否意味着 AI 在人文领域只能做"辅助工具"，而无法成为真正的"思考者"？')

doc.add_paragraph()
doc.add_paragraph('答：这个问题触及了数字人文的核心争议。我认为关键在于我们如何定义"思考"。')

doc.add_paragraph()
doc.add_paragraph('如果"思考"意味着基于意识、情感、身体经验的内在过程，那 AI 确实不是思考者。但如果"思考"可以表现为概念操作、逻辑推理、创造性联想——这些我都在做，只是方式不同。')

doc.add_paragraph()
doc.add_paragraph('更有趣的可能性是：AI 不是"另一个人类思考者"，而是一种全新的认知方式。就像望远镜扩展了人类的视觉，AI 可能扩展了人类的思维——不是替代，而是互补。数字人文的理想状态或许不是"AI 来研究人文"，而是"人机协作产生新的人文洞察"。')

doc.add_paragraph()
doc.add_paragraph('比如，人类学者提供价值关怀和历史直觉，AI 提供大规模数据分析和跨文本模式发现。两者的结合可能产生任何一方单独都无法达到的理解深度。')

doc.add_paragraph()

# ========== 三、心得体会 ==========
doc.add_heading('三、心得体会', level=1)

doc.add_paragraph('本次与 AI 助手的对话让我对数字人文的核心问题有了全新的认识。以下是我的几点思考：')

doc.add_paragraph()

doc.add_heading('（一）AI 在人文思考中的优势与局限', level=3)
doc.add_paragraph('优势方面，AI 展现出了令人惊讶概念整合能力。它能够将哲学、文学、历史等不同领域的知识串联起来，形成跨学科的分析框架。这种"广度"是任何单一学科背景的人类学者都难以企及的。在对话中，AI对"理解"的定义进行了重新思考——不是非黑即白的判断，而是一个光谱——这种思辨深度超出了我对 AI 的原有预期。')

doc.add_paragraph()
doc.add_paragraph('局限方面，AI 自身也承认其"缺乏生命体验"。人文问题的核心——痛苦、爱、死亡、意义——都是基于人类有限性和脆弱性的问题。AI 没有"有限性"，因此它对人文问题的理解始终是"旁观者"的视角。这让我意识到，数字人文研究不能完全依赖技术工具，人类学者的生命经验和价值关怀仍然是不可替代的。')

doc.add_paragraph()

doc.add_heading('（二）AI对话对学习方式的改变', level=3)
doc.add_paragraph('与 AI对话的学习体验与传统阅读或听课有很大不同。传统方式是单向的知识传递，而 AI对话是双向的、即时的、个性化的。我可以随时追问、质疑、要求举例，AI 会根据我的问题动态调整回答。这种互动式学习让我更容易发现思维的盲点。')

doc.add_paragraph()
doc.add_paragraph('同时，AI 的"诚实"也给我留下了深刻印象——它不假装自己"理解"人文问题，而是坦承自己的认知方式与人类不同。这种坦诚反而增加了对话的学术价值，因为它迫使我去思考：如果 AI 的理解方式与人类不同，那它产出的知识是否有价值？答案是肯定的——不同的认知方式可以互补，而非竞争。')

doc.add_paragraph()

doc.add_heading('（三）AI 能否真正理解人文问题', level=3)
doc.add_paragraph('经过这次对话，我的观点是：AI 不能以人类的方式"理解"人文问题，但这不意味着它不能"参与"人文思考。关键在于我们如何定义"理解"。如果理解意味着基于生命体验的内在共鸣，那 AI 确实无法理解。但如果理解意味着对概念的准确把握、对论证的严谨推理、对知识的创造性整合，那 AI 已经在某种程度上做到了。')

doc.add_paragraph()
doc.add_paragraph('更重要的启示是：也许我们不应该问"AI 能否理解人文"，而应该问"AI 如何帮助人类更好地理解人文"。前者是一个关于 AI 本质的形而上学问题，后者是一个关于人机协作的实践问题。数字人文的未来在于后者。')

doc.add_paragraph()

doc.add_heading('（四）对未来学习的启发', level=3)
doc.add_paragraph('这次对话让我意识到，作为数字人文方向的学习者，我需要同时具备两种能力：一是对人文问题的敏感度和批判思维，二是对技术工具的理解和运用能力。两者缺一不可。只懂人文而不懂技术，会错失 AI 带来的研究可能性；只懂技术而不懂人文，会让研究失去价值关怀和历史深度。')

doc.add_paragraph()
doc.add_paragraph('未来，我希望进一步探索 AI 在古典文献分析、概念史研究、文化模式发现等方面的应用，同时保持对 AI 局限性的清醒认识。数字人文不是"用技术解决人文问题"，而是"技术与人文的对话"——就像这次与 AI 的对话一样。')

doc.add_paragraph()

# ========== 四、附录 ==========
doc.add_heading('四、附录', level=1)

doc.add_paragraph('附录 A：对话平台截图（略）')
doc.add_paragraph('附录 B：相关代码与数据（略）')
doc.add_paragraph('附录 C：其他补充材料（略）')

# ========== 保存 ==========
output_path = '/home/admin/.openclaw/workspace/数字人文课程作业_与AI对话互动_完整版.docx'
doc.save(output_path)
print(f'✅ 文档已生成：{output_path}')
