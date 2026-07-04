#!/usr/bin/env python3
"""
叶畏兵毕业论文完善脚本 v4
修复：英文引号→中文引号、参考文献77/78合并、标点规范
"""
import docx
import re
import copy

INPUT = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
OUTPUT = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"

doc = docx.Document(INPUT)

# ==========================================
# 1. 分析
# ==========================================
print("=== 分析原文档 ===")

quote_count = 0
for p in doc.paragraphs:
    for ch in p.text:
        if ch in ('"', '"', '"', '"'):
            quote_count += 1
print(f"英文引号字符: {quote_count}")

# 找77/78合并
for i, p in enumerate(doc.paragraphs):
    if '[77]' in p.text and '[78]' in p.text:
        print(f"段落{i}存在77/78合并")

# ==========================================
# 2. 修复英文引号 → 中文引号
# ==========================================
print("\n=== 修复引号 ===")

# 收集所有需要修复的段落
fixes = []
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if not run.text:
            continue
        t = run.text
        if '"' in t or '"' in t or '"' in t or '"' in t:
            fixes.append(i)
            break

print(f"需要修复的段落: {len(fixes)}")

# 修复函数：智能替换引号
def fix_quotes(text):
    result = []
    i = 0
    # 先替换开引号
    while i < len(text):
        if text[i] == '"':
            result.append('"')
        elif text[i] == '"':
            result.append('"')
        else:
            result.append(text[i])
        i += 1
    return ''.join(result)

applied = 0
for p in doc.paragraphs:
    for run in p.runs:
        if run.text:
            old = run.text
            new = fix_quotes(old)
            if new != old:
                run.text = new
                applied += 1

print(f"已修复: {applied} 处")

# ==========================================
# 3. 修复77/78合并
# ==========================================
print("\n=== 修复参考文献合并 ===")

for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if '[77]' in p.text and '[78]' in p.text:
        split_pos = p.text.index('[78]')
        part77 = p.text[:split_pos].rstrip()
        part78 = p.text[split_pos:].strip()
        
        p.text = part77
        print(f"77: {part77[:70]}...")
        
        # 创建新段落
        new_p = doc.add_paragraph()
        # 复制格式
        if p.style:
            new_p.style = p.style
        # 设置内容
        if new_p.runs:
            new_p.runs[0].text = part78
        else:
            run = new_p.add_run(part78)
        
        # 移动新段落到正确位置（77之后）
        p._p.addnext(new_p._p)
        print(f"78: {part78[:70]}...")
        print("已拆分完成")
        break

# ==========================================
# 4. 保存
# ==========================================
doc.save(OUTPUT)
print(f"\n✅ 保存成功: {OUTPUT}")
