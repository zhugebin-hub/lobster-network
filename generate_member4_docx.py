#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成成员 4 的 Word 文档 - 符合任务要求版本"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_bullet(doc, text, size=11, bold=False):
    """添加项目符号"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font_size = Pt(size)
    r.bold = bold
    return p

def create_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    # 标题
    title = doc.add_heading('成员 4：房价下跌如何触发危机 + 真实案例', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 一、数据图说明
    doc.add_heading('一、数据图：房价与违约率双轴图', level=2)
    
    p = doc.add_paragraph()
    r = p.add_run('【图表说明】')
    r.bold = True
    
    doc.add_paragraph('上图展示了 2000-2008 年美国房价指数（左轴，2000 年=100）与次级贷款违约率（右轴，%）的变化趋势。')
    
    doc.add_heading('关键观察：', level=3)
    add_bullet(doc, '• 2000-2006 年：房价持续上涨，从 100 升至 189（+89%），违约率维持在低位（1.5%-5.5%）\n')
    add_bullet(doc, '• 2006 年：房价达到峰值 189，违约率开始上升至 5.5%\n')
    add_bullet(doc, '• 2007-2008 年：房价暴跌至 139（-26%），违约率飙升至 15.8%\n')
    add_bullet(doc, '• 剪刀差效应：房价下跌与违约率上升形成"剪刀差"，标志危机全面爆发', bold=True)
    
    # 二、真实案例：雷曼兄弟破产
    doc.add_heading('二、真实案例：雷曼兄弟破产（2008 年 9 月）', level=2)
    
    doc.add_heading('1. 雷曼兄弟公司简介', level=3)
    add_bullet(doc, '• 成立时间：1850 年，总部位于纽约\n')
    add_bullet(doc, '• 地位：美国第四大投资银行，全球性金融机构\n')
    add_bullet(doc, '• 业务：证券承销、交易、资产管理、固定收益产品\n')
    add_bullet(doc, '• 危机前资产规模：约 6390 亿美元')
    
    doc.add_heading('2. 雷曼为何会倒闭？', level=3)
    p = doc.add_paragraph()
    r = p.add_run('核心原因：')
    r.bold = True
    r = p.add_run('过度投资次级房贷相关证券（MBS、CDO）')
    
    add_bullet(doc, '• 2003-2007 年：雷曼大量收购房贷机构，扩大次级贷业务\n')
    add_bullet(doc, '• 持有大量 MBS 和 CDO 资产，总值超过 850 亿美元\n')
    add_bullet(doc, '• 杠杆率高达 31:1（每 1 美元自有资本对应 31 美元债务）\n')
    add_bullet(doc, '• 房价下跌后，MBS/CDO 价值暴跌，资产严重缩水')
    
    doc.add_heading('3. 破产传导路径', level=3)
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['阶段', '事件', '影响']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data = [
        ['① 房价下跌', '2006 年起美国房价开始下跌', '次级贷款违约率上升'],
        ['② MBS 贬值', '雷曼持有的 MBS/CDO 价值暴跌', '资产缩水 850 亿美元'],
        ['③ 流动性危机', '交易对手要求追加保证金', '现金短缺'],
        ['④ 寻求救助', '雷曼寻求买家或政府救助', '韩国产业银行退出谈判'],
        ['⑤ 破产申请', '2008 年 9 月 15 日申请破产保护', '史上最大破产案'],
        ['⑥ 全球冲击', '股市暴跌、信贷冻结', '金融危机全面爆发']
    ]
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_heading('4. 传导路径图示', level=3)
    p = doc.add_paragraph()
    r = p.add_run('房价下跌 → 次级贷违约 → MBS/CDO 贬值 → 雷曼资产缩水 → 流动性危机 → 破产 → 全球金融危机')
    r.font_size = Pt(12)
    r.bold = True
    
    doc.add_heading('5. 关键数据', level=3)
    add_bullet(doc, '• 破产日期：2008 年 9 月 15 日（星期一）\n')
    add_bullet(doc, '• 破产资产：6130 亿美元（美国历史上最大破产案）\n')
    add_bullet(doc, '• 裁员人数：约 2.5 万人失业\n')
    add_bullet(doc, '• 股市反应：道指当日下跌 4.4%，一周内下跌 10%\n')
    add_bullet(doc, '• 全球影响：多国股市暴跌，信贷市场冻结')
    
    # 三、替代案例：AIG 被接管
    doc.add_heading('三、替代案例：AIG 被政府接管（2008 年 9 月）', level=2)
    
    doc.add_heading('1. AIG 公司简介', level=3)
    add_bullet(doc, '• 全称：美国国际集团 (American International Group)\n')
    add_bullet(doc, '• 地位：全球最大保险公司之一\n')
    add_bullet(doc, '• 业务：保险、金融服务、资产管理')
    
    doc.add_heading('2. AIG 为何陷入危机？', level=3)
    p = doc.add_paragraph()
    r = p.add_run('核心原因：')
    r.bold = True
    r = p.add_run('大量出售 CDS（信用违约互换），为 MBS/CDO 提供"保险"')
    
    add_bullet(doc, '• AIG 金融产品部门出售了约 4400 亿美元的 CDS\n')
    add_bullet(doc, '• 假设房价不会下跌，CDS 不会赔付\n')
    add_bullet(doc, '• 房价下跌后，MBS/CDO 违约，AIG 需巨额赔付\n')
    add_bullet(doc, '• 2008 年 9 月：AIG 无法赔付，面临破产')
    
    doc.add_heading('3. AIG 传导路径', level=3)
    
    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Table Grid'
    headers2 = ['阶段', '事件', '影响']
    for i, header in enumerate(headers2):
        table2.rows[0].cells[i].text = header
        table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data2 = [
        ['① 出售 CDS', 'AIG 出售 4400 亿美元 CDS', '收取保费，承担风险'],
        ['② 房价下跌', 'MBS/CDO 大规模违约', 'CDS 触发赔付条件'],
        ['③ 赔付危机', 'AIG 需赔付数百亿美元', '现金不足'],
        ['④ 政府救助', '美联储提供 850 亿美元贷款', '政府接管 79.9% 股权'],
        ['⑤ 系统性风险', 'AIG 倒闭将引发连锁反应', '政府被迫救助']
    ]
    for i, row_data in enumerate(data2, 1):
        for j, cell_data in enumerate(row_data):
            table2.rows[i].cells[j].text = cell_data
    
    doc.add_heading('4. AIG 案例的启示', level=3)
    add_bullet(doc, '• CDS 等衍生品将风险从银行转移到保险公司\n')
    add_bullet(doc, '• 但风险并未消失，只是集中到 AIG 这样的系统性机构\n')
    add_bullet(doc, '• "太大而不能倒"：AIG 倒闭会引发全球金融系统崩溃\n')
    add_bullet(doc, '• 政府最终用纳税人钱救助，引发道德风险争议')
    
    # 四、总结
    doc.add_heading('四、总结：房价下跌如何触发危机', level=2)
    
    p = doc.add_paragraph()
    r = p.add_run('传导机制核心逻辑：')
    r.font_size = Pt(12)
    r.bold = True
    
    add_bullet(doc, '1. 房价下跌 → 次级贷款者无法 refinancing 或卖房还款\n')
    add_bullet(doc, '2. 房贷违约率飙升 → MBS/CDO 现金流中断\n')
    add_bullet(doc, '3. MBS/CDO 价值暴跌 → 持有这些资产的金融机构巨额亏损\n')
    add_bullet(doc, '4. 金融机构流动性危机 → 银行间停止拆借，信贷冻结\n')
    add_bullet(doc, '5. 雷曼破产/AIG 被接管 → 市场信心崩溃，全球金融危机爆发\n')
    add_bullet(doc, '6. 信贷紧缩 → 企业无法融资 → 裁员 → 消费下降 → 经济衰退')
    
    doc.add_heading('关键教训：', level=3)
    add_bullet(doc, '• 房价不会永远上涨，资产泡沫终将破裂\n')
    add_bullet(doc, '• 金融衍生品可能放大而非分散风险\n')
    add_bullet(doc, '• 高杠杆在顺周期时放大收益，在逆周期时放大损失\n')
    add_bullet(doc, '• 系统性风险需要监管干预，"太大而不能倒"是真实存在的')
    
    # 参考资料
    doc.add_heading('参考资料', level=2)
    add_bullet(doc, '• 美联储经济数据 (FRED): Case-Shiller 房价指数\n')
    add_bullet(doc, '• 美国财政部：《金融危机调查报告》(2011)\n')
    add_bullet(doc, '• 《大空头》(The Big Short) - Michael Lewis\n')
    add_bullet(doc, '• 《监守自盗》(Inside Job) 纪录片\n')
    add_bullet(doc, '• 维基百科：Lehman Brothers bankruptcy, AIG bailout')
    
    # 保存文档
    doc.save('/home/admin/.openclaw/workspace/成员 4-房价下跌触发危机与真实案例.docx')
    print('Word 文档已生成：成员 4-房价下跌触发危机与真实案例.docx')

if __name__ == '__main__':
    create_document()
