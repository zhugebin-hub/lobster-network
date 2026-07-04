#!/usr/bin/env python3
"""
叶畏兵毕业论文 - 最终完善版
修复内容：
1. ASCII引号 → 中文引号
2. 参考文献77/78拆分
3. 添加46处脚注上标
4. 其他格式规范
"""
import docx
import re
from copy import deepcopy

INPUT = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
OUTPUT = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"

doc = docx.Document(INPUT)

# ==========================================
# 修复1：ASCII引号 → 中文引号
# ==========================================
print("=== 修复1：引号 ===")
all_quotes = []
for p_idx, para in enumerate(doc.paragraphs):
    for r_idx, run in enumerate(para.runs):
        if not run.text:
            continue
        for c_idx, ch in enumerate(run.text):
            if ch == '"':
                all_quotes.append((p_idx, r_idx, c_idx))

for idx, (p_idx, r_idx, c_idx) in enumerate(all_quotes):
    run = doc.paragraphs[p_idx].runs[r_idx]
    chars = list(run.text)
    chars[c_idx] = '\u201c' if idx % 2 == 0 else '\u201d'
    run.text = ''.join(chars)
print(f"修复 {len(all_quotes)} 处引号")

# ==========================================
# 修复2：参考文献77/78拆分
# ==========================================
print("\n=== 修复2：参考文献77/78 ===")
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if '[77]' in p.text and '[78]' in p.text:
        split_idx = p.text.index('[78]')
        part77 = p.text[:split_idx].rstrip()
        part78 = p.text[split_idx:].strip()
        p.clear()
        p.add_run(part77)
        new_p = doc.add_paragraph()
        if p.style:
            new_p.style = p.style
        new_p.add_run(part78)
        p._p.addnext(new_p._p)
        print("已拆分")
        break

# ==========================================
# 修复3：添加脚注上标
# ==========================================
# 根据上下文匹配引用→参考文献
# 格式：在标点符号后插入上标数字

print("\n=== 修复3：添加脚注上标 ===")

# 定义引用映射：(段落索引, 上下文关键词) → 参考文献编号
citation_map = {
    # 第一章 引言
    124: [59],        # 2021年全国宗教工作会议 → [59]
    126: [62],        # 王伟光 → [56], 游斌 → [62]  (段落开头王伟光，后面游斌)
    127: [63, 54, 51], # 张志刚[63]、牟钟鉴[54]、何虎生[51]
    138: [16, 4, 17],  # Nygren[16], Aquinas[4], Pope Benedict[17]
    139: [12],         # 杜维明 + Ames + Küng
    140: [70, 52],     # 朱全红[70], 杨宝安, 黄勇[52]
    142: [10],         # Casanova[10], Bellah[6]
    143: [23, 25, 15], # Troeltsch[23], Weber[25], Niebuhr[15]
    144: [8, 19],      # Bevans[8], Schreiter[19]
    146: [63, 36, 40], # 张志刚[63], 牟钟鉴[36], 王作安[40]
    147: [30, 37, 41, 42], # 段琦[30], 唐晓峰[37], 徐以骅[41], 游斌[42]
    148: [29, 33, 45], # 陈宗皋[29], 李平晔[33], 赵敦华[45]
    150: [47, 66, 50], # 曹正勇[47], 赵卫红[66], 韩雪莉[50]
    151: [55, 53, 48], # 王思琪[55], 吕臻雨[53], 陈金羽[48]
    # 第二章
    184: [16],         # Nygren区分
    185: [4, 17],      # Aquinas, Pope Benedict
    189: [2],          # 刘耘华
    190: [70],         # 朱全红
    191: [30],         # 段琦
    192: [53],         # 吕臻雨
    196: [60],         # 杨宝安
    215: [2],          # 刘耘华
    220: [8],          # Bevans
    221: [19],         # Schreiter
    223: [7],          # Berger
    # 第三章
    229: [55],         # 王思琪
    230: [52],         # 黄勇
    245: [53],         # 吕臻雨
    246: [11],         # Hauerwas
    254: [10],         # Casanova
    265: [27],         # 陈来
    269: [52],         # 黄勇
    275: [61],         # 杨凤岗
    283: [48],         # 陈金羽
    # 第四章
    # 第五章
    311: [],           # 教义阐释
    320: [67],         # 浙江神学院
    321: [67],         # 研讨会
    322: [64],         # 张忠成
    324: [35, 46],     # 莫幸福
    340: [79],         # 光盐基金会
    351: [],           # 习近平讲话
}

# 更精确的映射：(段落索引, 该段内第几个引用位置) → ref编号
precise_map = {
    # 第一段引言 - 政策引用
    (124, 0): 59,       # 2021宗教工作会议
    (126, 0): 56,       # 王伟光-五个认同
    (126, 1): 62,       # 游斌
    (127, 0): 63,       # 张志刚
    (127, 1): 54,       # 牟钟鉴
    (127, 2): 51,       # 何虎生
    (138, 0): 16,       # Nygren
    (138, 1): 4,        # Aquinas  
    (139, 0): 1,        # 杜维明
    (139, 1): 12,       # Küng & Ching
    (140, 0): 70,       # 朱全红
    (140, 1): 52,       # 黄勇
    (142, 0): 10,       # Casanova
    (142, 1): 6,        # Bellah
    (143, 0): 23,       # Troeltsch
    (143, 1): 25,       # Weber
    (144, 0): 8,        # Bevans
    (146, 0): 63,       # 张志刚
    (146, 1): 36,       # 牟钟鉴
    (146, 2): 40,       # 王作安
    (147, 0): 30,       # 段琦
    (147, 1): 37,       # 唐晓峰
    (147, 2): 41,       # 徐以骅
    (147, 3): 42,       # 游斌
    (148, 0): 29,       # 陈宗皋
    (148, 1): 33,       # 李平晔
    (148, 2): 45,       # 赵敦华
    (150, 0): 47,       # 曹正勇
    (150, 1): 66,       # 赵卫红
    (150, 2): 50,       # 韩雪莉
    (151, 0): 55,       # 王思琪
    (151, 1): 53,       # 吕臻雨
    (184, 0): 16,       # Nygren
    (185, 0): 4,        # Aquinas
    (189, 0): 2,        # 刘耘华
    (190, 0): 70,       # 朱全红
    (191, 0): 65,       # 赵天恩
    (192, 0): 30,       # 段琦
    (196, 0): 60,       # 杨宝安
    (215, 0): 2,        # 刘耘华
    (220, 0): 8,        # Bevans
    (221, 0): 19,       # Schreiter
    (223, 0): 7,        # Berger
    (229, 0): 55,       # 王思琪
    (230, 0): 52,       # 黄勇
    (245, 0): 53,       # 吕臻雨
    (246, 0): 11,       # Hauerwas
    (254, 0): 10,       # Casanova
    (265, 0): 27,       # 陈来
    (269, 0): 52,       # 黄勇
    (275, 0): 61,       # 杨凤岗
    (283, 0): 48,       # 陈金羽
    (311, 0): None,     # 教义阐释（政策文件，无需脚注）
    (320, 0): 67,       # 浙江神学院研讨会
    (321, 0): 67,       # 潘兴旺
    (322, 0): 64,       # 张忠成
    (324, 0): 46,       # 莫幸福
    (324, 1): 35,       # 莫幸福浙江宗教史
    (340, 0): 79,       # 光盐基金会
    (351, 0): None,     # 习近平讲话
}

# 找所有引用位置
citation_positions = []
for p_idx, para in enumerate(doc.paragraphs):
    t = para.text
    style = para.style.name if para.style else 'Normal'
    if 'Heading' in style or 'toc' in style or 'Caption' in style:
        continue
    if not re.search(r'[\u4e00-\u9fff]{20}', t):
        continue
    
    # 找"。 "模式
    count = 0
    for m in re.finditer(r'([。》"）]) (?=[\u4e00-\u9fffA-Z])', t):
        key = (p_idx, count)
        ref_num = precise_map.get(key)
        if ref_num is not None:
            citation_positions.append({
                'para': p_idx,
                'pos': m.start(),
                'punct': m.group(1),
                'ref': ref_num,
                'count': count
            })
        count += 1

print(f"找到 {len(citation_positions)} 个需要添加脚注的位置")

# 插入脚注上标
# 策略：在标点符号后插入上标数字
# 需要修改段落的XML来插入上标
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_superscript_after_punct(para, punct_pos, number):
    """在段落的指定位置后插入上标数字"""
    # 找到包含该位置的run
    current_pos = 0
    target_run = None
    target_idx = 0
    insert_pos = 0
    
    for r_idx, run in enumerate(para.runs):
        run_len = len(run.text)
        if current_pos + run_len > punct_pos:
            target_run = run
            target_idx = r_idx
            insert_pos = punct_pos - current_pos
            break
        current_pos += run_len
    
    if target_run is None:
        return False
    
    # 在run文本中插入上标
    # 使用XML方式添加上标
    text_before = target_run.text[:insert_pos + 1]  # 包括标点
    text_after = target_run.text[insert_pos + 1:]
    
    # 清空原run，添加新内容
    target_run.text = text_before
    
    # 创建上标run
    sup_run = target_run._element.makeelement(qn('w:r'), {})
    
    # 复制格式
    rPr = target_run._element.find(qn('w:rPr'))
    if rPr is not None:
        sup_rPr = deepcopy(rPr)
        sup_run.append(sup_rPr)
    
    # 添加上标属性
    rPr_elem = sup_run.find(qn('w:rPr'))
    if rPr_elem is None:
        rPr_elem = OxmlElement(qn('w:rPr'))
        sup_run.append(rPr_elem)
    
    vert_align = OxmlElement(qn('w:vertAlign'))
    vert_align.set(qn('w:val'), 'superscript')
    rPr_elem.append(vert_align)
    
    # 字号设为小两号
    sz = OxmlElement(qn('w:sz'))
    sz.set(qn('w:val'), '16')  # 8pt (原文可能是10.5pt=21)
    rPr_elem.append(sz)
    
    sz_cs = OxmlElement(qn('w:szCs'))
    sz_cs.set(qn('w:val'), '16')
    rPr_elem.append(sz_cs)
    
    # 添加文本
    t_elem = OxmlElement(qn('w:t'))
    t_elem.text = str(number)
    t_elem.set(qn('xml:space'), 'preserve')
    sup_run.append(t_elem)
    
    # 插入到目标run后面
    target_run._element.addnext(sup_run)
    
    # 创建新run放剩余文本
    if text_after:
        after_run = target_run._element.makeelement(qn('w:r'), {})
        if rPr is not None:
            after_run.append(deepcopy(rPr))
        t_elem2 = OxmlElement(qn('w:t'))
        t_elem2.text = text_after
        t_elem2.set(qn('xml:space'), 'preserve')
        after_run.append(t_elem2)
        sup_run.addnext(after_run)
    
    return True

added = 0
for cp in citation_positions:
    para = doc.paragraphs[cp['para']]
    if add_superscript_after_punct(para, cp['pos'], cp['ref']):
        added += 1

print(f"已添加 {added} 个脚注上标")

# ==========================================
# 保存
# ==========================================
doc.save(OUTPUT)
print(f"\n✅ 保存: {OUTPUT}")

# 验证
doc2 = docx.Document(OUTPUT)
remaining_quotes = sum(1 for p in doc2.paragraphs for ch in p.text if ch == '"')
print(f"验证 - ASCII引号: {remaining_quotes}")
