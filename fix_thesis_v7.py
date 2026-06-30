#!/usr/bin/env python3
"""
叶畏兵毕业论文 - 全面完善脚本 v7
核心修复：
1. ASCII引号 (U+0022) → 中文引号 (U+201C / U+201D)
2. 参考文献77/78拆分
3. 生成脚注-参考文献对应表
"""
import docx
import re

INPUT = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
OUTPUT = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"

doc = docx.Document(INPUT)

# ==========================================
# 第一步：收集全文所有ASCII引号位置，全局配对
# ==========================================
print("=== 分析ASCII引号分布 ===")

# 收集所有包含ASCII引号的run
quote_runs = []  # (para_idx, run_idx, text)
for i, p in enumerate(doc.paragraphs):
    for j, run in enumerate(p.runs):
        if run.text and '"' in run.text:
            quote_runs.append((i, j))

print(f"包含ASCII引号的run数: {len(quote_runs)}")

# ==========================================
# 第二步：全文引号配对修复
# ==========================================
# 策略：按顺序遍历所有run中的引号，成对替换为 "" 和 ""

# 收集所有ASCII引号的 (para_idx, run_idx, char_position_in_run)
all_quotes = []
for i, p in enumerate(doc.paragraphs):
    for j, run in enumerate(p.runs):
        if not run.text:
            continue
        for k, ch in enumerate(run.text):
            if ch == '"':
                all_quotes.append((i, j, k))

print(f"ASCII引号总数: {len(all_quotes)}")

# 成对替换
for idx, (p_idx, r_idx, c_idx) in enumerate(all_quotes):
    run = doc.paragraphs[p_idx].runs[r_idx]
    chars = list(run.text)
    if idx % 2 == 0:
        chars[c_idx] = '"'  # 开引号
    else:
        chars[c_idx] = '"'  # 闭引号
    run.text = ''.join(chars)

print(f"已修复 {len(all_quotes)} 处引号")

# ==========================================
# 第三步：参考文献77/78拆分
# ==========================================
print("\n=== 拆分参考文献77/78 ===")
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if '[77]' in p.text and '[78]' in p.text:
        split_pos = p.text.index('[78]')
        part77 = p.text[:split_pos].rstrip()
        part78 = p.text[split_pos:].strip()
        
        p.text = part77
        
        new_p = doc.add_paragraph()
        if p.style:
            new_p.style = p.style
        new_p.add_run(part78)
        p._p.addnext(new_p._p)
        print(f"  ✅ 已拆分")
        print(f"  77: {part77[:70]}...")
        print(f"  78: {part78[:70]}...")
        break

# ==========================================
# 第四步：保存
# ==========================================
doc.save(OUTPUT)
print(f"\n✅ 保存: {OUTPUT}")

# 验证
doc2 = docx.Document(OUTPUT)
remaining = sum(1 for p in doc2.paragraphs for ch in p.text if ch == '"')
print(f"验证: 剩余ASCII引号 {remaining} 个")
