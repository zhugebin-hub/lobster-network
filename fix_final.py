#!/usr/bin/env python3
"""
叶畏兵毕业论文 - 完善脚本
1. ASCII引号→中文引号
2. 参考文献77/78拆分
3. 添加脚注上标（替换正文中的空格标记）
"""
import docx
import re
from copy import deepcopy
from lxml import etree
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
OUTPUT = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"

doc = docx.Document(INPUT)

# ==========================================
# Step 1: ASCII引号 → 中文引号
# ==========================================
all_quotes = []
for pi, p in enumerate(doc.paragraphs):
    for ri, run in enumerate(p.runs):
        if run.text:
            for ci, ch in enumerate(run.text):
                if ch == '"':
                    all_quotes.append((pi, ri, ci))

for idx, (pi, ri, ci) in enumerate(all_quotes):
    run = doc.paragraphs[pi].runs[ri]
    chars = list(run.text)
    chars[ci] = '\u201c' if idx % 2 == 0 else '\u201d'
    run.text = ''.join(chars)

print(f"Step 1: 修复 {len(all_quotes)} 处引号 ✅")

# ==========================================
# Step 2: 参考文献77/78拆分
# ==========================================
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if '[77]' in p.text and '[78]' in p.text:
        idx = p.text.index('[78]')
        p77, p78 = p.text[:idx].rstrip(), p.text[idx:].strip()
        p.clear()
        p.add_run(p77)
        np = doc.add_paragraph()
        if p.style: np.style = p.style
        np.add_run(p78)
        p._p.addnext(np._p)
        print("Step 2: 拆分77/78 ✅")
        break

# ==========================================
# Step 3: 添加脚注上标
# ==========================================
# 映射：段落号 → 该段内各引用位置对应的参考文献号
# (None表示该位置不需要脚注)
FM = {
    124: [59],
    126: [56, 62],
    127: [63, 54, 51],
    138: [16, 4],
    140: [70, 52],
    142: [10, 6],
    143: [23, 25, 15],
    144: [8, 19],
    146: [63, 36, 40],
    147: [30, 37, 41, 42],
    148: [29, 33, 45],
    150: [47, 66, 50],
    151: [55, 53, 48],
    184: [16],
    185: [4, 17],
    189: [2],
    190: [70],
    191: [65],
    192: [30, 53],
    196: [60],
    215: [2],
    216: [9],
    220: [8],
    221: [19],
    223: [7],
    229: [55],
    230: [52],
    245: [53],
    246: [11],
    254: [10],
    265: [27],
    269: [52],
    270: [9],
    275: [61],
    283: [48],
    293: [48],
    311: [2],
    320: [67],
    321: [67],
    322: [64],
    324: [46, 35],
    340: [79],
    351: [59],
}

def add_superscript(para, char_pos, num):
    """在指定字符位置后插入上标 [num]"""
    pos = 0
    for run in para.runs:
        if not run.text: continue
        rl = len(run.text)
        if pos <= char_pos < pos + rl:
            insert_at = char_pos - pos + 1
            before = run.text[:insert_at]
            after = run.text[insert_at:]
            run.text = before
            
            # 上标run
            r = etree.SubElement(run._element.getparent(), run._element.tag)
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                r.append(deepcopy(rPr))
            rp = r.find(qn('w:rPr'))
            if rp is None:
                rp = OxmlElement(qn('w:rPr'))
                r.append(rp)
            va = OxmlElement(qn('w:vertAlign')); va.set(qn('w:val'),'superscript'); rp.append(va)
            sz = OxmlElement(qn('w:sz')); sz.set(qn('w:val'),'16'); rp.append(sz)
            sz2 = OxmlElement(qn('w:szCs')); sz2.set(qn('w:val'),'16'); rp.append(sz2)
            t = OxmlElement(qn('w:t')); t.text = f'[{num}]'; t.set(qn('xml:space'),'preserve'); r.append(t)
            
            if after:
                r2 = etree.SubElement(r.getparent(), run._element.tag)
                if rPr is not None: r2.append(deepcopy(rPr))
                t2 = OxmlElement(qn('w:t')); t2.text = after; t2.set(qn('xml:space'),'preserve'); r2.append(t2)
            return True
        pos += rl
    return False

print("\nStep 3: 添加脚注上标")
added = 0
for pi, refs in FM.items():
    if pi >= len(doc.paragraphs): continue
    para = doc.paragraphs[pi]
    t = para.text
    positions = [m.start() for m in re.finditer(r'([。》"）]) (?=[\u4e00-\u9fffA-Z])', t)]
    for seq, pos in enumerate(positions):
        if seq < len(refs) and refs[seq]:
            if add_superscript(para, pos, refs[seq]):
                added += 1

print(f"  添加 {added} 个脚注上标 ✅")

# ==========================================
# 保存
# ==========================================
doc.save(OUTPUT)
print(f"\n✅ 保存: {OUTPUT}")

# 验证
doc2 = docx.Document(OUTPUT)
aq = sum(1 for p in doc2.paragraphs for ch in p.text if ch == '"')
print(f"验证: ASCII引号剩余 {aq} 个")
footnotes = sum(p.text.count('[') for p in doc2.paragraphs if re.search(r'\[\d+\]', p.text))
print(f"验证: 脚注标记 {footnotes} 处")
