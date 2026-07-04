#!/usr/bin/env python3
"""生成 QPPB 技术详解 PPT 大纲 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# === 全局样式设置 ===
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = Pt(20)
pf.space_after = Pt(6)

# 标题样式
for level in [1, 2, 3]:
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = '黑体'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1:
        h_style.font.size = Pt(18)
        h_style.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 2:
        h_style.font.size = Pt(16)
        h_style.font.color.rgb = RGBColor(0, 51, 102)
    else:
        h_style.font.size = Pt(14)
        h_style.font.color.rgb = RGBColor(0, 102, 153)


def add_cover_page(doc):
    """添加封面页"""
    for _ in range(4):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('QPPB 技术详解')
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = '黑体'
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('—— 基于 BGP 路由属性的 QoS 策略传播机制')
    run.font.size = Pt(16)
    run.font.name = '黑体'
    run.font.color.rgb = RGBColor(102, 102, 102)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph('')

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in ['汇报人：芦熠檑', '日期：2026 年 5 月']:
        run = info.add_run(line + '\n')
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()


def add_toc(doc):
    """添加目录页"""
    doc.add_heading('目  录', level=1)
    doc.add_paragraph('')

    toc_items = [
        ('第一部分', 'QPPB 概述与产生背景'),
        ('第二部分', 'QPPB 实现原理'),
        ('第三部分', 'QPPB 典型应用场景'),
    ]

    for part, title in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{part}：{title}')
        run.font.size = Pt(14)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p.paragraph_format.space_after = Pt(10)

    doc.add_page_break()


def add_part1(doc):
    """第一部分：QPPB 概述与产生背景"""
    doc.add_heading('第一部分：QPPB 概述与产生背景', level=1)

    # 1.1 什么是 QPPB
    doc.add_heading('1.1 什么是 QPPB？', level=2)

    items = [
        ('全称：', 'QoS Policy Propagation through BGP'),
        ('定义：', '一种特殊的复杂流分类方法，通过 BGP 路由属性对报文进行流分类'),
        ('核心思想：', 'BGP 路由发送者通过设置 BGP 属性预先对路由进行分类，BGP 路由接收者只需要配置合适的接收路由策略，匹配 BGP 路由属性后为 BGP 路由设置 QoS 参数，从而在转发数据时执行相应的 QoS 动作'),
        ('优势：', '在网络发生变化后，BGP 路由接收者可以不更改本地的配置，只是通过 BGP 路由发送者对发送的 BGP 路由属性做相关改动即可'),
    ]

    for label, content in items:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run = p.add_run(content)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 1.2 产生背景
    doc.add_heading('1.2 产生背景（痛点分析）', level=2)

    p = doc.add_paragraph()
    run = p.add_run('场景：')
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run = p.add_run('跨 AS 组网环境（如图 1-52 所示）')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('问题：')
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    problems = [
        '假设 AS400 是一个高优先级网络，所有往/返 AS400 的报文都需要重新设置报文的 IP Precedence，以保证其报文的优先转发',
        '对于 Node-C，只需要针对接口（连接 AS400 的接口）收发的流量进行重标记即可',
        '但对于 Node-A 或 Node-B，则需要执行针对目的地址为 AS400 内 IP 地址的流分类',
        '如果 AS400 内有大量的 IP 地址/地址段，则需要大量的流分类',
        '如果网络结构不稳定，需经常变化网络结构，则需要进行大量的配置修改',
    ]

    for prob in problems:
        p = doc.add_paragraph(prob, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.5)

    p = doc.add_paragraph()
    run = p.add_run('解决方案：')
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run = p.add_run('如果可以按照 AS 信息对报文分类，将使上述 Node-A 或 Node-B 的配置变得简单。为此，产生了 QPPB，一种可以针对 AS 信息、团体属性等聚类信息对报文进行流分类的技术。')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 1.3 QPPB 技术定位
    doc.add_heading('1.3 QPPB 技术定位', level=2)

    p = doc.add_paragraph()
    run = p.add_run('QPPB 是一种特殊的复杂流分类方法。')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('QPPB 技术顾名思义是指通过 BGP 传播 QoS 策略。应用 QPPB 技术可以由 BGP 路由发送者通过设置 BGP 属性预先对路由进行分类，BGP 路由接收者只需要配置合适的接收路由策略，匹配 BGP 路由属性后为 BGP 路由设置 QoS 参数，从而在转发数据时执行相应的 QoS 动作。')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()


def add_part2(doc):
    """第二部分：QPPB 实现原理"""
    doc.add_heading('第二部分：QPPB 实现原理', level=1)

    # 2.1 工作流程
    doc.add_heading('2.1 QPPB 工作流程（图 1-53）', level=2)

    steps = [
        'BGP 路由发送者（Node-C）在发送 BGP 路由时为 BGP 路由设置特定的属性（包括 AS_PATH、团体属性、扩展团体属性等）',
        'BGP 路由在通告过程中可以携带路由属性，这些特定的属性作为 BGP 路由分类的标识',
        'BGP 路由接收者（Node-A）预先在路由策略中设置路由属性匹配项，且对匹配上路由属性的路由设置关联的流动作。在收到 BGP 路由信息后，对匹配上路由属性的路由，在对应的 FIB 中设置 Behavior ID，不同的流动作设置了不同的 Behavior ID',
        '在数据转发过程中，Node-A 可以针对发送到目的网络的数据包，从 FIB 中获取对应的 Behavior ID，执行 Behavior ID 所对应的流动作',
    ]

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'步骤 {i}：')
        run.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run = p.add_run(step)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 2.2 关键技术点
    doc.add_heading('2.2 关键技术点', level=2)

    items = [
        ('Behavior ID：', '不同的流动作设置了不同的 Behavior ID，存储在 FIB 表项中。在数据转发过程中，设备可以根据目的网络从 FIB 中获取对应的 Behavior ID，执行相应的流动作'),
        ('QoS Local-ID：', 'QPPB 策略中绑定 qos-local-id 与 behavior，实现路由属性与 QoS 策略的关联'),
        ('策略传递机制：', '路由发送端通过 route-policy 设置 AS_PATH/Community/Ext-Community 等属性；路由接收端通过 route-policy import 匹配属性，apply qos-local-id'),
    ]

    for label, content in items:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run = p.add_run(content)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 重要说明
    p = doc.add_paragraph()
    run = p.add_run('重要说明：')
    run.bold = True
    run.font.name = '黑体'
    run.font.color.rgb = RGBColor(204, 0, 0)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run = p.add_run('QPPB 技术实际并没有在 BGP 路由信息中发送 QoS 策略，只是在路由发送方通过对通告的路由设置路由属性，在路由接收方根据目的网段的路由属性设置 QoS 策略。')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 2.3 上行 vs 下行 QPPB
    doc.add_heading('2.3 上行 vs 下行 QPPB', level=2)

    table = doc.add_table(rows=4, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['方向', '配置命令', '查表依据', '应用场景']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '黑体'
                run.font.size = Pt(11)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data = [
        ['上行（inbound）', 'qppb-policy policy source inbound', '根据源 IP 查路由表', '用户→ISP 流量计费'],
        ['下行（outbound）', 'qppb-policy policy outbound', '根据目的 IP 查路由表', 'ISP→用户 流量计费'],
        ['基于 IP 优先级', 'qppb-policy ip-precedence source', '根据源/目的地址', '按优先级分类'],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(11)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()


def add_part3(doc):
    """第三部分：QPPB 典型应用场景"""
    doc.add_heading('第三部分：QPPB 典型应用场景', level=1)

    # 3.1 典型应用一
    doc.add_heading('3.1 典型应用一：AS 域间流量分类（图 1-54）', level=2)

    p = doc.add_paragraph()
    run = p.add_run('场景：')
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run = p.add_run('使用 QPPB 可以方便地在 AS100 的边缘设备对 AS 域间的流量进行流分类。例如要在 Node-C 上对 AS200 和 AS400 之间的流量进行限速')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('配置方案：')
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    solutions = [
        '对于 AS200→AS400 方向的流量，在 Node-C 上的 AS100 域内所有接口使能针对源地址的 QPPB',
        '对于 AS400→AS200 方向的流量，在 Node-C 上与 AS400 相连的接口使能针对目的地址的 QPPB',
    ]

    for sol in solutions:
        p = doc.add_paragraph(sol, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.5)

    p = doc.add_paragraph()
    run = p.add_run('须知：')
    run.bold = True
    run.font.name = '黑体'
    run.font.color.rgb = RGBColor(204, 0, 0)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run = p.add_run('查 FIB 转发的是针对上行流量而不是下行流量，因此使能 QPPB 的接口是流量上行的接口。')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 3.2 典型应用二
    doc.add_heading('3.2 典型应用二：L3VPN 流量分类（图 1-55）', level=2)

    p = doc.add_paragraph()
    run = p.add_run('场景：')
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run = p.add_run('如图所示是 QPPB 技术在 BGP/MPLS L3VPN 组网环境中的应用。当 PE 连接多个 VPN 时，可以对某个 VPN-instance 在路由发布时设置 Community 等属性后，再将路由通告出去')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('配置方案：')
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    solutions = [
        'PE 连接多个 VPN 时，对某个 VPN-instance 在路由发布时设置 Community 等属性后，再将路由通告出去',
        '远端 PE 接收到路由信息后将路由及 QoS 等参数设置到 FIB 表项中',
        '使得从 CE 来的流量在转发时能执行相应的 QoS 动作',
        '这样，不同的 VPN 可获得不同的服务质量',
    ]

    for sol in solutions:
        p = doc.add_paragraph(sol, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.5)

    # 3.3 典型应用三
    doc.add_heading('3.3 典型应用三：用户→ISP 的流量计费（图 1-56）', level=2)

    p = doc.add_paragraph()
    run = p.add_run('场景：')
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run = p.add_run('如图所示是 QPPB 技术应用于用户到 ISP 的流量计费场景')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('配置方案：')
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    steps = [
        '通过 BGP 协议，发布路由时携带团体属性',
        '引入 BGP 路由时，匹配团体属性，在路由表中设置 Behavior ID',
        '配置 qppb-policy，匹配 qos-local-id，配置统计/CAR/Remark 等动作',
        '在流量入口方向使能基于目的地址的 QPPB',
        '在用户侧接口的 inbound 方向应用 qppb-policy',
        '报文在转发时，根据接口的 QPPB 使能配置，根据目的 IP 查路由表，获取路由表中的 Behavior ID（即 qos-local-id），再根据接口 qppb-policy 配置进行 qos-local-id 匹配，匹配上就进行统计/CAR/Remark 等动作处理',
    ]

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f'{i}. {step}')
        p.paragraph_format.left_indent = Cm(1.5)

    # 3.4 典型应用四
    doc.add_heading('3.4 典型应用四：ISP→用户的流量计费（图 1-57）', level=2)

    p = doc.add_paragraph()
    run = p.add_run('场景：')
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run = p.add_run('如图所示是 QPPB 技术应用于 ISP 到用户的流量计费场景')
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    run = p.add_run('配置方案：')
    run.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    steps = [
        '通过 BGP 协议，发布路由时携带团体属性',
        '引入 BGP 路由时，匹配团体属性，在路由表中设置 Behavior ID',
        '配置 qppb-policy，匹配 qos-local-id，配置统计/CAR/Remark 等动作',
        '在流量入接口方向使能基于源地址的 QPPB',
        '在用户侧接口的 outbound 方向应用 qppb-policy',
        '报文在转发时，根据接口的 QPPB 使能配置，根据源 IP 查路由表，获取路由表中的 Behavior ID（即 qos-local-id），qos-local-id 经过内部交换传递到流量出接口，再根据出接口的 qppb-policy 配置进行 qos-local-id 匹配，匹配上就进行统计/CAR/Remark 等动作处理',
    ]

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f'{i}. {step}')
        p.paragraph_format.left_indent = Cm(1.5)

    doc.add_page_break()


def add_end_page(doc):
    """添加结束页"""
    doc.add_paragraph('')
    doc.add_paragraph('')

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run('感谢聆听')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = '黑体'
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    end2 = doc.add_paragraph()
    end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end2.add_run('Q & A')
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run.font.color.rgb = RGBColor(102, 102, 102)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


# === 生成文档 ===
add_cover_page(doc)
add_toc(doc)
add_part1(doc)
add_part2(doc)
add_part3(doc)
add_end_page(doc)

output_path = '/home/admin/.openclaw/workspace/QPPB技术详解_大纲.docx'
doc.save(output_path)
print(f'✅ 文档已生成：{output_path}')
