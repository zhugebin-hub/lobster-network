#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '仿宋'
font.size = Pt(16)
font.color.rgb = RGBColor(0, 0, 0)
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 1.5
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(0)
paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符

# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('众生即菩萨净土')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '黑体'
run.font.color.rgb = RGBColor(0, 0, 0)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title.paragraph_format.space_after = Pt(6)
title.paragraph_format.line_spacing = 1.5

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('——依《维摩诘经》浅谈中国佛教众生观')
run.font.size = Pt(16)
run.font.name = '仿宋'
run.font.color.rgb = RGBColor(0, 0, 0)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
subtitle.paragraph_format.space_after = Pt(18)
subtitle.paragraph_format.line_spacing = 1.5

# 正文段落
body_text = """尊敬的各位评委、诸位善知识、同修大德：

阿弥陀佛！

末学是杭州佛学院学僧隆昱，承蒙学院悉心栽培，今日得此法缘，能与诸位讲经交流，无上荣幸。便依止《维摩诘经》的般若妙慧，就"众生观"这一话题，与各位作一探讨。

《维摩诘经·佛国品》有云："众生之类是菩萨佛土。所以者何？菩萨取于净国，皆为饶益诸众生故。"此语一出，颠覆常情。原来净土不在西天，而在众生身上；菩萨建设净土，不为自享清净，而为饶益众生。离开众生，没有净土；舍弃众生，不是菩萨。这便是中国佛教"众生观"的开显。

或许有人要问：净土宗讲往生西方，禅宗讲明心见性，为何《维摩诘经》偏说净土在众生身上？这正是不共之处。经中佛陀告螺髻梵王："我佛土净，而汝不见。"舍利弗所见娑婆充满丘陵坑坎、荆棘沙砾，梵王所见却如自在天宫般庄严。同一国土，所见迥异，关键不在国土本身，而在行者之心。僧肇大师注曰："土之净者，必由众生；众生之净，必因众行。"净土不是某个遥远的地理坐标，而是菩萨在度化众生的过程中，以悲愿与行动一点一滴建设起来的。没有众生这块"田地"，菩萨的悲愿向何处落实？没有众生的苦难作为道场，六度万行从何修起？

如何践行此众生观？回望近代，高僧大德用生命书写了最好的注脚。1937年抗战爆发，太虚大师通电号召全国佛教徒抗战护国，并组织救护队。面对"出家人为何参与战争"的质疑，大师开示道："侵略者破坏国家，伤害人民时，任何人皆负抵抗之义务。"并提出"佛必降魔，方能救世；僧应护国，乃可安禅。"圆瑛大师曾挥泪写下"出世犹垂忧国泪，居山恒作感时诗"，组织百余僧侣救护队，穿梭于枪林弹雨中救死扶伤。淞沪会战期间，僧侣救护队救护伤员及难民八千余人，掩埋遗骸一万多具。降魔与护生，本无差别；菩萨与众生，原是一体。

前贤虽已远去，但其"降魔救世"的精神，至今依然照耀着我们。今日世界，我国人民已走出风雨，但放眼全球，他国人民却处于水深火热之中。正如维摩诘大士所言："以一切众生病，是故我病。"我等佛子，深感悲切。

《入不二法门品》中，文殊师利问维摩诘何为不二，维摩诘默然无言。这一默然，超越言说分别，直契诸法实相。然"不二"非抹杀差别，维摩诘的默然，源于深谙佛理——照见众生本具佛性，了知自他不二，故能生起同体大悲。将此智慧运用于众生观，便知：众生之苦，即我之苦；众生之乐，即我之乐。菩萨视众生，如母视子，视一切众生为多生父母、未来诸佛。

维摩诘的伟大，不仅在于智慧，更在于方便。经中说他"虽为白衣，奉持沙门清净律行；虽处居家，不著三界"，出入酒肆市井，"若至博弈戏处，辄以度人"。他深入一切世间而不染，这正是《维摩诘经》为中国佛教奠定的精神底色：不离世间觉，不染世间相。六祖慧能大师承此精神，直言"佛法在世间，不离世间觉，离世觅菩提，恰如求兔角"，将中国佛教的入世品格推至高峰。

赵朴初居士曾言："中国佛教要发扬人间佛教思想，就要面向群众，面向社会，面向时代，面向世界。"总书记指出，要"坚持我国宗教中国化方向""积极引导宗教与社会主义社会相适应"。此二者看似一古一今、一内一外，实则理路相通。"中国化"不是外在的政治要求，而是佛教自身"契理契机"精神的当代表达——佛法传入中土两千年，从格义佛教到八宗并盛，从百丈清规到人间佛教，每一次本土化都是"应机施教"。今日之"机"，就是中国特色社会主义新时代；今日之"化"，就是让佛教教义与中国现实相融合、与社会主义核心价值观相贯通。总书记强调"以人民为中心"的发展思想，与菩萨"以众生为本"的悲愿，精神内核高度一致：佛法以众生为根，国家以人民本。

放眼当今，战云密布——俄乌硝烟未散，中东血泪又流。探究其源，虽是众生共业所感，实乃无明贪嗔所致。面对生灵涂炭，我等佛子当如何自处？是闭目一句"娑婆堪忍"，念佛求往生净土吗？若如此，则是错谬了菩萨本怀。经云："菩萨欲得净土，当净其心；随其心净，则佛土净。"是故，我等佛子应团结一心，以"爱国爱教"为核心，完善僧格，饶益有情，效仿维摩诘大士，以悲智双运回应时代呼唤；努力建设人间佛教，以佛教正法改良社会、促进人类进步、推动世界和平。

佛陀出生、修行、成佛，皆在人间。众生世界就是诸佛出处，人间苦难就是菩萨道场。建立人间净土，亦是诸佛菩萨的宏愿。作为后学的我们，要将这份宏愿化为当下每一步的坚实脚印：

以慈悲心，落实具体行动。同体大悲是无条件的悲悯，将每一次诵经、每一句佛号，皆回向一切众生，唤醒良知；随缘布施，救助同胞及难民；推动对话，促进跨文化交流，勇做和平使者。

以如实观，超越信息茧房。算法推送常遮蔽真相，当以多方求证、独立思考，不被情绪裹挟。如法发声，谴责一切暴行，不让无知与偏见成为仇恨的燃料。

以智慧眼，破除二元对立。网络信息纷杂，极端言论时有所闻。我等当以缘起智慧观照，莫被妄转。支持和平而不陷入对立，反对侵略而不仇视人民，须知战火下丧生的每一个生命，无论民族、立场，皆是未来诸佛。

以共命观，体认人类命运共同体。当今世界，一场战争波及全球，能源危机、粮食短缺、难民潮涌，无一国能够独善其身。总书记提出构建"人类命运共同体"的伟大理念，与佛教众生观在深层遥相呼应——缘起法告诉我们，一切法互为因缘，无一法孤立存在。地球尽为国土，人类悉为同胞。真正的众生观，必以平等心对待一切人类、一切生命。

昔日，维摩诘大士以一默答文殊；今我辈虽不能至，然心向往之。愿效法此"不二"之智，消弭世间对立，融通自他壁垒。

最后，让我们共同至诚祈愿：

愿战火纷飞处，早得清凉；
愿苦难流离者，速获安乐；
愿仇恨对立心，化为慈悲；
愿人类共命缘，结成和平；
愿佛教中国化，灯灯相续；
愿人间净土愿，步步实现。

阿弥陀佛！（合十）"""

for para in body_text.split('\n'):
    p = doc.add_paragraph()
    run = p.add_run(para)
    run.font.size = Pt(16)
    run.font.name = '仿宋'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # 首行缩进（除了空行和标题类行）
    if para.strip():
        p.paragraph_format.first_line_indent = Cm(0.74)

output_path = '/home/admin/.openclaw/workspace/众生即菩萨净土_隆昱.docx'
doc.save(output_path)
print(f"Done: {output_path}")
