#!/usr/bin/env python3
"""Manus 直播介绍 Word 文档 - 重新梳理直播内容介绍"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document('/home/admin/.openclaw/workspace/digital-employee-report/Manus智能体全攻略_直播介绍_最终版.docx')

# 找到包含"本次直播深度解读"的段落
for p in doc.paragraphs:
    if '本次直播' in p.text:
        for r in p.runs:
            r.text = ''
        r = p.runs[0] if p.runs else p.add_run()
        r.text = (
            '本次直播解读清华版《Manus智能体全攻略》，聚焦AI赋能高校教学的四大场景：'
            '课件智能生成（输入大纲一键出PPT）、论文协作写作（从选题到成稿全流程支持）、'
            '教学案例自动开发（3天完成全套数字化教学资源包）、实验数据可视化（一键生成可发表图表）。'
            '现场实操演示如何将2小时的教案制作压缩至10分钟。'
        )
        r.font.name = '微软雅黑'
        r.font.size = Pt(12)
        break

output = '/home/admin/.openclaw/workspace/digital-employee-report/Manus智能体全攻略_直播介绍_最终版.docx'
doc.save(output)
print(f'✅ 已更新: {output}')
