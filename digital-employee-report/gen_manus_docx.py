#!/usr/bin/env python3
"""将 Manus 直播材料转为 Word 文档"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ─── 全局样式 ───
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(12)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ─── 标题 ───
title = doc.add_heading('《Manus智能体全攻略》直播内容介绍', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
    run.font.name = '微软雅黑'
    run.font.size = Pt(22)

# ─── 基本信息 ───
doc.add_heading('一、基本信息', level=1)

basic_info = [
    ('书名', '《Manus智能体全攻略》'),
    ('出版社', '清华大学出版社'),
    ('主编', '洪金珠、李晓林、诸葛斌'),
    ('特色', '8个工程级案例源码 + 配套视频教程'),
]

for label, value in basic_info:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run_label = p.add_run(f'{label}：')
    run_label.bold = True
    run_label.font.size = Pt(12)
    run_label.font.name = '微软雅黑'
    run_value = p.add_run(value)
    run_value.font.size = Pt(12)
    run_value.font.name = '微软雅黑'

# ─── 直播文案 ───
doc.add_heading('二、直播文案', level=1)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run('🔥 清华大学出版社重磅新书《Manus智能体全攻略》直播专场！')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xD4, 0x2C, 0x00)
run.font.name = '微软雅黑'

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run('三位专家联袂主编，8个工程级案例源码 + 视频全配套！')
run.font.size = Pt(12)
run.font.name = '微软雅黑'

doc.add_paragraph()  # spacer

p = doc.add_paragraph()
run = p.add_run('🎯 直播间硬核演示：')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
run.font.name = '微软雅黑'

demo_items = [
    '📄 输入大纲，AI自动生成PPT课件',
    '📝 论文辅助写作全流程实操',
    '🏫 教学案例自动开发，3天搞定全套数字化教学资源包',
    '📊 实验数据一键生成可发表图表',
]
for item in demo_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'

doc.add_paragraph()  # spacer

p = doc.add_paragraph()
run = p.add_run('🎁 直播专属福利：')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xD4, 0x2C, 0x00)
run.font.name = '微软雅黑'

benefit_items = [
    '教材 + Manus教育版联合优惠',
    '教师完成培训可获 工信部认证证书',
]
for item in benefit_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('适合高校教师、AI教育从业者、对智能体应用感兴趣的开发者！')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
run.font.name = '微软雅黑'

# ─── 教育解决方案亮点 ───
doc.add_heading('三、教育解决方案亮点', level=1)

# 创建表格
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'

# 表头
headers = ['能力', '说明']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '微软雅黑'
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_data = [
    ('课件智能生成', '输入大纲自动输出PPT'),
    ('论文辅助写作', '从选题到成稿全流程支持'),
    ('教学案例开发', '3天完成全套数字化教学资源'),
    ('PDF转AR课件', '教材升级为沉浸式课件'),
    ('数据可视化', '实验数据自动生成可发表图表'),
]

for i, (cap, desc) in enumerate(rows_data):
    row = table.rows[i + 1]
    row.cells[0].text = cap
    row.cells[1].text = desc
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'

# 设置列宽
for row in table.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(11)

# ─── 适用人群 ───
doc.add_heading('四、适用人群', level=1)

audience = [
    '高校教师与教育工作者',
    'AI/智能体应用开发者',
    '教育数字化转型推动者',
    '对AI辅助教学感兴趣的从业者',
]
for item in audience:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'

# ─── 页脚信息 ───
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 本材料由 AI 辅助生成 —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = '微软雅黑'

# ─── 保存 ───
output = '/home/admin/.openclaw/workspace/digital-employee-report/Manus智能体全攻略_直播内容介绍.docx'
doc.save(output)
print(f'✅ Word 文档已生成: {output}')
