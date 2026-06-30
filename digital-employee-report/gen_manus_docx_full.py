#!/usr/bin/env python3
"""将 Manus 直播材料转为 Word 文档（完整版）"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ─── 页面设置 ───
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ─── 全局样式 ───
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ─── 标题样式 ───
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)

# Helper: 设置单元格字体
def set_cell_font(cell, size=11, bold=False, color=None):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.name = '微软雅黑'
            r.font.bold = bold
            if color:
                r.font.color.rgb = color

# ════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('《Manus智能体全攻略》')
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
r.font.name = '微软雅黑'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('直播内容介绍材料')
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
r.font.name = '微软雅黑'

for _ in range(3):
    doc.add_paragraph()

# 基本信息封面
info_items = [
    '出版社：清华大学出版社',
    '主编：洪金珠、李晓林、诸葛斌',
    '日期：2026年6月',
]
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(item)
    r.font.size = Pt(14)
    r.font.name = '微软雅黑'
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ════════════════════════════════════════════
# 一、基本信息
# ════════════════════════════════════════════
doc.add_heading('一、基本信息', level=1)

basic_info = [
    ('书名', '《Manus智能体全攻略》'),
    ('出版社', '清华大学出版社'),
    ('主编', '洪金珠、李晓林、诸葛斌'),
    ('内容定位', '系统介绍Manus智能体的基本原理、核心功能与实践应用'),
    ('章节结构', '全书共五章，涵盖智能体理论基础、开发路径、功能特性、使用指南与实战案例'),
    ('配套资源', '8个工程级案例源码 + 配套视频教程 + 案例回放链接 + 案例涉及文件 + 案例讲解录屏'),
]

for label, value in basic_info:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run(f'{label}：')
    run_label.bold = True
    run_label.font.size = Pt(12)
    run_label.font.name = '微软雅黑'
    run_value = p.add_run(value)
    run_value.font.size = Pt(12)
    run_value.font.name = '微软雅黑'

# ════════════════════════════════════════════
# 二、内容简介
# ════════════════════════════════════════════
doc.add_heading('二、内容简介', level=1)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run(
    '本书系统介绍了Manus智能体的基本原理、核心功能与实践应用，旨在帮助读者全面掌握这一前沿AI技术的使用方法与开发思路。全书共分为五章，涵盖智能体理论基础、开发路径、功能特性、使用指南与实战案例等多个维度，结构完整、内容丰富。'
)
run.font.size = Pt(12)
run.font.name = '微软雅黑'

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run(
    '本书内容详实、案例丰富，兼顾理论深度与实践指导，既适合作为高校人工智能等专业的教学参考书，也可作为企业技术团队、个人开发者以及AI产品设计者的实战指南。通过本书的学习，读者将能够掌握Manus智能体的核心技能，理解"连接思想与行动"的智能体范式，并将其应用于提升工作效率与创新能力的各类场景中。'
)
run.font.size = Pt(12)
run.font.name = '微软雅黑'

# ════════════════════════════════════════════
# 三、直播文案
# ════════════════════════════════════════════
doc.add_heading('三、直播文案', level=1)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run('🔥 清华大学出版社重磅新书《Manus智能体全攻略》直播专场！')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0xD4, 0x2C, 0x00)
r.font.name = '微软雅黑'

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run('三位专家联袂主编，8个工程级案例源码 + 视频全配套！')
r.font.size = Pt(13)
r.font.name = '微软雅黑'

# 硬核演示
p = doc.add_paragraph()
r = p.add_run('🎯 直播间硬核演示：')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
r.font.name = '微软雅黑'

demo_items = [
    '输入大纲，AI自动生成PPT课件',
    '论文辅助写作全流程实操',
    '教学案例自动开发，3天搞定全套数字化教学资源包',
    '实验数据一键生成可发表图表',
]
for item in demo_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'

# 专属福利
p = doc.add_paragraph()
r = p.add_run('🎁 直播专属福利：')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0xD4, 0x2C, 0x00)
r.font.name = '微软雅黑'

benefit_items = [
    '教材 + Manus教育版联合优惠',
    '教师完成培训可获 工信部认证证书',
]
for item in benefit_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'

p = doc.add_paragraph()
r = p.add_run('适合高校教师、AI教育从业者、对智能体应用感兴趣的开发者！')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
r.font.name = '微软雅黑'

# ════════════════════════════════════════════
# 四、教育解决方案亮点
# ════════════════════════════════════════════
doc.add_heading('四、教育解决方案亮点', level=1)

table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 表头
headers = ['能力', '说明']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_font(cell, size=12, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

rows_data = [
    ('课件智能生成', '输入大纲自动输出PPT'),
    ('论文辅助写作', '从选题到成稿全流程支持'),
    ('教学案例开发', '3天完成全套数字化教学资源'),
    ('PDF转AR课件', '教材升级为沉浸式课件'),
    ('数据可视化', '实验数据自动生成可发表图表'),
]

for i, (cap, desc) in enumerate(rows_data):
    row = table.rows[i + 1]
    row.cells[0].text = cap
    row.cells[1].text = desc
    set_cell_font(row.cells[0], size=11)
    set_cell_font(row.cells[1], size=11)

# 列宽
for row in table.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(11)

# ════════════════════════════════════════════
# 五、适用人群
# ════════════════════════════════════════════
doc.add_heading('五、适用人群', level=1)

audience = [
    '高校教师与教育工作者',
    'AI/智能体应用开发者',
    '教育数字化转型推动者',
    '对AI辅助教学感兴趣的从业者',
]
for item in audience:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'

# ════════════════════════════════════════════
# 六、作者简介（诸葛斌教授）
# ════════════════════════════════════════════
doc.add_heading('六、主编简介', level=1)

doc.add_heading('诸葛斌', level=2)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run(
    '浙江工商大学信息与电子工程学院/萨塞克斯人工智能学院教授。研究方向为互联网应用开发，国家线上一流课程负责人。'
    '主持国家自然科学基金1项、浙江省自然科学基金1项、浙江省科技计划项目1项，在国内外期刊上发表SCI及EI检索论文50余篇，'
    '申请发明专利50余项（授权30余项），主编教材10部。'
)
run.font.size = Pt(12)
run.font.name = '微软雅黑'

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run('主要贡献：')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6A)
run.font.name = '微软雅黑'

contributions = [
    '与袁非牛、洪金珠共同编著《数字人文：AI情感分析与文化生成》（上海教育出版社，2025年）',
    '与钉钉公司联合编写《钉钉AI助理应用实战》教材（清华大学出版社，已定稿提交）',
    '主编《Manus智能体全攻略》（清华大学出版社，已定稿提交）',
    '依托教育部-阿里云产学合作协同育人项目，联合编写《百炼智能体开发与应用》教材',
    '开设研究生DeepSeek主题课程，吸引线上线下900余人参与',
    '积极推广AI Agent技术与数字员工应用，在多个场合分享OpenClaw实践',
]
for item in contributions:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

doc.add_heading('洪金珠', level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run(
    '联合主编，与袁非牛、诸葛斌共同编著《数字人文：AI情感分析与文化生成》（上海教育出版社，2025年），'
    '在数字人文与AI情感分析领域有深入研究。'
)
run.font.size = Pt(12)
run.font.name = '微软雅黑'

doc.add_heading('李晓林', level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run('联合主编，与洪金珠、诸葛斌共同编著《Manus智能体全攻略》（清华大学出版社）。')
run.font.size = Pt(12)
run.font.name = '微软雅黑'

# ════════════════════════════════════════════
# 页脚
# ════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— 本材料由 AI 辅助生成 —')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.font.name = '微软雅黑'

# ─── 保存 ───
output = '/home/admin/.openclaw/workspace/digital-employee-report/Manus智能体全攻略_直播内容介绍.docx'
doc.save(output)
print(f'✅ Word 文档已生成: {output}')
print(f'   文件大小: {os.path.getsize(output)} bytes')
