#!/usr/bin/env python3
"""Manus 直播介绍 Word 文档 - 个人简介加入AI教育大会一等奖"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

template = Document('/home/admin/.openclaw/workspace/digital-employee-report/template_live_stream.docx')
doc = Document()

for section in doc.sections:
    section.top_margin = template.sections[0].top_margin
    section.bottom_margin = template.sections[0].bottom_margin
    section.left_margin = template.sections[0].left_margin
    section.right_margin = template.sections[0].right_margin

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

photo_path = '/home/admin/.openclaw/workspace/digital-employee-report/author_photo.png'

# ═══ 1. 视频要求 ═══
p = doc.add_paragraph()
r = p.add_run('视频要求：')
r.font.name = '微软雅黑'
r = p.add_run('格式支持MP4，大小不超过2G，建议视频分辨率为1280×720，时长60分钟左右为最佳；视频内容要求：即为授课视频，内容有关于直播所需要用到的图书即可')
r.font.name = '微软雅黑'

doc.add_paragraph()

# ═══ 2. 直播题目 ═══
p = doc.add_paragraph()
r = p.add_run('直播题目：')
r.font.name = '微软雅黑'
r = p.add_run('清华版《Manus智能体全攻略》教材解读——AI赋能高校教学新范式')
r.font.name = '微软雅黑'

doc.add_paragraph()

# ═══ 3. 所用图书 ═══
p = doc.add_paragraph()
r = p.add_run('所用图书：')
r.font.name = '微软雅黑'
r = p.add_run('《Manus智能体全攻略》')
r.font.name = '微软雅黑'
r = p.add_run('（')
r.font.name = '微软雅黑'
r = p.add_run('清华大学出版社')
r.font.name = '微软雅黑'
r = p.add_run('）')
r.font.name = '微软雅黑'

doc.add_paragraph()

# ═══ 4. 直播内容介绍 ═══
p = doc.add_paragraph()
r = p.add_run('直播内容介绍：')
r.font.name = '微软雅黑'
r = p.add_run('（100字以内为最佳）')
r.font.name = '微软雅黑'

p = doc.add_paragraph()
intro_text = (
    '本次直播解读清华版《Manus智能体全攻略》，'
    '涵盖五大教学场景：'
    '课件智能生成（"小龙虾三部曲"精品课件）、'
    '教学案例开发（烟草数据挖掘、网络课程动画、微信小程序等）、'
    '论文协作写作（从选题到IEEE成稿全流程）、'
    '教学视频制作（PPT逐页自动生成教学视频）、'
    '数据可视化（一键生成可发表图表），'
    '实操2小时教案制作压缩至10分钟。'
)
r = p.add_run(intro_text)
r.font.size = Pt(12)
r.font.name = '微软雅黑'
print(f'   直播内容介绍字数: {len(intro_text)}')

doc.add_paragraph()

# ═══ 5. 作者个人简介 ═══
p = doc.add_paragraph()
r = p.add_run('作者个人简介：')
r.font.name = '微软雅黑'
r = p.add_run('（100字以内为最佳）')
r.font.name = '微软雅黑'

p = doc.add_paragraph()
r = p.add_run(
    '诸葛斌，浙江工商大学信息与电子工程学院/萨塞克斯人工智能学院教授。'
    '研究方向为互联网应用开发与AI教育，获浙江省技术发明一等奖。'
    '2025全国高校人工智能教育大会优秀案例一等奖，'
    '"人工智能+"背景下基于阿里云的智能体教学实践与人才培养创新。'
    '联合钉钉撰写国内首本低代码开发教材，'
    '本次直播将实操演示如何用AI智能体将2小时的教案课件制作压缩至10分钟。'
)
r.font.size = Pt(12)
r.font.name = '微软雅黑'

doc.add_paragraph()

# ═══ 6. 直播时间 ═══
p = doc.add_paragraph()
r = p.add_run('直播时间：')
r.font.name = '微软雅黑'
r = p.add_run('6月24日（下周三）下午 3:00-4:00')
r.font.name = '微软雅黑'

doc.add_paragraph()

# ═══ 7. 作者照片 ═══
p = doc.add_paragraph()
r = p.add_run('（作者提供高清照片一张）')
r.font.name = '微软雅黑'

if os.path.exists(photo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(photo_path, width=Inches(2.0))

output = '/home/admin/.openclaw/workspace/digital-employee-report/Manus智能体全攻略_直播介绍_最终版.docx'
doc.save(output)
print(f'✅ 已更新: {output}')
print(f'   文件大小: {os.path.getsize(output)} bytes')
