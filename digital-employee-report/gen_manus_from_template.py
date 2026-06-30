#!/usr/bin/env python3
"""参考钉钉低代码直播模板，生成 Manus 直播介绍 Word 文档（含作者照片）"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# ─── 读取模板 ───
template = Document('/home/admin/.openclaw/workspace/digital-employee-report/template_live_stream.docx')

# ─── 创建新文档 ───
doc = Document()

# 保持模板的页面设置
for section in doc.sections:
    section.top_margin = template.sections[0].top_margin
    section.bottom_margin = template.sections[0].bottom_margin
    section.left_margin = template.sections[0].left_margin
    section.right_margin = template.sections[0].right_margin

# ─── 全局样式 ───
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

photo_path = '/home/admin/.openclaw/workspace/digital-employee-report/author_photo.png'

# ════════════════════════════════════════════
# 按照模板格式写入
# ════════════════════════════════════════════

# 1. 视频要求说明
p = doc.add_paragraph()
r = p.add_run('视频要求：')
r.font.name = '微软雅黑'
r = p.add_run('格式支持MP4，大小不超过2G，建议视频分辨率为1280×720，时长60分钟左右为最佳；视频内容要求：即为授课视频，内容有关于直播所需要用到的图书即可')
r.font.name = '微软雅黑'

# 空行
doc.add_paragraph()

# 2. 直播题目
p = doc.add_paragraph()
r = p.add_run('直播题目：')
r.font.name = '微软雅黑'
r = p.add_run('《Manus智能体全攻略》——AI智能体开发实战与应用')
r.font.name = '微软雅黑'

# 空行
doc.add_paragraph()

# 3. 所用图书
p = doc.add_paragraph()
r = p.add_run('所用图书：')
r.font.name = '微软雅黑'
r = p.add_run('《Manus智能体全攻略》')
r.font.name = '微软雅黑'
r = p.add_run('（')
r.font.name = '微软雅黑'
r = p.add_run('清华大学出版社')
r.font.name = '微软雅黑'
r = p.add_run('）')
r.font.name = '微软雅黑'

# 空行
doc.add_paragraph()

# 4. 直播内容介绍（100字以内）
p = doc.add_paragraph()
r = p.add_run('直播内容介绍：')
r.font.name = '微软雅黑'
r = p.add_run('（100字以内为最佳）')
r.font.name = '微软雅黑'

p = doc.add_paragraph()
r = p.add_run(
    '联合洪金珠、李晓林共同编著，已与清华大学出版社签订出版合同并已定稿提交。'
    '本书系统介绍Manus智能体的基本原理、核心功能与实践应用，涵盖智能体理论基础、开发路径、功能特性、'
    '使用指南与实战案例五大章节。提供8个工程级案例源码及配套视频教程，既适合作为高校人工智能专业的教学参考书，'
    '也可作为企业技术团队、个人开发者以及AI产品设计者的实战指南。'
    '通过本书学习，读者将掌握Manus智能体核心技能，理解"连接思想与行动"的智能体范式，'
    '并将其应用于提升工作效率与创新能力的各类场景中。'
)
r.font.size = Pt(12)
r.font.name = '微软雅黑'

# 空行
doc.add_paragraph()

# 5. 作者个人简介（100字以内）
p = doc.add_paragraph()
r = p.add_run('作者个人简介：')
r.font.name = '微软雅黑'
r = p.add_run('（100字以内为最佳）')
r.font.name = '微软雅黑'

p = doc.add_paragraph()
r = p.add_run(
    '诸葛斌，浙江工商大学信息与电子工程学院/萨塞克斯人工智能学院教授。'
    '研究方向为互联网应用开发，国家线上一流课程负责人。'
    '主持国家自然科学基金1项、浙江省自然科学基金1项，发表SCI及EI检索论文50余篇，'
    '申请发明专利50余项（授权30余项），主编教材10部。'
    '与钉钉公司联合编写《钉钉AI助理应用实战》教材，'
    '积极推广AI Agent技术与数字员工应用，在多个场合分享OpenClaw实践。'
)
r.font.size = Pt(12)
r.font.name = '微软雅黑'

# 空行
doc.add_paragraph()

# 6. 直播时间
p = doc.add_paragraph()
r = p.add_run('直播时间：')
r.font.name = '微软雅黑'
r = p.add_run('待定')
r.font.name = '微软雅黑'

# 空行
doc.add_paragraph()

# 7. 作者照片
p = doc.add_paragraph()
r = p.add_run('（作者提供高清照片一张）')
r.font.name = '微软雅黑'

# 插入照片
if os.path.exists(photo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(photo_path, width=Inches(2.0))

# ─── 保存 ───
output = '/home/admin/.openclaw/workspace/digital-employee-report/Manus智能体全攻略_直播介绍.docx'
doc.save(output)
print(f'✅ Word 文档已生成: {output}')
print(f'   文件大小: {os.path.getsize(output)} bytes')
