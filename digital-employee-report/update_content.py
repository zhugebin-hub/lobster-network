#!/usr/bin/env python3
"""Manus 直播介绍 Word 文档 - 修改直播内容介绍"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document('/home/admin/.openclaw/workspace/digital-employee-report/Manus智能体全攻略_直播介绍_最终版.docx')

# 找到包含"本次直播深度解读"的段落
for p in doc.paragraphs:
    if '本次直播深度解读' in p.text:
        for r in p.runs:
            r.text = ''
        r = p.runs[0] if p.runs else p.add_run()
        r.text = (
            '本次直播深度解读清华版《Manus智能体全攻略》四大核心价值：'
            '洪金珠、李晓林、诸葛斌三位专家联袂主编，含8个工程级案例源码；'
            '现场演示课件智能生成、论文协作写作、教学案例开发；'
            '提供3天完成数字化教学资源包等教育场景解决方案，'
            '实操演示如何用AI智能体将2小时的教案课件制作压缩至10分钟。'
        )
        r.font.name = '微软雅黑'
        r.font.size = Pt(12)
        break

output = '/home/admin/.openclaw/workspace/digital-employee-report/Manus智能体全攻略_直播介绍_最终版.docx'
doc.save(output)
print(f'✅ 已更新: {output}')
