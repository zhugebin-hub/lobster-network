#!/usr/bin/env python3
"""Manus 直播介绍 Word 文档 - 更新直播时间"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document('/home/admin/.openclaw/workspace/digital-employee-report/Manus智能体全攻略_直播介绍_优化版.docx')

# 找到"直播时间"段落并更新
for p in doc.paragraphs:
    if '直播时间' in p.text:
        # 清空所有 runs
        for r in p.runs:
            r.text = ''
        r = p.runs[0] if p.runs else p.add_run()
        r.text = '直播时间：6月25日（下周三）下午 3:00-4:00'
        r.font.name = '微软雅黑'
        r.font.size = Pt(12)
        break

output = '/home/admin/.openclaw/workspace/digital-employee-report/Manus智能体全攻略_直播介绍_最终版.docx'
doc.save(output)
print(f'✅ 已更新: {output}')
