"""
为薛瑾老师生成"闪光课堂教学节"授课证明和讲座证明
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_text(cell, text, font_name='仿宋', font_size=16, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    # 中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_certificate(doc, title, content, date='2026年6月1日'):
    # 标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.space_before = Pt(40)
    title_p.space_after = Pt(30)
    run = title_p.add_run(title)
    run.font.name = '方正小标宋简体'
    run.font.size = Pt(26)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')

    # 正文
    body_p = doc.add_paragraph()
    body_p.space_before = Pt(20)
    body_p.space_after = Pt(10)
    body_p.paragraph_format.line_spacing = 1.8
    
    # 首行缩进
    pPr = body_p._element.get_or_add_pPr()
    pPr.set(qn('w:ind'), '2em')
    
    run = body_p.add_run(content)
    run.font.name = '仿宋'
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 落款单位
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_p.space_before = Pt(40)
    run = sig_p.add_run('XX学校')
    run.font.name = '仿宋'
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 日期
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_p.add_run(date)
    run.font.name = '仿宋'
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    doc.add_paragraph()  # 空行

# ===== 创建文档 =====
doc = Document()

# 设置页面
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(3.5)
section.bottom_margin = Cm(3)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)

# ===== 证明一：公开课证明 =====
create_certificate(doc,
    '证  明',
    '薛瑾老师在我校小学部"闪光课堂教学节"活动中，执教五年级下册古诗词诵读《春夜喜雨》公开课一节。该课教学设计精巧，课堂呈现精彩，充分展现了教师扎实的教学功底和深厚的文学素养，获得了师生的一致好评。'
)

# 分页
doc.add_page_break()

# ===== 证明二：讲座证明 =====
create_certificate(doc,
    '证  明',
    '薛瑾老师在我校小学部"闪光课堂教学节"活动中，作了《小学中高年级古诗词教学中的高阶思维》专题讲座。讲座观点鲜明、内容充实，对古诗词教学中培养学生高阶思维提供了有益的实践思考和理论指引，得到了参会教师的充分肯定。'
)

# 保存
output_path = '/home/admin/.openclaw/workspace/薛瑾老师_闪光课堂教学节_证明.docx'
doc.save(output_path)
print(f'✅ 证明已生成：{output_path}')
