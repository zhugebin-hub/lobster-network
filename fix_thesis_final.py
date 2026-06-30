#!/usr/bin/env python3
"""
叶畏兵毕业论文完善 - 最终版
修复所有格式问题
"""
import docx
import re

INPUT = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
OUTPUT = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"

doc = docx.Document(INPUT)

# ==========================================
# Step 1: Fix all ASCII double quotes → Chinese quotes
# ==========================================
# Collect ALL ASCII quote positions globally
all_quote_positions = []  # (paragraph_index, run_index, char_index)
for p_idx, para in enumerate(doc.paragraphs):
    for r_idx, run in enumerate(para.runs):
        if not run.text:
            continue
        for c_idx, ch in enumerate(run.text):
            if ch == '"':
                all_quote_positions.append((p_idx, r_idx, c_idx))

print(f"Found {len(all_quote_positions)} ASCII quotes")

# Replace in pairs: even→open, odd→close
for idx, (p_idx, r_idx, c_idx) in enumerate(all_quote_positions):
    run = doc.paragraphs[p_idx].runs[r_idx]
    text_list = list(run.text)
    if idx % 2 == 0:
        text_list[c_idx] = '\u201c'  # "
    else:
        text_list[c_idx] = '\u201d'  # "
    run.text = ''.join(text_list)

# ==========================================
# Step 2: Fix reference 77/78 merge
# ==========================================
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if '[77]' in p.text and '[78]' in p.text:
        split_idx = p.text.index('[78]')
        part77 = p.text[:split_idx].rstrip()
        part78 = p.text[split_idx:].strip()
        
        # Update paragraph 77
        p.clear()
        p.add_run(part77)
        
        # Insert new paragraph for 78
        new_p = doc.add_paragraph()
        if p.style:
            new_p.style = p.style
        new_p.add_run(part78)
        
        # Move new_p right after p
        p._p.addnext(new_p._p)
        print(f"Split ref 77/78 successfully")
        break

# ==========================================
# Step 3: Fix mixed quote patterns like "" and ""
# ==========================================
# These are patterns where Chinese closing quote follows ASCII opening
for para in doc.paragraphs:
    for run in para.runs:
        if not run.text:
            continue
        old = run.text
        # Fix "" → ""
        new = old.replace('""', '""')
        # Fix "" → ""
        new = new.replace('""', '""')
        # Fix standalone "" at end of phrase
        new = new.replace('""', '""')
        if new != old:
            run.text = new

# ==========================================
# Step 4: Save
# ==========================================
doc.save(OUTPUT)
print(f"\nSaved to {OUTPUT}")

# Verify
doc2 = docx.Document(OUTPUT)
remaining = sum(1 for p in doc2.paragraphs for ch in p.text if ch == '"')
print(f"Remaining ASCII quotes: {remaining}")

# List any remaining issues
for i, p in enumerate(doc2.paragraphs):
    for ch in p.text:
        if ch == '"':
            print(f"  Para {i}: {p.text[:80]}")
            break
