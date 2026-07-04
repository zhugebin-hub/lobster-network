#!/usr/bin/env python3
"""生成优化后的百炼教材大纲 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
style.font.name = '仿宋'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ===== 辅助函数 =====
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        sizes = {0: Pt(22), 1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(11)}
        run.font.size = sizes.get(level, Pt(12))
    h.paragraph_format.space_before = Pt({0:0,1:12,2:10,3:8,4:6}.get(level, 8))
    h.paragraph_format.space_after = Pt({0:12,1:6,2:6,3:4,4:2}.get(level, 4))
    return h

def add_para(text, indent=True, bold_prefix=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = '仿宋'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '仿宋'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    return p

def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    indent = Cm(0.74 + level * 0.74)
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = '仿宋'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '仿宋'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = '仿宋'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = '仿宋'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def chapter_section(title, bullets):
    add_heading(title, level=3)
    for b in bullets:
        add_bullet(b)

# ===== 标题 =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('《阿里云百炼：智能体工程化与OpenClaw实战》\n优化后大纲结构（完整版）')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title.paragraph_format.space_after = Pt(12)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('——融合递归自主式分解与人机协作新范式')
run.font.size = Pt(16)
run.font.name = '楷体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
subtitle.paragraph_format.space_after = Pt(6)

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('诸葛斌  浙江工商大学 人工智能学院\n2026年5月')
run.font.size = Pt(12)
run.font.name = '楷体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
author.paragraph_format.space_after = Pt(18)

doc.add_paragraph('_' * 80).paragraph_format.space_after = Pt(12)

# ===== 优化说明 =====
add_heading('优化说明', level=1)

add_para(
    '本优化方案在原版大纲（第一部分百炼核心工程篇8章 + 第二部分OpenClaw高阶生态篇3章）基础上，'
    '新增第三部分"递归自主式人机协作篇"3章，形成14章完整结构。'
    '优化核心：将国家级智能体政策、软件定义价格（SDP）理论、递归自主式分解框架与百炼工程实践深度融合。'
)

add_table(
    ['优化维度', '原版', '优化后', '价值'],
    [
        ['案例连贯性', '10个独立案例', '1个贯穿案例 + 10个专题案例', '学完=完成真实项目'],
        ['理论高度', '工程实践为主', '递归分解框架 + SDP演进', '理论+实践双轮驱动'],
        ['人机协作', '隐含在案例中', '显性化决策表 + 能力画像', '明确何时用AI/何时用人'],
        ['政策对齐', '未显性对标', '第14章完整对标五大能力', '响应国家战略'],
        ['Hermes定位', '缺失', '第13章深度分析引擎', '双智能体协作闭环'],
        ['专著融合', '缺失', '第12章SDP现代演进', '十年理论延续'],
    ]
)

# ===== 第一部分 =====
add_heading('第一部分：百炼核心工程篇（案例集成，共8章）', level=1)

# 第1章
add_heading('第1章 范式重构：从大模型到代理式AI', level=2)
add_para('📍 贯穿案例定位：对应MB-001需求分析与方案设计，展示智能体能做什么。', indent=False, bold_prefix='')

add_heading('1.1 人工智能发展的新阶段：从生成到代理', level=3)
add_bullet('1.1.1 生成式大模型（LLM）的边际局限：无状态与被动响应')
add_bullet('1.1.2 代理式AI（Agentic AI）的崛起：自主行动与目标驱动')
add_bullet('1.1.3 国家级智能体政策解读：五大基础能力定义')

add_heading('1.2 智能体的核心运行机制', level=3)
add_bullet('1.2.1 感知系统：多模态输入与环境数据理解')
add_bullet('1.2.2 规划系统：思维链（CoT）与任务自主拆解')
add_bullet('1.2.3 决策执行：通过工具调用（Action）影响现实')

add_heading('1.3 人机协作决策矩阵', level=3)
add_bullet('1.3.1 执行者能力画像：小龙虾/Hermes/人三维量化')
add_bullet('1.3.2 动态匹配规则：安全→质量→预算→时间优先级')
add_bullet('1.3.3 四种人机协作模式详解')

add_heading('1.4 【案例1】实战案例：构建"智能售后分析专家"', level=3)
add_bullet('1.4.1 需求分析：业务场景识别与约束定义')
add_bullet('1.4.2 智能体搭建：Agent 2.0自主模式配置')
add_bullet('1.4.3 人机分工：AI收集数据 + 人审核决策')

add_para('📊 本章关键产出：智能体五大能力对标表、执行者能力画像、人机协作决策矩阵', indent=False, bold_prefix='')

# 第2章
add_heading('第2章 规划中枢：Agent 2.0的自主编排逻辑', level=2)
add_para('📍 贯穿案例定位：对应MB-003-A智能体搭建，核心能力配置。', indent=False, bold_prefix='')

add_heading('2.1 Agent 2.0自主模式深度解析', level=3)
add_bullet('2.1.1 开启"思考模式（Enable Thinking）"：提升逻辑上限')
add_bullet('2.1.2 ReAct（Reasoning and Acting）框架的工程化实现')
add_bullet('2.1.3 递归分解视角：ReAct = 分解→执行→反馈的闭环')

add_heading('2.2 可观测性：卡片流执行监控', level=3)
add_bullet('2.2.1 Thinking卡片：透视智能体的内部规划逻辑')
add_bullet('2.2.2 Tool Call与Observation：规划与反馈的闭环校验')
add_bullet('2.2.3 元业务执行日志：可追溯、可审计、可优化')

add_heading('2.3 从SDP到智能体：十年理论演进', level=3)
add_bullet('2.3.1 软件定义价格 → 元业务成本评估')
add_bullet('2.3.2 拍卖机制 → 动态匹配算法')
add_bullet('2.3.3 Multi-Agent协商 → 多智能体协作')

add_heading('2.4 【案例2】期末出卷智能体：基于提示词优化器的学术试卷生成', level=3)
add_bullet('2.4.1 任务分解：题型设计 → 难度控制 → 答案生成')
add_bullet('2.4.2 ReAct循环：生成→评估→修正→定稿')
add_bullet('2.4.3 人机协作：AI生成初稿 → 教师审核修改')

add_para('📊 本章关键产出：ReAct框架工程化实现指南、SDP→智能体8项理论映射表、元业务执行日志模板', indent=False, bold_prefix='')

# 第3章
add_heading('第3章 数据基石：数据中心与RAG记忆工程', level=2)
add_para('📍 贯穿案例定位：对应MB-002知识库构建，RAG记忆工程。', indent=False, bold_prefix='')

add_heading('3.1 百炼数据中心管理：从结构化到非结构化', level=3)
add_bullet('3.1.1 结构化数据：Text-to-SQL的逻辑支撑')
add_bullet('3.1.2 非结构化文档：语义碎片的深度加工')
add_bullet('3.1.3 元业务视角：文档收集→清洗→分类→入库')

add_heading('3.2 RAG检索增强生成的底层路径', level=3)
add_bullet('3.2.1 智能切片（Chunking）与向量化（Embedding）实践')
add_bullet('3.2.2 混合检索策略：语义索引与关键词过滤的动态平衡')
add_bullet('3.2.3 检索策略调优：自动化测试 + bad case分析')

add_heading('3.3 记忆系统分级：从会话记忆到长期知识', level=3)
add_bullet('3.3.1 短期记忆：会话上下文管理')
add_bullet('3.3.2 长期记忆：RAG知识库 + 历史案例复用')
add_bullet('3.3.3 元业务相似度匹配：基于向量的案例检索')

add_heading('3.4 【案例3】民典慧析：基于《民法典》知识库的法律条文匹配', level=3)
add_heading('3.5 【案例4】手搓实战：构建"百炼数码2025新品手册"', level=3)

add_heading('3.6 贯穿案例MB-002：知识库构建完整流程', level=3)
add_bullet('3.6.1 文档收集与清洗（小龙虾自动化，成本¥8，6h，质量0.85）')
add_bullet('3.6.2 智能切片与向量化（小龙虾自动化，成本¥15，8h，质量0.88）')
add_bullet('3.6.3 检索策略调优（小龙虾+Hermes协作，成本¥50，12h，质量0.90）')
add_bullet('3.6.4 知识审核（人工执行，安全L2，质量≥0.95）')

add_para('📊 本章关键产出：知识库构建SOP、检索策略调优方法论、MB-002元业务分解表（4个子任务）', indent=False, bold_prefix='')

# 第4章
add_heading('第4章 万能插座：MCP（模型上下文协议）深度集成', level=2)
add_para('📍 贯穿案例定位：对应MB-003-B MCP工具集成，能力扩展。', indent=False, bold_prefix='')

add_heading('4.1 MCP协议：智能体的能力倍增器', level=3)
add_bullet('4.1.1 标准化接口标准：像USB接口一样的万用连接')
add_bullet('4.1.2 兼容性探讨：OpenAPI REST与SSE长连接')
add_bullet('4.1.3 元业务标准接口：MCP = 元业务的"USB接口"')

add_heading('4.2 MCP广场：官方插件与自定义扩展', level=3)
add_bullet('4.2.1 官方核心：联网搜索、Python代码解释器、万相绘图')
add_bullet('4.2.2 业务对齐：将企业ERP/CRM系统封装为自定义工具')
add_bullet('4.2.3 自定义MCP开发：从API到智能体工具的封装流程')

add_heading('4.3 【案例5】基于百炼的导购助手：联动库存API与黄历工具调用', level=3)

add_heading('4.4 贯穿案例MB-003-B：MCP工具集成实战', level=3)
add_bullet('4.4.1 联网搜索MCP（实时查询）')
add_bullet('4.4.2 Python代码解释器MCP（数据计算）')
add_bullet('4.4.3 企业CRM/订单系统API封装（自定义MCP）')
add_bullet('4.4.4 万相绘图MCP（产品图片生成）')

add_para('📊 本章关键产出：MCP接口标准文档、自定义MCP开发指南、MB-003-B元业务分解表', indent=False, bold_prefix='')

# 第5章
add_heading('第5章 积木化开发：组件（Component）的构建与复用', level=2)
add_para('📍 贯穿案例定位：对应MB-003-C组件化封装，模块化开发。', indent=False, bold_prefix='')

add_heading('5.1 百炼组件设计哲学', level=3)
add_bullet('5.1.1 模块化封装：实现智能体功能的解耦与复用')
add_bullet('5.1.2 元能力标准接口：组件 = 元能力的模块化封装')
add_bullet('5.1.3 组件设计模式：意图识别、工单生成、情绪检测')

add_heading('5.2 AppFlow低代码集成与流程编排', level=3)
add_bullet('5.2.1 可视化画布：定义节点流转与变量传递逻辑')
add_bullet('5.2.2 场景适配：将模型能力封装为标准Web服务')
add_bullet('5.2.3 DAG依赖管理：元业务编排的工程实现')

add_heading('5.3 【案例6】营销物料一键生成：通义万相图像组件集成流程', level=3)

add_heading('5.4 贯穿案例MB-003-C：组件化封装实战', level=3)
add_bullet('5.4.1 意图识别组件（分类：咨询/投诉/售后）')
add_bullet('5.4.2 工单生成组件（结构化输出）')
add_bullet('5.4.3 情绪检测组件（敏感客户升级）')
add_bullet('5.4.4 AppFlow流程编排（可视化画布）')

add_para('📊 本章关键产出：组件设计规范、AppFlow流程编排模板、MB-003-C元业务分解表', indent=False, bold_prefix='')

# 第6章
add_heading('第6章 执行核心：沙箱运行与代码解释引擎', level=2)
add_para('📍 贯穿案例定位：对应MB-004系统集成与测试，沙箱自愈。', indent=False, bold_prefix='')

add_heading('6.1 代码解释器（Code Interpreter）运行机制', level=3)
add_bullet('6.1.1 隔离沙箱环境：执行安全性与资源配额管理')
add_bullet('6.1.2 运行时支持：内置Python库环境与环境变量配置')
add_bullet('6.1.3 元业务安全隔离：沙箱 = 元业务的执行安全边界')

add_heading('6.2 符号逻辑与计算增强', level=3)
add_bullet('6.2.1 结构化输出转换：从自然语言到数据对象的精确映射')
add_bullet('6.2.2 系统内部的JSON Schema校验与逻辑验证')

add_heading('6.3 自愈执行闭环：代码纠错循环', level=3)
add_bullet('6.3.1 运行时异常拦截与反馈修正')
add_bullet('6.3.2 长期运行任务的状态保持机制')
add_bullet('6.3.3 异常降级与重试：元业务执行容错机制')

add_heading('6.4 【案例7】企业经营数据可视化大屏：RDS数据库与DataV集成实战', level=3)

add_heading('6.5 贯穿案例MB-004：系统集成与测试', level=3)
add_bullet('6.5.1 功能测试：500个典型场景自动化测试（小龙虾，成本¥30，16h）')
add_bullet('6.5.2 性能测试：并发100/500/1000 QPS（小龙虾，成本¥15，8h）')
add_bullet('6.5.3 用户体验测试：20名测试用户NPS评分（人）')
add_bullet('6.5.4 问题修复：Hermes分析bad case + 小龙虾修复（成本¥200）')

add_para('📊 本章关键产出：沙箱执行安全规范、自愈执行闭环设计模式、MB-004元业务分解表（4个子任务）', indent=False, bold_prefix='')

# 第7章
add_heading('第7章 工程治理：生命周期管理与安全风控', level=2)
add_para('📍 贯穿案例定位：对应MB-003-E安全风控 + MB-005部署运营。', indent=False, bold_prefix='')

add_heading('7.1 账号权限与API Key的生命周期管理', level=3)
add_bullet('7.1.1 RAM细粒度授权：AliyunBailianFullAccess的按需分配')
add_bullet('7.1.2 安全合规底线：API Key轮转与加密存储策略')

add_heading('7.2 内容安全与版本控制', level=3)
add_bullet('7.2.1 版本快照与秒级回滚：保障业务连续性')
add_bullet('7.2.2 敏感词干预与风控拦截规则')

add_heading('7.3 政策合规与智能体标准', level=3)
add_bullet('7.3.1 五大能力对齐（感知/记忆/决策/交互/执行）')
add_bullet('7.3.2 安全合规框架（数据本地化/决策可追溯/人工兜底）')
add_bullet('7.3.3 能力分级与权限控制（L1-L4安全级别）')

add_heading('7.4 【案例8】项目沟通摘要助理：信息脱敏与重点摘要自动化', level=3)

add_heading('7.5 贯穿案例MB-003-E：安全与风控配置', level=3)
add_bullet('7.5.1 API Key加密存储与轮转')
add_bullet('7.5.2 敏感词干预规则配置')
add_bullet('7.5.3 客户数据脱敏处理')
add_bullet('7.5.4 版本快照与回滚机制')
add_bullet('7.5.5 人工升级通道（复杂问题转人工）')

add_heading('7.6 贯穿案例MB-005：上线部署与运营', level=3)
add_bullet('7.6.1 生产部署（人执行，安全L2）')
add_bullet('7.6.2 运营监控（小龙虾自动化）')
add_bullet('7.6.3 持续优化（Hermes分析 + 人决策 + 小龙虾执行）')

add_para('📊 本章关键产出：政策合规Checklist、安全分级与权限控制矩阵、MB-003-E和MB-005元业务分解表', indent=False, bold_prefix='')

# 第8章
add_heading('第8章 应用进阶：多模态感知与复杂集成', level=2)
add_para('📍 贯穿案例定位：对应MB-005多渠道部署运营。', indent=False, bold_prefix='')

add_heading('8.1 多模态感知模型（Qwen-VL）的工程应用', level=3)
add_bullet('8.1.1 图像/视频解析与视觉逻辑推导')
add_bullet('8.1.2 多模态元业务：图片诊断、文档理解、视频分析')

add_heading('8.2 跨渠道集成方案：Webhook与系统回调', level=3)
add_bullet('8.2.1 Webhook事件驱动：实时消息处理')
add_bullet('8.2.2 多渠道消息路由：钉钉/微信/Web/Telegram')
add_bullet('8.2.3 元业务跨渠道执行：同一元业务多平台适配')

add_heading('8.3 【案例9】深圳航空渠道违规行为自动巡检系统实战', level=3)
add_heading('8.4 【案例10】跨系统全链路营销智能客服集成方案', level=3)

add_heading('8.5 贯穿案例MB-005多渠道部署', level=3)
add_bullet('8.5.1 钉钉群机器人配置')
add_bullet('8.5.2 Web端客服界面集成')
add_bullet('8.5.3 多渠道消息统一路由')

add_para('📊 本章关键产出：多模态元业务设计规范、跨渠道集成架构、MB-005多渠道部署方案', indent=False, bold_prefix='')

# ===== 第二部分 =====
add_heading('第二部分：OpenClaw（小龙虾）高阶生态篇（共3章）', level=1)

add_heading('第9章 开放架构：OpenClaw 5层模型设计哲学', level=2)

add_heading('9.1 开放架构演进：百炼能力的生态延伸', level=3)
add_bullet('9.1.1 消息处理层：MsgContext对象的标准化流转')
add_bullet('9.1.2 Provider接入层：屏蔽异构模型与渠道差异')
add_bullet('9.1.3 5层模型与递归框架的映射关系')

add_heading('9.2 异步非阻塞架构与网关设计', level=3)
add_bullet('9.2.1 异步事件驱动：心跳轮询 + cron定时任务')
add_bullet('9.2.2 网关设计：消息路由、限流、熔断')
add_bullet('9.2.3 元业务执行引擎：DAG编排 + 并行调度')

add_heading('9.3 OpenClaw在递归框架中的定位', level=3)
add_bullet('9.3.1 "手脚"角色：自主执行层')
add_bullet('9.3.2 能力优势：主动触发、记忆持久化、多平台、技能系统')
add_bullet('9.3.3 与百炼的互补：百炼（规划+分析）+ OpenClaw（执行+协调）')

add_para('📊 本章关键产出：OpenClaw 5层模型架构图、元业务执行引擎设计、百炼+OpenClaw互补矩阵', indent=False, bold_prefix='')

add_heading('第10章 动态交互：OpenClaw记忆机制与车道调度', level=2)

add_heading('10.1 记忆分级流转：短期、长期与每日汇总', level=3)
add_bullet('10.1.1 短期记忆：会话上下文管理')
add_bullet('10.1.2 长期记忆：MEMORY.md + daily notes持久化')
add_bullet('10.1.3 Daily Summary：每日执行历史自动汇总')
add_bullet('10.1.4 记忆与RAG的融合：OpenClaw记忆 + 百炼知识库')

add_heading('10.2 车道机制（Lane）：高并发下的稳定性保障', level=3)
add_bullet('10.2.1 会话级车道 vs 全局级车道：确保上下文有序一致')
add_bullet('10.2.2 消息幂等性与去重策略')
add_bullet('10.2.3 元业务并发调度：车道 = 元业务的执行通道')

add_para('📊 本章关键产出：记忆分级流转架构图、车道机制配置指南、元业务并发调度策略', indent=False, bold_prefix='')

add_heading('第11章 群体智能：多智能体协作（MAS）设计模式', level=2)

add_heading('11.1 多智能体系统（MAS）协作范式', level=3)
add_bullet('11.1.1 中心化调度（Supervisor）vs 去中心化竞争协作模式')
add_bullet('11.1.2 动态路由：根据任务复杂度分配最佳智能体')
add_bullet('11.1.3 小龙虾与Hermes的协作模式：执行+分析双引擎')

add_heading('11.2 Hermes定位：深度分析引擎', level=3)
add_bullet('11.2.1 "大脑"角色：深度推理、结构化写作、方案对比')
add_bullet('11.2.2 能力画像：成本/速度/质量三维量化')
add_bullet('11.2.3 与小龙虾的分工边界与协作接口')

add_heading('11.3 未来前瞻', level=3)
add_bullet('11.3.1 端云协同：百炼模型与本地私有化算力的互补')
add_bullet('11.3.2 智能体交互标准的未来趋势')
add_bullet('11.3.3 从双智能体到多智能体生态')

add_para('📊 本章关键产出：MAS协作架构设计、小龙虾+Hermes分工协作矩阵、智能体交互标准展望', indent=False, bold_prefix='')

# ===== 第三部分（新增）=====
add_heading('第三部分：递归自主式人机协作篇（新增，共3章）', level=1)

add_heading('第12章 理论框架：递归自主式分解', level=2)

add_heading('12.1 元业务模型：属性、约束与分类', level=3)
add_bullet('12.1.1 元业务定义：可独立执行、可评估、可定价的最小业务单元')
add_bullet('12.1.2 约束属性：max_load、security_level、budget、deadline、quality_min、dependencies')
add_bullet('12.1.3 执行者候选：executor、cost、speed、quality、load')
add_bullet('12.1.4 元业务分类：数据型、分析型、执行型、决策型')

add_heading('12.2 递归分解算法：从复杂任务到原子操作', level=3)
add_bullet('12.2.1 分解规则：原子判断 → 方案生成 → 约束剪枝 → 最优选择 → 递归分解')
add_bullet('12.2.2 约束剪枝：时间/预算/安全/质量/依赖五维约束')
add_bullet('12.2.3 伪代码实现与工程化落地')
add_bullet('12.2.4 纳什均衡思想：局部最优 → 全局近似最优')

add_heading('12.3 动态匹配：小龙虾/Hermes/人的三维决策', level=3)
add_bullet('12.3.1 能力画像量化：成本/速度/质量/安全/创造力/可靠性/可解释性')
add_bullet('12.3.2 匹配决策矩阵：安全优先 → 质量次之 → 预算时间再次')
add_bullet('12.3.3 组合模式：纯AI / AI+人审 / 人主导 / 人机并行')

add_heading('12.4 执行协调：DAG编排与质量验证', level=3)
add_bullet('12.4.1 DAG依赖管理：有依赖按序执行，无依赖并行执行')
add_bullet('12.4.2 异常处理：降级、重试、人工接管')
add_bullet('12.4.3 结果聚合：多执行者结果合并与质量评估')
add_bullet('12.4.4 反馈闭环：质量评估 → 分解策略优化')

add_heading('12.5 从SDP到智能体：专著思想的现代演进', level=3)
add_bullet('12.5.1 软件定义价格 → 元业务成本评估（从资源价格→执行成本）')
add_bullet('12.5.2 拍卖机制 → 动态匹配算法（从资源竞争→能力互补）')
add_bullet('12.5.3 Multi-Agent协商 → 多智能体协作（从价格博弈→任务编排）')
add_bullet('12.5.4 业务聚类 → 元业务分类（从QoS聚类→多维属性聚类）')
add_bullet('12.5.5 流量分类传输 → 优先级调度（从网络流量→任务优先级）')
add_bullet('12.5.6 区块链优先级币 → 可信执行记录（从竞拍公正→执行可审计）')
add_bullet('12.5.7 VTN虚拟租户网络 → 元业务隔离（从网络隔离→执行隔离）')
add_bullet('12.5.8 MHM多跳模型 → DAG依赖管理（从网络路由→任务编排）')

add_para('📊 本章关键产出：元业务属性模型、递归分解算法伪代码、SDP→智能体8项理论映射表、动态匹配决策矩阵', indent=False, bold_prefix='')

add_heading('第13章 人机协作模式', level=2)

add_heading('13.1 能力画像：成本/速度/质量三维量化', level=3)
add_bullet('13.1.1 小龙虾能力画像：成本¥0.01-0.1/次，速度秒级，质量0.80-0.95')
add_bullet('13.1.2 Hermes能力画像：成本¥0.05-0.5/次，速度秒级，质量0.85-0.93')
add_bullet('13.1.3 人能力画像：成本¥5-50/次，速度分钟-小时级，质量0.90-0.99')
add_bullet('13.1.4 能力画像获取方法：基准测试 + 历史数据 + 动态更新')

add_heading('13.2 四种协作模式详解', level=3)
add_bullet('13.2.1 模式A：纯AI执行（小龙虾→Hermes→小龙虾）——常规报告、数据整理')
add_bullet('13.2.2 模式B：AI执行+人审核（小龙虾→人→小龙虾）——重要文档、对外沟通')
add_bullet('13.2.3 模式C：人主导+AI辅助（人→Hermes→人）——战略规划、创意设计')
add_bullet('13.2.4 模式D：人机并行（小龙虾+Hermes+人→聚合→人选择）——方案比选')
add_bullet('13.2.5 模式选择决策树：任务类型 → 安全级别 → 质量要求 → 预算时间')

add_heading('13.3 Hermes定位：深度分析引擎', level=3)
add_bullet('13.3.1 "大脑"角色：深度推理、结构化写作、方案对比、bad case分析')
add_bullet('13.3.2 与小龙虾的分工：小龙虾（手脚/执行）+ Hermes（大脑/分析）')
add_bullet('13.3.3 协作接口：标准化输入输出 + 结果传递 + 质量验证')

add_heading('13.4 任务看板与可视化', level=3)
add_bullet('13.4.1 钉钉集成任务看板：实时进度展示')
add_bullet('13.4.2 元业务执行状态：✅完成 / 🔄执行中 / ⏳等待 / ❌失败')
add_bullet('13.4.3 人工干预入口：[查看详情] [人工干预] [加速执行]')
add_bullet('13.4.4 质量评估展示：实际质量 vs 预期质量 vs 最低要求')

add_heading('13.5 安全分级与人工兜底', level=3)
add_bullet('13.5.1 安全级别L1-L4：公开/内部/机密/绝密')
add_bullet('13.5.2 人工兜底规则：L3/L4必须人执行，L1/L2可AI执行+人审核')
add_bullet('13.5.3 异常升级通道：AI无法处理 → 自动转人工')
add_bullet('13.5.4 审计追溯：所有AI执行过程结构化日志记录')

add_para('📊 本章关键产出：执行者能力画像表（7维度）、四种协作模式详解+适用场景、模式选择决策树、安全分级与人工兜底规则', indent=False, bold_prefix='')

add_heading('第14章 端到端实战', level=2)

add_heading('14.1 案例1：企业级智能客服系统（贯穿案例总结）', level=3)
add_bullet('14.1.1 任务约束：预算¥5000，截止30天，质量≥85%，安全L2')
add_bullet('14.1.2 一级分解：5个元业务（需求/知识库/开发/测试/部署）')
add_bullet('14.1.3 二级分解：20个元业务完整展开')
add_bullet('14.1.4 执行结果：成本¥2230，时间30天，质量0.94')
add_bullet('14.1.5 效果对比：成本降低26%，时间缩短50%，质量损失2%')

add_heading('14.2 案例2：行业研究报告（知识密集型）', level=3)
add_bullet('14.2.1 任务：完成"2026年AI智能体产业发展报告"')
add_bullet('14.2.2 分解：数据收集→文献综述→框架设计→深度分析→报告撰写→最终审核')
add_bullet('14.2.3 执行：小龙虾（数据）+ Hermes（分析）+ 人（决策）')
add_bullet('14.2.4 效果：成本¥58，时间10h，质量0.94 vs 纯人¥300+/20h+/0.96')

add_heading('14.3 案例3：运维自动化（执行密集型）', level=3)
add_bullet('14.3.1 任务：服务器巡检 + 异常诊断 + 自动修复')
add_bullet('14.3.2 分解：监控数据采集→异常检测→根因分析→修复执行→验证确认')
add_bullet('14.3.3 执行：小龙虾（监控+修复）+ Hermes（诊断）+ 人（审核）')
add_bullet('14.3.4 效果：故障响应从小时级降至分钟级，人工干预减少70%')

add_heading('14.4 政策合规对标', level=3)
add_bullet('14.4.1 五大能力对标：感知✅ 记忆✅ 决策✅ 交互✅ 执行✅')
add_bullet('14.4.2 安全合规对标：数据本地化✅ 决策可追溯✅ 人工兜底✅ 用户知情✅')
add_bullet('14.4.3 产业方向对标：智能制造✅ 金融风控✅ 智慧农业✅')
add_bullet('14.4.4 政策响应度评估：完整对齐国家级智能体政策要求')

add_heading('14.5 效果评估方法论', level=3)
add_bullet('14.5.1 效率指标：任务完成时间、并行度、资源利用率')
add_bullet('14.5.2 成本指标：总成本、成本节约率、ROI')
add_bullet('14.5.3 质量指标：实际质量、质量损失率、用户满意度')
add_bullet('14.5.4 安全指标：安全事件数、审计通过率、人工干预率')
add_bullet('14.5.5 持续优化：反馈闭环、能力画像更新、分解策略优化')

add_para('📊 本章关键产出：3个端到端案例完整分解、政策合规对标表、效果评估指标体系、持续优化方法论', indent=False, bold_prefix='')

# ===== 全书结构总览 =====
add_heading('全书结构总览', level=1)

add_table(
    ['部分', '章节', '章名', '核心内容', '贯穿案例对应', '状态'],
    [
        ['第一部分', '第1章', '范式重构', 'LLM→Agentic AI + 五大能力 + 人机决策', 'MB-001需求分析', '优化'],
        ['', '第2章', '规划中枢', 'Agent 2.0 + ReAct + SDP演进', 'MB-003-A智能体搭建', '优化'],
        ['', '第3章', '数据中心', 'RAG记忆 + 检索调优 + 记忆分级', 'MB-002知识库构建', '优化'],
        ['', '第4章', 'MCP集成', '标准化接口 + 自定义MCP', 'MB-003-B工具集成', '保留'],
        ['', '第5章', '组件化开发', '模块化封装 + AppFlow编排', 'MB-003-C组件封装', '优化'],
        ['', '第6章', '沙箱执行', '代码解释器 + 自愈闭环', 'MB-004测试验证', '优化'],
        ['', '第7章', '工程治理', '安全风控 + 政策合规', 'MB-003-E/MB-005', '新增7.3'],
        ['', '第8章', '多模态集成', 'Qwen-VL + 跨渠道', 'MB-005多渠道部署', '优化'],
        ['第二部分', '第9章', 'OpenClaw架构', '5层模型 + 异步网关', '执行协调层', '优化'],
        ['', '第10章', '记忆与车道', '记忆分级 + 车道调度', '记忆系统', '优化'],
        ['', '第11章', '多智能体协作', 'MAS + 小龙虾+Hermes分工', '动态匹配层', '新增11.2'],
        ['第三部分', '第12章', '递归分解框架', '元业务模型 + 递归算法 + SDP演进', '理论底座', '新增'],
        ['', '第13章', '人机协作模式', '能力画像 + 四种模式 + 安全分级', '执行层', '新增'],
        ['', '第14章', '端到端实战', '3个案例 + 政策对标 + 效果评估', '全书总结', '新增'],
    ]
)

# ===== 优化效果对比 =====
add_heading('优化效果对比', level=1)

add_table(
    ['指标', '原版', '优化后', '提升'],
    [
        ['总章节数', '11章', '14章', '+27%'],
        ['贯穿案例', '无', '1个完整案例（20个元业务）', '从0到1'],
        ['理论框架', '工程实践', '递归分解+SDP演进', '质的飞跃'],
        ['人机协作', '隐含', '显性化（决策表+能力画像+模式选择）', '从隐性到显性'],
        ['政策对齐', '无', '第14章完整对标', '从0到1'],
        ['Hermes定位', '缺失', '第13章深度分析引擎', '从缺失到完整'],
        ['专著融合', '缺失', '第12章8项理论映射', '从缺失到完整'],
        ['读者获得感', '学完各章技术', '学完=完成真实项目', '质的飞跃'],
    ]
)

# ===== 附录 =====
add_heading('附录：贯穿案例元业务分解总表', level=1)

add_table(
    ['元业务ID', '元业务名称', '二级数', '成本', '时间', '主要执行者', '对应章节'],
    [
        ['MB-001', '需求分析与方案设计', '4', '¥207', '5天', '人+Hermes', '第1章'],
        ['MB-002', '知识库构建', '4', '¥223', '10天', '小龙虾', '第3章'],
        ['MB-003', '智能体开发', '5', '¥1055', '20天', '小龙虾+人', '第1-5,7章'],
        ['MB-004', '系统集成与测试', '4', '¥545', '25天', '小龙虾+人', '第6-7章'],
        ['MB-005', '上线部署与运营', '3', '¥200', '30天', '人+小龙虾', '第7-8章'],
        ['合计', '', '20', '¥2230', '30天', '混合', '全书'],
    ]
)

# ===== 保存 =====
output_path = '/home/admin/.openclaw/workspace/优化后百炼教材大纲结构-递归自主式分解与人机协作新范式.docx'
doc.save(output_path)
print(f'✅ 优化后大纲已保存: {output_path}')
