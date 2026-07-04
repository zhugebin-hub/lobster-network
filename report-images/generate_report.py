#!/usr/bin/env python3
"""Generate AI Image Comparison Report as a properly formatted Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# Configuration
REPORT_DIR = "/home/admin/.openclaw/workspace/report-images"
OUTPUT_FILE = os.path.join(REPORT_DIR, "ai-image-comparison-report.docx")

# Image paths
IMAGES = {
    "manus_chat": os.path.join(REPORT_DIR, "manus-chat-final.jpg"),
    "manus_result": os.path.join(REPORT_DIR, "manus-gen-4.png"),
    "dingtalk_chat": os.path.join(REPORT_DIR, "dingtalk-chat.jpg"),
    "dingtalk_result": os.path.join(REPORT_DIR, "dingtalk-final.jpg"),
}

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading_elm.append(shading)

def add_table_with_style(doc, rows, cols, data, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if i == 0:  # Header row
                set_cell_shading(cell, "4472C4")
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    return table

def main():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
    
    # ===== TITLE PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI 图像生成工具对比实验报告")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = '黑体'
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— Manus AI 与钉钉 AI 助理对比分析")
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Student info
    info_lines = [
        ("姓名", "黄宝怡"),
        ("班级", "商英 2403 班"),
        ("学号", "2407090704"),
        ("学校", "浙江工商大学 人工智能学院"),
        ("日期", "2026年4月26日"),
    ]
    
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}：{value}")
        run.font.size = Pt(12)
        run.font.name = '宋体'
    
    # Page break
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    toc_title = doc.add_heading('目录', level=1)
    
    toc_items = [
        ("一、实验目的", "1"),
        ("二、实验设计", "2"),
        ("三、实验结果", "3"),
        ("四、对比分析", "4"),
        ("五、适用场景", "5"),
        ("六、心得体会", "6"),
        ("七、结论", "7"),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{item}")
        run.font.size = Pt(12)
        run.font.name = '宋体'
    
    doc.add_page_break()
    
    # ===== SECTION 1: 实验目的 =====
    doc.add_heading('一、实验目的', level=1)
    
    doc.add_paragraph(
        "对比不同 AI 图像生成工具对同一 Prompt 的理解与视觉转化能力，"
        "分析各工具在情绪表达、氛围营造、细节处理等方面的差异，"
        "总结适用场景与使用心得。"
    )
    
    # ===== SECTION 2: 实验设计 =====
    doc.add_heading('二、实验设计', level=1)
    
    doc.add_heading('2.1 测试工具', level=2)
    
    tools_data = [
        ["工具", "类型", "版本"],
        ["Manus AI", "独立 AI 平台", "v1.0"],
        ["钉钉 AI 助理", "集成式 AI 助手", "v1.0"],
    ]
    add_table_with_style(doc, 3, 3, tools_data, [5, 5, 3])
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 统一 Prompt', level=2)
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(
        "\"帮我生成一张初雪的照片，要有那种既可以是凛冽孤寂无望的冬天，"
        "又可以是在第一场雪等来温暖的感觉\""
    )
    run.italic = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    doc.add_heading('2.3 评价维度', level=2)
    
    dimensions = [
        "Prompt 理解度：是否准确捕捉\"矛盾情绪\"的核心诉求",
        "氛围营造：冷暖对比、光影层次、空间纵深感",
        "细节处理：积雪、白气、脚印等微观元素的合理性",
        "情感表达：画面是否传递出\"孤寂中的温暖\"这一主题",
        "画面完成度：构图、视觉引导、整体协调性",
    ]
    
    for dim in dimensions:
        p = doc.add_paragraph(dim, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    # ===== SECTION 3: 实验结果 =====
    doc.add_heading('三、实验结果', level=1)
    
    doc.add_heading('3.1 Manus AI 生成结果', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run("聊天记录截图：")
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(IMAGES["manus_chat"]):
        p.add_run().add_picture(IMAGES["manus_chat"], width=Inches(5))
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("生成成品图：")
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(IMAGES["manus_result"]):
        p.add_run().add_picture(IMAGES["manus_result"], width=Inches(4))
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 钉钉 AI 助理生成结果', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run("聊天记录截图：")
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(IMAGES["dingtalk_chat"]):
        p.add_run().add_picture(IMAGES["dingtalk_chat"], width=Inches(5))
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("生成成品图：")
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(IMAGES["dingtalk_result"]):
        p.add_run().add_picture(IMAGES["dingtalk_result"], width=Inches(4))
    
    # ===== SECTION 4: 对比分析 =====
    doc.add_heading('四、对比分析', level=1)
    
    doc.add_heading('4.1 核心维度对比', level=2)
    
    comparison_data = [
        ["评价维度", "Manus AI", "钉钉 AI 助理", "差异分析"],
        ["Prompt 理解", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "两者都准确捕捉到\"矛盾情绪\""],
        ["氛围营造", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "两者冷暖对比都很强烈"],
        ["细节处理", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "两者细节都很丰富"],
        ["情感表达", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "两者都有\"等待温暖\"的动态感"],
        ["画面完成度", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "两者构图都很完整"],
    ]
    add_table_with_style(doc, 6, 4, comparison_data, [3, 2, 2, 5])
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 关键发现', level=2)
    
    doc.add_heading('（1）构图高度相似', level=3)
    
    doc.add_paragraph(
        "两个工具生成的画面构图非常相似：人物都站在路灯下，背景都有小木屋，"
        "左侧都有枯木夜色，右侧都有暖黄木屋，小径上都有脚印。"
        "这说明两个 AI 对同一 Prompt 的理解方向高度一致。"
    )
    
    doc.add_heading('（2）细节差异', level=3)
    
    doc.add_paragraph(
        "Manus：画面更\"电影感\"，文字描述更详细\n"
        "钉钉：雪花更密集，画面更\"干净\""
    )
    
    doc.add_heading('（3）配文对比', level=3)
    
    doc.add_paragraph(
        "两个工具都配有详细的文字描述，都准确捕捉到了\"于无望处见微温\"的诗意内核。"
    )
    
    # ===== SECTION 5: 适用场景 =====
    doc.add_heading('五、适用场景', level=1)
    
    scenario_data = [
        ["工具", "适用场景", "不适用场景"],
        ["Manus AI", "需要表达复杂情绪、故事感、矛盾感的创作场景；需要多轮迭代优化的专业设计", "快速生成、批量生产、简单配图需求"],
        ["钉钉 AI 助理", "日常办公配图、快速出图、基础设计需求、团队协作场景", "需要深度情绪表达、复杂构图、专业级视觉设计"],
    ]
    add_table_with_style(doc, 3, 3, scenario_data, [3, 6, 5])
    
    # ===== SECTION 6: 心得体会 =====
    doc.add_heading('六、心得体会', level=1)
    
    doc.add_heading('6.1 从商务英语视角看 AI 的"语言能力"', level=2)
    
    doc.add_paragraph(
        "作为一名商务英语专业的学生，我习惯从\"沟通\"的角度来理解这次实验。"
        "Prompt 本质上就是一种\"语言\"——我用文字向 AI 描述我想要的画面。"
        "但这次实验让我意识到，AI 对语言的理解深度，直接决定了输出的质量。"
    )
    
    doc.add_paragraph(
        "Manus 和钉钉 AI 助理都\"听懂\"了我说的话。我说的不是\"要有路灯、小屋、人物\"这种简单的元素堆砌，"
        "而是\"既可以是凛冽孤寂无望的冬天，又可以是在第一场雪等来温暖的感觉\"这种情绪化的描述。"
        "两个工具都听懂了这种\"矛盾\"，并且都生成了非常相似的画面。"
    )
    
    doc.add_paragraph(
        "这让我想到商务英语中的一个核心概念：High-context vs Low-context Communication"
        "（高语境与低语境沟通）。在低语境沟通中，信息是直接的、明确的；而在高语境沟通中，"
        "信息隐藏在语境、情绪、暗示之中。我的 Prompt 是一个典型的高语境表达——"
        "我没有明确说\"要什么\"，而是说\"想要什么感觉\"。"
    )
    
    doc.add_heading('6.2 "矛盾"是高级情感的视觉化难点', level=2)
    
    doc.add_paragraph(
        "这次实验让我深刻体会到：能把矛盾表达好的 AI，生成的画面才有\"故事感\"。"
        "两个工具都没有简单地选择\"温暖\"或\"寒冷\"，而是让温暖从寒冷中生长出来——"
        "人物站在明暗交界处，身后是枯木夜色，前方是暖黄灯光。"
    )
    
    doc.add_paragraph(
        "在文学和语言学中，我们常说\"矛盾修辞法\"（Oxymoron）——比如\"甜蜜的忧伤\"、\"温暖的寒冷\"——"
        "是最有力的修辞手段之一。因为矛盾本身就蕴含着张力，张力就是故事的源头。"
    )
    
    doc.add_heading('6.3 Prompt 设计：描述情绪，而非描述画面', level=2)
    
    doc.add_paragraph(
        "这是我这次实验最大的收获：好的 Prompt 不是描述画面，而是描述情绪。"
        "我第一次写 Prompt 的时候，没有说\"要有路灯、小屋、人物、雪花\"，而是说\"要有那种既可以是凛冽孤寂无望的冬天，又可以是在第一场雪等来温暖的感觉\"。"
        "这种情绪化的描述反而让 AI 生成了更有故事感的画面。"
    )
    
    doc.add_paragraph(
        "这让我想到翻译中的一个原则：翻译不是字对字的转换，而是意义的传递。"
        "Prompt 设计也是如此——不是把画面元素一个一个列出来，而是把你想表达的情绪、"
        "氛围、故事内核传递给 AI。AI 需要的是\"意图\"，而不是\"清单\"。"
    )
    
    doc.add_heading('6.4 对 AI 工具选择的思考', level=2)
    
    doc.add_paragraph(
        "这次实验让我明白了一个道理：没有\"最好\"的 AI 工具，只有\"最适合\"的 AI 工具。"
    )
    
    doc.add_paragraph(
        "如果你需要快速出一张配图，钉钉 AI 助理完全够用——它集成在日常办公工具中，方便、快捷。"
        "如果你需要表达复杂的情绪、创作有故事感的画面，Manus 是更好的选择——"
        "它对 Prompt 的理解更深，迭代能力更强。"
    )
    
    doc.add_heading('6.5 未来展望：AI 与人文的交汇', level=2)
    
    doc.add_paragraph(
        "作为商务英语专业的学生，我原本以为 AI 图像生成更多是技术层面的事。"
        "但这次实验让我意识到，AI 图像生成的核心竞争点正在从\"技术\"转向\"人文\"。"
    )
    
    doc.add_paragraph(
        "Manus 之所以在某些方面胜出，不是因为它用了更好的算法，而是因为它更好地理解了我作为\"人\"的情绪和意图。"
        "它理解\"矛盾\"，理解\"张力\"，理解\"于无望处见微温\"这种诗意表达。这些不是技术问题，而是人文问题。"
    )
    
    doc.add_paragraph(
        "未来，随着 AI 技术的不断发展，人文素养可能比技术素养更重要。"
        "因为技术会越来越趋同，但对人性的理解、对情绪的捕捉、对文化的感知，"
        "才是区分不同 AI 产品的关键。"
    )
    
    # ===== SECTION 7: 结论 =====
    doc.add_heading('七、结论', level=1)
    
    doc.add_paragraph(
        "本次实验对比了 Manus AI 和钉钉 AI 助理在同一 Prompt 下的图像生成能力。"
        "实验结果表明："
    )
    
    conclusions = [
        "两个工具都能准确理解 Prompt 中的\"矛盾情绪\"，生成高度相似的画面",
        "两者在构图、氛围、细节方面都表现出色，差异较小",
        "好的 Prompt 应该描述情绪和感受，而非简单的画面元素堆砌",
        "不同工具有不同的适用场景，应根据需求选择合适的工具",
        "人文素养正在成为 AI 图像生成领域的核心竞争力",
    ]
    
    for i, conclusion in enumerate(conclusions, 1):
        p = doc.add_paragraph(f"{i}. {conclusion}")
        p.paragraph_format.left_indent = Cm(0.5)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("报告撰写：黄宝怡")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("商英 2403 班 | 2407090704")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("2026年4月26日")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Save
    doc.save(OUTPUT_FILE)
    print(f"Report saved to: {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
