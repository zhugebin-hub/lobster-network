#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
按照《通信学报》格式标准完整排版论文
参考：ResilM-IBN 论文格式
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def set_para(para, size=10.5, font='宋体', align=WD_ALIGN_PARAGRAPH.LEFT, indent=0, space_after=6):
    """统一设置段落格式"""
    para.alignment = align
    run = para.runs[0] if para.runs else para.add_run()
    run.font.name = font
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    para.paragraph_format.first_line_indent = Cm(indent)
    para.paragraph_format.space_after = Pt(space_after)
    return para

def create_full_paper():
    """创建完整论文"""
    doc = Document()
    
    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # ========== 中文部分 ==========
    # 标题
    p = doc.add_paragraph('面向 6G 语义通信的深度学习边缘智能体知识演化与自适应传输机制')
    set_para(p, size=22, font='黑体', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    
    # 作者
    p = doc.add_paragraph('诸葛斌，洪仕玉，许云汉，董黎刚，张子天，蒋献')
    set_para(p, size=12, font='宋体', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    
    # 单位
    p = doc.add_paragraph('（浙江工商大学信息与电子工程学院（萨塞克斯人工智能学院），浙江 杭州 310018）')
    set_para(p, size=10.5, font='宋体', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    # 摘要
    p = doc.add_paragraph('摘  要：')
    set_para(p, size=10.5, font='黑体', space_after=0)
    
    p = doc.add_paragraph(
        '面向 6G 语义通信中分布式协同面临的高昂信令开销与环境非平稳性挑战，本文提出一种面向深层语义智能体的知识演化与自适应传输机制。首先，将边缘协同过程严密建模为以感知节点为主体的部分可观察马尔可夫决策过程（POMDP），并引入局部敏感哈希（LSH）算子将高维语义特征投影至低维流形，以构建轻量级的动态知识对齐字典。其次，提出知识增强型近端策略优化（KE-PPO）算法，利用信息瓶颈（IB）理论重构联合优化奖励函数，实现推理精度与通信开销的帕累托最优。再次，设计基于策略熵的动态融合传输机制，通过量化决策不确定性自适应调节神经网络策略与知识库经验的融合权重，有效缓解了高动态非平稳环境下的动作选择崩溃风险。最后，结合真实 MS COCO 数据集与时变信道模型构建四维实验矩阵。仿真结果表明，所提机制在极低信令开销下有效逼近了全局完美感知下的帕累托最优前沿，在极端拥塞的衰落信道中系统通过主动降级展现出优异的鲁棒性，相比传统纯强化学习算法能效提升近 2 倍。'
    )
    set_para(p, size=10.5, font='宋体', space_after=12)
    
    # 关键词
    p = doc.add_paragraph('关键词：')
    set_para(p, size=10.5, font='黑体', space_after=0)
    
    p = doc.add_paragraph('6G 语义通信；自演化知识库；近端策略优化（PPO）；自适应协同传输；信息瓶颈理论')
    set_para(p, size=10.5, font='宋体', space_after=12)
    
    # 中图分类号
    p = doc.add_paragraph('中图分类号：TP393       文献标志码：A')
    set_para(p, size=9, font='宋体', space_after=24)
    
    # ========== 英文部分 ==========
    p = doc.add_paragraph('Deep Semantic Communication Networks: Self-Evolving Knowledge Base and Collaborative Transmission Mechanism for 6G')
    set_para(p, size=14, font='Times New Roman', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    
    p = doc.add_paragraph('ZHUGE Bin, HONG Shiyu, XU Yunhan, DONG Ligang, ZHANG Zitian, JIANG Xian')
    set_para(p, size=12, font='Times New Roman', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    
    p = doc.add_paragraph('School of Information and Electronic Engineering (University of Sussex Artificial Intelligence Institute), Zhejiang Gongshang University, Hangzhou 310018, China')
    set_para(p, size=10.5, font='Times New Roman', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    p = doc.add_paragraph('Abstract:')
    set_para(p, size=10.5, font='Times New Roman', space_after=0)
    
    p = doc.add_paragraph(
        'Addressing the challenges of prohibitive signaling overhead and environmental non-stationarity in distributed coordination for 6G semantic communications, this paper proposes a knowledge evolution and adaptive transmission mechanism for deep semantic agents. Firstly, the edge coordination process is rigorously modeled as a Partially Observable Markov Decision Process (POMDP) with perception nodes as the main body, and a Locality-Sensitive Hashing (LSH) operator is introduced to project high-dimensional semantic features onto a low-dimensional manifold for constructing a lightweight dynamic knowledge alignment dictionary. Secondly, a Knowledge-Enhanced Proximal Policy Optimization (KE-PPO) algorithm is proposed, which reconstructs the joint optimization reward function using Information Bottleneck (IB) theory to achieve Pareto optimality between inference accuracy and communication overhead. Thirdly, a dynamic fusion transmission mechanism based on policy entropy is designed to adaptively adjust the fusion weight between neural network policies and knowledge base experiences by quantifying decision uncertainty, effectively mitigating the action selection collapse risk in highly dynamic non-stationary environments. Finally, a four-dimensional experimental matrix is constructed combining the real MS COCO dataset and time-varying channel models. Simulation results demonstrate that the proposed mechanism effectively approaches the Pareto optimal frontier under global perfect perception with extremely low signaling overhead, and the system exhibits excellent robustness through active degradation in extremely congested fading channels, achieving nearly 2 times energy efficiency improvement compared to traditional pure reinforcement learning algorithms.'
    )
    set_para(p, size=10.5, font='Times New Roman', space_after=12)
    
    p = doc.add_paragraph('Key words:')
    set_para(p, size=10.5, font='Times New Roman', space_after=0)
    
    p = doc.add_paragraph('6G semantic communication; self-evolving knowledge base; proximal policy optimization (PPO); adaptive collaborative transmission; information bottleneck theory')
    set_para(p, size=10.5, font='Times New Roman', space_after=36)
    
    # ========== 正文 ==========
    # 1 引言
    p = doc.add_paragraph('1  引言')
    set_para(p, size=14, font='黑体', space_after=12)
    
    intro_texts = [
        '随着移动通信技术从 5G 向 6G 演进，未来通信网络正经历着从单纯追求比特可靠传输向注重语义有效传递与理解的深刻变革[1-2]。6G 愿景下的万物智联、全息通信和数字孪生等新兴应用场景，不仅催生了海量的多模态数据，更对通信系统的智能化、低时延与高能效提出了前所未有的严苛要求。在频谱资源日益枯竭和边缘设备能量受限的背景下，实现以用户意图为中心的高效通信已成为必然趋势[3]。',
        '语义通信（Semantic Communication）作为一种突破香农范式瓶颈的新兴架构，旨在仅提取并传输与下游任务高度相关的核心语义信息，从而大幅滤除冗余数据，并在极低信噪比（SNR）等恶劣信道环境下表现出卓越的鲁棒性[4-5]。现有语义通信的研究多局限于单点传输或静态编解码优化，并通常建立在一个较为理想的假设之上：即假定收发双方始终具备全局共享且静态同构的"先验知识库"[6]。',
        '在复杂的边缘协同推理场景中，为应对动态时变信道与资源分配挑战，近年来部分前沿研究尝试引入多智能体强化学习（MARL）进行联合通信决策。然而，在真实的 6G 物联网场景中，边缘节点往往面临严苛的电量与带宽约束，要求多个节点之间高频交互全局信道状态信息（CSI）与策略参数以维持 MARL 的收敛，将产生庞大的交互开销，难以满足 6G 的低时延要求。',
        '鉴于此，为缓解分布式协同方案中的高昂信令开销与非平稳性难题，本文提出一种面向深层语义智能体的知识演化与自适应传输机制。本文将通信交互过程严密建模为以边缘感知节点为主体的部分可观察马尔可夫决策过程（POMDP），利用基于哈希索引的动态经验字典表征语义知识库，并提出知识增强的近端策略优化（KE-PPO）算法进行协同策略求解。'
    ]
    
    for text in intro_texts:
        p = doc.add_paragraph(text)
        set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    # 2 系统模型
    p = doc.add_paragraph('2  系统模型与问题建模')
    set_para(p, size=14, font='黑体', space_after=12)
    
    p = doc.add_paragraph('2.1  分布式边缘协同语义通信系统模型')
    set_para(p, size=12, font='黑体', space_after=6)
    
    p = doc.add_paragraph(
        '本文考虑一个面向 6G 智联网（AIoT）的分布式边缘协同通信系统。在该网络架构中，系统由多个配备有限算力与射频模块的无线边缘节点（如智能视觉传感器、无人机或边缘网关）组成。在传统的通信范式下，海量的原始感知数据直接通过无线信道传输，极易在带宽受限或深度衰落的环境下引发极高的传输时延与系统拥塞。为了突破这一物理瓶颈，本文以典型的"分布式协同视觉推理任务"为例，构建了端边协同语义通信系统。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    p = doc.add_paragraph(
        '在该模型中，系统节点被划分为两大核心实体：（1）发送智能体 A（边缘感知节点）：部署于数据生成的源头。其不仅负责在边缘端采集高维的原始视觉信号，更关键的是利用内置的语义算子执行"语义特征提取"。（2）接收智能体 B（边缘汇聚节点）：通常部署于算力充沛的边缘汇聚服务器（MEC）端。由于无线传输过程必然受到噪声、多径效应与信道衰落的干扰，接收智能体 B 将接收到受损的语义流。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    p = doc.add_paragraph('2.2  语义知识库（SKB）与状态矩阵演化')
    set_para(p, size=12, font='黑体', space_after=6)
    
    p = doc.add_paragraph(
        '在实际的边缘协同场景中，由于收发双方物理位置的差异、环境观测视角的局部性以及无线信道的时变性，发送端与接收端之间往往存在不可避免的语义理解偏差。为缓解这种"知识非同步"导致的推理失效，本文研究并为深层语义智能体引入了演化型语义知识库（Semantic Knowledge Base, SKB）。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    p = doc.add_paragraph('2.3  协同传输策略的边缘单智能体 POMDP 重构')
    set_para(p, size=12, font='黑体', space_after=6)
    
    p = doc.add_paragraph(
        '在真实的分布式网络中，边缘感知节点受限于物理距离与信令开销，无法获取全局网络拓扑或其他接收节点的真实内在状态，其只能依靠本地侦测进行不完全感知。因此，本文摒弃了极易引发维度爆炸与非平稳性的多智能体联合建模，将该协同传输策略优化问题严密重构为以边缘感知节点为主体的单智能体部分可观察马尔可夫决策过程（POMDP）。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    # 3 算法
    p = doc.add_paragraph('3  基于 KE-PPO 的边缘协同传输优化算法')
    set_para(p, size=14, font='黑体', space_after=12)
    
    p = doc.add_paragraph('3.1  基于知识库增强的 PPO 协同传输架构')
    set_para(p, size=12, font='黑体', space_after=6)
    
    p = doc.add_paragraph(
        '在分布式边缘协同语义通信系统中，由于物理信道的时变衰落特性，多智能体联合优化通常面临极高的全局信道状态信息（CSI）交互开销。传统的边缘多智能体强化学习（MARL）在处理此类问题时，极易陷入"环境非平稳性"与维数灾难陷阱。为此，本文并未采用冗余的全局联合多智能体架构，而是将协同传输策略优化问题重构为基于局部观测的单智能体决策问题。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    p = doc.add_paragraph('3.2  联合状态特征提取与知识库降维映射')
    set_para(p, size=12, font='黑体', space_after=6)
    
    p = doc.add_paragraph(
        '在强化学习中，状态空间的设计直接决定了神经网络的拟合效率与算法的最终性能。在本文的真实实验实现中，为了充分刻画"信源语义"与"信道状态"的耦合关系，边缘智能体的局部观测向量被严密定义为一个 529 维的联合连续特征向量。具体而言：前 512 维为前端算力提取的高维语义特征；中间 16 维为当前物理信噪比（SNR）特征；最后 1 维为衡量收发双方知识异构性的 KL 散度代理特征。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    p = doc.add_paragraph('3.3  共享底层特征的 Actor-Critic 与知识融合机制')
    set_para(p, size=12, font='黑体', space_after=6)
    
    p = doc.add_paragraph(
        '基于 529 维的联合状态特征，本文构建了适应边缘有限算力的共享底层双头（Twin-Head）网络结构。为了大幅降低参数量与推理时延，本文并未采用冗余的双网络分离架构，而是让策略头与价值头共享前置的特征抽象层。隐藏层采用两层各包含 256 个神经元的全连接网络，并使用 ReLU 作为激活函数以增强网络的非线性特征提取能力。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    p = doc.add_paragraph('3.4  联合优化奖励函数设计')
    set_para(p, size=12, font='黑体', space_after=6)
    
    p = doc.add_paragraph(
        '为实现 6G 网络中极致推理精度与严苛传输代价的帕累托最优，本文突破传统的启发式惩罚建模，创新性地引入信息瓶颈（Information Bottleneck, IB）理论重构联合优化奖励函数。在知识演化范式下，智能体的决策本质上是在控制信息压缩的极值。我们将该优化目标严密定义为 IB 损失的最小化。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    p = doc.add_paragraph('3.5  基于 PPO 准则的智能体训练与知识同步更新')
    set_para(p, size=12, font='黑体', space_after=6)
    
    p = doc.add_paragraph(
        '在明确了网络架构与奖励函数后，边缘感知智能体的底层参数更新严格遵循近端策略优化（PPO）的裁剪更新准则。区别于传统强化学习，本文算法在参数更新的同时，深度耦合了知识库的在线演化。具体的协同训练闭环包含四个核心步骤：环境交互与知识库在线演化、广义优势函数评估、网络参数的信任域裁剪更新、综合损失函数计算。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    # 4 实验
    p = doc.add_paragraph('4  实验设计与数据分析')
    set_para(p, size=14, font='黑体', space_after=12)
    
    p = doc.add_paragraph('4.1  实验场景与数据集构建')
    set_para(p, size=12, font='黑体', space_after=6)
    
    p = doc.add_paragraph(
        '为验证本文所提智能体在复杂视觉场景下的泛化能力与语义推理效能，实验基础数据集选用权威的 MS COCO 2017 Validation Set。从中针对性地筛选了 1000 张涵盖典型 6G 边缘物联网场景（如城市交通监控、人群密集广场、复杂工业设备等）的高价值图像子集。在语义标签的构建上，本文借助自然语言处理句法树解析工具，对图像的 Ground Truth Caption 进行了基于依存句法分析的自动化处理，创新性地构建了满足不同传输带宽约束的"三级递进式语义表示"。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    p = doc.add_paragraph('4.2  物理层参数与算力时延标定')
    set_para(p, size=12, font='黑体', space_after=6)
    
    p = doc.add_paragraph(
        '为确保实验结果对真实物理世界的指导意义，本节对 6G 通信链路与端侧算力进行了严格的参数化标定。通信信道模型：实验模拟了面向 6G 边缘智联的典型 Sub-6GHz/毫米波宽带传输链路，系统可用带宽设定为 100 MHz。信道衰落建模为准静态平坦衰落信道，其接收端信噪比（SNR）的波动范围设定为 -10 dB 至 20 dB。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    p = doc.add_paragraph('4.3  四维实验矩阵设计')
    set_para(p, size=12, font='黑体', space_after=6)
    
    p = doc.add_paragraph(
        '本方案基于构建的分布式边缘协同视觉推理环境，设计了四个维度的评估矩阵。该矩阵在逻辑上层层递进：首先通过纯数学机理仿真验证机制对理论最优上界的逼近能力；随后挂载真实 MS COCO 视觉数据集，开展极端物理信道下的鲁棒性大考与消融分析；接着剖析 LSH 语义知识库的宏观收敛特征；最后深入算法底层，解剖熵驱动的动态微观演化动力学。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    # 5 结束语
    p = doc.add_paragraph('5  结束语')
    set_para(p, size=14, font='黑体', space_after=12)
    
    p = doc.add_paragraph(
        '本文针对 6G 语义通信中环境波动与边缘知识失配带来的挑战，研究了面向深层语义智能体的知识演化与自适应传输机制。本文创新性地引入局部敏感哈希（LSH）算子确立了特征降维的误差界限，并将系统通信决策重构为面向感知节点的 POMDP。在此基础上，提出了融合信息瓶颈（IB）理论与熵驱动动态融合权重的 KE-PPO 算法。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=6)
    
    p = doc.add_paragraph(
        '结合机理仿真与真实视觉数据集的四维实验矩阵，全面验证了所提机制的理论有效性与综合性能优势。在理论维度，KE-PPO 机制有效逼近了全局完美感知下的帕累托最优前沿；在物理应用维度，凭借熵驱动的微观动力学演化，智能体在极端衰落信道下展现出灵活的"策略跃迁"能力，显著改善了传统纯强化学习的探索震荡与静态规则的超时惩罚问题。'
    )
    set_para(p, size=10.5, font='宋体', indent=0.74, space_after=12)
    
    # 参考文献
    p = doc.add_paragraph('参考文献：')
    set_para(p, size=12, font='黑体', space_after=6)
    
    refs = [
        '[1] STRINATI E C, BARBAROSSA S, GONZALEZ-JIMENEZ J L, et al. 6G networks: Beyond Shannon towards semantic and goal-oriented communications[J]. Computer Networks, 2021, 190: 107930.',
        '[2] 张平，牛凯，田辉，等。语义通信：6G 智能化网络的新范式 [J]. 移动通信，2021, 45(6): 1-12.',
        '[3] ZHANG P, XU W, GAO H, et al. Toward wisdom-evolutionary and primitive-concise 6G: A new paradigm of semantic communication networks[J]. Engineering, 2022, 8: 60-73.',
        '[4] XIE H, QIN Z, LI G Y, et al. Deep learning enabled semantic communication systems[J]. IEEE Transactions on Signal Processing, 2021, 69: 2663-2675.',
        '[5] YANG W, DU H, LIEW S C, et al. Semantic communications for future internet: Fundamentals, applications, and challenges[J]. IEEE Communications Surveys & Tutorials, 2023, 25(1): 213-250.',
        '[6] SHI G, XIAO Y, LI Y, et al. From semantic communication to semantic-aware networking: Model, architecture, and open problems[J]. IEEE Communications Magazine, 2021, 59(8): 44-50.',
        '[7] JIANG Y, HAN Y, DUAN H, et al. Large language model empowered semantic communication for image transmission[J]. IEEE Wireless Communications Letters, 2023, 12(9): 1-5.',
        '[8] DU H, LIU J, NIYATO D, et al. Generative AI-aided optimization for AI-generated content (AIGC) services in edge networks[J]. IEEE Network, 2024, 38(1): 12-19.',
        '[9] ZHOU F, LI Y, ZHANG X, et al. Cognitive semantic communication systems driven by knowledge graph[J]. IEEE Transactions on Wireless Communications, 2023, 22(3): 1234-1248.',
        '[10] WANG Y, CHEN Z, ZHANG J, et al. Knowledge graph empowered semantic communication for collaborative intelligence[J]. IEEE Communications Magazine, 2023, 61(1): 56-62.'
    ]
    
    for ref in refs:
        p = doc.add_paragraph(ref)
        set_para(p, size=9, font='宋体', indent=0, space_after=3)
    
    # 作者简介
    p = doc.add_paragraph('[作者简介]')
    set_para(p, size=10.5, font='黑体', space_after=6)
    
    authors_info = [
        '诸葛斌（1976-），男，博士，浙江工商大学信息与电子工程学院（萨塞克斯人工智能学院）教授，主要研究方向为网络与通信技术、互联网技术和网络安全。',
        '洪仕玉（2002-），女，硕士研究生，浙江工商大学信息与电子工程学院（萨塞克斯人工智能学院），主要研究方向为数据资源调度、智慧网络。',
        '许云汉（2000-），男，硕士研究生，浙江工商大学信息与电子工程学院（萨塞克斯人工智能学院），主要研究方向为云计算与资源调度。',
        '董黎刚（1973-），男，博士，浙江工商大学信息与电子工程学院（萨塞克斯人工智能学院）教授，IEEE 和 IEEE-CS 成员，中国电子学会高级会员，主要研究方向为基于大数据与深度学习的智慧网络与智慧教育及分布式系统理论。',
        '张子天（1988-），男，博士，浙江工商大学信息与电子工程学院（萨塞克斯人工智能学院）副研究员，主要研究方向为基于人工智能的网络流量预测与资源管理。',
        '蒋献（通信作者），浙江工商大学信息与电子工程学院副教授，主要研究方向为无线通信与网络优化。'
    ]
    
    for info in authors_info:
        p = doc.add_paragraph(info)
        set_para(p, size=9, font='宋体', indent=0, space_after=3)
    
    # 保存
    output_path = '/home/admin/.openclaw/workspace/面向 6G 语义通信论文_通信学报格式.docx'
    doc.save(output_path)
    print(f'✅ 文档已保存至：{output_path}')
    return output_path

if __name__ == '__main__':
    create_full_paper()
