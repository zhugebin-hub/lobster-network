#!/usr/bin/env python3
"""
婚姻工资溢价实验报告 - 简洁版（无图片，精简文字，保留数据表格）
"""

import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = "/home/admin/.openclaw/workspace/stata_experiment2/output"

with open(os.path.join(OUTPUT_DIR, "regression_results.json"), 'r', encoding='utf-8') as f:
    results = json.load(f)
with open(os.path.join(OUTPUT_DIR, "descriptive_stats.json"), 'r', encoding='utf-8') as f:
    desc_stats = json.load(f)

doc = Document()

def sf(run, size=12, name='仿宋_GB2312', bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.name = name
    run.bold = bold
    run.italic = italic
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)

style = doc.styles['Normal']
style.font.name = '仿宋_GB2312'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

def add_header_row(table, headers):
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                sf(r, 10, bold=True)

def fill_table(table, data, first_col_left=False):
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            c = table.rows[ri].cells[ci]
            c.text = val
            for p in c.paragraphs:
                if first_col_left and ci == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    sf(r, 10)

# ========== 标题 ==========
title = doc.add_heading('《婚姻工资溢价的计量分析》实验报告', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    sf(r, 18, '黑体', bold=True)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
for t in ['课程：计量经济学\n', '姓名：陈政道\n', '日期：2026年6月9日']:
    sf(info.add_run(t))
doc.add_paragraph('')

# ========== 一、摘要 ==========
doc.add_heading('一、摘要', level=2)
doc.add_paragraph(
    '本研究基于WAGE1数据集（526个观测值），使用扩展的明瑟收入方程分析婚姻状态对工资的影响。'
    '结果表明，在控制教育、经验、性别和任期后，已婚男性工资平均高出约19.5%（p<0.001）。'
    '同时验证了教育回报率约9.3%、经验边际回报递减、以及约27.4%的性别工资差距。'
)

# ========== 二、引言 ==========
doc.add_heading('二、引言', level=2)
doc.add_paragraph(
    '已婚男性平均工资显著高于未婚男性，称为"婚姻工资溢价"。'
    '本研究旨在检验：在控制教育、经验、性别和任期后，该溢价是否依然存在。'
)

# ========== 三、数据与变量 ==========
doc.add_heading('三、数据与变量', level=2)
doc.add_paragraph('数据来源：WAGE1（Wooldridge教材配套数据），526个个体观测值。')

doc.add_heading('3.1 变量定义与描述性统计', level=3)
var_table = doc.add_table(rows=8, cols=6)
var_table.style = 'Table Grid'
var_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(var_table, ['变量', '定义', '均值', '标准差', '最小值', '最大值'])
var_data = [
    ('ln_wage', '对数小时工资', '2.010', '0.552', '0.209', '3.689'),
    ('educ', '受教育年限（年）', '11.99', '2.97', '5', '21'),
    ('exper', '工作经验（年）', '17.40', '8.30', '0', '45'),
    ('female', '性别（女=1）', '0.460', '0.499', '0', '1'),
    ('married', '婚姻（已婚=1）', '0.614', '0.487', '0', '1'),
    ('tenure', '工作任期（年）', '4.68', '3.55', '0', '19'),
    ('hours', '年工作小时', '39.63', '8.30', '20', '69'),
]
fill_table(var_table, var_data)
doc.add_paragraph('')

doc.add_heading('3.2 婚姻分组比较', level=3)
grp_table = doc.add_table(rows=5, cols=3)
grp_table.style = 'Table Grid'
grp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(grp_table, ['指标', '已婚组 (n=323)', '未婚组 (n=203)'])
grp_data = [
    ('平均工资（$/小时）', '9.54', '7.35'),
    ('对数工资均值', '2.101', '1.864'),
    ('受教育年限（年）', '12.1', '11.7'),
    ('工作经验（年）', '17.7', '16.9'),
]
fill_table(grp_table, grp_data)
doc.add_paragraph('')
doc.add_paragraph('原始婚姻工资差距：23.8%。已婚组教育和经验略高，提示需控制混杂因素。')

doc.add_heading('3.3 理论模型', level=3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p.add_run(
    'ln(wage) = β₀ + β₁·married + β₂·educ + β₃·exper + β₄·exper² + β₅·female + β₆·tenure + β₇·tenure² + ε'
), 11, bold=True)

# ========== 四、实证策略 ==========
doc.add_heading('四、实证策略', level=2)
doc.add_paragraph(
    'OLS估计，逐步加入控制变量：'
    '模型1（仅婚姻）→ 模型2（+教育）→ 模型3（+经验）→ 模型4（+性别+任期）。'
)

# ========== 五、回归结果 ==========
doc.add_heading('五、回归结果', level=2)

reg_table = doc.add_table(rows=15, cols=5)
reg_table.style = 'Table Grid'
reg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(reg_table, ['变量', '模型1', '模型2', '模型3', '模型4'])

reg_rows = [
    ('married', '0.2376***', '0.2018***', '0.1936***', '0.1948***'),
    ('', '(0.0484)', '(0.0417)', '(0.0399)', '(0.0379)'),
    ('educ', '—', '0.0933***', '0.0944***', '0.0930***'),
    ('', '', '(0.0068)', '(0.0066)', '(0.0063)'),
    ('exper', '—', '—', '0.0391***', '0.0404***'),
    ('', '', '', '(0.0079)', '(0.0075)'),
    ('exper_sq', '—', '—', '-0.0007***', '-0.0007***'),
    ('', '', '', '(0.0002)', '(0.0002)'),
    ('female', '—', '—', '—', '-0.2739***'),
    ('', '', '', '', '(0.0370)'),
    ('tenure', '—', '—', '—', '-0.0251*'),
    ('', '', '', '', '(0.0140)'),
    ('const', '1.8637***', '0.7678***', '0.3319***', '0.5287***'),
    ('', '(0.0379)', '(0.0867)', '(0.1065)', '(0.1074)'),
]
fill_table(reg_table, reg_rows, first_col_left=True)
doc.add_paragraph('')

doc.add_heading('5.1 统计量汇总', level=3)
stat_table = doc.add_table(rows=6, cols=5)
stat_table.style = 'Table Grid'
stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(stat_table, ['统计量', '模型1', '模型2', '模型3', '模型4'])
stat_data = [
    ('观测值', '526', '526', '526', '526'),
    ('R-squared', '0.044', '0.295', '0.357', '0.424'),
    ('Adjusted R²', '0.042', '0.292', '0.352', '0.416'),
    ('F统计量', '24.13', '109.34', '72.21', '54.45'),
    ('Prob > F', '0.0000', '0.0000', '0.0000', '0.0000'),
]
fill_table(stat_table, stat_data, first_col_left=True)

note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(note_p.add_run('注：* p<0.10, ** p<0.05, *** p<0.01；括号内为标准误'), 9, italic=True)

# ========== 六、结论 ==========
doc.add_heading('六、结论', level=2)
doc.add_paragraph(
    '（1）婚姻工资溢价显著：控制其他因素后，已婚者工资高出约19.5%（p<0.001）。'
)
doc.add_paragraph(
    '（2）教育回报率稳健：约9.3%，各模型基本一致。'
)
doc.add_paragraph(
    '（3）经验回报递减：经验系数为正，平方项为负，符合理论预期。'
)
doc.add_paragraph(
    '（4）性别差距显著：女性工资低约27.4%（p<0.001）。'
)
doc.add_paragraph(
    '（5）模型解释力良好：R²=0.424，解释约42%的工资变异。'
)

doc.add_heading('6.1 局限性', level=3)
doc.add_paragraph(
    '（1）内生性：婚姻非随机分配，可能存在选择偏差。\n'
    '（2）遗漏变量：能力、性格等不可观测因素未控制。\n'
    '（3）横截面数据：无法区分生产力效应和选择效应。'
)

doc.add_heading('6.2 改进方向', level=3)
doc.add_paragraph(
    '（1）使用面板数据+固定效应模型。\n'
    '（2）工具变量法解决内生性。\n'
    '（3）异质性分析（分性别、教育水平等）。'
)

# ========== 附录 ==========
doc.add_page_break()
doc.add_heading('附录：Stata Do-file 代码', level=2)

do_file_path = "/home/admin/.openclaw/workspace/stata_experiment2/marriage_wage_analysis.do"
with open(do_file_path, 'r', encoding='utf-8') as f:
    do_code = f.read()

p = doc.add_paragraph()
sf(p.add_run(do_code), 9, 'Courier New')

output_path = os.path.join(OUTPUT_DIR, "婚姻工资溢价实验报告_简洁版.docx")
doc.save(output_path)
print(f"✓ 简洁版报告已保存: {output_path}")
print(f"✓ 共 {len(doc.paragraphs)} 段落, {len(doc.tables)} 个表格")
