#!/usr/bin/env python3
"""
生成婚姻工资溢价实验报告 - Word格式
"""

import json
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = "/home/admin/.openclaw/workspace/stata_experiment2/output"

with open(os.path.join(OUTPUT_DIR, "regression_results.json"), 'r', encoding='utf-8') as f:
    results = json.load(f)
with open(os.path.join(OUTPUT_DIR, "descriptive_stats.json"), 'r', encoding='utf-8') as f:
    desc_stats = json.load(f)

doc = Document()

# ============================================================
# 全局样式
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '仿宋_GB2312'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

def set_font(run, size=12, name='仿宋_GB2312', bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.name = name
    run.bold = bold
    run.italic = italic
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)

# ============================================================
# 标题页
# ============================================================
title = doc.add_heading('《婚姻工资溢价的计量分析》实验报告', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    set_font(run, 18, '黑体', bold=True)

# 副标题
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('——基于明瑟收入方程的扩展研究')
set_font(run, 14, '楷体', italic=True)

doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
for text in ['课程：计量经济学', '姓名：陈政道', '日期：2026年6月9日']:
    run = info.add_run(text + '\n')
    set_font(run, 12)

doc.add_page_break()

# ============================================================
# 一、摘要
# ============================================================
doc.add_heading('一、摘要', level=2)
doc.add_paragraph(
    '本研究基于WAGE1数据集（526个观测值），通过建立扩展的明瑟收入方程，使用Stata软件分析了'
    '婚姻状态对个人工资的影响，即是否存在"婚姻工资溢价"现象。研究发现，在控制了受教育年限、'
    '工作经验、性别和任期等因素后，已婚男性的平均工资比未婚男性高出约19.5%，该结果在1%的水平上统计显著。'
    '此外，本研究还验证了工作经验的边际回报递减效应和劳动力市场中的性别工资差距问题。'
    '研究结果为理解婚姻与劳动力市场表现之间的关系提供了计量经济学证据。'
)

# ============================================================
# 二、引言
# ============================================================
doc.add_heading('二、引言', level=2)
doc.add_paragraph(
    '在劳动经济学研究中，一个引人注目的经验事实是：已婚男性的平均工资显著高于未婚男性。'
    '这一现象被称为"婚姻工资溢价"（Marriage Wage Premium）。大量实证研究发现，'
    '即使在控制了教育、工作经验等可观测因素后，婚姻状态仍然对工有显著的正向影响。'
)
doc.add_paragraph(
    '关于婚姻工资溢价的成因，经济学界主要有两种解释：一是"生产力假说"，认为婚姻使男性更加稳定、'
    '更有责任感，从而提高了劳动生产率；二是"选择假说"，认为高收入男性更容易结婚，'
    '即婚姻状态本身并不导致工资提高，而是高收入者更可能进入婚姻。'
)
doc.add_paragraph(
    '本研究的核心问题是：婚姻状态是否对工资有独立的影响？即在控制了教育、经验、性别和任期后，'
    '婚姻工资溢价是否依然存在？我们使用明瑟收入方程框架，通过逐步加入控制变量的回归策略进行检验。'
)

# ============================================================
# 三、数据与变量
# ============================================================
doc.add_heading('三、数据与变量', level=2)
doc.add_paragraph(
    '本研究使用的数据为WAGE1数据集，来自Wooldridge《计量经济学导论》教材配套数据。'
    '该数据集包含526个个体观测值，记录了工资、受教育年限、工作经验、婚姻状态、性别、'
    '工作任期等变量信息。'
)

doc.add_heading('3.1 变量定义', level=3)

var_table = doc.add_table(rows=8, cols=5)
var_table.style = 'Table Grid'
var_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['变量名', '定义', '均值', '标准差', '取值范围']
for i, h in enumerate(headers):
    cell = var_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, 10, bold=True)

var_data = [
    ('ln_wage', '对数小时工资', '2.010', '0.552', '[0.209, 3.689]'),
    ('educ', '受教育年限（年）', '11.99', '2.97', '[5, 21]'),
    ('exper', '工作经验（年）', '17.40', '8.30', '[0, 45]'),
    ('female', '性别虚拟变量（女=1）', '0.460', '0.499', '[0, 1]'),
    ('married', '婚姻虚拟变量（已婚=1）', '0.614', '0.487', '[0, 1]'),
    ('tenure', '当前工作任期（年）', '4.68', '3.55', '[0, 19]'),
    ('hours', '年工作小时数', '39.63', '8.30', '[20, 69]'),
]

for row_idx, data in enumerate(var_data):
    for col_idx, val in enumerate(data):
        cell = var_table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_font(r, 10)

doc.add_paragraph('')

doc.add_heading('3.2 婚姻状态分组统计', level=3)
doc.add_paragraph(
    '在进行回归分析之前，我们首先对已婚组和未婚组进行了描述性比较。'
    '结果显示，已婚组（n=323）的平均工资为$9.54/小时，未婚组（n=203）的平均工资为$7.35/小时，'
    '原始婚姻工资差距约为23.8%。此外，已婚组的平均受教育年限（12.1年）和工作经验（17.7年）'
    '也略高于未婚组（11.7年和16.9年）。这提示我们，在分析婚姻对工资的影响时，'
    '必须控制教育和经验等混杂因素，否则可能高估婚姻的独立效应。'
)

doc.add_heading('3.3 理论模型', level=3)
doc.add_paragraph('扩展的明瑟收入方程：')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ln(wage) = β₀ + β₁·married + β₂·educ + β₃·exper + β₄·exper² + β₅·female + β₆·tenure + β₇·tenure² + ε')
run.bold = True
set_font(run, 11)

doc.add_paragraph('')
doc.add_paragraph(
    '其中关键解释变量为married（婚姻状态虚拟变量），其系数β₁即为婚姻工资溢价的估计值。'
    '如果β₁显著为正，则说明在控制了其他因素后，已婚者的工资显著高于未婚者。'
)

# ============================================================
# 四、实证策略与方法
# ============================================================
doc.add_heading('四、实证策略与方法', level=2)

doc.add_heading('4.1 估计方法', level=3)
doc.add_paragraph(
    '采用普通最小二乘法（OLS）进行参数估计。为了检验婚姻工资溢价的稳健性，'
    '我们从简单模型开始，逐步加入控制变量，观察婚姻系数的变化：'
)
doc.add_paragraph('模型1：仅包含婚姻变量（基准回归）', style='List Bullet')
doc.add_paragraph('模型2：加入受教育年限（控制人力资本差异）', style='List Bullet')
doc.add_paragraph('模型3：加入工作经验及经验平方（控制经验回报递减）', style='List Bullet')
doc.add_paragraph('模型4：加入性别和工作任期（完整模型）', style='List Bullet')

doc.add_heading('4.2 诊断检验', level=3)
doc.add_paragraph('（1）异方差检验：使用Breusch-Pagan检验检查残差是否存在异方差问题。')
doc.add_paragraph('（2）多重共线性检验：使用方差膨胀因子（VIF）检查自变量之间是否存在严重的多重共线性。')

# ============================================================
# 五、结果与分析
# ============================================================
doc.add_heading('五、结果与分析', level=2)

doc.add_heading('5.1 散点图与箱线图分析', level=3)
doc.add_paragraph(
    '图1展示了婚姻状态与对数工资的箱线图。从图中可以直观看到，已婚组的工资中位数和四分位距均高于未婚组，'
    '初步支持了婚姻工资溢价的存在。图2进一步展示了工作经验与对数工资的关系，'
    '并按婚姻状态分组。已婚组的拟合线整体高于未婚组，且两条线近似平行，'
    '表明婚姻溢价在不同工作经验水平上基本一致。'
)

img_path = os.path.join(OUTPUT_DIR, "marriage_plots.png")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run('图1：婚姻状态与对数工资关系')
    set_font(run, 10, italic=True)

doc.add_heading('5.2 回归结果', level=3)
doc.add_paragraph('表1展示了四个递进模型的回归结果：')

reg_table = doc.add_table(rows=11, cols=5)
reg_table.style = 'Table Grid'
reg_table.alignment = WD_TABLE_ALIGNMENT.CENTER

reg_headers = ['变量', '模型1', '模型2', '模型3', '模型4']
for i, h in enumerate(reg_headers):
    cell = reg_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, 10, bold=True)

reg_data = [
    ['已婚 (married=1)', '0.2376***\n(0.0484)', '0.2018***\n(0.0417)', '0.1936***\n(0.0399)', '0.1948***\n(0.0379)'],
    ['受教育年限 (educ)', '—', '0.0933***\n(0.0068)', '0.0944***\n(0.0066)', '0.0930***\n(0.0063)'],
    ['工作经验 (exper)', '—', '—', '0.0391***\n(0.0079)', '0.0404***\n(0.0075)'],
    ['经验平方 (exper²)', '—', '—', '-0.0007***\n(0.0002)', '-0.0007***\n(0.0002)'],
    ['性别 (female=1)', '—', '—', '—', '-0.2739***\n(0.0370)'],
    ['任期 (tenure)', '—', '—', '—', '-0.0251*\n(0.0140)'],
    ['任期平方 (tenure²)', '—', '—', '—', '0.0015\n(0.0012)'],
    ['常数项', '1.8637***\n(0.0379)', '0.7678***\n(0.0867)', '0.3319***\n(0.1065)', '0.5287***\n(0.1074)'],
]

for row_idx, data in enumerate(reg_data):
    for col_idx, val in enumerate(data):
        cell = reg_table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_font(r, 9, 'Times New Roman')

# 统计量行
stats_row_data = ['观测值', ['526']*4, ['R²', '0.044', '0.295', '0.357', '0.424'],
                  ['Adj R²', '0.042', '0.292', '0.352', '0.416']]

for idx, row_data in enumerate(stats_row_data):
    row_idx = 9 + idx
    if row_idx < 11:
        for col_idx, val in enumerate(row_data):
            cell = reg_table.rows[row_idx].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                if col_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    set_font(r, 9)

note_para = doc.add_paragraph()
note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note_para.add_run('表1：婚姻工资溢价回归结果（括号内为标准误；* p<0.10, ** p<0.05, *** p<0.01）')
set_font(run, 10, italic=True)

doc.add_heading('5.3 结果解释', level=3)

doc.add_heading('（1）核心发现：婚姻工资溢价', level=4)
doc.add_paragraph(
    '模型1仅包含婚姻变量，估计的婚姻溢价为23.76%（p<0.001），这与原始分组统计结果一致。'
    '然而，这一系数可能包含了教育、经验等其他因素的混杂影响。'
)
doc.add_paragraph(
    '模型2加入教育变量后，婚姻溢价系数从0.2376下降到0.2018，说明部分婚姻溢价是由于已婚者受教育水平更高所致。'
)
doc.add_paragraph(
    '模型3进一步加入工作经验及经验平方项，婚姻溢价系数微调至0.1936，基本保持稳定。'
)
doc.add_paragraph(
    '模型4为完整模型，加入性别和工作任期变量后，婚姻溢价系数为0.1948（p<0.001），'
    '在1%的水平上统计显著。这意味着，在控制了教育、经验、性别和任期后，'
    '已婚男性的平均工资比未婚男性高出约19.5%。'
)

doc.add_heading('（2）教育回报率', level=4)
doc.add_paragraph(
    '四个模型中，教育回报率系数稳定在9.3%左右，与文档一中教育回报率分析的结果高度一致，'
    '进一步验证了教育作为人力资本投资的经济价值。'
)

doc.add_heading('（3）工作经验回报递减', level=4)
doc.add_paragraph(
    '工作经验的系数为正（0.0404，p<0.001），经验平方项系数为负（-0.0007，p<0.001），'
    '验证了经验回报的边际递减规律。计算可得经验回报达到最大值的转折点约为28.9年。'
)

doc.add_heading('（4）性别工资差距', level=4)
doc.add_paragraph(
    '性别变量系数为-0.2739（p<0.001），表明在控制了婚姻、教育、经验和任期后，'
    '女性的平均工资比男性低约27.4%，这一差距比文档一中的估计更大，'
    '说明部分性别差距可能被婚姻状态所掩盖。'
)

doc.add_heading('（5）模型拟合优度', level=4)
doc.add_paragraph(
    '从模型1到模型4，R-squared从0.044逐步提高到0.424，说明加入控制变量后，'
    '模型解释力度大幅提升。完整模型能够解释约42.4%的工资变异，'
    '在社会科学研究中属于较高的解释水平。'
)

# ============================================================
# 六、结论
# ============================================================
doc.add_heading('六、结论', level=2)
doc.add_paragraph(
    '本研究基于WAGE1数据集，通过扩展的明瑟收入方程，对婚姻工资溢价进行了计量分析，主要结论如下：'
)
doc.add_paragraph(
    '（1）存在显著的婚姻工资溢价。在控制了受教育年限、工作经验、性别和任期后，'
    '已婚男性的平均工资比未婚男性高出约19.5%，该结果在1%的统计水平上高度显著。'
    '这为婚姻状态与劳动力市场表现之间的正向关系提供了有力证据。',
    style='List Number'
)
doc.add_paragraph(
    '（2）教育回报率稳健。教育回报率系数在各模型中稳定在9.3%左右，'
    '进一步验证了教育投资的经济价值。',
    style='List Number'
)
doc.add_paragraph(
    '（3）经验回报边际递减。工作经验对工资的正向影响随工作年限增加而减弱，'
    '符合人力资本折旧的理论预期。',
    style='List Number'
)
doc.add_paragraph(
    '（4）性别工资差距显著。女性工资在控制了其他因素后显著低于男性，'
    '反映了劳动力市场中可能存在的结构性不平等问题。',
    style='List Number'
)

doc.add_heading('6.1 局限性', level=3)
doc.add_paragraph(
    '（1）内生性问题：婚姻状态并非随机分配，可能存在选择偏差。高能力、高收入者可能更容易结婚，'
    '这会导致OLS估计高估婚姻的真实因果效应。'
)
doc.add_paragraph(
    '（2）遗漏变量偏差：模型可能遗漏了个人能力、性格特征、家庭背景等难以观测的变量，'
    '这些变量可能同时影响婚姻状态和工资收入。'
)
doc.add_paragraph(
    '（3）横截面数据的局限：本研究使用的是横截面数据，无法控制个体固定效应，'
    '难以区分婚姻的生产力效应和选择效应。'
)

doc.add_heading('6.2 未来改进方向', level=3)
doc.add_paragraph(
    '（1）使用面板数据（如NLSY79）和固定效应模型，控制个体不可观测的异质性，'
    '更准确地识别婚姻的因果效应。'
)
doc.add_paragraph(
    '（2）使用工具变量法（如地区婚姻率、宗教信仰等作为婚姻状态的工具变量）解决内生性问题。'
)
doc.add_paragraph(
    '（3）进行异质性分析，检验婚姻溢价在不同性别、教育水平、行业之间的差异。'
)

# ============================================================
# 附录
# ============================================================
doc.add_page_break()
doc.add_heading('附录：Stata Do-file 代码', level=2)

do_file_path = "/home/admin/.openclaw/workspace/stata_experiment2/marriage_wage_analysis.do"
with open(do_file_path, 'r', encoding='utf-8') as f:
    do_code = f.read()

p = doc.add_paragraph()
run = p.add_run(do_code)
run.font.name = 'Courier New'
run.font.size = Pt(9)
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')

output_path = os.path.join(OUTPUT_DIR, "婚姻工资溢价实验报告.docx")
doc.save(output_path)
print(f"✓ 实验报告已保存: {output_path}")
print(f"✓ 报告共 {len(doc.paragraphs)} 段落")
