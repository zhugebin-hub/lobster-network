#!/usr/bin/env python3
"""Generate Word document for 叶畏兵 paper on tears theology - revised version."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# --- Default font ---
style = doc.styles['Normal']
font = style.font
font.name = '仿宋_GB2312'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
pf = style.paragraph_format
pf.first_line_indent = Cm(0.74)
pf.alignment = 2
pf.line_spacing = Pt(25)
pf.space_before = Pt(0)
pf.space_after = Pt(0)

def add_blank_line():
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def add_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = 1
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_author(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = 1
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '仿宋_GB2312'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

def add_meta_line(label, content):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.alignment = 2
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(25)
    run1 = p.add_run(label)
    run1.font.size = Pt(12)
    run1.font.bold = True
    run1.font.name = '黑体'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run2 = p.add_run(content)
    run2.font.size = Pt(12)
    run2.font.name = '仿宋_GB2312'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

def add_heading_1(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.alignment = 2
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(25)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_heading_2(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.alignment = 2
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(25)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_paragraph(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.74)
    pf.alignment = 2
    pf.line_spacing = Pt(25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '仿宋_GB2312'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

# ============================================================
# TITLE & META
# ============================================================
add_title('眼泪的神学：从尼希米的"坐下哭泣"看基督徒家国情感的正当性')
add_author('叶畏兵')

add_meta_line('摘\u3000要：', '本文关注中国教会中一个长期存在却鲜被正视的现象：信徒为国族苦难或同胞困境流泪时，常被劝勉"要刚强""要有信心"，仿佛哭泣本身等同于灵性软弱或信心不足。本文以《尼希米记》第一章"坐下哭泣"的叙事为切入点，结合旧约哀歌传统与新约耶稣、保罗的眼泪叙事，论证眼泪在圣经传统中并非信心的对立面，而是一种有力量的信仰表达。在此基础上，本文反思中国教会"刚强叙事"的历史与文化根源，提出眼泪应当被恢复为正当的神学范畴。进一步地，本文以尼希米的家国牵挂为范式，论证中国基督徒对民族命运的情感关切具有圣经根据。最后，本文从讲台信息和牧养关怀两个维度提出实践建议。')

add_meta_line('关键词：', '眼泪神学；尼希米记；家国情感；情感神学；基督教中国化')

add_meta_line('作者简介：', '叶畏兵，台州市椒江基督教堂教牧人员。')

add_blank_line()

# ============================================================
# READ MD FILE AND CONVERT
# ============================================================
with open('/home/admin/.openclaw/workspace/叶畏兵_眼泪神学_去AI修订版.md', 'r') as f:
    content = f.read()

# Skip title/meta part, start from first heading
lines = content.split('\n')

in_main = False
for line in lines:
    stripped = line.strip()
    if not stripped:
        continue
    if stripped == '# 眼泪的神学：从尼希米的"坐下哭泣"看基督徒家国情感的正当性':
        continue
    if stripped == '叶畏兵':
        continue
    if stripped.startswith('摘'):
        continue
    if stripped.startswith('关'):
        continue
    if stripped.startswith('作者'):
        continue
    if stripped.startswith('# '):
        continue

    if stripped.startswith('## '):
        in_main = True
        heading_text = stripped[3:].strip()
        if heading_text.startswith('一、') or heading_text.startswith('二、') or heading_text.startswith('三、') or heading_text.startswith('四、') or heading_text.startswith('五、') or heading_text == '结语' or heading_text == '参考文献':
            add_heading_1(heading_text)
        else:
            # check if it's a subheading
            add_heading_2(heading_text)
    elif in_main and stripped:
        add_paragraph(stripped)

# --- References ---
refs = [
    '[1]《圣经》（和合本），南京：中国基督教协会，2000年。',
    '[2]Brueggemann, Walter. The Message of the Psalms: A Theological Commentary. Minneapolis: Augsburg Publishing House, 1984.',
    '[3]Williamson, H.G.M. Ezra, Nehemiah. Word Biblical Commentary, Vol. 16. Waco: Word Books, 1985.',
    '[4]Yamauchi, Edwin M. "Nehemiah." In The Expositor\'s Bible Commentary, Vol. 3, edited by Frank E. Gaebelein, 637-724. Grand Rapids: Zondervan, 1990.',
    '[5]Roberts, J.J.M. "Nehemiah." In The New Interpreter\'s Bible, Vol. III, 815-918. Nashville: Abingdon Press, 1999.',
    '[6]丁光训：《丁光训文集》，南京：译林出版社，1998年。',
    '[7]丁光训：《神学思考的足迹》，北京：宗教文化出版社，2005年。',
    '[8]赵晓阳编：《基督教中国化研究》，北京：社会科学文献出版社，2013年。',
    '[9]卓新平：《当代西方新教神学》，上海：上海三联书店，1998年。',
    '[10]许志伟：《基督教神学基础》，北京：宗教文化出版社，2006年。',
    '[11]尤思德：《和而不同：基督教神学与中国文化的对话》，香港：汉语圣经协会，2009年。',
]

for ref in refs:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.alignment = 2
    pf.line_spacing = Pt(25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(ref)
    run.font.size = Pt(10.5)
    run.font.name = '仿宋_GB2312'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

output_path = '/home/admin/.openclaw/workspace/叶畏兵_眼泪的神学_修订版.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
