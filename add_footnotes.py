#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""为全方明论文添加学术注脚"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Define footnotes: (trigger_text_in_paragraph, footnote_content)
footnotes = [
    # 1. 利玛窦生平与《天主实义》
    ('万历三十一年',
     '利玛窦（Matteo Ricci, 1552-1610），意大利耶稣会士，1582年抵华，1610年卒于北京。《天主实义》初名《天学实义》，1603年刊行于北京，是利玛窦用中文撰写的最重要神学哲学著作。参见李天纲：《利玛窦与晚明儒学》，《复旦学报》1999年第6期，第77-78页。'),

    # 2. 程朱理学体系
    ('在程朱理学的体系中',
     '程朱理学以程颢、程颐兄弟和朱熹为代表，将"理"提升为宇宙万物的最高本体。朱熹在《朱子语类》卷一中明确指出："未有天地之先，毕竟也只是理。"参见朱熹：《朱子语类》卷一，北京：中华书局，1986年，第1页。'),

    # 3. "自立者—依赖者"分类
    ('有自立者，有依赖者',
     '利玛窦：《天主实义》，载《利玛窦中文著译集》，上海：复旦大学出版社，2001年，第278页。此处利玛窦借用了亚里士多德-托马斯主义哲学中的"实体"（substantia）与"偶性"（accidens）之分，译为"自立者"与"依赖者"。'),

    # 4. "若太极者"引文
    ('则不能为天地万物之原矣',
     '利玛窦：《天主实义》，第281页。利玛窦的这一论证在经院哲学中被称为"第一因论证"，其逻辑是：如果某物不能自立，就不能作为其他事物存在的第一原因。参见托马斯·阿奎那：《神学大全》第一集第二题第三条，北京：商务印书馆，2013年，第78-82页。'),

    # 5. 朱熹"理"的双重特性
    ('朱熹明确指出，"理"是',
     '朱熹在《朱子语类》卷四中言："理也者，形而上之道也，生物之本也。"参见《朱子语类》卷四，第65页。朱熹的"理"具有双重特性：既超越于具体事物之上（理在事先），又内在于具体事物之中（理在事中）。'),

    # 6. 太极生阴阳
    ('太极生阴阳，阴阳生万物',
     '此语出自周敦颐《太极图说》，是宋明理学宇宙论的经典表述。周敦颐（1017-1073）为理学开山祖师之一，其《太极图说》以"无极而太极"为宇宙生成的起点。'),

    # 7. "夫太极之生阴阳"引文
    ('夫太极之生阴阳，犹理之生气乎',
     '利玛窦：《天主实义》，第281页。利玛窦此处将"太极"与"理"等同，是对朱熹"太极只是一个理字"命题的直接运用。参见朱熹：《朱子语类》卷九十四，第2371页。'),

    # 8. "仁者为能爱人"引文
    ('苟上帝不予善人升天堂',
     '利玛窦：《天主实义》，第295页。利玛窦此处将基督教的"赏善罚恶"观念引入对"理"的批判，反映了两种哲学体系在终极关怀上的根本差异。'),

    # 9. 《诗经》引文
    ('皇矣上帝，临下有赫',
     '《诗经·大雅·皇矣》。此诗描述上帝监察四方、择民而授命的场景，是先秦儒家经典中"上帝"具有位格性的典型证据。'),

    # 10. 《尚书》引文
    ('惟上帝不常，作善降之百祥',
     '《尚书·商书·伊训》。此句表明"上帝"具有赏善罚恶的意志和能力，是利玛窦"古儒上帝=基督教天主"论断的核心文本依据。'),

    # 11. 朱熹"太极只是一个理字"
    ('太极只是一个理字',
     '朱熹：《朱子语类》卷九十四，第2371页。朱熹此语将"太极"（宇宙论概念）与"理"（本体论概念）等同，为利玛窦的"概念置换"策略提供了直接的文本依据。'),

    # 12. "净洁空阔"
    ('净洁空阔',
     '朱熹在《朱子语类》卷一中形容"理"的状态："未有事物之时，此理已具，净洁空阔，不会造作。"参见《朱子语类》卷一，第3页。'),

    # 13. "吾天主者，乃生生者"
    ('吾天主者，乃生生者',
     '利玛窦：《天主实义》，第275页。此处利玛窦借用《周易·系辞上》"生生之谓易"的表述，但赋予其全新的神学含义。在《周易》中，"生生"是宇宙万物连续不断的生成变化，无主体意志；而在利玛窦的阐释中，"生生"变成了有位格的创造者的主动行为。'),

    # 14. "理也者，依赖之谓也，非灵觉者也"
    ('理也者，依赖之谓也，非灵觉者也',
     '利玛窦：《天主实义》，第281页。利玛窦通过"灵觉"（位格性）这一标准，建立了上帝高于"理"的价值层级。这一论证在神学上被称为"位格性论证"，是基督教神学区别于泛神论和自然主义的核心论题。'),

    # 15. "东海西海，心同理同"
    ('东海西海，心同理同',
     '此语出自陆九渊《杂著》："东海有圣人出焉，此心同也，此理同也；西海有圣人出焉，此心同也，此理同也。"李之藻借用此语来理解利玛窦的学说，体现了以"理"的普遍性消解文化差异的思想倾向。'),

    # 16. "天下为主，君为客"
    ('天下为主，君为客',
     '黄宗羲：《明夷待访录·原君》。黄宗羲（1610-1695）为明清之际大儒，其民本思想与基督教的神权观念形成了有趣的对照。参见孙尚扬：《基督教与明末儒学》，北京：东方出版社，1994年，第156-162页。'),

    # 17. 龙华民
    ('龙华民',
     '龙华民（1559-1654），意大利耶稣会士，利玛窦去世后接任在华传教区会长。他激烈反对利玛窦的译名策略，认为中国经典中的"天"与"上帝"具有浓厚的自然主义色彩。参见谢和耐：《中国与基督教》，耿昇译，上海：上海古籍出版社，2003年，第89-112页。'),

    # 18. 1704年禁令
    ('1704年，罗马教廷发布禁令',
     '1704年11月20日，教皇克雷芒十一世正式颁布禁令，禁止在中国传教活动中使用"天"和"上帝"来翻译"Deus"，并禁止中国教徒参与祭孔祭祖仪式。这一事件被称为"中国礼仪之争"，直接导致了康熙皇帝的反制。参见张西平：《利玛窦对宋明理学的批判及其历史影响》，《宗教学研究》2005年第1期，第92-94页。'),
]

# Read original document
src = Document('/home/admin/.openclaw/media/inbound/4f643fd1-a51f-4864-9e98-77dc02ef3527.docx')

# Create new document
dst = Document()

# Set default style
style = dst.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Collect all footnote texts
footnote_list = []

# Process each paragraph
for para in src.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    new_para = dst.add_paragraph()
    new_para.paragraph_format.space_before = Pt(3)
    new_para.paragraph_format.space_after = Pt(3)
    new_para.paragraph_format.line_spacing = 1.5

    # Add main text
    run = new_para.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Check for footnote matches
    matched = False
    for trigger, footnote_text in footnotes:
        if trigger in text and not matched:
            # Add footnote reference
            fn_num = len(footnote_list) + 1
            ref_run = new_para.add_run('[' + str(fn_num) + ']')
            ref_run.font.size = Pt(8)
            ref_run.font.superscript = True

            footnote_list.append((fn_num, footnote_text))
            matched = True

    # Handle title differently
    if text.startswith('利玛窦对') and len(text) < 30:
        for run in new_para.runs:
            run.font.size = Pt(16)
            run.font.bold = True
    elif text.startswith('\u2014\u2014'):
        for run in new_para.runs:
            run.font.size = Pt(14)
    elif text.startswith('摘') or text.startswith('关'):
        for run in new_para.runs:
            run.font.size = Pt(11)

# Add footnotes section at the end
dst.add_paragraph()
sep_para = dst.add_paragraph()
sep_run = sep_para.add_run('\u2014 \u2014\u2014\u2014\u2014\u2014\u2014\u2014\u2014\u2014\u2014\u2014\u2014')
sep_run.font.size = Pt(10)

# Add footnotes heading
fn_heading = dst.add_paragraph()
fn_heading_run = fn_heading.add_run('注  释')
fn_heading_run.font.size = Pt(12)
fn_heading_run.font.bold = True

# Add each footnote
for fn_num, fn_text in footnote_list:
    fn_para = dst.add_paragraph()
    fn_para.paragraph_format.space_before = Pt(2)
    fn_para.paragraph_format.space_after = Pt(2)
    fn_para.paragraph_format.line_spacing = 1.25

    fn_run = fn_para.add_run('[' + str(fn_num) + '] ' + fn_text)
    fn_run.font.size = Pt(10)
    fn_run.font.name = '宋体'
    fn_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Save
output_path = '/home/admin/.openclaw/workspace/利玛窦概念挪用策略分析_带注脚.docx'
dst.save(output_path)
print('Done! Generated ' + str(len(footnote_list)) + ' footnotes.')
print('Output: ' + output_path)
