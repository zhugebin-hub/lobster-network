#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Fix references in the US-China Trade War report."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document('/home/admin/.openclaw/workspace/中美贸易战读写译报告.docx')

# Remove the old references section (everything from "参考文献" onward)
new_paras = []
skip = False
for p in doc.paragraphs:
    if p.text.strip() == '参考文献':
        skip = True
        continue
    if skip:
        continue
    new_paras.append(p)

# We need to rebuild the document properly
# Actually, let's just remove trailing paragraphs after "参考文献"
para_count = len(doc.paragraphs)
ref_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '参考文献':
        ref_idx = i
        break

if ref_idx is not None:
    # Delete all paragraphs from ref_idx to end
    for i in range(len(doc.paragraphs) - 1, ref_idx - 1, -1):
        p = doc.paragraphs[i]
        # Remove the paragraph from its parent element
        p_elem = p._element
        p_elem.getparent().remove(p_elem)

# Add corrected references
refs = [
    '[1] 商务部. 关于中美经贸摩擦的事实与中方立场（白皮书）[R]. 北京: 国务院新闻办公室, 2018.',
    '[2] 余淼杰. 中美贸易战的经济效应分析[J]. 国际经济评论, 2018(5): 9-23.',
    '[3] 刘劲松. 中美经贸摩擦的根源与前景[J]. 世界经济与政治, 2019(4): 4-22.',
    '[4] 鞠建东, 侯雪雁. 全球价值链视角下的中美贸易摩擦[J]. 管理世界, 2019, 35(5): 36-48.',
    '[5] Amiti M, Redding S J, Weinstein D E. The Impact of the 2018 Trade War on U.S. Prices and Welfare[J]. Journal of Economic Perspectives, 2019, 33(4): 187-210.',
    '[6] Bown C P. The US-China Trade War and Phase One Agreement[R]. Peterson Institute for International Economics, Policy Brief 21-3, 2021.',
    '[7] WTO. World Trade Statistical Review 2020[R]. Geneva: World Trade Organization, 2020.',
    '[8] UNCTAD. World Investment Report 2022[R]. Geneva: United Nations Conference on Trade and Development, 2022.',
    '[9] 国家统计局. 2020年国民经济和社会发展统计公报[R]. 北京: 中国统计出版社, 2021.',
    '[10] 中国现代国际关系研究院. 国际战略与安全形势评估（2020/2021）[M]. 北京: 时事出版社, 2021.',
    '[11] 张宇燕, 徐秀军. 中美经贸关系与国际秩序变革[J]. 世界经济与政治, 2019(8): 4-25.',
    '[12] Irwin D A. Clashing over Commerce: A History of US Trade Policy[M]. Chicago: University of Chicago Press, 2017.',
    '[13] 中国海关总署. 2019年中国对外贸易统计数据[EB/OL]. [2025-06-01]. http://www.customs.gov.cn.',
    '[14] 彼得森国际经济研究所. US-China Trade War: Tariffs and Trade Data[EB/OL]. [2025-06-01]. https://www.piie.com/research/trade-war.',
]

# Add "参考文献" heading
from docx.text.paragraph import Paragraph

def add_head(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    return p

def add_para(doc, text, indent=True, left_indent=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if left_indent:
        p.paragraph_format.left_indent = left_indent
        p.paragraph_format.first_line_indent = Cm(0)
    _add_mixed_run(p, text)
    return p

def _add_mixed_run(para, text):
    buf = ''; is_ascii = None
    for ch in text:
        a = ord(ch) < 128
        if is_ascii is None: is_ascii = a
        if a != is_ascii:
            _flush(para, buf, is_ascii); buf = ''; is_ascii = a
        buf += ch
    if buf: _flush(para, buf, is_ascii)

def _flush(para, text, is_ascii):
    r = para.add_run(text); r.font.size = Pt(12)
    if is_ascii: r.font.name = 'Times New Roman'
    else: r.font.name = '宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Fix heading style too
for lv in (1, 2):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = '黑体'
    hs.font.size = Pt(14) if lv == 1 else Pt(12)
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()  # spacer
add_head(doc, '参考文献', level=1)

for ref in refs:
    add_para(doc, ref, indent=False)

out = '/home/admin/.openclaw/workspace/中美贸易战读写译报告_修正版.docx'
doc.save(out)
print(f'Done: {out}')
