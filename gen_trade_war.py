#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate Word doc: US-China Trade War report."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(3.17)
sec.right_margin = Cm(3.17)

sty = doc.styles['Normal']
sty.font.name = '宋体'
sty.font.size = Pt(12)
sty.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = sty.paragraph_format
pf.line_spacing = Pt(20)
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.first_line_indent = Cm(0.74)

for lv in (1, 2, 3):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = '黑体'
    hs.font.size = Pt(14) if lv == 1 else Pt(12)
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hpf = hs.paragraph_format
    hpf.line_spacing = Pt(24) if lv == 1 else Pt(20)
    hpf.space_before = Pt(12) if lv == 1 else Pt(6)
    hpf.space_after = Pt(6)
    hpf.first_line_indent = None

def add_para(text, indent=True, left_indent=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if left_indent:
        p.paragraph_format.left_indent = left_indent
        p.paragraph_format.first_line_indent = Cm(0)
    _add_mixed_run(p, text)
    return p

def add_head(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    return p

def _add_mixed_run(para, text):
    buf = ''; is_ascii = None
    for ch in text:
        a = ord(ch) < 128
        if is_ascii is None: is_ascii = a
        if a != is_ascii:
            _flush(para, buf, is_ascii); buf = ''; is_ascii = a
        buf += ch
    if buf: _flush(para, buf, is_ascii)

def _flush(para, text, is_ascii):
    r = para.add_run(text); r.font.size = Pt(12)
    if is_ascii: r.font.name = 'Times New Roman'
    else: r.font.name = '宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def items(*paras):
    for p in paras:
        add_para(p, left_indent=Cm(0.74))

def s1(t, *ps):
    add_head(t, level=1)
    for p in ps: add_para(p)

def s2(t, *ps):
    add_head(t, level=2)
    for p in ps: add_para(p)

# Title
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run('中美贸易战的成因、影响与对策分析')
tr.font.name = '黑体'; tr.font.size = Pt(22); tr.bold = True
tr.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
tp.paragraph_format.line_spacing = Pt(36)
tp.paragraph_format.space_before = Pt(36)
tp.paragraph_format.space_after = Pt(18)
doc.add_paragraph()

# ============================================================
# Content
# ============================================================

s1('一、引言',
    '2018年3月，美国政府依据《1974年贸易法》第301条款，宣布对中国输美商品加征关税，标志着中美贸易战的正式爆发。此后，两国经历了多轮关税升级、贸易谈判和协议签署，中美经贸关系进入了一个充满不确定性的新阶段。',
    '中美贸易战不仅是两国之间的经贸摩擦，更是全球经贸格局深度调整的重要信号。作为全球第一大和第二大经济体，中美两国的经贸关系直接影响全球产业链布局、国际贸易秩序和多边贸易体制的稳定。据世界贸易组织（WTO）统计，中美贸易战导致2019年全球货物贸易量增速降至1.2%，为2009年以来的最低水平（WTO, 2020）。',
    '本文旨在系统分析中美贸易战的成因、影响，并提出应对策略，为理解中美经贸关系的未来走向提供参考。',
)

s1('二、中美贸易战的成因分析',
    '中美贸易战的爆发并非偶然事件，而是多种因素长期积累的结果。从经济、政治和战略三个维度来看，其成因主要包括以下几个方面：',
)

s2('（一）贸易失衡与结构性矛盾',
    '中美贸易失衡是贸易战最直接的导火索。自2001年中国加入世界贸易组织（WTO）以来，中国对美贸易顺差持续扩大。据中国海关总署数据，2017年中美贸易顺差达到2,758亿美元，占美国全球贸易逆差的34%。',
    '美方认为，贸易失衡的根源在于中国经济的结构性问题：一是产业政策倾斜，中国政府通过补贴、税收优惠和信贷支持等方式扶持重点产业，被认为构成了不公平竞争；二是市场准入限制，部分行业（如金融、电信、互联网）对外资设置较高门槛；三是知识产权保护不足，美方指责中国通过强制技术转让等方式获取美国企业的核心技术。',
)

s2('（二）科技竞争与战略焦虑',
    '中美贸易战的深层原因是科技领域的竞争。近年来，中国在5G通信、人工智能、半导体、量子计算等前沿技术领域快速发展，对美国的科技领先地位构成了挑战。',
    '2018年美国商务部对中兴通讯实施出口管制，随后将华为列入\u201c实体清单\u201d，禁止美国企业向其出售关键技术产品。这些举措反映了美国对中国科技崛起的战略焦虑。2019年5月，美国进一步限制中国企业在美国的半导体采购，标志着中美科技战全面升级。',
    '美国担心，如果中国在关键核心技术领域实现突破，将削弱其在全球产业链中的主导地位，进而影响其国家安全利益和全球战略优势。',
)

s2('（三）国际经济秩序的重塑需求',
    '二战后建立的国际经济秩序以美国为主导，WTO、IMF和世界银行等国际机构在规则制定上体现了发达国家的利益偏好。然而，随着中国等新兴市场国家的崛起，现有国际经济秩序的公平性和代表性受到质疑。',
    '中国主张构建更加公正合理的国际经济新秩序，推动多边贸易体制改革和全球经济治理变革。这种诉求与美国维护既有国际秩序主导地位的战略目标产生了结构性矛盾，成为中美贸易战的重要制度背景。',
)

s1('三、中美贸易战的影响分析',
    '中美贸易战对全球经济产生了广泛而深远的影响，涉及贸易、投资、产业链、金融市场等多个领域。',
)

s2('（一）对中美两国经济的影响',
    '从美国方面来看，加征关税虽然减少了从中国的进口，但并未显著改善美国的贸易逆差。据彼得森国际经济研究所（PIIE）的研究，美国对华贸易逆差从2017年的2,758亿美元下降至2019年的2,200亿美元，但同时增加了对越南、墨西哥等国的进口，贸易转移效应明显（Bown, 2021）。',
    '此外，关税成本最终由美国消费者和企业承担。研究表明，美国对华加征关税导致美国国内物价上涨，消费者每年额外支出约510亿美元（Amiti et al., 2019）。部分依赖中国供应链的美国制造业企业也面临成本上升和生产中断的风险。',
    '从中国方面来看，贸易战对出口导向型产业造成了一定冲击。2019年中国对美出口同比下降约12.5%，部分劳动密集型产业（如纺织、电子组装）出现了产能向东南亚转移的趋势。然而，中国通过扩大内需、推动产业升级和深化与其他国家和地区的经贸合作，有效对冲了贸易战的部分负面影响。2020年中国GDP仍实现了2.3%的正增长，成为全球唯一实现正增长的主要经济体（国家统计局, 2021）。',
)

s2('（二）对全球产业链和供应链的影响',
    '中美贸易战加速了全球产业链的重构。跨国公司为了降低关税风险和供应链集中度过高的风险，纷纷实施\u201c中国+1\u201d战略，将部分产能转移至越南、印度、墨西哥等国家。',
    '据联合国贸易和发展会议（UNCTAD）数据，2019年至2022年间，越南吸收的外国直接投资（FDI）增长了约30%，部分得益于中美贸易战带来的产能转移效应。然而，产业链转移并非简单的\u201c脱钩\u201d，而是呈现\u201c多元化\u201d和\u201c区域化\u201d趋势。中国仍然是全球制造业的核心节点，许多转移至东南亚的工厂仍高度依赖中国的中间品供应。',
)

s2('（三）对多边贸易体制的影响',
    '中美贸易战对以WTO为核心的多边贸易体制造成了严重冲击。美国绕过WTO争端解决机制，单方面依据国内法对中国加征关税，违反了WTO最惠国待遇和关税约束原则。这一行为削弱了WTO的权威性和有效性。',
    '2019年12月，WTO上诉机构因法官任命受阻而停止运作，多边贸易争端解决机制陷入瘫痪。中美贸易战在一定程度上加剧了多边贸易体制的危机，推动了各国寻求区域贸易协定（如RCEP、CPTPP）和双边贸易安排。',
)

s1('四、应对中美贸易战的对策建议',
    '面对中美贸易战的长期化和复杂化，中国需要采取综合性的应对策略，既要维护自身核心利益，也要推动全球经贸秩序的稳定与发展。',
)

s2('（一）坚持多边主义，推动WTO改革',
    '中国应坚定支持多边贸易体制，积极参与WTO改革进程。具体而言，可以推动以下改革措施：',
)

items('WTO改革建议',
    '1. 恢复上诉机构正常运作，推动成员就法官任命达成共识。',
    '2. 完善争端解决机制，提高裁决执行效率。',
    '3. 推动电子商务、投资便利化等新议题谈判，使WTO规则适应数字经济时代的需求。',
)

s2('（二）深化改革开放，增强经济韧性',
    '应对贸易战的长远之策在于深化改革开放，提升中国经济的核心竞争力：',
)

items('深化改革方向',
    '1. 扩大市场准入：进一步缩减外资准入负面清单，放宽金融、教育、医疗等领域的市场准入。',
    '2. 加强知识产权保护：完善知识产权法律体系，加大执法力度，营造公平竞争的市场环境。',
    '3. 推动产业升级：加大对高端制造业、半导体、人工智能等关键领域的研发投入，突破\u201c卡脖子\u201d技术瓶颈。',
    '4. 扩大内需：通过收入分配改革、社会保障体系完善等措施，释放国内消费潜力，降低对外部市场的依赖。',
)

s2('（三）拓展多元化国际市场',
    '中国应积极拓展与美国以外的国家和地区的经贸合作，分散贸易风险：',
)

items('多元化市场策略',
    '1. 深化RCEP合作：充分利用RCEP框架下的关税减让和贸易便利化措施，扩大与东盟、日本、韩国的贸易规模。',
    '2. 推进中欧投资协定：推动中欧投资协定早日生效，深化与欧洲国家的经贸联系。',
    '3. 加强南南合作：扩大与非洲、拉美等发展中国家的经贸合作，推动共建\u201c一带一路\u201d高质量发展。',
)

s2('（四）加强国际沟通与舆论引导',
    '中美贸易战不仅是经济博弈，也是舆论和话语权的较量。中国应加强国际沟通，向世界讲清楚中国改革开放的诚意和维护多边贸易体制的决心。通过发布白皮书、参与国际论坛、加强与国际媒体合作等方式，塑造客观公正的国际舆论环境，减少误解和偏见。',
)

s1('五、结论',
    '中美贸易战是21世纪最重要的国际经济事件之一，其成因涉及贸易失衡、科技竞争和国际秩序重塑等多重因素。贸易战对中美两国经济、全球产业链和多边贸易体制造成了深远影响。',
    '从长期来看，中美经贸关系的健康发展符合两国和全球的共同利益。中国应坚持改革开放，增强经济韧性，同时积极参与全球治理改革，推动构建更加公正合理的国际经济新秩序。中美两国也应通过对话协商管控分歧，在竞争中寻求合作，共同维护全球经济的稳定与繁荣。',
)

# ============================================================
# References
# ============================================================
doc.add_paragraph()
add_head('参考文献', level=1)

refs = [
    '[1] 商务部国际贸易经济合作研究院. 中美经贸摩擦回顾与展望[M]. 北京: 中国商务出版社, 2021.',
    '[2] 余淼杰. 中美贸易战的经济学分析[J]. 经济研究, 2019, 54(6): 4-18.',
    '[3] 王孝松, 谢申书. 中美贸易战对中国出口企业的影响研究[J]. 世界经济, 2020, 43(2): 3-27.',
    '[4] 鞠建东, 王璐. 全球价值链重构与中美贸易摩擦[J]. 国际经济评论, 2019(4): 93-109.',
    '[5] Amiti M, Redding S J, Weinstein D E. The Impact of the 2018 Trade War on U.S. Prices and Welfare[J]. Journal of Economic Perspectives, 2019, 33(4): 187-210.',
    '[6] Bown C P. The US-China Trade War and Phase One Agreement[J]. Journal of Policy Modeling, 2021, 43(4): 805-843.',
    '[7] WTO. World Trade Report 2020: Government Policies to Promote Innovation in the Digital Age[R]. Geneva: World Trade Organization, 2020.',
    '[8] UNCTAD. World Investment Report 2022[R]. Geneva: United Nations Conference on Trade and Development, 2022.',
    '[9] 国家统计局. 2020年国民经济和社会发展统计公报[R]. 北京: 中国统计出版社, 2021.',
    '[10] 中国社会科学院世界经济与政治研究所. 国际形势和中国外交蓝皮书（2021）[M]. 北京: 社会科学文献出版社, 2021.',
    '[11] 张宇燕. 中美贸易战的国际政治经济学分析[J]. 外交评论, 2019, 36(3): 1-21.',
    '[12] Baldwin R, Evenett S J. COVID-19 and Trade Policy: Why Turning Inward Will Not Work[R]. London: CEPR Press, 2020.',
    '[13] 中国海关总署. 2019年中国对外贸易统计数据[EB/OL]. [2025-06-01]. http://www.customs.gov.cn.',
    '[14] 彼得森国际经济研究所. 中美贸易摩擦追踪报告[EB/OL]. [2025-05-20]. https://www.piie.com.',
]

for ref in refs:
    add_para(ref, indent=False)

out = '/home/admin/.openclaw/workspace/中美贸易战读写译报告.docx'
doc.save(out)
print(f'Done: {out}')
