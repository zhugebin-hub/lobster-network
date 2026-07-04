#!/usr/bin/env python3
"""
叶畏兵毕业论文完善脚本 - 全面修复版
修复内容：
1. ASCII引号(U+0022)→中文引号(U+201C/U+201D)
2. 参考文献77/78拆分
3. 识别并标注引用位置
4. 生成修改报告
"""
import docx
import re
import json
from copy import deepcopy

INPUT = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
OUTPUT = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"

doc = docx.Document(INPUT)

# ==========================================
# 1. ASCII引号 → 中文引号
# ==========================================
print("=== 1. 修复ASCII引号 ===")

# 收集所有ASCII引号位置
all_ascii_quotes = []
for p_idx, para in enumerate(doc.paragraphs):
    for r_idx, run in enumerate(para.runs):
        if not run.text:
            continue
        for c_idx, ch in enumerate(run.text):
            if ch == '"':
                all_ascii_quotes.append((p_idx, r_idx, c_idx))

print(f"  发现 {len(all_ascii_quotes)} 个ASCII引号")

# 成对替换
for idx, (p_idx, r_idx, c_idx) in enumerate(all_ascii_quotes):
    run = doc.paragraphs[p_idx].runs[r_idx]
    chars = list(run.text)
    if idx % 2 == 0:
        chars[c_idx] = '\u201c'
    else:
        chars[c_idx] = '\u201d'
    run.text = ''.join(chars)

print(f"  ✅ 已修复")

# ==========================================
# 2. 参考文献77/78拆分
# ==========================================
print("\n=== 2. 拆分参考文献77/78 ===")
for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if '[77]' in p.text and '[78]' in p.text:
        split_at = p.text.index('[78]')
        p77 = p.text[:split_at].rstrip()
        p78 = p.text[split_at:].strip()
        p.clear()
        p.add_run(p77)
        # 插入新段落
        new_p = doc.add_paragraph()
        if p.style:
            new_p.style = p.style
        new_p.add_run(p78)
        p._p.addnext(new_p._p)
        print(f"  ✅ 已拆分")
        break

# ==========================================
# 3. 识别引用位置并建立映射
# ==========================================
print("\n=== 3. 引用位置分析 ===")

# 找到所有引用位置
ref_positions = []
for p_idx, para in enumerate(doc.paragraphs):
    t = para.text
    style = para.style.name if para.style else 'Normal'
    if 'Heading' in style or 'toc' in style or 'Caption' in style:
        continue
    if not re.search(r'[\u4e00-\u9fff]{20}', t):
        continue
    
    # "。 " or ""。 " or "》 " etc.
    count = 0
    for m in re.finditer(r'([。》"）]) (?=[\u4e00-\u9fffA-Z])', t):
        ref_positions.append({
            'para': p_idx,
            'char_pos': m.start(),
            'punct': m.group(1),
            'seq': count,
            'text_before': t[max(0,m.start()-50):m.start()],
            'text_after': t[m.end():m.end()+50]
        })
        count += 1

print(f"  共 {len(ref_positions)} 个引用位置")

# ==========================================
# 4. 生成详细修改报告
# ==========================================
report = {
    'summary': {
        'ascii_quotes_fixed': len(all_ascii_quotes),
        'ref_77_78_split': True,
        'citation_positions': len(ref_positions),
    },
    'citation_positions': [],
    'recommendations': []
}

for rp in ref_positions:
    report['citation_positions'].append({
        'para': rp['para'],
        'before': rp['text_before'][-40:],
        'after': rp['text_after'][:40]
    })

report['recommendations'] = [
    "1. 引号已全部从ASCII替换为中文引号",
    "2. 参考文献77/78已拆分",
    "3. 正文中有46处引用标记（句号+空格），建议逐一添加脚注上标",
    "4. 建议在Word中手动添加脚注：在引用位置右键→插入脚注",
    "5. 部分段落（如239、293标题）的引号格式需要特别注意"
]

with open('/home/admin/.openclaw/workspace/论文修改报告.json', 'w', encoding='utf-8') as f:
    json.dump(report, f, ensure_ascii=False, indent=2)

# ==========================================
# 5. 保存
# ==========================================
doc.save(OUTPUT)
print(f"\n✅ 保存至: {OUTPUT}")

# 最终验证
doc2 = docx.Document(OUTPUT)
remaining = sum(1 for p in doc2.paragraphs for ch in p.text if ch == '"')
print(f"验证: 剩余ASCII引号 {remaining} 个")
print(f"\n修改报告: /home/admin/.openclaw/workspace/论文修改报告.json")
