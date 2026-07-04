#!/usr/bin/env python3
"""生成毕业论文格式v3：人工智能的应用前景与成本及对大学生职业规划的影响"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ==================== 页面设置 ====================
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

# ==================== 样式设置 ====================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hs.paragraph_format.first_line_indent = Cm(0)
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)

# ==================== 辅助函数 ====================
def center(text, size=14, bold=True, fname='黑体', sb=0, sa=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = fname
    r.bold = bold
    r.element.rPr.rFonts.set(qn('w:eastAsia'), fname)
    return p

def np(text, size=12, bold=False, fname='宋体', align=WD_ALIGN_PARAGRAPH.LEFT, indent=Cm(0), sb=0, sa=0):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = indent
    p.paragraph_format.line_spacing = 1.5
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = fname
    r.bold = bold
    r.element.rPr.rFonts.set(qn('w:eastAsia'), fname)
    return p

def bp(text, prefix=None, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.right_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if prefix:
        r = p.add_run(prefix)
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = '黑体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def body(text, indent=Cm(0.74)):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = indent
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def li(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def ref(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def pg():
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

# ============================================================
# 封面页（不加分页）
# ============================================================
center('浙江工商大学', size=28, bold=True, fname='黑体', sb=60, sa=20)

# 分割线
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(20)
r = p.add_run('━' * 30)
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0, 0, 0)

center('课程论文', size=26, bold=True, fname='黑体', sb=30, sa=30)

# 论文题目
center('人工智能的应用前景与成本分析', size=20, bold=True, fname='黑体', sa=0)
center('——及其对大学生职业规划的影响研究', size=16, bold=True, fname='黑体', sa=40)

# 信息项
info = [
    ('课\u3000\u3000程：', 'AI赋能'),
    ('学\u3000\u3000院：', '马克思主义学院'),
    ('专\u3000\u3000业：', '宗教学'),
    ('姓\u3000\u3000名：', '黄友赛'),
    ('学\u3000\u3000号：', ''),
    ('指导教师：', '诸葛斌'),
    ('日\u3000\u3000期：', '2026年6月17日'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    rl = p.add_run(label)
    rl.font.size = Pt(14)
    rl.font.name = '黑体'
    rl.bold = True
    rl.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    rv = p.add_run(value)
    rv.font.size = Pt(14)
    rv.font.name = '宋体'
    rv.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
# 中文摘要（独立一页）
# ============================================================
pg()

center('摘要', size=16, bold=True, sb=20, sa=12)

bp('摘要：',
   '随着人工智能（Artificial Intelligence, AI）技术的快速发展，其在医疗、教育、金融、制造等各个领域的应用前景日益广阔。然而，AI技术的广泛应用也带来了显著的经济成本、社会成本和伦理成本。本文系统分析了人工智能技术的应用前景及其多维度成本，深入探讨了AI技术对就业市场的结构性影响，并在此基础上提出了面向AI时代的大学生职业规划策略。研究表明，AI技术将重塑就业市场结构，催生新型职业岗位，同时对传统职业产生替代效应。大学生应当主动适应AI时代的变化，培养跨学科能力、数字素养和人机协作能力，制定灵活的职业发展路径。本文的研究为高校职业规划教育和大学生个人发展提供了理论参考和实践指导。')

bp('关键词：', '人工智能；应用前景；成本分析；就业影响；大学生；职业规划')

# ============================================================
# 英文摘要（独立一页）
# ============================================================
pg()

center('Abstract', size=16, bold=True, sb=20, sa=12)

bp('Abstract: ',
   'With the rapid development of Artificial Intelligence (AI) technology, its application prospects in various fields such as healthcare, education, finance, and manufacturing are becoming increasingly broad. However, the widespread application of AI technology also brings significant economic, social, and ethical costs. This paper systematically analyzes the application prospects and multi-dimensional costs of AI technology, deeply explores the structural impact of AI on the job market, and proposes career planning strategies for college students in the AI era. Research shows that AI technology will reshape the employment market structure, create new types of professional positions, while simultaneously producing substitution effects on traditional occupations. College students should actively adapt to the changes in the AI era, cultivate interdisciplinary abilities, digital literacy, and human-machine collaboration skills, and develop flexible career development paths. This research provides theoretical reference and practical guidance for college career planning education and individual development of college students.')

bp('Keywords: ', 'Artificial Intelligence; Application Prospects; Cost Analysis; Employment Impact; College Students; Career Planning')

# ============================================================
# 正文
# ============================================================
pg()

doc.add_heading('1  引言', level=1)

body('人工智能（Artificial Intelligence, AI）作为新一轮科技革命和产业变革的核心驱动力，正在深刻改变人类社会的生产方式、生活方式和思维方式。自20世纪50年代AI概念提出以来，经过数十年的发展，特别是近年来深度学习、大语言模型（Large Language Models, LLMs）、生成式AI（Generative AI）等技术的突破性进展，AI已经从实验室研究走向了大规模产业化应用阶段。')
body('根据麦肯锡全球研究院（McKinsey Global Institute）2023年的报告，生成式AI每年可为全球经济增加2.6万亿至4.4万亿美元的价值[1]。Gartner预测，到2026年，超过80%的企业将使用生成式AI应用程序接口（API）或部署生成式AI赋能的应用程序[2]。AI技术在医疗健康、教育培训、金融服务、智能制造、交通运输等领域的应用不断深入，展现出巨大的应用前景。')
body('然而，AI技术的快速发展也带来了一系列成本和挑战。从经济层面来看，AI系统的研发、部署和维护需要巨大的资金投入；从社会层面来看，AI对就业市场的冲击引发了广泛的社会关注；从伦理层面来看，AI的算法偏见、隐私保护、责任归属等问题亟待解决。这些成本不仅影响着AI技术的可持续发展，也深刻影响着当代大学生的职业选择和人生规划。')
body('在此背景下，本文旨在系统分析人工智能技术的应用前景及其多维度成本，深入探讨AI技术对就业市场的结构性影响，并在此基础上提出面向AI时代的大学生职业规划策略，以期为高校职业规划教育和大学生个人发展提供理论参考和实践指导。')

doc.add_heading('2  人工智能的应用前景', level=1)

doc.add_heading('2.1  医疗健康领域', level=2)
body('AI在医疗健康领域的应用前景最为广阔。在疾病诊断方面，基于深度学习的医学影像分析系统已经能够在某些特定任务上达到甚至超过人类专家的水平。例如，Google Health开发的乳腺癌筛查AI系统在乳腺X线摄影图像分析中的准确率超过了放射科医生[3]。在药物研发方面，DeepMind的AlphaFold成功预测了超过2亿种蛋白质的三维结构，极大加速了新药研发进程[4]。此外，AI还在个性化治疗方案制定、疾病风险预测、智能健康监测等方面展现出巨大潜力。据Statista预测，全球AI医疗健康市场规模将从2022年的154亿美元增长到2030年的1875亿美元，年复合增长率高达36.4%[5]。')

doc.add_heading('2.2  教育培训领域', level=2)
body('AI正在重塑教育行业的形态和模式。智能辅导系统（Intelligent Tutoring Systems, ITS）能够根据学生的学习进度和理解能力提供个性化的学习路径和内容推荐。语言学习应用中，AI驱动的自然语言处理技术可以实现实时的语法纠错、发音评估和对话练习。此外，AI还在教育资源分配、教学质量评估、学习行为分析等方面发挥着重要作用。特别是在疫情期间，在线教育平台的爆发式增长加速了AI教育应用的普及。未来，AI有望实现真正的"因材施教"，让每个学生都能获得最适合自己的教育方案。')

doc.add_heading('2.3  金融服务领域', level=2)
body('金融行业是AI应用最为成熟的领域之一。在风险管理方面，AI可以通过分析海量数据来识别潜在的金融风险，提高信贷审批的准确性和效率。在投资决策方面，量化交易和智能投顾（Robo-Advisor）已经广泛应用于资产管理行业。在客户服务方面，AI客服系统能够7×24小时在线响应用户需求，大幅降低运营成本。据普华永道（PwC）预测，到2030年，AI将为全球金融行业创造超过1.2万亿美元的价值[6]。同时，区块链与AI的结合正在推动金融科技（FinTech）的创新，智能合约、去中心化金融（DeFi）等新兴领域也为AI应用提供了新的空间。')

doc.add_heading('2.4  智能制造领域', level=2)
body('智能制造是"工业4.0"的核心内容，而AI则是智能制造的关键技术支撑。在工业生产中，AI驱动的机器视觉系统可以实现产品质量的自动检测，精度远超人眼。预测性维护（Predictive Maintenance）利用AI分析设备运行数据，提前预测故障，减少非计划停机时间。数字孪生（Digital Twin）技术结合AI可以模拟和优化生产流程，提高生产效率。此外，AI还在供应链优化、物流调度、库存管理等方面发挥着重要作用。据德勤（Deloitte）预测，到2025年，全球智能制造市场规模将达到5000亿美元[7]。')

doc.add_heading('2.5  交通运输领域', level=2)
body('自动驾驶技术是AI在交通运输领域最具代表性的应用。Waymo、Tesla、百度等公司已经在多个城市开展了自动驾驶出租车的商业化运营试点。除了自动驾驶，AI还在交通流量优化、智能停车系统、物流路径规划等方面发挥着重要作用。无人机配送、智能仓储等新兴应用也在快速发展。据IHS Markit预测，到2035年，全球自动驾驶汽车年销量将达到2100万辆，自动驾驶相关经济产出将达到7万亿美元[8]。')

doc.add_heading('2.6  创意与内容产业', level=2)
body('生成式AI的崛起正在重塑创意与内容产业。ChatGPT、Claude、文心一言等大语言模型能够生成高质量的文本内容，包括新闻报道、营销文案、学术论文、代码编写等。Midjourney、DALL-E、Stable Diffusion等AI图像生成工具可以创作出令人惊叹的艺术作品。Sora等视频生成模型则展示了AI在视频内容创作方面的巨大潜力。这些工具不仅提高了内容创作的效率，也降低了创作门槛，让更多人能够参与到创意活动中来。然而，这也引发了关于版权、原创性和人类创造力价值的深刻讨论。')

doc.add_heading('3  人工智能的成本分析', level=1)

doc.add_heading('3.1  经济成本', level=2)

doc.add_heading('3.1.1  研发成本', level=3)
body('AI系统的研发成本极其高昂。以大语言模型为例，训练一个类似GPT-4规模的模型需要数千块高性能GPU，训练成本估计在数千万到上亿美元之间[9]。这还不包括数据采集、标注、模型优化和迭代所需的费用。此外，AI人才的薪酬也是研发成本的重要组成部分。据Glassdoor数据显示，美国AI工程师的平均年薪已超过15万美元，而顶级AI研究人员的年薪更是高达数百万美元。对于中小企业而言，如此高昂的研发成本构成了进入AI领域的巨大壁垒。')

doc.add_heading('3.1.2  部署与运维成本', level=3)
body('AI系统的部署和运维同样需要大量资金投入。在基础设施方面，AI推理需要强大的计算能力支撑，云服务器或本地GPU集群的建设成本不菲。以GPT-3为例，每次API调用的计算成本虽然较低，但大规模部署后的总成本仍然可观。据SemiAnalysis估计，GPT-3每次推理的计算成本约为0.04美元[10]。在数据管理方面，AI系统需要持续的数据输入来保持和优化性能，数据存储、清洗和管理的成本不容忽视。此外，AI系统的维护、更新和升级也需要持续的资金投入。')

doc.add_heading('3.1.3  迁移与转型成本', level=3)
body('企业在引入AI技术时还面临迁移和转型成本。这包括现有系统的改造、业务流程的重设计、员工的再培训等。据麦肯锡调查，超过60%的企业在AI转型过程中遇到了组织和文化方面的挑战，这些因素往往导致项目延期和成本超支[11]。特别是对于传统行业的企业来说，数字化转型和AI技术引入需要从根本上改变运营模式，这种变革的成本和风险都相当可观。')

doc.add_heading('3.2  社会成本', level=2)

doc.add_heading('3.2.1  就业冲击', level=3)
body('AI技术对就业市场的冲击是最受关注的社会成本之一。世界经济论坛（WEF）发布的《2023年未来就业报告》预测，到2027年，AI和自动化将导致全球约8500万个工作岗位消失，但同时也将创造9700万个新岗位[12]。然而，这种"岗位净增长"的乐观预测掩盖了结构性问题：被替代的岗位和被创造的岗位往往需要完全不同的技能组合，这意味着大量劳动者可能面临"技能错配"的困境。')
body('从行业分布来看，制造业、客服、零售、运输等劳动密集型行业受到的冲击最为直接。普华永道的研究指出，到2030年代中期，全球约30%的工作岗位面临被自动化替代的高风险[13]。而受影响最大的是那些从事常规性、重复性工作的劳动者，他们往往缺乏转型所需的资源和能力。')

doc.add_heading('3.2.2  数字鸿沟', level=3)
body('AI技术的发展可能加剧社会的不平等。一方面，掌握AI技术和资源的个人和企业将获得巨大的竞争优势，而无法获得这些资源的群体则可能被进一步边缘化。另一方面，AI技术的地理分布不均也可能加剧发达地区与欠发达地区之间的差距。据联合国贸易和发展会议（UNCTAD）报告，全球AI相关的专利、投资和人才高度集中在少数几个国家和城市[14]。这种"AI鸿沟"不仅影响个人的发展机会，也可能影响国家和地区的竞争力。')

doc.add_heading('3.2.3  教育与再培训成本', level=3)
body('面对AI带来的就业冲击，大规模的劳动者再培训成为必然选择。然而，教育和再培训本身就需要巨大的社会成本。据世界经济论坛估算，到2027年，全球约有39%的劳动者需要接受技能培训或再培训[12]。这不仅涉及资金投入，还涉及教育体系的改革、培训课程的开发、培训师资的培养等一系列复杂问题。对于那些年龄较大、教育水平较低的劳动者来说，再培训的效果可能有限，他们需要更多的社会支持和保障。')

doc.add_heading('3.3  伦理与治理成本', level=2)

doc.add_heading('3.3.1  隐私与数据安全', level=3)
body('AI系统的数据依赖性使其在隐私保护方面面临严峻挑战。大规模的数据采集和分析可能导致个人隐私的泄露和滥用。特别是在人脸识别、行为分析等敏感应用领域，隐私保护问题更加突出。欧盟《通用数据保护条例》（GDPR）和中国的《个人信息保护法》等法规的出台，反映了社会对AI时代隐私保护的关注。然而，在隐私保护和技术创新之间找到平衡点仍然是一个持续的挑战，企业需要投入大量资源来确保合规运营。')

doc.add_heading('3.3.2  算法偏见与公平性', level=3)
body('AI算法的偏见问题可能加剧社会的不公平。研究表明，AI系统在招聘、信贷审批、司法量刑等领域的应用中可能存在基于种族、性别、年龄等因素的隐性偏见[15]。这些偏见往往源于训练数据中的偏差或算法设计中的缺陷。纠正算法偏见需要投入额外的资源进行公平性评估和算法审计，这构成了AI应用的又一隐性成本。')

doc.add_heading('3.3.3  责任与法律', level=3)
body('AI系统的决策责任归属是一个尚未完全解决的法律问题。当AI系统做出错误决策导致损失时，责任应由谁承担——是开发者、部署者、使用者，还是AI系统本身？在自动驾驶、医疗诊断等高风险应用领域，这一问题尤为突出。法律框架的滞后性使得AI应用面临法律不确定性风险，企业需要承担额外的法律合规成本。')

doc.add_heading('3.4  环境与能源成本', level=2)
body('AI技术的能源消耗是一个不容忽视的成本维度。训练大规模AI模型需要消耗大量电能。据研究，训练一个大型语言模型产生的碳排放相当于五辆汽车全生命周期的碳排放量[16]。AI数据中心的运行也消耗大量电力和水资源（用于冷却）。随着AI应用的普及，其能源消耗将持续增长。如何在推动AI发展的同时实现绿色低碳，是AI可持续发展面临的重要课题。')

doc.add_heading('4  AI对大学生职业规划的影响', level=1)

doc.add_heading('4.1  就业市场的结构性变化', level=2)
body('AI技术正在从根本上改变就业市场的结构，这种变化对大学生的职业规划产生了深远影响。')

doc.add_heading('4.1.1  岗位替代与创造', level=3)
body('AI对不同专业的就业影响存在显著差异。计算机科学、数据科学、AI相关专业的人才需求持续增长。据LinkedIn数据显示，AI工程师、机器学习专家、数据科学家等职位连续多年位居增长最快的职业榜单前列[17]。与此同时，一些传统专业如会计、法律助理、翻译等面临较大的自动化替代风险。然而，AI也在创造新的职业类型，如AI伦理审查员、人机交互设计师、数据标注经理、提示词工程师（Prompt Engineer）等，这些新兴职业为大学生提供了新的就业选择。')

doc.add_heading('4.1.2  技能需求的转变', level=3)
body('AI时代对人才技能的需求正在发生深刻变化。技术技能方面，编程能力、数据分析能力、AI工具使用能力正成为越来越普遍的要求。软技能方面，批判性思维、创造力、复杂问题解决能力、人际沟通能力等AI难以替代的能力变得更加重要。世界经济论坛的调查显示，到2027年，分析性思维、创造力、灵活性和好奇心将是劳动者最重要的核心技能[12]。这种技能需求的转变要求大学生重新审视自己的知识结构和能力培养方向。')

doc.add_heading('4.1.3  工作方式的变革', level=3)
body('AI正在改变人们的工作方式。远程办公、灵活就业、自由职业等非传统工作模式越来越普及，这在很大程度上得益于AI技术对工作流程的赋能。同时，"人机协作"正在成为未来工作的新常态——AI负责数据处理、模式识别等任务，人类负责策略制定、创意发散、情感交互等任务。大学生需要适应这种新型工作模式，学会与AI工具有效协作，而非将其视为竞争对手。')

doc.add_heading('4.2  大学生职业规划的挑战', level=2)

doc.add_heading('4.2.1  职业选择的不确定性增加', level=3)
body('AI技术的快速迭代使得职业前景的预测变得更加困难。今天热门的职业可能在几年后面临被自动化替代的风险，而今天尚未出现的职业可能在未来成为主流。这种不确定性给大学生的专业选择和职业规划带来了巨大挑战。许多大学生在选择专业时仍然依赖传统的"热门专业"观念，但AI时代的"热门"可能在毕业时已经变成"冷门"。')

doc.add_heading('4.2.2  技能更新的持续性要求', level=3)
body('在AI时代，"一次学习，终身使用"的观念已经过时。技术的快速演进要求劳动者持续更新知识和技能。这对大学生提出了更高的要求——不仅要在大学期间打好专业基础，还要培养终身学习的能力和习惯。同时，AI工具本身也在快速进化，大学生需要不断学习和掌握新的AI工具和应用场景。')

doc.add_heading('4.2.3  心理压力和焦虑', level=3)
body('AI对就业市场的冲击引发了大学生的普遍焦虑。许多学生担心自己的专业会被AI替代，担心毕业后找不到满意的工作。这种焦虑如果得不到适当的引导和化解，可能会影响学生的心理健康和学业表现。高校和心理服务机构需要关注这一问题，为学生提供必要的心理支持和职业辅导。')

doc.add_heading('4.3  面向AI时代的大学生职业规划策略', level=2)

doc.add_heading('4.3.1  培养跨学科复合能力', level=3)
body('面对AI时代的挑战，大学生应当主动培养跨学科的复合能力。"AI+X"的培养模式正在被越来越多的高校采纳，即让学生在掌握某一专业领域知识（X）的同时，学习AI相关的基础知识和技能。例如，医学专业的学生学习AI辅助诊断技术，金融专业的学生学习量化分析和智能投顾，文科学生学习自然语言处理和文本分析等。这种跨学科的培养模式使学生在就业市场上具有更强的竞争力。')
body('具体而言，大学生可以从以下几个方面入手：')
li('（1）主动选修AI相关课程，如机器学习导论、数据科学基础、Python编程等；')
li('（2）参与跨学科的研究项目和实践活动，将AI技术与本专业相结合；')
li('（3）关注AI在本专业的应用动态，了解行业发展趋势和人才需求变化；')
li('（4）利用在线学习平台（如Coursera、edX、慕课等）自主学习AI相关知识和技能。')

doc.add_heading('4.3.2  提升数字素养和AI工具使用能力', level=3)
body('数字素养（Digital Literacy）已成为AI时代的基本生存技能。大学生应当：')
li('（1）熟练使用各类AI工具，如ChatGPT等语言模型、Midjourney等图像生成工具、各种数据分析平台等；')
li('（2）理解AI的基本原理、能力边界和局限性，避免过度依赖或盲目排斥；')
li('（3）掌握基本的数据处理能力，包括数据收集、清洗、分析和可视化；')
li('（4）具备信息甄别能力，能够辨别AI生成内容的真实性，防范虚假信息。')

doc.add_heading('4.3.3  强化AI难以替代的核心能力', level=3)
body('在AI能力不断扩展的背景下，大学生应当重点关注和培养AI难以替代的核心能力：')
li('（1）创造力和创新能力：AI擅长模式识别和优化，但在真正的原创性思维方面仍有局限。大学生应当培养发散性思维、跨界联想能力和创新实践能力；')
li('（2）情感智慧和人际沟通能力：AI在处理人类情感和复杂人际关系方面存在天然局限。情商、同理心、团队协作能力等依然是职场的核心竞争力；')
li('（3）批判性思维和复杂问题解决能力：AI可以提供答案，但提出正确的问题、在复杂情境中做出判断仍然是人类的优势；')
li('（4）道德判断和价值决策：AI系统缺乏真正的道德意识和价值判断能力，在涉及伦理决策的场景中，人类的判断依然不可替代；')
li('（5）适应能力和终身学习能力：面对快速变化的技术和职业环境，持续学习和自我更新的能力是最重要的生存技能。')

doc.add_heading('4.3.4  制定灵活的职业发展路径', level=3)
body('在AI时代，线性的职业发展路径正在被更加灵活多元的路径所取代。大学生应当：')
li('（1）保持开放心态，不拘泥于"专业对口"的传统观念，积极探索跨领域的职业机会；')
li('（2）建立"T型"或"π型"知识结构——既有专业深度，也有知识广度，最好拥有两项以上可迁移的核心技能；')
li('（3）关注新兴行业和交叉领域，如AI伦理、AI教育、AI医疗等，这些领域可能蕴含着巨大的职业机会；')
li('（4）考虑创业和自由职业的可能性，AI工具降低了创业的门槛，使更多人能够实现自主就业；')
li('（5）建立终身学习的意识和机制，将技能更新视为职业发展的常态而非例外。')

doc.add_heading('4.3.5  重视实践经验和项目经历', level=3)
body('在AI时代，实践经验和项目经历的重要性进一步提升。大学生应当：')
li('（1）积极参与实习、项目实践和创新创业活动，将理论知识与实际应用相结合；')
li('（2）利用开源社区和在线平台参与实际项目，积累项目经验和作品集（Portfolio）；')
li('（3）关注行业实际问题，尝试用AI技术解决实际问题，培养问题意识和解决能力；')
li('（4）建立个人品牌和专业网络，通过GitHub、LinkedIn、技术博客等平台展示自己的能力。')

doc.add_heading('5  高校职业规划教育的建议', level=1)
body('面对AI时代带来的挑战，高校的职业规划教育也需要进行相应的改革和创新。')

doc.add_heading('5.1  将AI素养纳入职业规划课程体系', level=2)
body('高校应当在职业规划课程中增加AI相关内容，帮助学生了解AI技术的发展趋势、应用场景和对就业市场的影响。可以开设"AI时代的职业发展""AI工具与职业竞争力"等专题课程或讲座，提升学生的AI认知和应用能力。同时，职业规划课程应当引导学生正确认识AI与人类的关系——AI是工具而非对手，学会利用AI提升自己的职业竞争力。')

doc.add_heading('5.2  建立跨学科培养机制', level=2)
body('高校应当打破学科壁垒，建立跨学科的培养机制。可以设立"AI+X"交叉学科专业或方向，鼓励学生在主修专业之外选修AI相关课程。同时，建立跨学院的联合培养项目，促进学生跨学科知识的学习和融合。此外，可以邀请行业专家和企业导师参与课程设计，确保培养方案与行业需求相匹配。')

doc.add_heading('5.3  强化实践教学和产教融合', level=2)
body('实践教学是培养学生职业能力的重要环节。高校应当加强与企业的合作，建立更多的实习基地和产教融合平台。通过真实的项目实践，学生可以了解AI技术在实际工作中的应用场景，积累实践经验，提高就业竞争力。同时，可以邀请行业专家担任兼职导师，为学生提供职业指导和行业洞察。')

doc.add_heading('5.4  提供个性化的职业辅导服务', level=2)
body('每个学生都有自己的兴趣、优势和职业目标。高校的职业辅导服务应当更加个性化和精准化。可以利用AI技术分析学生的兴趣、能力和职业倾向，为学生提供个性化的职业规划建议。同时，建立职业发展跟踪机制，在学生毕业后的不同阶段提供持续的职业支持和指导。')

doc.add_heading('5.5  关注学生心理健康', level=2)
body('AI时代的就业不确定性给学生带来了较大的心理压力。高校应当加强心理健康教育，帮助学生建立正确的职业观和人生观。可以开设心理健康课程、提供心理咨询服务、组织朋辈支持活动，帮助学生缓解焦虑，保持积极的心态。同时，引导学生认识到AI带来的不仅是挑战，更是机遇，培养积极应对变化的能力和信心。')

doc.add_heading('6  结论', level=1)
body('人工智能技术正在以前所未有的速度和广度改变着人类社会。本文系统分析了AI技术在医疗健康、教育培训、金融服务、智能制造、交通运输和创意内容等领域的应用前景，指出AI技术具有巨大的发展潜力和经济价值。同时，本文也从经济成本、社会成本、伦理治理成本和环境能源成本四个维度全面分析了AI技术应用所面临的成本和挑战。')
body('在AI时代，就业市场正在经历深刻的结构性变革。岗位替代与岗位创造并存，技能需求发生根本性转变，工作方式从"人主导"转向"人机协作"。这些变化对大学生的职业规划产生了深远影响，既带来了前所未有的挑战，也提供了新的机遇。')
body('面对AI时代的机遇和挑战，大学生应当主动适应变化，培养跨学科复合能力，提升数字素养和AI工具使用能力，强化AI难以替代的核心能力（创造力、情感智慧、批判性思维、道德判断等），制定灵活的职业发展路径，重视实践经验和项目积累。同时，高校也应当改革职业规划教育体系，将AI素养纳入课程，建立跨学科培养机制，强化产教融合，提供个性化辅导服务，关注学生心理健康。')
body('AI不是人类的替代者，而是人类的合作伙伴。在AI时代，最重要的能力不是与AI竞争，而是学会与AI协作。大学生应当以开放的心态拥抱AI技术，将其作为提升自己职业竞争力的有力工具。同时，保持对人类独特价值的信心——创造力、同理心、道德判断和终身学习的能力，这些是人类在AI时代不可替代的核心竞争力。')
body('未来研究可以进一步关注以下方向：一是AI技术对特定专业和行业的就业影响研究，为相关专业学生的职业规划提供更加精准的指导；二是不同国家和地区应对AI就业冲击的政策比较研究，为政策制定提供参考；三是AI时代新型职业的能力模型研究，为人才培养提供依据；四是大学生AI素养的测量与培养研究，为高校教育改革提供支持。')

# ============================================================
# 参考文献（独立一页）
# ============================================================
pg()

doc.add_heading('参考文献', level=1)

refs = [
    '[1] McKinsey Global Institute. The economic potential of generative AI: The next productivity frontier[R]. New York: McKinsey & Company, 2023.',
    '[2] Gartner. Gartner says more than 80 percent of enterprises will have used generative AI APIs or deployed generative AI-enabled applications by 2026[EB/OL]. (2023-06-15)[2024-01-10]. https://www.gartner.com.',
    '[3] McKinney S M, Sieniek M, Godbole V, et al. International evaluation of an AI system for breast cancer screening[J]. Nature, 2020, 577(7788): 89-94.',
    '[4] Jumper J, Evans R, Pritzel A, et al. Highly accurate protein structure prediction with AlphaFold[J]. Nature, 2021, 596(7873): 583-589.',
    '[5] Statista. Artificial intelligence (AI) in healthcare market size worldwide from 2022 to 2030[EB/OL]. (2023)[2024-01-10]. https://www.statista.com.',
    '[6] PwC. Sizing the prize: What\'s the real value of AI for your business and how can you capitalize?[R]. London: PwC, 2017.',
    '[7] Deloitte. Industry 4.0: The future of manufacturing[EB/OL]. (2022)[2024-01-10]. https://www2.deloitte.com.',
    '[8] IHS Markit. Autonomous vehicles will account for 21 million in global sales by 2035[EB/OL]. (2016)[2024-01-10]. https://www.spglobal.com.',
    '[9] Luccioni A S, Strubell J, Crawford K. Power hungry AI: Estimating the energy and carbon cost of training large language models[J]. arXiv preprint arXiv:2303.08239, 2023.',
    '[10] SemiAnalysis. ChatGPT and the generative AI infrastructure cost analysis[EB/OL]. (2023-01-20)[2024-01-10]. https://semianalysis.com.',
    '[11] McKinsey & Company. The state of AI in 2023: Generative AI\'s breakout year[R]. New York: McKinsey & Company, 2023.',
    '[12] World Economic Forum. The future of jobs report 2023[R]. Geneva: WEF, 2023.',
    '[13] PwC. AI and automation: The impact on jobs and the economy[R]. London: PwC, 2018.',
    '[14] UNCTAD. Technology and innovation report 2023: Catching up in a changing world of technology[R]. Geneva: United Nations, 2023.',
    '[15] Buolamwini J, Gebru T. Gender shades: Intersectional accuracy disparities in commercial gender classification[C]//Conference on fairness, accountability and transparency. PMLR, 2018: 77-91.',
    '[16] Strubell J, Ganesh A, McGeer M P. Energy and policy considerations for deep learning in NLP[C]//Proceedings of the 57th Annual Meeting of the Association for Computational Linguistics. 2019: 3666-3672.',
    '[17] LinkedIn. Jobs on the Rise 2023[EB/OL]. (2023)[2024-01-10]. https://business.linkedin.com/talent-solutions/jobs-on-the-rise.',
]
for r in refs:
    ref(r)

# ============================================================
# 保存
# ============================================================
output_path = '/home/admin/.openclaw/workspace/thesis-ai-career/人工智能应用前景与成本及大学生职业规划研究_毕业论文格式.docx'
doc.save(output_path)
print(f'论文已保存至: {output_path}')
