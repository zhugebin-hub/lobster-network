#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成道长版问卷Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_questionnaire_docx():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 设置文档边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
    
    # 标题
    title = doc.add_heading('道教道长心理服务需求调查问卷', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题信息
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('（道长版 V1.0）')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # 问卷说明
    doc.add_paragraph('')  # 空行
    
    # 问卷说明框
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = info_para.add_run('调查目的：')
    run.bold = True
    run = info_para.add_run('了解道长对心理疏导的认知、实践经验和培训需求')
    
    run = info_para.add_run('\n适用对象：')
    run.bold = True
    run = info_para.add_run('道教教职人员（道长/道长/坤道）')
    
    run = info_para.add_run('\n预计填写时间：')
    run.bold = True
    run = info_para.add_run('15-20分钟')
    
    run = info_para.add_run('\n匿名填写，数据仅用于学术研究')
    
    doc.add_paragraph('')  # 空行
    
    # 问卷说明正文
    doc.add_heading('问卷说明', level=2)
    
    instructions = [
        '尊敬的道长：',
        '您好！本问卷旨在了解道教教职人员对心理疏导服务的认知、实践经验和发展需求，为道教心理服务的理论研究和实践探索提供数据支持。',
        '• 本问卷采用匿名方式，所有数据仅用于学术研究，绝不泄露个人隐私',
        '• 问卷填写约需15-20分钟',
        '• 除标注"多选"的题目外，其余均为单选题',
        '• 如有任何疑问，请联系研究者',
        '感谢您的支持与参与！',
        '',
        '研究者：择白',
        '浙江省双通班学员',
        '调查时间：2026年'
    ]
    
    for text in instructions:
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph('')  # 空行
    
    # A部分：基本信息
    doc.add_heading('A部分：基本信息', level=2)
    
    questions_a = [
        {
            'q': 'A1. 您的道派：（必答）',
            'options': ['全真派', '正一派', '其他：_______']
        },
        {
            'q': 'A2. 您的出家/皈依年限：（必答）',
            'options': ['3年以下', '3-5年', '5-10年', '10-20年', '20年以上']
        },
        {
            'q': 'A3. 您在道观中的主要职务：（必答）',
            'options': ['住持/监院', '知客', '经师', '执事', '普通道士', '其他：_______']
        },
        {
            'q': 'A4. 您所在的道观类型：（必答）',
            'options': ['大型道观（常驻道士20人以上）', '中型道观（常驻道士10-20人）', '小型道观（常驻道士10人以下）', '宫观/小庙']
        },
        {
            'q': 'A5. 您所在的道观所在地区：（必答）',
            'options': ['城市中心', '城市郊区', '县城', '乡镇/农村']
        },
        {
            'q': 'A6. 您的性别：（必答）',
            'options': ['男', '女']
        },
        {
            'q': 'A7. 您的年龄段：（必答）',
            'options': ['25岁以下', '25-35岁', '36-45岁', '46-55岁', '56-65岁', '65岁以上']
        }
    ]
    
    for item in questions_a:
        p = doc.add_paragraph()
        run = p.add_run(item['q'])
        run.bold = True
        for opt in item['options']:
            p = doc.add_paragraph(f'□ {opt}')
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph('')  # 空行
    
    # B部分：心理疏导实践经验
    doc.add_heading('B部分：心理疏导实践经验', level=2)
    
    questions_b = [
        {
            'q': 'B1. 您是否曾为信众提供过心理疏导/心灵慰藉？（必答）',
            'options': ['经常提供（每周多次）', '偶尔提供（每月几次）', '很少提供（每季度几次）', '从未提供（跳至第B5题）']
        },
        {
            'q': 'B2. 您通常为信众疏导哪些类型的困扰？（最多选3项）（多选，必答）',
            'options': ['家庭矛盾/婚姻问题', '子女教育/代际冲突', '事业/财运不顺', '健康问题/疾病焦虑', '人际关系困扰', '情绪问题（焦虑、抑郁、愤怒）', '人生意义/价值困惑', '信仰困惑/信仰危机', '死亡恐惧/临终关怀', '其他：_______']
        },
        {
            'q': 'B3. 您通常采用哪些方式进行疏导？（可多选）（多选，必答）',
            'options': ['讲经解惑——讲解道教经典，以经义开导', '因果开示——以因果报应、善恶有报的道理劝导', '命理分析——通过八字、风水等分析运势，缓解焦虑', '冥想/静坐指导——教信众坐忘、心斋等修炼方法', '科仪法事——通过祈福、消灾等仪式给予心理安慰', '日常陪伴——闲聊、喝茶、散步中的自然开导', '推荐修行——建议信众诵经、拜忏、持咒等', '转介建议——建议信众寻求专业心理咨询', '其他：_______']
        },
        {
            'q': 'B4. 您认为自己的心理疏导效果如何？（必答）',
            'options': ['效果非常好，信众反馈积极', '效果比较好，大部分有帮助', '效果一般，视情况而定', '效果不太好，经常感到力不从心', '不确定效果如何']
        },
        {
            'q': 'B5. 您平均每次与信众交流的时间大约是：（必答）',
            'options': ['15分钟以内', '15-30分钟', '30-60分钟', '60分钟以上', '不提供心理疏导（跳至C部分）']
        },
        {
            'q': 'B6. 您每月大约为多少位信众提供心理疏导？（必答）',
            'options': ['0位（不提供）', '1-5位', '6-15位', '16-30位', '30位以上']
        }
    ]
    
    for item in questions_b:
        p = doc.add_paragraph()
        run = p.add_run(item['q'])
        run.bold = True
        for opt in item['options']:
            p = doc.add_paragraph(f'□ {opt}')
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph('')  # 空行
    
    # C部分：道教心理服务认知
    doc.add_heading('C部分：道教心理服务认知', level=2)
    
    questions_c = [
        {
            'q': 'C1. 您是否认为道观应该提供心理疏导服务？（必答）',
            'options': ['非常应该，这是道观的重要功能', '比较应该，可以适当开展', '一般，看情况', '不太应该，道观主要功能是修行和祭祀', '完全不应该，这不是道观的职责']
        },
        {
            'q': 'C2. 您认为道教心理疏导与专业心理咨询相比，优势在于？（可多选）（多选，必答）',
            'options': ['信仰力量——信众对道长和道教有天然信任', '文化根基——道教智慧更契合中国人的文化心理', '身心同调——修炼方法（冥想、呼吸等）可直接调节身心', '社群支持——道观提供信仰社群，增强归属感', '仪式力量——科仪法事有独特的心理疗愈功能', '道德框架——提供明确的价值导向和人生意义', '可及性——道观分布广，信众容易接触', '成本低——随缘乐助，经济门槛低', '没有明显优势', '其他：_______']
        },
        {
            'q': 'C3. 您认为道教心理疏导的局限性在于？（可多选）（多选，必答）',
            'options': ['专业性不足——缺乏系统的心理学训练', '边界不清——宗教辅导与心理咨询的界限模糊', '效果评估难——缺乏科学的疗效评估方法', '个案复杂——严重心理问题超出道长能力范围', '时间精力——道长日常事务多，无暇深入疏导', '缺乏场地——道观没有专门的疏导空间', '社会认知——公众对道教心理服务缺乏了解', '缺乏规范——没有行业标准和服务规范', '没有明显局限', '其他：_______']
        },
        {
            'q': 'C4. 您是否了解以下心理学流派/方法？（多选，必答）',
            'options': ['认知行为疗法（CBT）', '精神分析/心理动力学', '人本主义/来访者中心疗法', '正念/冥想疗法（MBSR/MBCT）', '接纳承诺疗法（ACT）', '存在主义心理治疗', '荣格分析心理学', '焦点解决短期治疗（SFBT）', '哀伤辅导/危机干预', '以上都不了解']
        },
        {
            'q': 'C5. 您认为道教修炼方法与心理学方法有哪些相通之处？（可多选）（多选，必答）',
            'options': ['坐忘/心斋 ≈ 正念冥想——都强调觉察当下、不评判', '无为 ≈ 接纳——都强调放下执着、顺应自然', '齐物 ≈ 认知解离——都强调超越二元对立', '养生 ≈ 身心医学——都关注身心互动', '修道 ≈ 自我实现——都追求人格完善', '科仪 ≈ 仪式治疗——都通过仪式行为促进疗愈', '师承 ≈ 治疗关系——都强调师徒/咨访关系的重要性', '说不清相通之处', '其他：_______']
        }
    ]
    
    for item in questions_c:
        p = doc.add_paragraph()
        run = p.add_run(item['q'])
        run.bold = True
        for opt in item['options']:
            p = doc.add_paragraph(f'□ {opt}')
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph('')  # 空行
    
    # D部分：培训需求与发展意愿
    doc.add_heading('D部分：培训需求与发展意愿', level=2)
    
    questions_d = [
        {
            'q': 'D1. 您是否希望接受心理学相关培训？（必答）',
            'options': ['非常希望', '比较希望', '一般', '不太希望', '完全不需要']
        },
        {
            'q': 'D2. 您最希望学习哪些内容？（最多选3项）（多选，必答）',
            'options': ['基础心理学知识——心理学流派、人的心理发展规律', '咨询技巧——倾听、共情、提问、反馈等基本技能', '危机干预——自杀预防、哀伤辅导、创伤处理', '道教心理学——道教经典中的心理学思想', '冥想/正念训练——坐忘、心斋等方法的心理学解读', '伦理与边界——宗教辅导与心理咨询的伦理差异', '案例督导——实际案例讨论与专业督导', '转介知识——何时转介、如何转介给专业机构', '其他：_______']
        },
        {
            'q': 'D3. 您希望的培训形式是？（可多选）（多选，必答）',
            'options': ['短期集中培训——3-5天封闭式培训', '系统课程——8-12次系统学习（每周一次）', '工作坊——实操性强的体验式学习', '线上学习——录播课程+直播答疑', '案例督导小组——定期案例讨论', '跟岗实习——到心理咨询机构实习观摩', '自学资料——书籍、文章、视频资料', '其他：_______']
        },
        {
            'q': 'D4. 您能接受的培训时长是？（必答）',
            'options': ['1-3天（短期）', '1-2周（中期）', '1-3个月（长期系统）', '半年以上（深度培训）']
        },
        {
            'q': 'D5. 您能接受的培训费用是？（必答）',
            'options': ['免费（道协/政府资助）', '500元以内', '500-1000元', '1000-3000元', '3000元以上（高质量培训值得投入）']
        },
        {
            'q': 'D6. 如果道观开设心理疏导室，您是否愿意参与？（必答）',
            'options': ['非常愿意，这是我的兴趣所在', '比较愿意，可以尝试', '看情况，如果有培训和支持的话', '不太愿意，我还有其他职责', '不愿意']
        },
        {
            'q': 'D7. 您认为道长开展心理疏导服务需要哪些支持？（可多选）（多选，必答）',
            'options': ['专业培训——系统的心理学和咨询技能培训', '专业督导——定期案例督导，防止反移情和职业倦怠', '转介渠道——与心理咨询机构建立合作关系', '场地支持——道观内设置专用疏导室', '政策支持——道教协会的政策指导和规范', '学术支持——与高校/研究机构合作研究', '资金支持——培训经费、场地建设经费', '社会认可——提高公众对道教心理服务的认知', '其他：_______']
        }
    ]
    
    for item in questions_d:
        p = doc.add_paragraph()
        run = p.add_run(item['q'])
        run.bold = True
        for opt in item['options']:
            p = doc.add_paragraph(f'□ {opt}')
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph('')  # 空行
    
    # E部分：道教心理服务发展
    doc.add_heading('E部分：道教心理服务发展', level=2)
    
    questions_e = [
        {
            'q': 'E1. 您是否认为需要建立"道教心理疏导"的专业规范？（必答）',
            'options': ['非常需要，应该有行业标准和伦理规范', '比较需要，可以有指导性文件', '一般，看发展情况', '不太需要，保持灵活性更好', '完全不需要，道教本就无法规范']
        },
        {
            'q': 'E2. 您认为道教心理服务的理想模式是？（可多选）（多选，必答）',
            'options': ['道观内设疏导室——道观内设立专用心理疏导空间', '道长+心理咨询师合作——道长与专业咨询师共同服务', '社区心理服务——道观与社区合作，提供心理服务', '线上心理服务——通过网络提供远程疏导', '修炼营/工作坊——定期举办道教修炼体验活动', '讲经+疏导结合——讲经法会中融入心理疏导', '道教心理医院——建立专门的道教心理医疗机构', '其他：_______']
        },
        {
            'q': 'E3. 您是否愿意将个人疏导经验整理为案例供研究使用？（必答）',
            'options': ['非常愿意，匿名处理即可', '比较愿意，但需要审核', '看情况', '不太愿意，涉及信众隐私', '完全不愿意']
        },
        {
            'q': 'E4. 您认为道教心理服务面临的最大挑战是？（必答）',
            'options': ['道长专业能力不足', '缺乏社会认知和信任', '缺乏政策和资金支持', '宗教辅导与心理咨询的边界问题', '严重心理问题的处理能力不足', '道长时间和精力有限', '缺乏效果评估方法', '其他：_______']
        }
    ]
    
    for item in questions_e:
        p = doc.add_paragraph()
        run = p.add_run(item['q'])
        run.bold = True
        for opt in item['options']:
            p = doc.add_paragraph(f'□ {opt}')
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph('')  # 空行
    
    # F部分：开放性问题
    doc.add_heading('F部分：开放性问题', level=2)
    
    open_questions = [
        'F1. 请分享一次您印象最深的心理疏导经历（成功或困难的都可以）：（选答）',
        'F2. 您认为道教经典中哪些思想对心理疏导最有价值？（选答）',
        'F3. 您对道教心理服务的未来发展有什么建议或期望？（选答）'
    ]
    
    for q in open_questions:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        # 添加答题线
        for _ in range(3):
            p = doc.add_paragraph('_____________________________________________________________________________')
            p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph('')  # 空行
    
    # 结束语
    doc.add_heading('问卷结束', level=2)
    
    closing_text = [
        '感恩道长的耐心填写！您的每一份回答都将为道教心理服务的理论研究和实践探索提供宝贵数据。',
        '',
        '如您愿意参与后续的深度访谈或案例研究，请留下联系方式（选填）：',
        '',
        '道号/姓名：_______',
        '所在道观：_______',
        '电话/微信：_______',
        '',
        '再次感谢您的支持与参与！福生无量天尊！'
    ]
    
    for text in closing_text:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
    
    # 保存文档
    output_path = '/home/admin/.openclaw/workspace/双通班论文_道长版问卷V1.0.docx'
    doc.save(output_path)
    print(f'Word文档已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    create_questionnaire_docx()
