#!/usr/bin/env python3
"""生成书册目录框架Word文档"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = "/home/admin/.openclaw/workspace"

# 书册信息
BOOK_TITLE = "AI赋能：青年教师教学创新实践探索"
BOOK_SUBTITLE = "——基于AI技术赋能课堂教学、作业设计、专业提升与技术应用的实践探索"

# 四个分类
categories = [
    {
        "编": "第一编",
        "四字主标题": "智启课堂",
        "副标题": "AI赋能教学新范式",
        "描述": "聚焦课堂教学创新",
        "sub_titles": [
            '一、情境创设：从"静态呈现"到"沉浸体验"',
            '二、思维激发：从"被动接受"到"主动探究"',
            '三、互动对话：从"单向讲授"到"多维碰撞"',
            '四、深度学习：从"浅层理解"到"深度建构"',
        ],
        "articles": [
            '问题引领，AI赋能——初中科学《日地月的相对运动》教学实践',
            'AI融教启思 破局素养教学——基于"嫌疑人\'k\'的现身"的课堂实践',
            '情境・情感・评价：三维赋能初中语文教学路径研究',
            'AI赋能历史课堂教学实践 — 以《辛亥革命》为例',
            'AI赋能初中文言文情境化教学与深度学习实践',
            '基于AI的情境可视化与兴趣激发实践',
            'AI赋能"中国梦"主题教学的智能体辅助探究实践',
            '基于AI实时图像生成的批判性思维培育教学实践',
            '与"豆包"结伴而行',
        ],
    },
    {
        "编": "第二编",
        "四字主标题": "智创作业",
        "副标题": "AI赋能设计新路径",
        "描述": "聚焦作业设计与精准教学",
        "sub_titles": [
            '一、分层设计：从"统一发放"到"量体裁衣"',
            '二、智能批改：从"机械重复"到"精准诊断"',
            '三、错题追踪：从"零散记录"到"系统分析"',
            '四、个性推送：从"题海战术"到"靶向训练"',
        ],
        "articles": [
            '基于数智作业的初中数学学情诊断与分层教学实践',
            'AI应用：初中英语"听说+数智作业"精准教学实践',
            'AI智能批阅与个性化推题赋能初三数学精准教学的实践案例',
            '"一核三阶"：基于科大讯飞AI的作文智能评改',
            '基于大语言模型的初中语文人物传记习作分层批改与精准反馈实践',
            '基于智慧作业平台和错题归因的初中数学精准教学实践',
            '数智赋能，科学增效——基于数智作业平台的初中科学教育"双减"加法实践',
            '基于AI的初中历史个性化作业与错题精准辅导',
        ],
    },
    {
        "编": "第三编",
        "四字主标题": "智驭技术",
        "副标题": "AI赋能应用新探索",
        "描述": "聚焦技术工具与场景应用",
        "sub_titles": [
            '一、工具选择：从"盲目跟风"到"理性适配"',
            '二、场景拓展：从"浅层尝试"到"深度融合"',
            '三、跨科融合：从"学科孤岛"到"协同创新"',
            '四、实验探究：从"抽象讲解"到"可视化呈现"',
        ],
        "articles": [
            'AI听说课堂赋能初中英语语法课的教学实践案例',
            '英语AI听说课堂赋能初一英语"人人开口"——基于实时语音评测系统的互动教学实践',
            '让数学动起来——AI赋能教学案例',
            'AI赋能初中数学相似三角形复习',
            '《等式的基本性质》教学案例分析',
            'AI赋能地理教学实践案例',
            '利用gemini3快速制作几何题配套模型一些实践尝试',
            'AI赋能的科学实验教学',
            'AI赋能微观可视化：摩擦起电的电子转移探究教学实践',
            'AI赋能初中科学凸透镜成像精准教学实践案例',
            'AI赋能初中科学精准教学—从课堂到课后的全场景实践与成效',
            '数据跑起来，教学更明白 ——智慧操场赋能精准教学',
            'AI赋能初中社会学科教学的实践探索',
            'AI赋能心理健康教育教学应用案例',
            'AI赋能教育教学应用的实践案例',
            '双轨并行：信息化工具与DeepSeek智能体在初中数学实验教学中的融合应用研究',
        ],
    },
    {
        "编": "第四编",
        "四字主标题": "智育良师",
        "副标题": "AI赋能成长新生态",
        "描述": "聚焦教师专业发展",
        "sub_titles": [
            '一、教研研修：从"闭门造车"到"云端共研"',
            '二、教学反思：从"感性认知"到"理性建构"',
            '三、专业阅读：从"碎片浏览"到"深度萃取"',
            '四、教学相长：从"单向传授"到"协同共进"',
        ],
        "articles": [
            '巧用AI绘图点亮细节描写',
            '基于人工智能通识教育的学科教学实践——以初中道德与法治《坚守公平》为例',
            'AI赋能下"入境—入心—入情"三阶深度阅读的教学探索——以《桃花源记》为例',
            'AI赋能：让"笨拙"的背影"触手可及" ——以《背影》一课的情感体验深化为例',
        ],
    },
]


def create_word_outline():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading(BOOK_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x8a)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(BOOK_SUBTITLE)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x5a, 0x6c, 0x7d)
    
    doc.add_paragraph('')
    
    # 总览表
    doc.add_heading('一、书册框架总览', level=1)
    
    table = doc.add_table(rows=5, cols=4, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['编', '四字主标题', '副标题', '文章数量']
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for i, cat in enumerate(categories):
        row = table.rows[i+1]
        row.cells[0].text = cat['编']
        row.cells[1].text = cat['四字主标题']
        row.cells[2].text = cat['副标题']
        row.cells[3].text = f"{len(cat['articles'])}篇"
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    # 各编详细目录
    for cat in categories:
        doc.add_heading(f'{cat["编"]}：{cat["四字主标题"]} · {cat["副标题"]}', level=1)
        
        p = doc.add_paragraph()
        run = p.add_run(f'📌 聚焦：{cat["描述"]}')
        run.font.color.rgb = RGBColor(0x2e, 0x86, 0xc1)
        
        # 小目录标题
        doc.add_heading('小目录标题（对仗）', level=2)
        for sub in cat['sub_titles']:
            p = doc.add_paragraph()
            run = p.add_run(sub)
            run.font.bold = True
            run.font.size = Pt(13)
        
        # 对应文章
        doc.add_heading('对应文章', level=2)
        for i, article in enumerate(cat['articles'], 1):
            p = doc.add_paragraph()
            run = p.add_run(f'{i}. {article}')
            run.font.size = Pt(11)
        
        doc.add_paragraph('')
    
    # 对仗结构说明
    doc.add_heading('二、对仗结构说明', level=1)
    
    table = doc.add_table(rows=5, cols=4, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['编', '四字主标题', '副标题', '小标题模式']
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    patterns = ['从A到B（对比式）', '从A到B（对比式）', '从A到B（对比式）', '从A到B（对比式）']
    for i, cat in enumerate(categories):
        row = table.rows[i+1]
        row.cells[0].text = cat['编']
        row.cells[1].text = cat['四字主标题']
        row.cells[2].text = cat['副标题']
        row.cells[3].text = patterns[i]
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    # 四字主标题对仗说明
    doc.add_heading('四字主标题对仗', level=2)
    p = doc.add_paragraph()
    p.add_run('四字主标题对仗：')
    run = p.add_run('启（开启）· 创（创造）· 驭（驾驭）· 育（培育）')
    run.font.bold = True
    p.add_run('——四个动词，词性一致，意境递进。')
    
    word_path = os.path.join(OUTPUT_DIR, '书册目录框架_AI赋能青年教师教学创新实践探索.docx')
    doc.save(word_path)
    print(f"Word文档已生成：{word_path}")
    return word_path


if __name__ == '__main__':
    word_path = create_word_outline()
    print(f"\n✅ 文件生成完成！")
    print(f'Word: {word_path}')
