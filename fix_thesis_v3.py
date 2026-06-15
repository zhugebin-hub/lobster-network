#!/usr/bin/env python3
"""
叶畏兵毕业论文完善 - 全面修复脚本
"""
import docx
import re
import copy
from docx.oxml.ns import qn

INPUT = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
OUTPUT = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"

doc = docx.Document(INPUT)

# ==========================================
# 第一步：统计和分析
# ==========================================
print("=== 分析原文档 ===")

# 找出所有包含英文引号的段落
quote_paragraphs = []
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '"' in t or '"' in t:
        quote_paragraphs.append((i, t))

print(f"包含英文引号的段落: {len(quote_paragraphs)} 处")

# 检查77/78合并
merged_77 = None
for i, p in enumerate(doc.paragraphs):
    if '[77]' in p.text and '[78]' in p.text:
        merged_77 = (i, p.text)
        print(f"发现77/78合并: 段落{i}")

# ==========================================
# 第二步：修复英文引号
# ==========================================
print("\n=== 修复英文引号 ===")

fix_count = 0

for p in doc.paragraphs:
    for run in p.runs:
        if not run.text:
            continue
        
        text = run.text
        new_text = []
        quote_stack = []
        
        i = 0
        while i < len(text):
            ch = text[i]
            if ch == '"' or ch == '"':
                # 开引号
                new_text.append('"')
            elif ch == '"' or ch == '"':
                # 闭引号
                new_text.append('"')
            else:
                new_text.append(ch)
            i += 1
        
        fixed = ''.join(new_text)
        if fixed != text:
            run.text = fixed
            fix_count += 1

print(f"引号修复: {fix_count} 处")

# ==========================================
# 第三步：修复77/78合并
# ==========================================
print("\n=== 修复参考文献合并 ===")

for i in range(len(doc.paragraphs)):
    p = doc.paragraphs[i]
    t = p.text
    # 检查77包含78
    if '[77]' in t and '[78]' in t:
        # 拆分
        idx_78 = t.index('[78]')
        text_77 = t[:idx_78].strip()
        text_78 = t[idx_78:].strip()
        
        # 更新77
        p.text = text_77
        
        # 在77后面插入78
        new_p_elem = copy.deepcopy(p._p)
        # 清除原有run内容
        for run_elem in new_p_elem.findall('.//' + qn('w:r')):
            new_p_elem.remove(run_elem)
        # 添加新内容
        new_run = docx.oxml.OxmlElement(qn('w:r'))
        new_run_props = docx.oxml.OxmlElement(qn('w:rPr'))
        new_run.append(new_run_props)
        new_run_text = docx.oxml.OxmlElement(qn('w:t'))
        new_run_text.text = text_78
        new_run_text.set(qn('xml:space'), 'preserve')
        new_run.append(new_run_text)
        new_p_elem.append(new_run)
        
        p._p.addnext(new_p_elem)
        print(f"已拆分77: {text_77[:60]}...")
        print(f"已拆分78: {text_78[:60]}...")
        break

# ==========================================
# 第四步：保存
# ==========================================
doc.save(OUTPUT)
print(f"\n✅ 保存完成: {OUTPUT}")
