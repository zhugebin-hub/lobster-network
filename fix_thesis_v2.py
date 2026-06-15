#!/usr/bin/env python3
"""
完善叶畏兵毕业论文 v2
修复内容：
1. 英文引号 → 中文引号
2. 修复参考文献77/78合并
3. 其他标点规范化
4. 生成脚注-参考文献对应表
"""

import docx
import re
import copy

INPUT = "/home/admin/.openclaw/media/inbound/8466908b-15d1-4924-b3d7-7dafcce89578.docx"
OUTPUT = "/home/admin/.openclaw/workspace/叶畏兵_论文_完善版.docx"

doc = docx.Document(INPUT)

# 统计原始状态
total_paragraphs = len(doc.paragraphs)
total_tables = len(doc.tables)
total_runs = sum(len(p.runs) for p in doc.paragraphs)
print(f"段落数: {total_paragraphs}")
print(f"表格数: {total_tables}")
print(f"runs总数: {total_runs}")

# ========== 分析引用标记 ==========
# 找到所有可能的引用位置（段末空格）
citation_candidates = []
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if t and t.endswith(' ') and len(t) > 10:
        # 检查是否在正文区域（非标题、非摘要、非参考文献）
        style = p.style.name if p.style else ''
        if 'Heading' not in style and 'Caption' not in style and 'toc' not in style:
            citation_candidates.append({
                'idx': i,
                'text': t[:120],
                'style': style
            })

print(f"\n引用标记候选: {len(citation_candidates)} 处")
for c in citation_candidates[:5]:
    print(f"  [{c['idx']}] ...{c['text'][-60:]}")

# ========== 修复策略 ==========
# 策略1: 在run级别替换英文引号为中文引号
# 策略2: 处理段落级问题

fixes_applied = 0
quote_fixes = []

for p_idx, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if not run.text:
            continue
        
        original = run.text
        fixed = original
        
        # 1. 英文双引号 → 中文双引号（智能成对替换）
        if '"' in fixed or '"' in fixed:
            result = []
            depth = 0
            for ch in fixed:
                if ch == '"' or ch == '"':
                    if depth % 2 == 0:
                        result.append('"')  # 开引号
                    else:
                        result.append('')  # 闭引号
                    depth += 1
                else:
                    result.append(ch)
            new_text = ''.join(result)
            if new_text != fixed:
                quote_fixes.append({
                    'para': p_idx,
                    'from': original[:80],
                    'to': new_text[:80]
                })
                fixed = new_text
                fixes_applied += 1
        
        # 2. 半角逗号/分号/冒号 → 全角（仅中文段落）
        has_chinese = bool(re.search(r'[\u4e00-\u9fff]', fixed))
        if has_chinese:
            # 中文后的半角标点
            old = fixed
            fixed = re.sub(r'([\u4e00-\u9fff\uff01-\uff5f]):', r'\1：', fixed)
            fixed = re.sub(r'([\u4e00-\u9fff\uff01-\uff5f]);', r'\1；', fixed)
            fixed = re.sub(r'([\u4e00-\u9fff\uff01-\uff5f]),(?=[\u4e00-\u9fff])', r'\1，', fixed)
            if fixed != old:
                fixes_applied += 1
        
        # 3. 半角括号 → 全角（中文段落中）
        if has_chinese:
            old = fixed
            fixed = fixed.replace('（', '(').replace('）', ')')  # 先不替换括号，保留原样
            fixed = old  # 括号暂时不动
        
        # 4. 特殊错误修复
        if '"' in fixed:
            fixed = fixed.replace('"', '"')
        if '"' in fixed:
            fixed = fixed.replace('"', '"')
        
        run.text = fixed

# ========== 修复参考文献表格/段落 ==========
# 找到参考文献部分，检查77/78合并
ref_start = None
ref_end = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '参考文献':
        ref_start = i
    if p.text.strip().startswith('附 录') and ref_start:
        ref_end = i
        break

print(f"\n参考文献范围: 段落 {ref_start} - {ref_end}")

# 检查77/78
for i in range(ref_start or 0, ref_end or len(doc.paragraphs)):
    t = doc.paragraphs[i].text
    if '[77]' in t and '[78]' in t:
        print(f"\n发现合并问题! 段落{i}:")
        print(f"  {t}")
        # 拆分：在[78]处分割
        if '[78]' in t:
            parts = t.split('[78]')
            doc.paragraphs[i].text = parts[0].strip()
            # 插入新段落作为78
            new_p = copy.deepcopy(doc.paragraphs[i])
            new_p.text = '[78]' + parts[1].strip()
            # 在i+1处插入
            i_elem = doc.paragraphs[i]._p
            i_elem.addnext(new_p._p)
            fixes_applied += 1
            print(f"  已拆分!")
            break

# ========== 保存 ==========
doc.save(OUTPUT)
print(f"\n✅ 修复完成!")
print(f"  引号替换: {len(quote_fixes)} 处")
print(f"  总修复: {fixes_applied} 处")
print(f"  输出: {OUTPUT}")

# ========== 生成脚注对应表 ==========
# 分析正文中引用与参考文献的对应关系
print("\n" + "="*60)
print("📋 建议的脚注-参考文献对应表（部分示例）")
print("="*60)

# 读取关键引用段落
key_citations = [
    ("2016年全国宗教工作会议", "习近平. 在全国宗教工作会议上的讲话[N]. 人民日报，2016-04-24（002）."),
    ("2021年全国宗教工作会议", "习近平. 在全国宗教工作会议上的讲话[N]. 人民日报，2021-12-05（002）."),
    ("王伟光-五个认同", "王伟光. 不断增进'五个认同' 深入推进我国宗教中国化[N]. 人民政协报，2022-04-14（010）."),
    ("游斌-宗教中国化", "游斌. 深入推进我国宗教中国化，增进'五个认同'[N]. 中国民族报，2022-04-08（005）."),
    ("张志刚-宗教中国化", "张志刚. '宗教中国化'义理沉思[J]. 世界宗教研究，2016（3）：1-12."),
    ("牟钟鉴-宗教中国化", "牟钟鉴. 如何深入理解'坚持中国化方向'[J]. 世界宗教研究，2016（3）：13-18."),
    ("何虎生-发展历程", "何虎生，胡竞方. 新中国·新时期·新时代坚持我国宗教中国化的发展历程研究[J]. 世界宗教研究，2020（1）：1-12."),
    ("Nygren-Agape", "Nygren, Anders. Agape and Eros: The Christian Idea of Love. Chicago: University of Chicago Press, 1982."),
    ("Aquinas-神学大全", "Aquinas, Thomas. Summa Theologica. II-II, QQ. 23-46."),
    ("Bhabha-第三空间", "Bhabha, Homi K. The Location of Culture. London: Routledge, 1994."),
    ("刘耘华-跨文化诠释", "刘耘华. 诠释的圆环[M]. 北京：北京大学出版社，2005."),
    ("Rawls-交叠共识", "Rawls, John. Political Liberalism. New York: Columbia University, 1993."),
    ("Bevans-处境化神学", "Bevans, Stephen B. Models of Contextual Theology. Maryknoll: Orbis Books, 2002."),
    ("Schreiter-本土神学", "Schreiter, Robert J. Constructing Local Theologies. Maryknoll: Orbis Books, 1985."),
    ("Casanova-公共宗教", "Casanova, José. Public Religions in the Modern World. Chicago: University of Chicago Press, 1994."),
    ("段琦-本色化", "段琦. 奋进的历程：中国基督教的本色化[M]. 北京：商务印书馆，2004."),
    ("杨凤岗-叠合身份", "杨凤岗. 中国基督徒的社会认同与'叠合身份'困境[J]. 社会，2012，32（4）：1-22."),
]

for mention, ref in key_citations:
    print(f"  提及: {mention}")
    print(f"  对应: {ref}")
    print()
