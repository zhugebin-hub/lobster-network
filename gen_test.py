# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)
style.paragraph_format.line_spacing = 1.5

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('浙教版七年级下册信息科技综合测试卷')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'

# 信息栏
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('（考试时间：90分钟  满分：100分）')
run.font.size = Pt(12)

# 学生信息
info2 = doc.add_paragraph()
info2.paragraph_format.space_before = Pt(10)
run = info2.add_run('班级：__________    姓名：__________    得分：__________')
run.font.size = Pt(12)

doc.add_paragraph('_' * 70)

# ========== 第一部分：选择题 ==========
p = doc.add_paragraph()
run = p.add_run('一、选择题（每题3分，共30题，共90分）')
run.bold = True
run.font.size = Pt(14)

questions = [
    ("1. 物联网的英文缩写是（  ）",
     "A. IT      B. IoT      C. AI      D. ERP",
     "B"),
    ("2. 物联网被定义为（  ）",
     "A. 计算机与计算机相连的网络\nB. 物和物相连的互联网络\nC. 手机与手机相连的网络\nD. 人与物相连的网络",
     "B"),
    ("3. 物联网的雏形之一是20世纪90年代的（  ）",
     "A. 智能冰箱    B. 特洛伊咖啡壶    C. 电子宠物    D. 掌上电脑",
     "B"),
    ("4. 以下不属于物联网特征的是（  ）",
     "A. 全面感知    B. 可靠传递    C. 智能处理    D. 独立运行",
     "D"),
    ("5. 物联网一般由几层组成？（  ）",
     "A. 三层    B. 四层    C. 五层    D. 两层",
     "B"),
    ("6. 物联网架构中，相当于人的眼、耳等器官的是（  ）",
     "A. 应用层    B. 网络层    C. 感知层    D. 平台层",
     "C"),
    ("7. 以下属于物联网网络层设备的是（  ）",
     "A. 传感器    B. 摄像头    C. Wi-Fi    D. 云服务器",
     "C"),
    ("8. 物联网架构中，云平台属于（  ）",
     "A. 感知层    B. 网络层    C. 平台层    D. 应用层",
     "C"),
    ("9. 以下不属于物联网应用领域的是（  ）",
     "A. 智慧城市    B. 智慧医疗    C. 传统手工业    D. 智慧物流",
     "C"),
    ("10. 大数据的特征不包括（  ）",
     "A. 数据量大    B. 速度快    C. 类型多    D. 价值密度高",
     "D"),
    ("11. 人工智能与物联网结合被称为（  ）",
     "A. IoB      B. AIoT      C. IaaS      D. SaaS",
     "B"),
    ("12. 感知技术包含传感技术和（  ）",
     "A. 识别技术    B. 云计算    C. 数据库技术    D. 编程技术",
     "A"),
    ("13. 传感器在物联网中被比喻为（  ）",
     "A. 人的大脑    B. 人的神经网络    C. 人的感觉器官    D. 人的四肢",
     "C"),
    ("14. 二维码使用什么图案表示二进制数据？（  ）",
     "A. 黑白圆形图案    B. 黑白矩形图案    C. 彩色圆形图案    D. 彩色矩形图案",
     "B"),
    ("15. 高速公路ETC使用的是（  ）",
     "A. 条形码识别技术    B. 二维码识别技术    C. 射频识别技术（RFID）    D. 人脸识别技术",
     "C"),
    ("16. 传感器的工作流程是（  ）",
     "A. 检测元件→被测的量→输出电信号\nB. 被测的量→检测元件→输出电信号\nC. 输出电信号→检测元件→被测的量\nD. 被测的量→输出电信号→检测元件",
     "B"),
    ("17. 光敏传感器对应人类的（  ）",
     "A. 听觉    B. 触觉    C. 视觉    D. 嗅觉",
     "C"),
    ("18. 以下不属于按照输出信号类型分类的传感器交互方式的是（  ）",
     "A. 数字信号    B. 模拟信号    C. 光信号    D. 以上都是交互方式",
     "C"),
    ("19. 蓝牙技术中，一个主设备可以匹配（  ）",
     "A. 只能一个从设备\nB. 一个或多个从设备\nC. 最多两个从设备\nD. 不能匹配从设备",
     "B"),
    ("20. Wi-Fi属于（  ）",
     "A. 有线通信技术    B. 无线局域网通信技术\nC. 卫星通信技术    D. 红外通信技术",
     "B"),
    ("21. NFC的全称是（  ）",
     "A. 近场通信（非接触式识别和互联技术）\nB. 远程通信技术\nC. 蓝牙通信技术    D. 卫星通信技术",
     "A"),
    ("22. MQTT协议最初的设计用途是（  ）",
     "A. 网页浏览\nB. 将石油管道上的传感器与卫星相连接\nC. 电子邮件传输    D. 文件下载",
     "B"),
    ("23. MQTT协议是一种（  ）",
     "A. 请求/响应协议    B. 发布/订阅协议    C. 点对点协议    D. 广播协议",
     "B"),
    ("24. MQTT协议包含三种角色，分别是（  ）",
     "A. 服务器、客户端、路由器\nB. 中介、发布者、订阅者\nC. 发送者、接收者、转发者\nD. 主机、从机、网关",
     "B"),
    ("25. 信息系统由硬件、软件、通信网络、数据和（  ）构成",
     "A. 服务器    B. 路由器    C. 用户    D. 数据库",
     "C"),
    ("26. 物联网数据的汇集处理方式中，将接收到的数据立刻进行处理并保存的方式称为（  ）",
     "A. 批处理    B. 流处理    C. 缓存处理    D. 异步处理",
     "B"),
    ("27. 暂时保存数据，定期进行集中化处理的方式称为（  ）",
     "A. 流处理    B. 批处理    C. 实时处理    D. 缓存处理",
     "B"),
    ("28. 根据控制主体的不同，控制形式可分为（  ）",
     "A. 远程控制和本地控制\nB. 手动控制和自动控制\nC. 开环控制和闭环控制\nD. 集中控制和分散控制",
     "B"),
    ("29. 在自动浇灌系统中，当土壤湿度小于阈值时，应该（  ）",
     "A. 停止水泵    B. 启动水泵    C. 不做任何操作    D. 报警",
     "B"),
    ("30. 物联系统原型搭建的流程顺序正确的是（  ）",
     "A. 需求分析→硬件搭建→架构设计→软件开发→系统测试→原型验证\nB. 需求分析→架构设计→硬件搭建→软件开发→系统测试→原型验证\nC. 架构设计→需求分析→硬件搭建→软件开发→原型验证→系统测试\nD. 硬件搭建→软件开发→需求分析→架构设计→系统测试→原型验证",
     "B"),
]

for i, (q, options, answer) in enumerate(questions):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(q)
    run.font.size = Pt(11)
    
    lines = options.split('\n')
    for line in lines:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.0)
        run2 = p2.add_run('    ' + line)
        run2.font.size = Pt(11)

# ========== 第二部分：判断题 ==========
doc.add_paragraph('_' * 70)

p = doc.add_paragraph()
run = p.add_run('二、判断题（每题1分，共10题，共10分）')
run.bold = True
run.font.size = Pt(14)

tf_questions = [
    ("1. 物联网就是互联网，两者没有区别。（  ）", "×"),
    ("2. 物联网的感知层相当于人的神经网络。（  ）", "×"),
    ("3. 特洛伊咖啡壶是物联网的雏形之一。（  ）", "√"),
    ("4. 条形码和二维码都是利用光电转换设备进行识别的技术。（  ）", "√"),
    ("5. 压敏传感器对应人类的听觉。（  ）", "×"),
    ("6. 蓝牙技术中，一个主设备只能匹配一个从设备。（  ）", "×"),
    ("7. MQTT协议比HTTP协议更节约网络带宽和系统资源。（  ）", "√"),
    ("8. 物联网服务器的主要作用是收发数据、处理数据和存储数据。（  ）", "√"),
    ("9. 自动控制是指人通过计算机、手机等终端发送指令控制设备。（  ）", "×"),
    ("10. 物联系统原型调试时，可以分块进行验证。（  ）", "√"),
]

for i, (q, answer) in enumerate(tf_questions):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(q)
    run.font.size = Pt(11)

# ========== 答案页 ==========
doc.add_page_break()

p = doc.add_paragraph()
run = p.add_run('参考答案')
run.bold = True
run.font.size = Pt(16)
run.font.name = '黑体'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 选择题答案
p = doc.add_paragraph()
run = p.add_run('一、选择题答案（每题3分，共90分）')
run.bold = True
run.font.size = Pt(13)

answers = [q[2] for q in questions]
col1 = answers[0:10]
col2 = answers[10:20]
col3 = answers[20:30]

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
for j in range(10):
    run = p.add_run('  第{}题：{}      '.format(j+1, col1[j]))
    run.font.size = Pt(11)
p = doc.add_paragraph()
for j in range(10):
    run = p.add_run('  第{}题：{}      '.format(j+11, col2[j]))
    run.font.size = Pt(11)
p = doc.add_paragraph()
for j in range(10):
    run = p.add_run('  第{}题：{}      '.format(j+21, col3[j]))
    run.font.size = Pt(11)

# 判断题答案
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(15)
run = p.add_run('二、判断题答案（每题1分，共10分）')
run.bold = True
run.font.size = Pt(13)

tf_answers = [q[1] for q in tf_questions]
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
for j in range(10):
    run = p.add_run('  第{}题：{}      '.format(j+1, tf_answers[j]))
    run.font.size = Pt(11)

# 保存
output_path = '/home/admin/.openclaw/workspace/浙教版七下信息科技综合测试卷.docx'
doc.save(output_path)
print('OK: ' + output_path)
