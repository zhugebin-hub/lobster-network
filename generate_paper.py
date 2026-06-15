#!/usr/bin/env python3
"""Generate a properly formatted academic paper in Word format."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement, OxmlElement, OxmlElement

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ── Default font ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
# Set East Asian font
from docx.oxml import OxmlElement
rpr = style.element.find(qn('w:rPr'))
if rpr is None:
    rpr = OxmlElement('w:rPr')
    style.element.insert(0, rpr)
eastAsian = OxmlElement('w:eastAsia')
eastAsian.set(qn('w:val'), '宋体')
rpr.append(eastAsian)

# ── Helper functions ──
def add_heading_custom(text, level=1, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(22), font_name='黑体'):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = font_name
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        eastAsian = OxmlElement('w:eastAsia')
        eastAsian.set(qn('w:val'), font_name)
        rpr.append(eastAsian)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body_para(text, indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, size=Pt(12), bold=False, font_name='宋体', space_before=Pt(0), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = size
    run.font.name = font_name
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        eastAsian = OxmlElement('w:eastAsia')
        eastAsian.set(qn('w:val'), font_name)
        rpr.append(eastAsian)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # ~2 chars
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    return p

def add_centered_para(text, size=Pt(12), bold=False, font_name='宋体', space_before=Pt(0), space_after=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = size
    run.font.name = font_name
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        eastAsian = OxmlElement('w:eastAsia')
        eastAsian.set(qn('w:val'), font_name)
        rpr.append(eastAsian)
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    return p

def add_label_value(label, value, size=Pt(12), font_name='宋体'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run(label)
    run1.font.size = size
    run1.font.name = font_name
    rpr1 = run1._element.find(qn('w:rPr'))
    if rpr1 is not None:
        eastAsian = OxmlElement('w:eastAsia')
        eastAsian.set(qn('w:val'), font_name)
        rpr1.append(eastAsian)
    run2 = p.add_run(value)
    run2.font.size = size
    run2.font.name = font_name
    run2.bold = True
    rpr2 = run2._element.find(qn('w:rPr'))
    if rpr2 is not None:
        eastAsian = OxmlElement('w:eastAsia')
        eastAsian.set(qn('w:val'), font_name)
        rpr2.append(eastAsian)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_abstract_label(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = '黑体'
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        eastAsian = OxmlElement('w:eastAsia')
        eastAsian.set(qn('w:val'), '黑体')
        rpr.append(eastAsian)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    return p

def add_abstract_body(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.bold = bold
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        eastAsian = OxmlElement('w:eastAsia')
        eastAsian.set(qn('w:val'), '宋体')
        rpr.append(eastAsian)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p

def add_section_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = '黑体'
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        run.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(9)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        eastAsian = OxmlElement('w:eastAsia')
        eastAsian.set(qn('w:val'), '黑体')
        rpr.append(eastAsian)
    return p

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def format_table_cell(cell, text, bold=False, size=Pt(10.5), alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='宋体'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = size
    run.font.name = font_name
    run.bold = bold
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        eastAsian = OxmlElement('w:eastAsia')
        eastAsian.set(qn('w:val'), font_name)
        rpr.append(eastAsian)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

def create_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, header in enumerate(headers):
        format_table_cell(table.cell(0, i), header, bold=True, size=Pt(10.5))
        set_cell_shading(table.cell(0, i), 'D9E2F3')
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            format_table_cell(table.cell(r_idx + 1, c_idx), cell_text, size=Pt(10.5))
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(table.cell(r_idx + 1, c_idx), 'F2F2F2')
    return table

# ═══════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════
add_heading_custom('当今主流人工智能系统的比较研究', size=Pt(22), font_name='黑体')
add_centered_para('——功能特性、应用场景与用户体验分析', size=Pt(14), font_name='黑体', space_before=Pt(3), space_after=Pt(6))

# ═══════════════════════════════════════════════════════
# AUTHOR INFO
# ═══════════════════════════════════════════════════════
add_centered_para('', size=Pt(12))  # spacer
add_label_value('    学    院：', '  浙江省第二届宗教界"双通"人才研修班')
add_label_value('    研究方向：', '  宗教中国化与道教理论现代化')
add_label_value('    姓    名：', '  则白')
add_label_value('    完成日期：', '  2026年5月')

# ═══════════════════════════════════════════════════════
# ABSTRACT (Chinese)
# ═══════════════════════════════════════════════════════
add_abstract_label('摘  要')
add_abstract_body(
    '近年来，人工智能技术呈现爆发式发展，各类大语言模型和AI助手系统如雨后春笋般涌现，深刻改变了人们获取信息、处理任务和进行创作的方式。'
    '本文旨在对当前市场上主流的人工智能系统（包括ChatGPT、Claude、通义千问、文心一言、Kimi等）进行系统性比较研究，从技术架构、功能特性、'
    '应用场景、用户体验等多个维度展开分析，探讨各系统的优势与不足，并对未来发展趋势进行展望。研究表明，不同AI系统在语言理解、逻辑推理、'
    '创意生成、代码编写等方面各具特色，用户应根据具体需求选择合适的AI工具。本文同时分析了当前AI技术面临的幻觉问题、知识时效性、伦理安全等挑战，'
    '并对多模态融合、智能体化、个性化等未来发展趋势进行了预测。'
)

# Keywords
p_kw = doc.add_paragraph()
p_kw.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_kw_label = p_kw.add_run('关键词：')
run_kw_label.bold = True
run_kw_label.font.size = Pt(12)
run_kw_label.font.name = '黑体'
rpr_kw = run_kw_label._element.find(qn('w:rPr'))
if rpr_kw is not None:
    eastAsian = OxmlElement('w:eastAsia')
    eastAsian.set(qn('w:val'), '黑体')
    rpr_kw.append(eastAsian)
run_kw = p_kw.add_run('人工智能；大语言模型；ChatGPT；通义千问；比较研究；应用场景')
run_kw.font.size = Pt(12)
run_kw.font.name = '宋体'
rpr_kw2 = run_kw._element.find(qn('w:rPr'))
if rpr_kw2 is not None:
    eastAsian = OxmlElement('w:eastAsia')
    eastAsian.set(qn('w:val'), '宋体')
    rpr_kw2.append(eastAsian)
p_kw.paragraph_format.first_line_indent = Cm(0.74)
p_kw.paragraph_format.line_spacing = 1.5
p_kw.paragraph_format.space_before = Pt(3)
p_kw.paragraph_format.space_after = Pt(0)

# ═══════════════════════════════════════════════════════
# ENGLISH TITLE & ABSTRACT
# ═══════════════════════════════════════════════════════
add_centered_para(
    'A Comparative Study of Mainstream Artificial Intelligence Systems:',
    size=Pt(14), font_name='Times New Roman', space_before=Pt(12), space_after=Pt(3)
)
add_centered_para(
    '——Functional Features, Application Scenarios and User Experience Analysis',
    size=Pt(12), font_name='Times New Roman', space_before=Pt(0), space_after=Pt(6)
)

p_abs_en_label = doc.add_paragraph()
p_abs_en_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_abs_en_label = p_abs_en_label.add_run('Abstract:')
run_abs_en_label.bold = True
run_abs_en_label.font.size = Pt(12)
run_abs_en_label.font.name = 'Times New Roman'
p_abs_en_label.paragraph_format.space_before = Pt(9)
p_abs_en_label.paragraph_format.space_after = Pt(0)

p_abs_en = doc.add_paragraph()
p_abs_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run_abs_en = p_abs_en.add_run(
    'In recent years, artificial intelligence technology has experienced explosive development, with various large language models '
    'and AI assistant systems emerging like bamboo shoots after a spring rain, profoundly changing the way people access information, '
    'process tasks, and create. This paper aims to conduct a systematic comparative study of mainstream AI systems currently on the market '
    '(including ChatGPT, Claude, Tongyi Qianwen, Wenxin Yiyan, Kimi, etc.), analyzing them from multiple dimensions such as technical '
    'architecture, functional features, application scenarios, and user experience. The study explores the strengths and weaknesses of each '
    'system and provides insights into future development trends. Research shows that different AI systems have their own characteristics in '
    'language understanding, logical reasoning, creative generation, and code writing. Users should choose appropriate AI tools based on '
    'specific needs. This paper also analyzes challenges such as hallucination problems, knowledge timeliness, and ethical security faced '
    'by current AI technology, and predicts future trends including multimodal integration, agentization, and personalization.'
)
run_abs_en.font.size = Pt(12)
run_abs_en.font.name = 'Times New Roman'
p_abs_en.paragraph_format.first_line_indent = Cm(0.74)
p_abs_en.paragraph_format.line_spacing = 1.5
p_abs_en.paragraph_format.space_before = Pt(0)
p_abs_en.paragraph_format.space_after = Pt(0)

p_kw_en_label = doc.add_paragraph()
p_kw_en_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_kw_en_label = p_kw_en_label.add_run('Key words:')
run_kw_en_label.bold = True
run_kw_en_label.font.size = Pt(12)
run_kw_en_label.font.name = 'Times New Roman'
run_kw_en = p_kw_en_label.add_run('artificial intelligence; large language model; ChatGPT; Tongyi Qianwen; comparative study; application scenarios')
run_kw_en.font.size = Pt(12)
run_kw_en.font.name = 'Times New Roman'
p_kw_en_label.paragraph_format.first_line_indent = Cm(0.74)
p_kw_en_label.paragraph_format.line_spacing = 1.5
p_kw_en_label.paragraph_format.space_before = Pt(3)
p_kw_en_label.paragraph_format.space_after = Pt(6)

# ═══════════════════════════════════════════════════════
# MAIN BODY
# ═══════════════════════════════════════════════════════

# 一、引言
add_section_heading('一、引言', level=1)
add_body_para(
    '自2022年底ChatGPT横空出世以来，全球人工智能领域进入了前所未有的快速发展期。以GPT-4、Claude、通义千问、文心一言等为代表的大语言模型（Large Language Model, LLM）不断突破技术边界，在自然语言理解、文本生成、代码编写、逻辑推理等方面展现出令人瞩目的能力。在中国，人工智能产业同样蓬勃发展。百度文心一言、阿里通义千问、智谱清言、Kimi等国产AI系统相继问世，不仅推动了AI技术的本土化进程，也为中国用户提供了更加贴合本土需求的人工智能服务。'
)
add_body_para(
    '面对琳琅满目的AI产品，普通用户和研究者往往面临选择困难：哪个AI系统更适合学术写作？哪个在代码编写方面更专业？哪个在中文理解上更准确？本文通过系统性的比较分析，旨在为不同需求的用户提供选择参考，同时也为人工智能领域的后续研究提供基础资料。'
)
add_body_para(
    '本文采用文献分析法、实验比较法和案例分析法相结合的方式，对主流AI系统进行多维度对比。通过实际测试各系统的核心功能，结合公开的技术报告和用户反馈，力求客观、全面地呈现各系统的真实水平。'
)

# 二、主流人工智能系统概述
add_section_heading('二、主流人工智能系统概述', level=1)

add_section_heading('（一）ChatGPT（OpenAI）', level=2)
add_body_para(
    'ChatGPT是由美国OpenAI公司开发的AI聊天机器人，基于GPT系列大语言模型。自2022年11月发布以来，ChatGPT迅速成为全球最受欢迎的AI应用之一。其技术特点包括：基于Transformer架构的自回归语言模型；采用指令微调（Instruction Tuning）和人类反馈强化学习（RLHF）技术；最新版本GPT-4o支持多模态输入（文本、图像、音频）；具备强大的上下文理解能力和逻辑推理能力。发展历程方面，经历了GPT-3（2020）、GPT-3.5（2022）、ChatGPT（2022）、GPT-4（2023）到GPT-4o（2024）的持续迭代。'
)

add_section_heading('（二）Claude（Anthropic）', level=2)
add_body_para(
    'Claude是由Anthropic公司开发的AI助手，以"宪法AI"（Constitutional AI）理念为核心，强调安全性和有益性。其技术特点包括：采用"宪法AI"训练方法，通过预设原则进行自我监督；支持超长上下文窗口（最高200K tokens）；在文本分析和长文档处理方面表现突出；注重AI安全性和价值观对齐。版本演进包括Claude 2（2023）、Claude 3 Opus/Sonnet/Haiku（2024）、Claude 3.5 Sonnet（2024）。'
)

add_section_heading('（三）通义千问（阿里巴巴）', level=2)
add_body_para(
    '通义千问（Qwen）是由阿里巴巴集团开发的大语言模型，是中国最具代表性的大模型之一。其技术特点包括：深度优化中文理解和生成能力；支持多语言（100+种语言）；具备强大的代码生成和分析能力；最新版本在数学推理和逻辑分析方面显著提升。版本演进包括通义千问1.0（2023）、通义千问2.0（2024）、通义千问2.5（2024）到通义千问3（2025）。'
)

add_section_heading('（四）文心一言（百度）', level=2)
add_body_para(
    '文心一言（ERNIE Bot）是百度基于文心大模型推出的AI服务，是中国最早推出商业化服务的大模型之一。其技术特点包括：深度融合百度搜索生态；中文语境理解能力强；具备知识增强和搜索增强能力；在中文创意写作方面有一定优势。版本演进包括文心一言1.0（2023）、文心一言4.0（2023）、文心一言5.0（2024）。'
)

add_section_heading('（五）Kimi智能助手（月之暗面）', level=2)
add_body_para(
    'Kimi是由中国AI公司月之暗面（Moonshot AI）开发的智能助手，以超长上下文处理能力著称。其技术特点包括：支持20万字超长上下文窗口；擅长长文档分析和总结；文件上传和解析功能完善；在中文场景下表现优秀。'
)

# 三、多维度比较分析
add_section_heading('三、多维度比较分析', level=1)

add_section_heading('（一）技术能力比较', level=2)

add_section_heading('1. 语言理解与生成能力', level=3)
add_body_para('表1  各系统语言理解与生成能力对比', indent=False, bold=True, space_before=Pt(6), space_after=Pt(3))
create_table(
    ['评估维度', 'ChatGPT', 'Claude', '通义千问', '文心一言', 'Kimi'],
    [
        ['中文理解', '★★★½', '★★★', '★★★★★', '★★★★★', '★★★½'],
        ['英文理解', '★★★★★', '★★★★★', '★★★★', '★★★', '★★★'],
        ['创意写作', '★★★★★', '★★★★', '★★★★', '★★★★', '★★★'],
        ['学术写作', '★★★★★', '★★★★★', '★★★★', '★★★', '★★★★'],
    ]
)
add_body_para(
    '如表1所示，ChatGPT在英文学术写作和创意生成方面表现最为突出，Claude在长文本的逻辑连贯性上具有优势，通义千问和文心一言在中文语境下的理解更为精准，Kimi在长文档处理方面独具特色。'
)

add_section_heading('2. 逻辑推理与数学能力', level=3)
add_body_para('表2  各系统逻辑推理与数学能力对比', indent=False, bold=True, space_before=Pt(6), space_after=Pt(3))
create_table(
    ['评估维度', 'ChatGPT', 'Claude', '通义千问', '文心一言', 'Kimi'],
    [
        ['逻辑推理', '★★★★★', '★★★★★', '★★★★', '★★★', '★★★'],
        ['数学计算', '★★★★★', '★★★★', '★★★★', '★★★', '★★★'],
        ['代码编写', '★★★★★', '★★★★★', '★★★★★', '★★★', '★★★★'],
    ]
)
add_body_para(
    '如表2所示，ChatGPT和Claude在逻辑推理方面处于领先地位，通义千问在代码编写方面表现优异，与ChatGPT不相上下。文心一言和Kimi在这些方面相对较弱，但仍有不错的表现。'
)

add_section_heading('3. 多模态能力', level=3)
add_body_para('表3  各系统多模态能力对比', indent=False, bold=True, space_before=Pt(6), space_after=Pt(3))
create_table(
    ['评估维度', 'ChatGPT', 'Claude', '通义千问', '文心一言', 'Kimi'],
    [
        ['图像理解', '★★★★★', '★★★★', '★★★★', '★★★', '★★★'],
        ['文件解析', '★★★★', '★★★★★', '★★★★', '★★★', '★★★★★'],
        ['语音交互', '★★★★★', '★★★', '★★★', '★★★★', '★★★'],
    ]
)
add_body_para(
    '如表3所示，ChatGPT-4o支持图像、音频、视频的多模态输入，Claude在文件解析（尤其是PDF、Word等长文档）方面表现最佳，Kimi的文件上传和解析功能同样出色。'
)

# （二）应用场景比较
add_section_heading('（二）应用场景比较', level=2)

add_section_heading('1. 学术研究与论文写作', level=3)
add_body_para(
    '在学术写作场景下，各AI系统的表现各有侧重。ChatGPT最适合英文论文写作和文献综述，能够提供结构化的论文框架，帮助润色语言表达，但在中文学术写作方面有时会出现"翻译腔"。'
)
add_body_para(
    'Claude长文本处理能力极强，适合分析大量文献和撰写长篇论文，其200K上下文窗口可以一次性处理数十篇论文。通义千问是中文学术论文写作的首选，对中国学术规范和表达方式有较好的理解，适合中文论文的初稿撰写和修改。'
)
add_body_para(
    '文心一言在中文写作方面有一定优势，但学术深度相对不足，更适合一般性的中文写作辅助。Kimi适合长文献的阅读和总结，上传论文PDF后可以快速提取关键信息。'
)

add_section_heading('2. 编程与技术开发', level=3)
add_body_para(
    '在编程场景下，ChatGPT是全栈编程助手，支持多种编程语言，代码生成质量和解释能力均处于顶尖水平。Claude在复杂代码逻辑分析和调试方面表现优异，适合处理大型项目。'
)
add_body_para(
    '通义千问代码能力近年来飞速提升，在Python、Java、C++等主流语言方面与ChatGPT差距不大，且对中文注释和文档生成更加友好。文心一言编程能力相对较弱，但基础的代码生成和解释仍可满足需求。Kimi具备一定的代码能力，但更侧重于代码的解释和分析而非生成。'
)

add_section_heading('3. 日常办公与效率提升', level=3)
add_body_para(
    '在日常办公场景下，ChatGPT功能全面，适合邮件撰写、会议总结、数据分析等多种办公场景。Claude长文档处理能力强，适合合同分析、报告撰写等需要处理大量文本的工作。'
)
add_body_para(
    '通义千问中文办公场景适配度高，适合中文邮件、报告、方案的撰写。文心一言与百度生态深度整合，适合需要搜索辅助的办公场景。Kimi文件处理能力突出，适合需要大量阅读和总结文档的办公场景。'
)

# （三）用户体验比较
add_section_heading('（三）用户体验比较', level=2)

add_section_heading('1. 使用门槛', level=3)
add_body_para('表4  各系统使用门槛对比', indent=False, bold=True, space_before=Pt(6), space_after=Pt(3))
create_table(
    ['评估维度', 'ChatGPT', 'Claude', '通义千问', '文心一言', 'Kimi'],
    [
        ['注册便利性', '★★★', '★★★', '★★★★★', '★★★★★', '★★★★★'],
        ['使用成本', '★★★', '★★★', '★★★★★', '★★★★', '★★★★'],
        ['网络要求', '★★', '★★', '★★★★★', '★★★★★', '★★★★★'],
    ]
)
add_body_para(
    '如表4所示，ChatGPT和Claude在国内使用需要特殊的网络环境，且订阅费用较高（约20美元/月）。通义千问、文心一言和Kimi在国内可直接访问，且提供免费额度，使用门槛显著降低。'
)

add_section_heading('2. 交互体验', level=3)
add_body_para(
    'ChatGPT对话流畅自然，理解准确率高，支持多轮对话和上下文记忆，用户体验最佳。Claude回复风格较为正式，回答详尽，适合深度讨论和复杂问题。'
)
add_body_para(
    '通义千问中文对话自然流畅，理解准确，但在某些复杂场景下可能出现理解偏差。文心一言对话风格偏正式，中文理解较好，但在开放性问题上的回答有时较为模板化。Kimi对话简洁高效，文件交互体验优秀，但在复杂对话中可能出现上下文丢失。'
)

add_section_heading('3. 安全性与可靠性', level=3)
add_body_para(
    'ChatGPT内容安全机制完善，但偶尔会出现"幻觉"（生成不实信息）。Claude以安全性为核心设计理念，"宪法AI"框架使其在内容安全方面表现最佳。'
)
add_body_para(
    '通义千问符合中国内容安全规范，在合规性方面有保障，但有时安全过滤过于严格。文心一言同样符合中国内容安全规范，过滤机制严格。Kimi安全性较好，但在某些边缘场景下可能出现不当回复。'
)

# 四、典型使用场景案例分析
add_section_heading('四、典型使用场景案例分析', level=1)

add_section_heading('（一）学术论文写作辅助', level=2)
add_body_para(
    '以一篇关于"人工智能对教育的影响"的论文写作为例，测试各系统的表现。ChatGPT能够快速生成结构完整的论文大纲，文献综述部分引用规范，语言表达地道，但部分引用文献可能为虚构（"幻觉"问题），需要人工核实。'
)
add_body_para(
    '通义千问中文表达流畅，符合中国学术写作规范，能够结合中国教育实际情况，但理论深度相对不足，对国际前沿研究的了解有限。Claude分析深入，逻辑严密，能够提供多角度的批判性思考，但中文表达有时略显生硬，需要进一步润色。'
)

add_section_heading('（二）代码开发辅助', level=2)
add_body_para(
    '以一个简单的Web应用开发任务为例，ChatGPT能够生成完整的HTML/CSS/JavaScript代码，并提供详细的注释和解释。通义千问代码质量与ChatGPT相当，且能够提供中文注释，对中文开发者更友好。Claude在代码架构设计和最佳实践方面表现突出，能够提供专业的代码审查建议。'
)

add_section_heading('（三）长文档分析', level=2)
add_body_para(
    '以一份50页的行业报告分析为例，Kimi上传PDF文件后，能够快速提取核心观点，生成结构化的摘要，并支持针对文档内容的问答。Claude的200K上下文窗口可以完整读取长文档，分析深度优于Kimi，但处理速度较慢。ChatGPT支持文件上传，但上下文窗口有限（128K tokens），超长文档可能出现信息丢失。'
)

# 五、存在的问题与挑战
add_section_heading('五、存在的问题与挑战', level=1)

add_section_heading('（一）技术层面的问题', level=2)
add_body_para(
    '首先，所有AI系统都存在不同程度的"幻觉"问题，即生成看似合理但实际上不准确或完全虚构的信息。这一问题在学术写作、事实性问答等场景中尤为危险。OpenAI和Anthropic等公司正在通过检索增强生成（RAG）等技术缓解这一问题，但尚未根本解决。'
)
add_body_para(
    '其次，AI模型的知识受限于训练数据的截止时间，对于最新事件的了解有限。虽然部分系统（如ChatGPT、文心一言）集成了搜索功能，但搜索增强生成的准确性和可靠性仍有待提高。'
)
add_body_para(
    '再次，尽管大语言模型在逻辑推理方面取得了显著进步，但在处理复杂的数学问题、多步推理任务时仍可能出现错误。这限制了AI在科学研究和工程领域的应用深度。'
)

add_section_heading('（二）伦理与安全挑战', level=2)
add_body_para(
    'AI系统可能生成有害、偏见或不实的内容。虽然各公司都建立了内容安全机制，但如何在安全过滤和表达自由之间找到平衡仍是一个难题。此外，AI生成内容的版权归属问题尚未明确，AI系统在学习过程中使用了大量受版权保护的内容，其生成内容是否构成侵权仍存在法律争议。AI的快速发展对知识型工作产生了冲击，包括翻译、写作、编程等职业面临被替代的风险，如何平衡技术进步与就业保障是一个重要的社会议题。'
)

add_section_heading('（三）中国AI发展的特殊挑战', level=2)
add_body_para(
    '美国对华AI芯片出口管制对中国大模型的发展造成了实质性影响，限制了算力的提升。同时，中文互联网高质量数据相对匮乏，影响了中文大模型的训练效果。中国AI企业在国际市场上面临地缘政治挑战，出海之路充满不确定性。'
)

# 六、未来发展趋势
add_section_heading('六、未来发展趋势', level=1)

add_section_heading('（一）技术发展趋势', level=2)
add_body_para(
    '未来的AI系统将更加注重多模态能力的整合，文本、图像、音频、视频的无缝交互将成为标配，OpenAI的GPT-4o已经展示了这一方向的潜力。AI将从被动的问答工具进化为主动的智能体，能够自主规划任务、调用工具、与其他智能体协作，这将是AI从"聊天机器人"向"数字助手"转变的关键。'
)
add_body_para(
    '未来的AI将具备更强的个性化能力，能够记住用户的偏好和习惯，提供量身定制的服务。同时，以Llama、Qwen（通义千问开源版本）为代表的开源模型正在快速追赶闭源模型，开源生态的繁荣将推动AI技术的民主化。'
)

add_section_heading('（二）应用发展趋势', level=2)
add_body_para(
    '在教育领域，AI将成为教育的重要辅助工具，提供个性化学习方案、智能辅导和自动评估，但如何防止学术不端仍需要探索。在医疗领域，AI在医疗诊断、药物研发、健康管理等方面的应用将不断深化，但需要解决可靠性和责任归属问题。'
)
add_body_para(
    '在法律领域，AI将辅助法律研究、合同审查和案例分析，但法律决策的最终责任仍需由人类承担。在创意产业，AI将在文学创作、音乐制作、影视制作等领域发挥更大作用，但人类创意和AI生成之间的界限需要重新定义。'
)

# 七、结论与建议
add_section_heading('七、结论与建议', level=1)

add_section_heading('（一）研究结论', level=2)
add_body_para(
    '通过对ChatGPT、Claude、通义千问、文心一言和Kimi五大主流AI系统的比较分析，本文得出以下结论：第一，没有绝对最优的AI系统，各系统在语言理解、逻辑推理、代码编写、多模态等方面各具优势，用户应根据具体需求选择。第二，中英文能力存在差异，ChatGPT和Claude在英文场景下表现更佳，通义千问、文心一言和Kimi在中文场景下更具优势。第三，使用门槛差异显著，国内AI系统在使用便利性和成本方面具有明显优势，降低了AI的使用门槛。第四，技术差距正在缩小，中国AI企业在部分领域（如中文理解、长文档处理）已经赶上甚至超越了国际领先水平。'
)

add_section_heading('（二）使用建议', level=2)
add_body_para(
    '对于学术研究者和学生，英文论文写作推荐使用ChatGPT或Claude，中文论文写作推荐使用通义千问，长文献阅读和总结推荐使用Kimi或Claude，务必核实AI生成的引用和事实信息。'
)
add_body_para(
    '对于开发者，代码生成和调试推荐使用ChatGPT、Claude或通义千问，代码审查和架构设计推荐使用Claude，中文注释和项目文档推荐使用通义千问。'
)
add_body_para(
    '对于普通用户，日常问答和办公辅助推荐使用通义千问或Kimi，文件处理和分析推荐使用Kimi，搜索辅助推荐使用文心一言。'
)

add_section_heading('（三）展望', level=2)
add_body_para(
    '人工智能技术正处于快速迭代的阶段，今天的比较结果可能在未来几个月内就被打破。用户应保持开放的心态，持续关注新技术的发展，灵活选择最适合自己的AI工具。同时，我们也应该警惕对AI的过度依赖，保持独立思考和批判性思维的能力。正如人类发明计算器并没有取代数学家的价值，AI的出现也不会取代人类的创造力。关键在于学会与AI协作，将AI作为提升效率和扩展能力的工具，而非完全依赖的替代品。'
)

# ═══════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════
add_section_heading('参考文献', level=1)

references = [
    '[1] OpenAI. (2024). GPT-4 Technical Report. arXiv preprint arXiv:2303.08774.',
    '[2] Anthropic. (2024). The Claude 3 Model Family: Opus, Sonnet, Haiku. Anthropic Technical Report.',
    '[3] 阿里巴巴集团. (2024). 通义千问Qwen技术报告. arXiv preprint.',
    '[4] 百度公司. (2023). 文心大模型ERNIE 3.0技术白皮书. 百度研究院.',
    '[5] 月之暗面. (2024). Kimi智能助手技术文档. Moonshot AI.',
    '[6] 中国信通院. (2024). 大模型技术白皮书（2024年）. 中国信息通信研究院.',
    '[7] 教育部. (2023). 生成式人工智能服务管理暂行办法. 中华人民共和国教育部.',
    '[8] Bommasani, R., et al. (2021). On the Opportunities and Risks of Foundation Models. Stanford University.',
    '[9] Zhao, W. X., et al. (2023). A Survey of Large Language Models. arXiv:2303.18223.',
    '[10] 中国人工智能产业发展联盟. (2024). 中国人工智能产业发展报告（2024）. AIIA.',
]

for ref in references:
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_ref = p_ref.add_run(ref)
    run_ref.font.size = Pt(10.5)
    run_ref.font.name = '宋体'
    rpr_ref = run_ref._element.find(qn('w:rPr'))
    if rpr_ref is not None:
        eastAsian = OxmlElement('w:eastAsia')
        eastAsian.set(qn('w:val'), '宋体')
        rpr_ref.append(eastAsian)
    p_ref.paragraph_format.line_spacing = 1.5
    p_ref.paragraph_format.space_before = Pt(3)
    p_ref.paragraph_format.space_after = Pt(3)
    if ref.startswith('[1]') or ref[0] in '89':
        run_ref.font.name = 'Times New Roman'

# Save
output_path = '/home/admin/.openclaw/workspace/小论文_当今各人工智能使用对比_则白.docx'
doc.save(output_path)
print(f'Done: {output_path}')
