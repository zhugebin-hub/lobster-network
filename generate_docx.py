#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成成员 4 的 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    # 标题
    title = doc.add_heading('成员 4：危机形成原因与传导机制', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 一、核心术语解释
    doc.add_heading('一、核心术语解释', level=2)
    
    doc.add_paragraph('1. MBS（Mortgage-Backed Security）- 抵押贷款支持证券', style='Heading 3')
    doc.add_paragraph('通俗解释：银行把成千上万个房贷打包成一个"大包裹"，然后把这个包裹分成小份卖给投资者。购房者每月还的房贷，就变成了投资者的收益。')
    
    doc.add_paragraph('2. ABS（Asset-Backed Security）- 资产支持证券', style='Heading 3')
    doc.add_paragraph('通俗解释：把各种能产生现金流的资产（如信用卡欠款、汽车贷款、学生贷款等）打包成证券出售。MBS 是 ABS 的一种特殊形式。')
    
    doc.add_paragraph('3. CDO（Collateralized Debt Obligation）- 担保债务凭证', style='Heading 3')
    doc.add_paragraph('通俗解释：把不同风险等级的 MBS 或其他债务再次打包、分层。就像把不同品质的水果混在一起做成水果拼盘，再按品质分档出售。')
    
    doc.add_paragraph('4. CDS（Credit Default Swap）- 信用违约互换', style='Heading 3')
    doc.add_paragraph('通俗解释：一种"保险"。买方定期付保费，如果某个债务违约，卖方赔钱。但和保险不同的是，任何人都可以买，哪怕你根本没有这个债务（相当于给别人的房子买火灾保险）。')
    
    # 二、美国住房按揭风险为何不断累积
    doc.add_heading('二、美国住房按揭风险为何不断累积', level=2)
    
    doc.add_paragraph('1. 宽松信贷政策（2001-2006 年）', style='Heading 3')
    p = doc.add_paragraph()
    p.add_run('• 美联储连续降息，联邦基金利率从 6.5% 降至 1%\n')
    p.add_run('• 银行放贷标准大幅降低\n')
    p.add_run('• 次级贷款（Subprime Loan）大量发放：向信用差、收入低、无稳定工作的人群放贷')
    
    doc.add_paragraph('2. 诱人性贷款产品', style='Heading 3')
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['贷款类型', '特点', '风险']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data = [
        ['可调利率抵押贷款 (ARM)', '前 2-3 年低利率，之后大幅上调', '还款额突然翻倍'],
        ['只还利息贷款', '前期只还利息，不还本金', '后期还款压力巨大'],
        ['负摊销贷款', '月供低于利息，欠款越还越多', '债务不降反升']
    ]
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph('3. 房价持续上涨的幻觉', style='Heading 3')
    doc.add_paragraph('• 1997-2006 年，美国房价平均上涨约 124%\n• 人们相信"房价只涨不跌"\n• 即使还不起贷款，也可以卖房获利')
    
    # 三、风险传导机制
    doc.add_heading('三、风险传导机制：从分散到放大', level=2)
    
    doc.add_paragraph('风险传导链条：', style='Heading 3')
    doc.add_paragraph('购房者违约 → MBS 违约 → CDO 违约 → 金融机构巨额亏损 → 信贷紧缩 → 经济衰退')
    
    doc.add_paragraph('1. 为什么风险看似被分散，结果却被放大？', style='Heading 3')
    
    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    headers2 = ['问题', '说明']
    for i, header in enumerate(headers2):
        table2.rows[0].cells[i].text = header
        table2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data2 = [
        ['底层资产质量差', '大量次级贷款本身就有高违约风险'],
        ['模型错误', '评级机构假设房价不会全国性下跌'],
        ['相关性被低估', '房价下跌时，所有地区的房贷同时违约'],
        ['杠杆过高', '金融机构用借来的钱大量投资 CDO'],
        ['CDS 连锁反应', 'AIG 等公司卖出大量 CDS，无力赔付']
    ]
    for i, row_data in enumerate(data2, 1):
        for j, cell_data in enumerate(row_data):
            table2.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph('2. 房价下跌如何传导到整个金融市场', style='Heading 3')
    doc.add_paragraph('① 房价下跌 (2006 年起)\n 次级贷款者无法 refinancing 或卖房\n③ 房贷违约率飙升 (从 2% 升至 10%+)\n MBS 价值暴跌\n⑤ 持有 MBS/CDO 的金融机构巨额亏损\n⑥ 银行间互不信任，停止拆借\n⑦ 信贷紧缩 → 企业无法融资 → 裁员 → 消费下降\n⑧ 经济衰退')
    
    # 四、关键数据
    doc.add_heading('四、关键数据', level=2)
    
    doc.add_paragraph('1. 美国房价指数变化（2000-2012）', style='Heading 3')
    table3 = doc.add_table(rows=7, cols=3)
    table3.style = 'Table Grid'
    headers3 = ['年份', '房价指数 (2000=100)', '变化率']
    for i, header in enumerate(headers3):
        table3.rows[0].cells[i].text = header
        table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data3 = [
        ['2000', '100', '-'],
        ['2004', '154', '+54%'],
        ['2006 (峰值)', '189', '+89%'],
        ['2008', '157', '-17%'],
        ['2010', '139', '-26%'],
        ['2012 (谷底)', '134', '-29%']
    ]
    for i, row_data in enumerate(data3, 1):
        for j, cell_data in enumerate(row_data):
            table3.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph('2. 次级贷款违约率', style='Heading 3')
    table4 = doc.add_table(rows=5, cols=2)
    table4.style = 'Table Grid'
    headers4 = ['时间', '违约率']
    for i, header in enumerate(headers4):
        table4.rows[0].cells[i].text = header
        table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data4 = [
        ['2005 年', '2.5%'],
        ['2007 年 Q1', '5.5%'],
        ['2008 年 Q1', '10.2%'],
        ['2009 年 Q1', '15.8%']
    ]
    for i, row_data in enumerate(data4, 1):
        for j, cell_data in enumerate(row_data):
            table4.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph('3. 主要金融机构损失', style='Heading 3')
    table5 = doc.add_table(rows=5, cols=2)
    table5.style = 'Table Grid'
    headers5 = ['机构', '损失金额 (亿美元)']
    for i, header in enumerate(headers5):
        table5.rows[0].cells[i].text = header
        table5.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data5 = [
        ['美林证券', '520'],
        ['花旗集团', '480'],
        ['美国银行', '350'],
        ['AIG', '1800 (含 CDS 赔付)']
    ]
    for i, row_data in enumerate(data5, 1):
        for j, cell_data in enumerate(row_data):
            table5.rows[i].cells[j].text = cell_data
    
    # 五、结论与启示
    doc.add_heading('五、结论与启示', level=2)
    
    doc.add_paragraph('1. 为什么会发生？', style='Heading 3')
    doc.add_paragraph('• 根本原因：宽松信贷 + 金融创新失控 + 监管缺位\n• 直接触发：房价下跌导致次级贷款大规模违约\n• 放大机制：复杂的金融衍生品将风险扩散到全球')
    
    doc.add_paragraph('2. 造成了什么影响？', style='Heading 3')
    doc.add_paragraph('• 全球金融危机，多国经济衰退\n• 美国失业率最高达 10%\n• 全球股市蒸发约 30 万亿美元\n• 雷曼兄弟破产，多家银行被救助')
    
    doc.add_paragraph('3. 给我们的启示', style='Heading 3')
    table6 = doc.add_table(rows=6, cols=2)
    table6.style = 'Table Grid'
    headers6 = ['启示', '说明']
    for i, header in enumerate(headers6):
        table6.rows[0].cells[i].text = header
        table6.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data6 = [
        ['金融创新需要监管', '复杂衍生品可能掩盖真实风险'],
        ['房价不会永远上涨', '资产泡沫终将破裂'],
        ['风险管理至关重要', '杠杆过高会放大损失'],
        ['信用评级不可盲信', '评级机构存在利益冲突'],
        ['系统性风险需要防范', '个别机构倒闭可能引发连锁反应']
    ]
    for i, row_data in enumerate(data6, 1):
        for j, cell_data in enumerate(row_data):
            table6.rows[i].cells[j].text = cell_data
    
    # 六、参考资料
    doc.add_heading('六、参考资料', level=2)
    doc.add_paragraph('• 《大空头》(The Big Short) 电影/书籍\n• 《监守自盗》(Inside Job) 纪录片\n• 美联储经济数据 (FRED)\n• 美国财政部危机报告 (2011)')
    
    # 保存文档
    doc.save('/home/admin/.openclaw/workspace/成员 4-危机形成原因与传导机制.docx')
    print('Word 文档已生成：成员 4-危机形成原因与传导机制.docx')

if __name__ == '__main__':
    create_document()
