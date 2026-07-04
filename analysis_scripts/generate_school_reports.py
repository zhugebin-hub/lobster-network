#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成学校错题报告 - Excel + Word 格式
"""

import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 题目内容
questions = {
    1: {"content": "少先队是（ ）创立和领导的", "answer": "D", "category": "少先队组织"},
    2: {"content": "队旗红旗象征（ ）", "answer": "C", "category": "少先队标志"},
    3: {"content": "少先队集会时由谁带领呼号", "answer": "C", "category": "少先队礼仪"},
    4: {"content": "队会中呼号的口号是", "answer": "C", "category": "少先队礼仪"},
    5: {"content": "10 月 13 日纪念什么", "answer": "B", "category": "少先队知识"},
    6: {"content": "小队一般由几人组成", "answer": "B", "category": "少先队组织"},
    7: {"content": "什么是少先队的生命", "answer": "B", "category": "少先队活动"},
    8: {"content": "不属于少先队作风的是", "answer": "B", "category": "少先队作风"},
    9: {"content": "加入少先队的年龄", "answer": "A", "category": "少先队知识"},
    10: {"content": "红领巾代表什么的一角", "answer": "A", "category": "少先队标志"},
    11: {"content": "维护国家荣誉是", "answer": "A", "category": "公民责任"},
    12: {"content": "社会主义核心价值观个人层面", "answer": "C", "category": "价值观"},
    13: {"content": "不是传统美德的是", "answer": "C", "category": "传统美德"},
    14: {"content": "属于社会公德表现的是", "answer": "B", "category": "社会公德"},
    15: {"content": "小组合作正确做法", "answer": "B", "category": "团队合作"},
    16: {"content": "看到乱涂乱画怎么做", "answer": "D", "category": "社会公德"},
    17: {"content": "答应同学的事情应该", "answer": "C", "category": "诚信"},
    18: {"content": "不属于传统节日的是", "answer": "B", "category": "传统文化"},
    19: {"content": "升国旗奏国歌时应该", "answer": "D", "category": "礼仪规范"},
    20: {"content": "班规制定正确说法", "answer": "B", "category": "民主参与"},
}

# 读取数据
input_file = '/home/admin/.openclaw/media/inbound/262cd4ab-93a6-44e8-b1fc-5e5b3bd7131a.xlsx'
df = pd.read_excel(input_file)
df_clean = df[df['得分'] > 0].copy()

# 转换对/错为 1/0
question_cols = [i for i in range(1, 21)]
for q in question_cols:
    df_clean[q] = df_clean[q].map({'对': 1, '错': 0})

output_dir = Path('/home/admin/.openclaw/workspace/stock-reports')
output_dir.mkdir(parents=True, exist_ok=True)

schools = df_clean['学校'].unique()

# ============ 1. 生成综合错题统计表（Excel） ============
excel_data = []

for school in schools:
    school_df = df_clean[df_clean['学校'] == school]
    for q in question_cols:
        error_count = (school_df[q] == 0).sum()
        total = len(school_df)
        error_rate = error_count / total * 100
        correct_rate = 100 - error_rate
        
        excel_data.append({
            '学校': school,
            '题号': f'第{q}题',
            '知识点': questions[q]['category'],
            '题目内容': questions[q]['content'],
            '正确答案': questions[q]['answer'],
            '错误人数': error_count,
            '总人数': total,
            '错误率': f'{error_rate:.2f}%',
            '正确率': f'{correct_rate:.2f}%',
            '重点关注': '是' if error_rate > 50 else '否'
        })

excel_df = pd.DataFrame(excel_data)
excel_file = output_dir / '三校错题统计总表.xlsx'
excel_df.to_excel(excel_file, index=False)
print(f"✓ 错题统计总表已保存：{excel_file}")

# ============ 2. 生成每所学校的 Word 报告 ============

for school in schools:
    school_df = df_clean[df_clean['学校'] == school]
    total_students = len(school_df)
    avg_score = school_df['得分'].mean()
    max_score = school_df['得分'].max()
    min_score = school_df['得分'].min()
    
    # 创建 Word 文档
    doc = Document()
    
    # 标题
    title = doc.add_heading(f'{school} 少先队知识测试错题分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_heading('一、基本信息', level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    basic_info = [
        ('学校名称', school),
        ('参测人数', str(total_students)),
        ('平均分', f'{avg_score:.2f}分'),
        ('最高分', f'{max_score}分'),
        ('最低分', f'{min_score}分'),
    ]
    
    for i, (label, value) in enumerate(basic_info):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
    
    # 成绩分布
    doc.add_heading('二、成绩分布', level=1)
    scores = school_df['得分']
    dist = {
        '90 分以上': (scores >= 90).sum(),
        '80-89 分': ((scores >= 80) & (scores < 90)).sum(),
        '70-79 分': ((scores >= 70) & (scores < 80)).sum(),
        '60-69 分': ((scores >= 60) & (scores < 70)).sum(),
        '60 分以下': (scores < 60).sum(),
    }
    
    dist_table = doc.add_table(rows=6, cols=2)
    dist_table.style = 'Table Grid'
    dist_table.cell(0, 0).text = '分数段'
    dist_table.cell(0, 1).text = '人数'
    
    for i, (range_name, count) in enumerate(dist.items(), 1):
        dist_table.cell(i, 0).text = range_name
        dist_table.cell(i, 1).text = str(count)
    
    # 错题统计
    doc.add_heading('三、错题统计（按正确率排序）', level=1)
    
    error_stats = []
    for q in question_cols:
        error_count = (school_df[q] == 0).sum()
        error_rate = error_count / total_students * 100
        correct_rate = 100 - error_rate
        error_stats.append({
            '题号': q,
            '错误人数': error_count,
            '错误率': error_rate,
            '正确率': correct_rate
        })
    
    # 按错误率从高到低排序
    error_stats.sort(key=lambda x: x['错误率'], reverse=True)
    
    error_table = doc.add_table(rows=1, cols=6)
    error_table.style = 'Table Grid'
    header_cells = error_table.rows[0].cells
    headers = ['题号', '知识点', '题目内容', '错误人数', '错误率', '重点关注']
    for i, h in enumerate(headers):
        header_cells[i].text = h
    
    for stat in error_stats:
        q = stat['题号']
        row = error_table.add_row().cells
        row[0].text = f'第{q}题'
        row[1].text = questions[q]['category']
        row[2].text = questions[q]['content']
        row[3].text = str(stat['错误人数'])
        row[4].text = f'{stat["错误率"]:.2f}%'
        row[5].text = '★' if stat['错误率'] > 50 else ''
    
    # 重点错题分析
    doc.add_heading('四、重点错题分析（错误率>50%）', level=1)
    
    focus_errors = [s for s in error_stats if s['错误率'] > 50]
    
    if focus_errors:
        for stat in focus_errors:
            q = stat['题号']
            para = doc.add_paragraph()
            q_content = questions[q]['content']
            para.add_run(f'第{q}题：{q_content}').bold = True
            doc.add_paragraph(f'知识点：{questions[q]["category"]}')
            doc.add_paragraph(f'错误人数：{stat["错误人数"]}人（{stat["错误率"]:.2f}%）')
            doc.add_paragraph(f'正确答案：{questions[q]["answer"]}')
            
            # 教学建议
            suggestion_para = doc.add_paragraph()
            suggestion_para.add_run('教学建议：').bold = True
            if questions[q]['category'] == '少先队活动':
                suggestion_para.add_run('强调"活动是少先队的生命"这一核心概念，通过实际活动体验加深理解。')
            elif questions[q]['category'] == '少先队标志':
                suggestion_para.add_run('展示队旗、红领巾实物，讲解象征意义，组织辨认练习。')
            elif questions[q]['category'] == '少先队礼仪':
                suggestion_para.add_run('现场演示呼号流程，反复练习，形成肌肉记忆。')
            elif questions[q]['category'] == '价值观':
                suggestion_para.add_run('结合生活实例讲解社会主义核心价值观，制作记忆卡片。')
            else:
                suggestion_para.add_run('课堂重点讲解，布置相关练习题巩固。')
            
            doc.add_paragraph()  # 空行
    else:
        doc.add_paragraph('本校无错误率超过 50% 的题目，整体掌握情况良好。')
    
    # 知识点掌握情况
    doc.add_heading('五、知识点掌握情况', level=1)
    
    # 按知识点分组
    category_stats = {}
    for q, info in questions.items():
        cat = info['category']
        if cat not in category_stats:
            category_stats[cat] = []
        correct_rate = (school_df[q] == 1).sum() / total_students * 100
        category_stats[cat].append((q, correct_rate))
    
    cat_table = doc.add_table(rows=1, cols=4)
    cat_table.style = 'Table Grid'
    cat_headers = ['知识点', '题号', '平均正确率', '掌握程度']
    for i, h in enumerate(cat_headers):
        cat_table.rows[0].cells[i].text = h
    
    for cat, q_rates in category_stats.items():
        avg_rate = np.mean([r[1] for r in q_rates])
        q_nums = ','.join([str(r[0]) for r in q_rates])
        
        if avg_rate >= 90:
            level = '优秀'
        elif avg_rate >= 70:
            level = '良好'
        elif avg_rate >= 60:
            level = '合格'
        else:
            level = '待加强'
        
        row = cat_table.add_row().cells
        row[0].text = cat
        row[1].text = q_nums
        row[2].text = f'{avg_rate:.2f}%'
        row[3].text = level
    
    # 改进措施
    doc.add_heading('六、改进措施与建议', level=1)
    
    measures = [
        '1. 针对错误率高的题目，安排专题复习课进行重点讲解。',
        '2. 建立学生错题本，要求学生整理错题并定期复习。',
        '3. 开展少先队知识竞赛，以赛促学，提高学习兴趣。',
        '4. 组织主题班会，强化少先队基础知识和礼仪规范。',
        '5. 家校联动，向家长反馈测试情况，共同督促学习。',
        '6. 定期测评，每月进行一次小测，跟踪学习效果。',
    ]
    
    for m in measures:
        doc.add_paragraph(m)
    
    # 页脚
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run(f'报告生成时间：2026 年 4 月 21 日\n').italic = True
    footer_para.add_run('数据分析师：小龙虾 AI 助手').italic = True
    
    # 保存 Word 文档
    word_file = output_dir / f'{school}_错题分析报告.docx'
    doc.save(word_file)
    print(f"✓ {school} 报告已保存：{word_file}")

print("\n===== 所有报告生成完成 =====")
print(f"输出目录：{output_dir}")
