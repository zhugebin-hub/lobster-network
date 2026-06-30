#!/usr/bin/env python3
"""生成计算机网络101计划培育团队对比分析Word文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement
import datetime

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_east_asia(run, font_name='宋体'):
    """为run设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_cell_font(cell, text, font_name='宋体', size=Pt(10), bold=False, align=None):
    """设置单元格文本和格式"""
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = size
    set_east_asia(run, font_name)
    run.font.bold = bold
    return run

# ========== 标题 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('计算机网络"101计划"培育团队对比分析')
run.font.size = Pt(22)
run.font.bold = True
set_east_asia(run, '黑体')

# 生成日期
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
date_run.font.size = Pt(10)
set_east_asia(date_run, '宋体')

doc.add_paragraph()  # 空行

# ========== 一、团队基本信息 ==========
h1 = doc.add_heading('一、团队基本信息', level=1)
for run in h1.runs:
    set_east_asia(run, '黑体')

# 团队基本信息表格
table1 = doc.add_table(rows=7, cols=3, style='Table Grid')
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
headers = ['对比维度', '诸葛斌团队', '陈翔团队']
for i, h in enumerate(headers):
    cell = table1.rows[0].cells[i]
    set_cell_font(cell, h, '黑体', Pt(11), True, WD_ALIGN_PARAGRAPH.CENTER)
    shading_el = OxmlElement('w:shd')
    shading_el.set(qn('w:fill'), 'D9E2F3')
    shading_el.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_el)

data1 = [
    ['对比维度', '诸葛斌团队', '陈翔团队'],
    ['所属单位', '绍兴文理学院', '未明确'],
    ['研究方向', '计算机网络、AI教育、云计算', '计算机网络、教学改革'],
    ['职称/职务', '副教授/硕士生导师', '未明确'],
    ['在线课程', 'MOOC + 省级一流实验课程', '省级一流课程'],
    ['合作范围', '覆盖22所高校', '未明确'],
    ['AI教育', 'AI原生教学范式、阿里云合作', 'AI教师培养项目'],
]

for row_idx, row_data in enumerate(data1, 0):
    for col_idx, text in enumerate(row_data):
        if row_idx == 0:
            continue  # 已处理表头
        cell = table1.rows[row_idx].cells[col_idx]
        set_cell_font(cell, text)

doc.add_paragraph()

# ========== 二、科研项目对比 ==========
h2 = doc.add_heading('二、科研项目对比', level=1)
for run in h2.runs:
    set_east_asia(run, '黑体')

p = doc.add_paragraph()
run = p.add_run('优势分析：')
run.font.bold = True
set_east_asia(run, '黑体')

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('诸葛斌团队：国家级项目经验丰富，参与973、863等重大科研项目，科研实力更强')
set_east_asia(run, '宋体')

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('陈翔团队：教改项目丰富，主持多项教育部协同育人项目，教学改革经验丰富')
set_east_asia(run, '宋体')

doc.add_paragraph()

# ========== 三、教学成果对比 ==========
h3 = doc.add_heading('三、教学成果对比', level=1)
for run in h3.runs:
    set_east_asia(run, '黑体')

p = doc.add_paragraph()
run = p.add_run('优势分析：')
run.font.bold = True
set_east_asia(run, '黑体')

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('诸葛斌团队：在线课程建设成果突出，MOOC + 省级一流实验课程，覆盖22所高校')
set_east_asia(run, '宋体')

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('陈翔团队：教学成果奖级别更高（省级一等奖），学科竞赛指导成果显著')
set_east_asia(run, '宋体')

doc.add_paragraph()

# ========== 四、学科竞赛与指导学生 ==========
h4 = doc.add_heading('四、学科竞赛与指导学生', level=1)
for run in h4.runs:
    set_east_asia(run, '黑体')

p = doc.add_paragraph()
run = p.add_run('优势分析：')
run.font.bold = True
set_east_asia(run, '黑体')

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('陈翔团队：学科竞赛指导经验丰富，成果显著，连续10余年获奖')
set_east_asia(run, '宋体')

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('诸葛斌团队：研究生培养体系完善，有硕士生导师资格')
set_east_asia(run, '宋体')

doc.add_paragraph()

# ========== 五、AI教育应用对比 ==========
h5 = doc.add_heading('五、AI教育应用对比', level=1)
for run in h5.runs:
    set_east_asia(run, '黑体')

p = doc.add_paragraph()
run = p.add_run('优势分析：')
run.font.bold = True
set_east_asia(run, '黑体')

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('诸葛斌团队：AI教育应用走在前列，提出AI原生教学范式，与阿里云深度合作')
set_east_asia(run, '宋体')

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('陈翔团队：有AI教师培养项目经验，但在AI教学应用方面成果较少')
set_east_asia(run, '宋体')

doc.add_paragraph()

# ========== 六、综合对比总结 ==========
h6 = doc.add_heading('六、综合对比总结', level=1)
for run in h6.runs:
    set_east_asia(run, '黑体')

table2 = doc.add_table(rows=7, cols=3, style='Table Grid')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

summary_headers = ['对比维度', '诸葛斌团队', '陈翔团队']
for i, h in enumerate(summary_headers):
    cell = table2.rows[0].cells[i]
    set_cell_font(cell, h, '黑体', Pt(11), True, WD_ALIGN_PARAGRAPH.CENTER)
    shading_el = OxmlElement('w:shd')
    shading_el.set(qn('w:fill'), 'D9E2F3')
    shading_el.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_el)

summary_data = [
    ['科研实力', '★★★★★ (国家级项目)', '★★★☆☆ (教改项目为主)'],
    ['教学成果', '★★★★☆ (MOOC+一流课程)', '★★★★★ (省级一等奖)'],
    ['在线课程', '★★★★★ (覆盖22所高校)', '★★★☆☆'],
    ['学科竞赛', '★★★☆☆', '★★★★★ (10余年获奖)'],
    ['AI教育', '★★★★★ (AI原生范式)', '★★★☆☆'],
    ['综合评价', '科研+AI驱动型', '教学+竞赛驱动型'],
]

for row_idx, row_data in enumerate(summary_data, 1):
    for col_idx, text in enumerate(row_data):
        cell = table2.rows[row_idx].cells[col_idx]
        set_cell_font(cell, text)

doc.add_paragraph()

# ========== 七、合作建议 ==========
h7 = doc.add_heading('七、合作建议', level=1)
for run in h7.runs:
    set_east_asia(run, '黑体')

p = doc.add_paragraph()
run = p.add_run('互补优势：')
run.font.bold = True
set_east_asia(run, '黑体')

items = [
    '科研 + 教学：诸葛斌团队科研实力强，陈翔团队教学成果优',
    '在线 + 竞赛：诸葛斌团队在线课程丰富，陈翔团队竞赛指导强',
    'AI + 传统：诸葛斌团队AI教育领先，陈翔团队传统教学扎实',
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    set_east_asia(run, '宋体')

p = doc.add_paragraph()
run = p.add_run('合作方向：')
run.font.bold = True
set_east_asia(run, '黑体')

directions = [
    '课程共建：联合开发计算机网络课程，共享MOOC资源',
    '竞赛指导：诸葛斌团队提供AI技术支持，陈翔团队提供竞赛指导经验',
    '教材编写：联合编写AI+计算机网络教材，发挥各自优势',
    '教改研究：共同申报国家级教改项目，结合AI教育与传统教学',
]
for d in directions:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(d)
    set_east_asia(run, '宋体')

doc.add_paragraph()

# ========== 结论 ==========
h8 = doc.add_heading('结论', level=1)
for run in h8.runs:
    set_east_asia(run, '黑体')

p = doc.add_paragraph()
run = p.add_run('两个团队各有优势，诸葛斌团队在科研实力、在线课程、AI教育应用方面领先；陈翔团队在教学成果、学科竞赛指导方面突出。建议加强合作，优势互补，共同打造计算机网络"101计划"标杆课程。')
set_east_asia(run, '宋体')
run.font.size = Pt(12)

# 保存
output_path = '/home/admin/.openclaw/workspace/计算机网络101计划培育团队对比分析.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
